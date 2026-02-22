from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.button import Button
from kivy.uix.screenmanager import ScreenManager, Screen
from kivy.uix.scrollview import ScrollView
from kivy.uix.gridlayout import GridLayout
from kivy.uix.textinput import TextInput
from kivy.uix.popup import Popup
from kivy.uix.label import Label
from kivy.uix.floatlayout import FloatLayout
from kivy.uix.widget import Widget
from kivy.properties import (NumericProperty, BooleanProperty, 
                          ObjectProperty, ListProperty, StringProperty)
from kivy.core.window import Window
from kivy.graphics import Color, Rectangle, Line, RoundedRectangle
from kivy.clock import Clock
from kivy.animation import Animation
from kivy.core.text import LabelBase
from kivy.utils import get_color_from_hex
from kivy.config import Config
from kivy.core.image import Image as CoreImage
import sqlite3
import pandas as pd
import os
import datetime
import re
import json
import shutil
import glob
from openpyxl import load_workbook, Workbook
from tkinter import Tk, filedialog
from molodniki_extended import ExtendedMolodnikiTableScreen
from new_taxation_menu import TaxationPopup
import statistics

# Monkey patch to fix Kivy TextInput mode issue
import kivy.uix.textinput

def patched_insert_text(self, substring, from_undo=False):
    if self.readonly:
        return
    if from_undo:
        self._undo.append((self.cursor, self.text))
    # Ensure substring is properly handled without calling mode
    if not substring:
        return
    cc, cr = self.cursor
    ci = self.cursor_index()
    text = self.text
    len_str = len(substring)
    if not from_undo:
        self._undo.append([substring, (cc, cr), (cc + len_str, cr), text, -1])
    self.text = text[:ci] + substring + text[ci:]
    self.cursor = self.get_cursor_from_index(ci + len_str)
    self._refresh_text(self.text)
    self._trigger_update_graphics()

kivy.uix.textinput.TextInput.insert_text = patched_insert_text

LabelBase.register(name='Roboto', 
                 fn_regular='fonts/Roboto-Medium.ttf',
                 fn_bold='fonts/Roboto-Bold.ttf')

class ThemeManager:
    def __init__(self):
        self.themes = []
        self.current_theme_index = 0
        self.themes_dir = 'themes'
        os.makedirs(self.themes_dir, exist_ok=True)
        self.load_themes()
        self.load_config()
        
    def load_themes(self):
        self.themes = []
        self.themes.extend([
            {
                'type': 'color',
                'name': 'light',
                'background': get_color_from_hex('#FEF7FF'),
                'text_color': '#1C1B1F'
            },
            {
                'type': 'color',
                'name': 'dark',
                'background': get_color_from_hex('#1C1B1F'),
                'text_color': '#E6E1E5'
            }
        ])
        
        for file in os.listdir(self.themes_dir):
            if file.lower().endswith(('.png', '.jpg', '.jpeg')):
                self.themes.append({
                    'type': 'image',
                    'name': os.path.splitext(file)[0],
                    'background': os.path.join(self.themes_dir, file),
                    'text_color': '#FFFFFF'
                })
    
    def add_theme(self, image_path):
        dest = os.path.join(self.themes_dir, os.path.basename(image_path))
        shutil.copy(image_path, dest)
        self.load_themes()
    
    @property
    def current_theme(self):
        return self.themes[self.current_theme_index]
    
    def save_config(self):
        config = {
            'theme_index': self.current_theme_index,
            'themes_dir': self.themes_dir
        }
        with open('theme_config.json', 'w') as f:
            json.dump(config, f)
    
    def load_config(self):
        try:
            with open('theme_config.json', 'r') as f:
                config = json.load(f)
                self.current_theme_index = config.get('theme_index', 0)
                self.themes_dir = config.get('themes_dir', 'themes')
        except (FileNotFoundError, json.JSONDecodeError):
            self.current_theme_index = 0

class ModernButton(Button):
    bg_color = ListProperty([1, 1, 1, 1])
    no_shadow = BooleanProperty(False)

    def __init__(self, **kwargs):
        self.no_shadow = kwargs.pop('no_shadow', False)
        super().__init__(**kwargs)
        self.background_color = (0, 0, 0, 0)
        self.font_name = 'Roboto'
        self.font_size = '16sp'
        self.bold = False
        self.size_hint = (None, None)
        self.height = 50
        self.padding = (20, 5)

        with self.canvas.before:
            if not self.no_shadow:
                Color(rgba=(0, 0, 0, 0.1))
                self.shadow = RoundedRectangle(
                    pos=(self.x+3, self.y-3),
                    size=self.size,
                    radius=[10]
                )
            self.bg_color_instruction = Color(rgba=self.bg_color)
            self.background = RoundedRectangle(
                pos=self.pos,
                size=self.size,
                radius=[10]
            )

        self.bind(pos=self.update_graphics, size=self.update_graphics)
        self.bind(text=self.update_width)

    def update_width(self, instance, value):
        self.width = self.texture_size[0] + 60

    def update_graphics(self, *args):
        self.background.pos = self.pos
        if not self.no_shadow:
            self.shadow.pos = (self.x+3, self.y-3)
            self.shadow.size = self.size
        self.background.size = self.size

    def on_touch_down(self, touch):
        if self.collide_point(*touch.pos):
            Animation(rgba=[c * 0.9 for c in self.bg_color], d=0.1).start(self.bg_color_instruction)
        return super().on_touch_down(touch)

    def on_touch_up(self, touch):
        Animation(rgba=self.bg_color, d=0.2).start(self.bg_color_instruction)
        return super().on_touch_up(touch)

class AutoCompleteTextInput(TextInput):
    next_widget = ObjectProperty(None)
    prev_widget = ObjectProperty(None)
    row_index = NumericProperty(0)
    col_index = NumericProperty(0)
    suggestions = ListProperty([])

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.bind(text=self.show_suggestions)
        self.popup = None

    def keyboard_on_key_down(self, window, keycode, text, modifiers):
        key = keycode[1]
        if key == 'down':
            self.focus_next('down')
        elif key == 'up':
            self.focus_previous('up')
        elif key == 'right':
            self.focus_next('right')
        elif key == 'left':
            self.focus_previous('left')
        else:
            super().keyboard_on_key_down(window, keycode, text, modifiers)
        return True

    def show_suggestions(self, instance, value):
        if not value or len(value) < 3:
            return

        conn = sqlite3.connect('forest_data.db')
        cursor = conn.cursor()
        cursor.execute('''
            SELECT value FROM suggestions
            WHERE column_index = ? AND value LIKE ?
            ORDER BY LENGTH(value) ASC, value ASC
            LIMIT 1
        ''', (self.col_index, f'{value}%'))
        results = cursor.fetchall()
        conn.close()

        if results:
            self.text = results[0][0]

    def get_table_screen(self):
        return App.get_running_app().root.get_screen('table')

    def focus_next(self, direction):
        table_screen = self.get_table_screen()
        if direction == 'right' and self.next_widget:
            self.next_widget.focus = True
        elif direction == 'down':
            next_row = self.row_index + 1
            if next_row < len(table_screen.inputs):
                table_screen.inputs[next_row][self.col_index].focus = True

    def focus_previous(self, direction):
        table_screen = self.get_table_screen()
        if direction == 'left' and self.prev_widget:
            self.prev_widget.focus = True
        elif direction == 'up':
            prev_row = self.row_index - 1
            if prev_row >= 0:
                table_screen.inputs[prev_row][self.col_index].focus = True

class TreeDataInputPopup(Popup):
    def __init__(self, table_screen, row_index, **kwargs):
        super().__init__(
            title='Ввод данных дерева',
            size_hint=(0.7, 0.8),
            separator_height=0,
            background_color=(0.5, 0.5, 0.5, 1),  # Серый фон
            overlay_color=(0, 0, 0, 0.5),
            **kwargs
        )
        self.table_screen = table_screen
        self.row_index = row_index
        self.fields = [
            ('Порода', 1),
            ('ж/ф', 2),
            ('шт/либо лет', 3),
            ('D, см', 4),
            ('H, м', 5),
            ('Сост-е', 6),
            ('Модель', 7),
            ('Примечания', 8)
        ]
        self.data = {}
        self.create_ui()

    def create_ui(self):
        content = FloatLayout()

        label = Label(
            text='Введите данные дерева:',
            font_name='Roboto',
            font_size='18sp',
            color=(1, 1, 1, 1),
            pos_hint={'center_x': 0.5, 'center_y': 0.9},
            size_hint=(None, None),
            size=(250, 50)
        )

        scroll = ScrollView(size_hint=(0.9, 0.7), pos_hint={'center_x': 0.5, 'center_y': 0.5})
        layout = GridLayout(cols=1, spacing=10, size_hint_y=None)
        layout.bind(minimum_height=layout.setter('height'))

        with layout.canvas.before:
            Color(rgba=(0, 0, 0, 0))

        self.input_fields = []
        for field_name, col_index in self.fields:
            field_layout = BoxLayout(orientation='vertical', size_hint_y=None, height=70, spacing=5)
            field_label = Label(
                text=field_name,
                font_name='Roboto',
                font_size='16sp',
                color=(1, 1, 1, 1),
                size_hint_y=None,
                height=20
            )
            input_field = AutoCompleteTextInput(
                multiline=False,
                size_hint_y=None,
                height=40,
                background_color=(1, 1, 1, 0.8),
                col_index=col_index,
                font_name='Roboto'
            )
            self.input_fields.append(input_field)
            field_layout.add_widget(field_label)
            field_layout.add_widget(input_field)
            layout.add_widget(field_layout)

        scroll.add_widget(layout)

        btn_box = BoxLayout(
            orientation='horizontal',
            spacing=10,
            size_hint=(1, None),
            height=40,
            pos_hint={'center_x': 0.5, 'center_y': 0.1}
        )
        save_btn = ModernButton(
            text='Сохранить',
            bg_color=get_color_from_hex('#00FF00'),
            size_hint=(0.5, None),
            height=40,
            no_shadow=True
        )
        save_btn.bind(on_press=self.save_data)
        exit_btn = ModernButton(
            text='Выйти',
            bg_color=get_color_from_hex('#FF0000'),
            size_hint=(0.5, None),
            height=40,
            no_shadow=True
        )
        exit_btn.bind(on_press=self.dismiss)
        btn_box.add_widget(save_btn)
        btn_box.add_widget(exit_btn)

        content.add_widget(label)
        content.add_widget(scroll)
        content.add_widget(btn_box)

        self.content = content
        self.open()

    def save_to_suggestions(self, col_index, value):
        try:
            conn = sqlite3.connect('forest_data.db')
            cursor = conn.cursor()
            cursor.execute('''
                INSERT OR IGNORE INTO suggestions (column_index, value)
                VALUES (?, ?)
            ''', (col_index, value))
            conn.commit()
            conn.close()
        except Exception as e:
            print(f"Error saving suggestion: {e}")

    def save_data(self, instance):
        for i, (field_name, col_index) in enumerate(self.fields):
            value = self.input_fields[i].text.strip()
            if value:
                self.data[col_index] = value
                # Save to suggestions
                self.save_to_suggestions(col_index, value)

        # Fill the row in the table
        for col_index, value in self.data.items():
            if col_index < len(self.table_screen.inputs[self.row_index]):
                self.table_screen.inputs[self.row_index][col_index].text = value

        # Save to page_data
        self.table_screen.save_current_page()

        # Auto-fill next tree numbers
        base_number = self.table_screen.current_page * self.table_screen.rows_per_page + self.row_index + 1
        for row_idx in range(self.row_index + 1, len(self.table_screen.inputs)):
            tree_number = base_number + (row_idx - self.row_index)
            self.table_screen.inputs[row_idx][0].text = str(tree_number)

        # Show success
        self.table_screen.show_success("Данные дерева сохранены!")
        self.dismiss()


class ExitConfirmPopup(Popup):
    def __init__(self, **kwargs):
        super().__init__(
            title='',
            separator_height=0,
            size_hint=(0.4, 0.3),
            background='atlas://data/images/defaulttheme/modalview-background',
            overlay_color=(0, 0, 0, 0.5)
        )

        content = FloatLayout()

        close_btn = ModernButton(
            text='X',
            size_hint=(None, None),
            size=(40, 40),
            pos_hint={'right': 0.95, 'top': 0.95},
            bg_color=(1, 0, 0, 1),
            no_shadow=True
        )
        close_btn.bind(on_press=self.dismiss)

        label = Label(
            text='Вы уверены, что хотите выйти?',
            font_name='Roboto',
            font_size='18sp',
            color=(0.2, 0.2, 0.2, 1),
            pos_hint={'center_x': 0.5, 'center_y': 0.65},
            size_hint=(None, None),
            size=(250, 50)
        )

        btn_box = BoxLayout(
            orientation='horizontal',
            spacing=15,
            size_hint=(None, None),
            size=(200, 45),
            pos_hint={'center_x': 0.5, 'center_y': 0.25}
        )
        yes_btn = ModernButton(
            text='Выход',
            bg_color=get_color_from_hex('#FF0000'),
            color=get_color_from_hex('#000000'),
            size_hint=(0.5, None),
            height=45,
            no_shadow=True
        )
        no_btn = ModernButton(
            text='Отмена',
            bg_color=get_color_from_hex('#00FF00'),
            color=get_color_from_hex('#000000'),
            size_hint=(0.5, None),
            height=45,
            no_shadow=True
        )
        yes_btn.bind(on_press=lambda x: App.get_running_app().stop())
        no_btn.bind(on_press=self.dismiss)

        btn_box.add_widget(yes_btn)
        btn_box.add_widget(no_btn)

        content.add_widget(close_btn)
        content.add_widget(label)
        content.add_widget(btn_box)

        self.content = content

class MainMenu(Screen):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.bg_image = None
        self.bg_rect = None
        self.setup_ui()
        
    def setup_ui(self):
        self.clear_widgets()
        theme = App.get_running_app().theme_manager.current_theme
        
        with self.canvas.before:
            if theme['type'] == 'color':
                Color(rgba=theme['background'])
                self.bg_rect = Rectangle(pos=self.pos, size=self.size)
            else:
                self.bg_image = CoreImage(theme['background']).texture
                self.bg_rect = Rectangle(
                    texture=self.bg_image,
                    pos=self.pos,
                    size=self.size
                )
            
        self.bind(pos=self.update_bg, size=self.update_bg)
        
        main_layout = BoxLayout(orientation='vertical', padding=[50, 20, 20, 20])

        center_layout = BoxLayout(
            orientation='vertical',
            size_hint=(0.8, None),
            height=400,
            pos_hint={'center_x': 0.5, 'center_y': 0.75},
            spacing=15
        )
        
        title = Label(
            text='Фанаты Пихты',
            font_size='24sp',
            font_name='Roboto',
            size_hint_y=None,
            height=60
        )
        center_layout.add_widget(title)
        
        buttons = [
            ('Перечётная ведомость', '#FFA500', self.show_add_section),
            ('РУМ (Молодняки)', '#00BFFF', self.show_add_molodniki_section),
            ('Темы', '#FFFF00', self.show_theme_chooser),
            ('Выход', '#FF0000', self.confirm_exit)
        ]
        
        for text, color, callback in buttons:
            btn = ModernButton(
                text=text,
                bg_color=get_color_from_hex(color),
                color=get_color_from_hex('#000000'),
                size_hint=(None, None),
                size=(250, 50),
                pos_hint={'center_x': 0.5}
            )
            btn.bind(on_press=callback)
            center_layout.add_widget(btn)
        
        main_layout.add_widget(center_layout)
        
        footer = Label(
            text='by forestcrypt®',
            font_size='12sp',
            color=(1, 1, 1, 1),
            size_hint=(1, None),
            height=30,
            pos_hint={'right': 0.98, 'y': 0.02}
        )
        main_layout.add_widget(footer)
        
        self.add_widget(main_layout)
        
    def update_bg(self, *args):
        if self.bg_rect:
            self.bg_rect.pos = self.pos
            self.bg_rect.size = self.size
            
    def change_theme(self, instance):
        themes = App.get_running_app().theme_manager.themes
        current_idx = App.get_running_app().theme_manager.current_theme_index
        new_idx = (current_idx + 1) % len(themes)
        App.get_running_app().theme_manager.current_theme_index = new_idx
        App.get_running_app().theme_manager.save_config()
        App.get_running_app().reload_theme()

    def show_theme_chooser(self, instance):
        ThemeChooser().open()

    def show_add_section(self, instance):
        content = FloatLayout()

        close_btn = ModernButton(
            text='X',
            size_hint=(None, None),
            size=(40, 40),
            pos_hint={'right': 0.95, 'top': 0.95},
            bg_color=(1, 0, 0, 1),
            no_shadow=True
        )
        close_btn.bind(on_press=lambda x: self.section_popup.dismiss())

        # Поля для адреса
        scroll = ScrollView(size_hint=(0.9, 0.8), pos_hint={'center_x': 0.5, 'center_y': 0.5})
        layout = BoxLayout(orientation='vertical', spacing=10, size_hint_y=None, padding=[10, 10])
        layout.bind(minimum_height=layout.setter('height'))

        # Номер участка
        section_layout = BoxLayout(orientation='vertical', size_hint_y=None, height=70, spacing=5)
        section_label = Label(
            text='Номер участка:',
            font_name='Roboto',
            font_size='16sp',
            color=(0.2, 0.2, 0.2, 1),
            size_hint_y=None,
            height=25
        )
        self.section_number_input = TextInput(
            hint_text="Введите номер участка",
            multiline=False,
            size_hint_y=None,
            height=40,
            background_color=(1, 1, 1, 0.8),
            font_name='Roboto'
        )
        section_layout.add_widget(section_label)
        section_layout.add_widget(self.section_number_input)
        layout.add_widget(section_layout)

        # Квартал
        quarter_layout = BoxLayout(orientation='vertical', size_hint_y=None, height=70, spacing=5)
        quarter_label = Label(
            text='Квартал:',
            font_name='Roboto',
            font_size='16sp',
            color=(0.2, 0.2, 0.2, 1),
            size_hint_y=None,
            height=25
        )
        self.quarter_input = TextInput(
            hint_text="Введите номер квартала",
            multiline=False,
            size_hint_y=None,
            height=40,
            background_color=(1, 1, 1, 0.8),
            font_name='Roboto'
        )
        quarter_layout.add_widget(quarter_label)
        quarter_layout.add_widget(self.quarter_input)
        layout.add_widget(quarter_layout)

        # Выдел
        plot_layout = BoxLayout(orientation='vertical', size_hint_y=None, height=70, spacing=5)
        plot_label = Label(
            text='Выдел:',
            font_name='Roboto',
            font_size='16sp',
            color=(0.2, 0.2, 0.2, 1),
            size_hint_y=None,
            height=25
        )
        self.plot_input = TextInput(
            hint_text="Введите номер выдела",
            multiline=False,
            size_hint_y=None,
            height=40,
            background_color=(1, 1, 1, 0.8),
            font_name='Roboto'
        )
        plot_layout.add_widget(plot_label)
        plot_layout.add_widget(self.plot_input)
        layout.add_widget(plot_layout)

        # Лесничество
        forestry_layout = BoxLayout(orientation='vertical', size_hint_y=None, height=70, spacing=5)
        forestry_label = Label(
            text='Лесничество:',
            font_name='Roboto',
            font_size='16sp',
            color=(0.2, 0.2, 0.2, 1),
            size_hint_y=None,
            height=25
        )
        self.forestry_input = TextInput(
            hint_text="Введите название лесничества",
            multiline=False,
            size_hint_y=None,
            height=40,
            background_color=(1, 1, 1, 0.8),
            font_name='Roboto'
        )
        forestry_layout.add_widget(forestry_label)
        forestry_layout.add_widget(self.forestry_input)
        layout.add_widget(forestry_layout)

        # Участковое лесничество
        district_forestry_layout = BoxLayout(orientation='vertical', size_hint_y=None, height=70, spacing=5)
        district_forestry_label = Label(
            text='Участковое лесничество:',
            font_name='Roboto',
            font_size='16sp',
            color=(0.2, 0.2, 0.2, 1),
            size_hint_y=None,
            height=25
        )
        self.district_forestry_input = TextInput(
            hint_text="Введите название участкового лесничества",
            multiline=False,
            size_hint_y=None,
            height=40,
            background_color=(1, 1, 1, 0.8),
            font_name='Roboto'
        )
        district_forestry_layout.add_widget(district_forestry_label)
        district_forestry_layout.add_widget(self.district_forestry_input)
        layout.add_widget(district_forestry_layout)

        scroll.add_widget(layout)

        btn_box = BoxLayout(
            orientation='horizontal',
            spacing=10,
            size_hint=(1, None),
            height=50,
            pos_hint={'center_x': 0.5, 'center_y': 0.1}
        )
        buttons = [
            ('Сохранить', '#00FF00', self.save_section),
            ('Загрузить', '#0000FF', self.show_load_popup),
            ('Отмена', '#FF0000', lambda x: self.section_popup.dismiss())
        ]

        for text, color, callback in buttons:
            btn = ModernButton(
                text=text,
                bg_color=get_color_from_hex(color),
                size_hint=(0.33, None),
                height=45,
                no_shadow=False
            )
            btn.bind(on_press=callback)
            btn_box.add_widget(btn)

        content.add_widget(close_btn)
        content.add_widget(scroll)
        content.add_widget(btn_box)

        self.section_popup = Popup(
            title="Управление участками",
            content=content,
            size_hint=(0.7, 0.8),
            separator_height=0,
            background='atlas://data/images/defaulttheme/modalview-background',
            overlay_color=(0, 0, 0, 0.5)
        )
        self.section_popup.open()

    def add_section(self, instance):
        try:
            conn = sqlite3.connect('forest_data.db')
            cursor = conn.cursor()
            cursor.execute('INSERT INTO sections DEFAULT VALUES')
            conn.commit()
            self.show_success("Новый участок создан!")
        except Exception as e:
            self.show_error(f"Ошибка: {str(e)}")
        finally:
            conn.close()

    def save_section(self, instance):
        section_number = self.section_number_input.text.strip()
        quarter = self.quarter_input.text.strip()
        plot = self.plot_input.text.strip()
        forestry = self.forestry_input.text.strip()
        district_forestry = self.district_forestry_input.text.strip()

        if not section_number:
            self.show_error("Введите номер участка!")
            return

        try:
            conn = sqlite3.connect('forest_data.db')
            cursor = conn.cursor()
            cursor.execute('''
                UPDATE sections
                SET section_number = ?, quarter = ?, plot = ?, forestry = ?, district_forestry = ?
                WHERE id = (SELECT MAX(id) FROM sections)
            ''', (section_number, quarter, plot, forestry, district_forestry))
            conn.commit()
            self.section_popup.dismiss()
            table_screen = App.get_running_app().root.get_screen('table')
            table_screen.current_section = section_number
            # Можно добавить хранение адреса в table_screen если нужно
            table_screen.update_section_label()
            App.get_running_app().root.current = 'table'
        except Exception as e:
            self.show_error(f"Ошибка сохранения: {str(e)}")
        finally:
            conn.close()

    def show_add_molodniki_section(self, instance):
        # Сначала создаем новую запись в таблице
        try:
            conn = sqlite3.connect('forest_data.db')
            cursor = conn.cursor()
            cursor.execute('INSERT INTO molodniki_sections DEFAULT VALUES')
            conn.commit()
        except Exception as e:
            pass  # Игнорируем ошибку, если таблица не существует
        finally:
            conn.close()

        content = FloatLayout()

        close_btn = ModernButton(
            text='X',
            size_hint=(None, None),
            size=(40, 40),
            pos_hint={'right': 0.95, 'top': 0.95},
            bg_color=(1, 0, 0, 1),
            no_shadow=True
        )
        close_btn.bind(on_press=lambda x: self.molodniki_popup.dismiss())

        label = Label(
            text='Введите номер участка молодняков:',
            font_name='Roboto',
            font_size='18sp',
            color=(0.2, 0.2, 0.2, 1),
            pos_hint={'center_x': 0.5, 'center_y': 0.7},
            size_hint=(None, None),
            size=(300, 50)
        )

        self.molodniki_section_input = TextInput(
            hint_text="Введите номер участка молодняков",
            multiline=False,
            size_hint=(None, None),
            size=(300, 40),
            pos_hint={'center_x': 0.5, 'center_y': 0.55},
            background_color=(1, 1, 1, 0.8)
        )

        btn_box = BoxLayout(
            orientation='horizontal',
            spacing=10,
            size_hint=(None, None),
            size=(300, 50),
            pos_hint={'center_x': 0.5, 'center_y': 0.3}
        )
        buttons = [
            ('Загрузить', '#00FF00', self.load_molodniki_json),
            ('Сохранить', '#0000FF', self.save_molodniki_section),
            ('Отмена', '#FF0000', lambda x: self.molodniki_popup.dismiss())
        ]

        for text, color, callback in buttons:
            btn = ModernButton(
                text=text,
                bg_color=get_color_from_hex(color),
                size_hint=(0.33, None),
                height=45,
                no_shadow=False
            )
            btn.bind(on_press=callback)
            btn_box.add_widget(btn)

        content.add_widget(close_btn)
        content.add_widget(label)
        content.add_widget(self.molodniki_section_input)
        content.add_widget(btn_box)

        self.molodniki_popup = Popup(
            title="Управление участками молодняков",
            content=content,
            size_hint=(0.6, 0.5),
            separator_height=0,
            background='atlas://data/images/defaulttheme/modalview-background',
            overlay_color=(0, 0, 0, 0.5)
        )
        self.molodniki_popup.open()

    def load_molodniki_json(self, instance):
        self.molodniki_popup.dismiss()
        App.get_running_app().root.current = 'molodniki'
        Clock.schedule_once(lambda dt: self.show_molodniki_json_popup(), 0.1)

    def show_molodniki_json_popup(self):
        molodniki_screen = App.get_running_app().root.get_screen('molodniki')
        molodniki_screen.load_section_popup()

    def add_molodniki_section(self, instance):
        try:
            conn = sqlite3.connect('forest_data.db')
            cursor = conn.cursor()
            cursor.execute('INSERT INTO molodniki_sections DEFAULT VALUES')
            conn.commit()
            self.show_success("Новый участок молодняков создан!")
        except Exception as e:
            self.show_error(f"Ошибка: {str(e)}")
        finally:
            conn.close()

    def save_molodniki_section(self, instance):
        section_number = self.molodniki_section_input.text.strip()
        if not section_number:
            self.show_error("Введите номер участка!")
            return

        try:
            conn = sqlite3.connect('forest_data.db')
            cursor = conn.cursor()
            # Используем INSERT OR REPLACE, чтобы создать или обновить запись с данным номером участка
            cursor.execute('''
                INSERT OR REPLACE INTO molodniki_sections (id, section_number, created_at)
                VALUES ((SELECT id FROM molodniki_sections WHERE section_number = ?), ?, CURRENT_TIMESTAMP)
            ''', (section_number, section_number))
            conn.commit()
            self.molodniki_popup.dismiss()
            App.get_running_app().root.get_screen('molodniki').current_section = section_number
            # Убираем вызов update_section_label, так как section_label больше не существует
            App.get_running_app().root.current = 'molodniki'
        except Exception as e:
            self.show_error(f"Ошибка сохранения: {str(e)}")
        finally:
            conn.close()

    def confirm_exit(self, instance):
        ExitConfirmPopup().open()

    def show_taxation_popup(self, instance):
        """Открыть popup таксационных показателей"""
        TaxationPopup().open()

    def show_success(self, message):
        content = FloatLayout()
        label = Label(
            text=message,
            color=(0, 0.5, 0, 1),
            font_name='Roboto',
            pos_hint={'center_x': 0.5, 'center_y': 0.5}
        )
        content.add_widget(label)
        Popup(
            title='Успешно',
            content=content,
            size_hint=(0.6, 0.3),
            background='atlas://data/images/defaulttheme/modalview-background',
            overlay_color=(0, 0, 0, 0.5)
        ).open()

    def show_error(self, message):
        content = FloatLayout()
        label = Label(
            text=message,
            color=(1, 0, 0, 1),
            font_name='Roboto',
            pos_hint={'center_x': 0.5, 'center_y': 0.5}
        )
        content.add_widget(label)
        Popup(
            title='Ошибка',
            content=content,
            size_hint=(0.6, 0.3),
            background='atlas://data/images/defaulttheme/modalview-background',
            overlay_color=(0, 0, 0, 0.5)
        ).open()

    def show_load_popup(self, instance):
        conn = sqlite3.connect('forest_data.db')
        cursor = conn.cursor()
        cursor.execute('SELECT section_number FROM sections WHERE section_number IS NOT NULL AND section_number != "" ORDER BY id DESC')
        sections = cursor.fetchall()
        conn.close()
        if not sections:
            self.show_error("Нет сохраненных участков!")
            return
        content = FloatLayout()

        close_btn = ModernButton(
            text='X',
            size_hint=(None, None),
            size=(40, 40),
            pos_hint={'right': 0.95, 'top': 0.95},
            bg_color=(1, 0, 0, 1),
            no_shadow=True
        )
        close_btn.bind(on_press=lambda x: self.load_popup.dismiss())

        scroll = ScrollView(size_hint=(1, 0.9), pos_hint={'center_x': 0.5, 'center_y': 0.45})
        layout = GridLayout(cols=1, spacing=5, size_hint_y=None)
        layout.bind(minimum_height=layout.setter('height'))
        for section in sections:
            btn = ModernButton(
                text=section[0],
                size_hint_y=None,
                height=40,
                bg_color=(0, 1, 0, 1),
                color=(0, 0, 0, 1),
                no_shadow=True
            )
            btn.bind(on_release=lambda b, s=section[0]: self.load_saved_section(s))
            layout.add_widget(btn)
        scroll.add_widget(layout)

        content.add_widget(close_btn)
        content.add_widget(scroll)
        self.load_popup = Popup(
            title="Выберите участок",
            content=content,
            size_hint=(0.5, 0.6),
            separator_height=0,
            background_color=(0.5, 0.5, 0.5, 1),
            overlay_color=(0, 0, 0, 0.5)
        )
        self.load_popup.open()

    def load_saved_section(self, section_number):
        files = glob.glob(os.path.join(App.get_running_app().root.get_screen('table').reports_dir, f"{section_number}_*.xlsx"))
        if files:
            latest_file = max(files, key=os.path.getctime)
            table_screen = App.get_running_app().root.get_screen('table')
            try:
                df = pd.read_excel(latest_file)
                data = df.values.tolist()
                
                table_screen.current_section = section_number
                table_screen.update_section_label()
                table_screen.page_data.clear()
                
                for page_num in range(0, len(df), table_screen.rows_per_page):
                    page = page_num // table_screen.rows_per_page
                    page_data = df.iloc[page_num:page_num+table_screen.rows_per_page].values.tolist()
                    table_screen.page_data[page] = page_data
                
                table_screen.current_page = 0
                table_screen.load_page_data()
                table_screen.update_pagination()
                self.show_success("Данные участка загружены!")
                App.get_running_app().root.current = 'table'
            except Exception as e:
                self.show_error(f"Ошибка загрузки: {str(e)}")
        else:
            self.show_error("Файл участка не найден!")
        self.load_popup.dismiss()

    def show_load_molodniki_popup(self, instance):
        conn = sqlite3.connect('forest_data.db')
        cursor = conn.cursor()
        cursor.execute('SELECT section_number FROM molodniki_sections WHERE section_number IS NOT NULL AND section_number != "" ORDER BY id DESC')
        sections = cursor.fetchall()
        conn.close()
        if not sections:
            self.show_error("Нет сохраненных участков молодняков!")
            return
        content = FloatLayout()

        close_btn = ModernButton(
            text='X',
            size_hint=(None, None),
            size=(40, 40),
            pos_hint={'right': 0.95, 'top': 0.95},
            bg_color=(1, 0, 0, 1),
            no_shadow=True
        )
        close_btn.bind(on_press=lambda x: self.load_molodniki_popup.dismiss())

        scroll = ScrollView(size_hint=(1, 0.9), pos_hint={'center_x': 0.5, 'center_y': 0.45})
        layout = GridLayout(cols=1, spacing=5, size_hint_y=None)
        layout.bind(minimum_height=layout.setter('height'))
        for section in sections:
            btn = ModernButton(
                text=section[0],
                size_hint_y=None,
                height=40,
                bg_color=(0, 1, 0, 1),
                color=(0, 0, 0, 1),
                no_shadow=True
            )
            btn.bind(on_release=lambda b, s=section[0]: self.load_saved_molodniki_section(s))
            layout.add_widget(btn)
        scroll.add_widget(layout)

        content.add_widget(close_btn)
        content.add_widget(scroll)
        self.load_molodniki_popup = Popup(
            title="Выберите участок молодняков",
            content=content,
            size_hint=(0.5, 0.6),
            separator_height=0,
            background_color=(0.5, 0.5, 0.5, 1),
            overlay_color=(0, 0, 0, 0.5)
        )
        self.load_molodniki_popup.open()

    def load_saved_molodniki_section(self, section_number):
        molodniki_screen = App.get_running_app().root.get_screen('molodniki')
        files = glob.glob(os.path.join(molodniki_screen.reports_dir, f"Молодняки_расширенный_{section_number}_*.xlsx"))
        if files:
            latest_file = max(files, key=os.path.getctime)
            try:
                df = pd.read_excel(latest_file)

                molodniki_screen.current_section = section_number
                molodniki_screen.update_section_label()
                molodniki_screen.page_data.clear()

                for page_num in range(0, len(df), molodniki_screen.rows_per_page):
                    page = page_num // molodniki_screen.rows_per_page
                    page_data = df.iloc[page_num:page_num+molodniki_screen.rows_per_page].values.tolist()
                    # Дополняем до 29 столбцов если нужно
                    for row in page_data:
                        while len(row) < 29:
                            row.append('')
                    molodniki_screen.page_data[page] = page_data

                molodniki_screen.current_page = 0
                molodniki_screen.load_page_data()
                molodniki_screen.update_pagination()
                self.molodniki_popup.dismiss()
                self.show_success("Данные участка молодняков загружены!")
                App.get_running_app().root.current = 'molodniki'
            except Exception as e:
                self.show_error(f"Ошибка загрузки: {str(e)}")
        else:
            self.show_error("Файл участка молодняков не найден!")



class TableScreen(Screen):
    current_page = NumericProperty(0)
    total_pages = NumericProperty(1)
    unsaved_changes = BooleanProperty(False)
    focused_cell = ListProperty([0, 0])
    edit_mode = BooleanProperty(False)
    current_section = StringProperty("")
    MAX_PAGES = 200

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.theme_manager = App.get_running_app().theme_manager
        self.reports_dir = "reports"
        os.makedirs(self.reports_dir, exist_ok=True)
        self.db_name = 'forest_data.db'
        self.rows_per_page = 50
        self.page_data = {}
        self.default_column_names = ['№ дерева*', 'Порода*', 'ж/ф', 'шт/либо лет',
                                   'D, см*', 'H, м', 'Сост-е', 'Модель', 'Примечания']
        self.column_names = None
        self.header_buttons = []
        self.load_column_config()
        if not self.column_names:
            self.column_names = self.default_column_names.copy()
        self.setup_database()
        self.create_ui()
        self.load_existing_data()
        Window.bind(on_key_down=self.key_action)

    def load_column_config(self):
        try:
            with open('table_config.json', 'r', encoding='utf-8') as f:
                config = json.load(f)
                self.column_names = config.get('column_names')
        except (FileNotFoundError, json.JSONDecodeError, KeyError):
            self.column_names = None

    def save_column_config(self):
        config = {
            'column_names': self.column_names
        }
        with open('table_config.json', 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=4)

    def key_action(self, window, key, scancode, codepoint, modifier):
        if key == 115 and 'ctrl' in modifier:
            self.save_current_page()

    def setup_database(self):
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        cursor.execute('''CREATE TABLE IF NOT EXISTS trees (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        tree_number INTEGER,
                        species TEXT,
                        age TEXT,
                        count TEXT,
                        diameter REAL,
                        height REAL,
                        condition TEXT,
                        model TEXT,
                        notes TEXT,
                        section_id INTEGER,
                        FOREIGN KEY(section_id) REFERENCES sections(id))''')
        
        cursor.execute('''CREATE TABLE IF NOT EXISTS sections (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        section_number TEXT UNIQUE,
                        quarter TEXT,
                        plot TEXT,
                        forestry TEXT,
                        district_forestry TEXT,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')
        
        cursor.execute('''CREATE TABLE IF NOT EXISTS suggestions (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        column_index INTEGER,
                        value TEXT,
                        UNIQUE(column_index, value))''')
        
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_suggestions ON suggestions (column_index, value)')
        conn.commit()
        conn.close()

    def create_ui(self):
        main_layout = BoxLayout(orientation='horizontal', padding=10, spacing=10)
        
        with self.canvas.before:
            self.bg_color = Color(1, 1, 1, 1)
            self.bg_rect = Rectangle(pos=self.pos, size=self.size)
            self.bind(pos=self._update_bg, size=self._update_bg)
        
        self._update_background(self.theme_manager.current_theme)

        # Табличная часть (левая панель)
        table_panel = BoxLayout(orientation='vertical', size_hint_x=0.7)
        
        # Заголовок участка
        header_layout = BoxLayout(orientation='vertical', size_hint=(1, None), height=30)
        self.section_label = Label(
            text=f"Участок: {self.current_section}", 
            font_name='Roboto',
            size_hint=(1, None),
            height=30,
            color=self._get_text_color()
        )
        header_layout.add_widget(self.section_label)
        table_panel.add_widget(header_layout)
        
        # Пагинация
        pagination = BoxLayout(size_hint_y=None, height=40, spacing=5)
        self.page_label = Label(
            text=f'Страница {self.current_page+1} из {self.total_pages}', 
            size_hint_x=0.4, 
            font_name='Roboto',
            color=self._get_text_color()
        )
        prev_btn = ModernButton(
            text='← Предыдущая', 
            size_hint_x=0.3,
            bg_color=get_color_from_hex('#00FF00'),
            color=self._get_text_color()
        )
        prev_btn.bind(on_press=lambda x: self.change_page(-1))
        next_btn = ModernButton(
            text='Следующая →', 
            size_hint_x=0.3,
            bg_color=get_color_from_hex('#00FF00'),
            color=self._get_text_color()
        )
        next_btn.bind(on_press=lambda x: self.change_page(1))
        pagination.add_widget(prev_btn)
        pagination.add_widget(self.page_label)
        pagination.add_widget(next_btn)
        table_panel.add_widget(pagination)
        
        # Основная таблица
        scroll = ScrollView(do_scroll_x=True, do_scroll_y=True, bar_width=10)
        self.table = GridLayout(cols=9, size_hint=(None, None), spacing=2)
        self.table.bind(minimum_height=self.table.setter('height'), 
                       minimum_width=self.table.setter('width'))
        
        # Заголовки столбцов
        self.header_buttons = []
        for i, column_name in enumerate(self.column_names):
            btn = ModernButton(
                text=column_name,
                bg_color=get_color_from_hex('#00FF00'),
                size_hint_y=None,
                height=30,
                size_hint_x=None,
                width=100,
                color=get_color_from_hex('#000000'),
                font_name='Roboto',
                bold=True,
                halign='center',
                no_shadow=True
            )
            btn.bind(on_press=lambda x, idx=i: self.edit_column_name(idx))
            self.header_buttons.append(btn)
            self.table.add_widget(btn)
        
        # Создаем строки таблицы
        self.inputs = []
        for row_idx in range(self.rows_per_page):
            row = []
            for col_idx in range(9):
                inp = AutoCompleteTextInput(multiline=False, size_hint_y=None, height=30, size_hint_x=None, width=150 if col_idx == 1 else 100)
                inp.row_index = row_idx
                inp.col_index = col_idx
                inp.bind(focus=self.update_focus)
                inp.font_name = 'Roboto'

                if col_idx > 0:
                    inp.prev_widget = row[col_idx-1] if row else None
                    if row:
                        row[col_idx-1].next_widget = inp
                if row_idx > 0:
                    inp.prev_widget = self.inputs[row_idx-1][col_idx] if self.inputs else None
                    if self.inputs:
                        self.inputs[row_idx-1][col_idx].next_widget = inp

                if col_idx in [0,4,5]:
                    inp.input_filter = 'float' if col_idx in [4,5] else 'int'
                if col_idx == 0:
                    inp.bind(focus=self.show_tree_popup)
                row.append(inp)
                self.table.add_widget(inp)
            self.inputs.append(row)

        # Добавляем кнопку "Итого" по середине после строки итогов
        # Пустая строка для разделения
        spacer = BoxLayout(orientation='horizontal', size_hint_y=None, height=10)
        self.table.add_widget(spacer)

        # Кнопка "Итого" по середине
        button_container = BoxLayout(orientation='horizontal', size_hint_y=None, height=50, size_hint_x=1)
        button_spacer = BoxLayout(size_hint_x=0.25)  # Спейсер слева
        button_container.add_widget(button_spacer)

        self.total_summary_button = ModernButton(
            text='Итого',
            bg_color=get_color_from_hex('#00FF00'),  # Зеленый цвет
            size_hint=(None, None),
            size=(200, 50),
            font_size='18sp',
            bold=True
        )
        self.total_summary_button.bind(on_press=self.show_total_summary_popup)
        button_container.add_widget(self.total_summary_button)

        button_spacer2 = BoxLayout(size_hint_x=0.25)  # Спейсер справа
        button_container.add_widget(button_spacer2)

        self.table.add_widget(button_container)

        scroll.add_widget(self.table)
        table_panel.add_widget(scroll)
        main_layout.add_widget(table_panel)
        
        # Правая панель управления
        control_panel = BoxLayout(
            orientation='vertical', 
            size_hint_x=0.3,
            spacing=15,
            padding=[0, 10, 0, 0]
        )
        
        # Основные кнопки управления
        controls = BoxLayout(
            orientation='vertical',
            size_hint_y=None,
            height=420,
            spacing=10,
            pos_hint={'top': 1}
        )
        
        button_handlers = {
            'Сохранить отчет': self.show_save_dialog,
            'Сохранить страницу': self.save_current_page,
            'Загрузить участок': self.load_section,
            'Открыть папку': self.open_excel_file,
            'Очистить данные': self.clear_table_data,
            'В главное меню': self.go_back
        }

        button_colors = {
            'Сохранить отчет': '#00FF00',
            'Сохранить страницу': '#00FFFF',
            'Загрузить участок': '#006400',
            'Открыть папку': '#0000FF',
            'Очистить данные': '#800000',
            'В главное меню': '#FF0000'
        }
        
        for text, color in button_colors.items():
            btn = ModernButton(
                text=text,
                bg_color=get_color_from_hex(color),
                size_hint=(None, None),
                size=(220, 45),
                pos_hint={'center_x': 0.5}
            )
            btn.bind(on_press=button_handlers[text])
            controls.add_widget(btn)
        
        control_panel.add_widget(controls)


        
        main_layout.add_widget(control_panel)
        self.add_widget(main_layout)

    def _update_bg(self, instance, value):
        self.bg_rect.pos = self.pos
        self.bg_rect.size = self.size

    def _update_background(self, theme):
        if theme['type'] == 'image':
            try:
                self.bg_color.rgba = (1, 1, 1, 1)
                self.bg_rect.texture = CoreImage(theme['background']).texture
            except Exception as e:
                print(f"Error loading background image: {str(e)}")
        else:
            self.bg_color.rgba = theme['background']
            self.bg_rect.texture = None

    def _get_text_color(self):
        theme = self.theme_manager.current_theme
        if theme['type'] == 'image':
            return get_color_from_hex('#FFFFFF')
        else:
            return get_color_from_hex(theme['text_color'])

    def update_section_label(self):
        self.section_label.text = f"Участок: {self.current_section}"

    def toggle_edit_mode(self, instance):
        self.edit_mode = not self.edit_mode
        instance.bg_color = get_color_from_hex('#FFA500' if self.edit_mode else '#00FF00')

    def update_focus(self, instance, value):
        if value:
            self.focused_cell = [instance.row_index, instance.col_index]

    def show_tree_popup(self, instance, value):
        if value and instance.text.strip():
            if not self.edit_mode:
                # In normal mode, only show popup if other columns are empty
                if not any(inp.text.strip() for inp in self.inputs[instance.row_index][1:]):
                    self.save_suggestion(0, instance.text.strip())
                    TreeDataInputPopup(self, instance.row_index).open()
            else:
                # In edit mode, always show popup for editing existing data
                self.save_suggestion(0, instance.text.strip())
                TreeDataInputPopup(self, instance.row_index).open()



    def save_suggestion(self, col_index, value):
        try:
            conn = sqlite3.connect(self.db_name)
            cursor = conn.cursor()
            cursor.execute('''
                INSERT OR IGNORE INTO suggestions (column_index, value)
                VALUES (?, ?)
            ''', (col_index, value))
            conn.commit()
            conn.close()
        except Exception as e:
            print(f"Error saving suggestion: {e}")



    def save_current_page(self, instance=None):
        page_data = []
        for row in self.inputs:
            page_data.append([inp.text for inp in row])
        self.page_data[self.current_page] = page_data

        # Валидация данных после сохранения
        self.validate_page_data()

    def validate_page_data(self):
        """Валидация данных на странице"""
        warnings = []
        for row_idx, row in enumerate(self.page_data.get(self.current_page, [])):
            if len(row) >= 6:
                tree_num = row[0].strip()
                species = row[1].strip()
                diameter_str = row[4].strip()
                height_str = row[5].strip()

                if tree_num and not species:
                    warnings.append(f"Строка {row_idx+1}: Нет породы для дерева №{tree_num}")
                if species and not tree_num:
                    warnings.append(f"Строка {row_idx+1}: Нет номера дерева для породы {species}")

                # Проверка числовых значений
                if diameter_str:
                    try:
                        d = float(diameter_str)
                        if d <= 0 or d > 500:
                            warnings.append(f"Строка {row_idx+1}: Подозрительный диаметр {d} см")
                    except ValueError:
                        warnings.append(f"Строка {row_idx+1}: Некорректный диаметр '{diameter_str}'")

                if height_str:
                    try:
                        h = float(height_str)
                        if h <= 0 or h > 100:
                            warnings.append(f"Строка {row_idx+1}: Подозрительная высота {h} м")
                    except ValueError:
                        warnings.append(f"Строка {row_idx+1}: Некорректная высота '{height_str}'")

        if warnings:
            self.show_error("Предупреждения в данных:\n" + "\n".join(warnings[:5]))  # Показываем первые 5

    def show_save_dialog(self, instance=None):
        content = FloatLayout()

        close_btn = ModernButton(
            text='X',
            size_hint=(None, None),
            size=(40, 40),
            pos_hint={'right': 0.95, 'top': 0.95},
            bg_color=(1, 0, 0, 1),
            no_shadow=True
        )
        close_btn.bind(on_press=lambda x: self.save_popup.dismiss())

        label = Label(
            text="Введите имя файла:",
            font_name='Roboto',
            font_size='18sp',
            color=(0.2, 0.2, 0.2, 1),
            pos_hint={'center_x': 0.5, 'center_y': 0.7},
            size_hint=(None, None),
            size=(200, 50)
        )

        self.filename_input = TextInput(
            hint_text="Имя файла",
            multiline=False,
            size_hint=(None, None),
            size=(250, 40),
            pos_hint={'center_x': 0.5, 'center_y': 0.5},
            background_color=(1, 1, 1, 0.8)
        )
        default_name = f"{self.current_section}_{datetime.datetime.now().strftime('%Y-%m-%d_%H-%M')}"
        self.filename_input.text = default_name

        btn_box = BoxLayout(
            orientation='horizontal',
            spacing=5,
            size_hint=(None, None),
            size=(250, 50),
            pos_hint={'center_x': 0.5, 'center_y': 0.3}
        )
        ok_btn = ModernButton(
            text="Сохранить",
            bg_color=get_color_from_hex('#00FF00'),
            size_hint=(0.5, None),
            height=40,
            no_shadow=True
        )
        cancel_btn = ModernButton(
            text="Отмена",
            bg_color=get_color_from_hex('#FF0000'),
            size_hint=(0.5, None),
            height=40,
            no_shadow=True
        )
        btn_box.add_widget(ok_btn)
        btn_box.add_widget(cancel_btn)

        content.add_widget(close_btn)
        content.add_widget(label)
        content.add_widget(self.filename_input)
        content.add_widget(btn_box)

        self.save_popup = Popup(
            title="Сохранение отчета",
            content=content,
            size_hint=(0.6, 0.5),
            separator_height=0,
            background='',
            overlay_color=(0, 0, 0, 0.5)
        )
        ok_btn.bind(on_press=self.save_all_formats)
        cancel_btn.bind(on_press=self.save_popup.dismiss)
        self.save_popup.open()

    def save_all_formats(self, instance):
        """Сохранить данные во всех форматах сразу"""
        success_messages = []
        error_messages = []

        # Проверка наличия данных
        if not self.page_data:
            error_messages.append("Нет данных для сохранения!")
            self.show_error("Ошибки сохранения:\n" + "\n".join(error_messages))
            return

        # Проверка наличия участка
        if not self.current_section:
            error_messages.append("Не указан номер участка!")
            self.show_error("Ошибки сохранения:\n" + "\n".join(error_messages))
            return

        # Проверка существования папки reports
        if not os.path.exists(self.reports_dir):
            try:
                os.makedirs(self.reports_dir, exist_ok=True)
            except Exception as e:
                error_messages.append(f"Не удалось создать папку reports: {str(e)}")
                self.show_error("Ошибки сохранения:\n" + "\n".join(error_messages))
                return

        try:
            # Расчет итогов
            totals_data = self.calculate_totals()
            print(f"Totals calculated: {totals_data}")  # Debug

            # Сохранение в JSON
            result, error = self.save_to_json(totals_data)
            if result:
                success_messages.append(result)
            else:
                error_messages.append(f"JSON: {error}")

            # Сохранение в Excel
            result, error = self.save_to_excel_without_dialog(totals_data)
            if result:
                success_messages.append(result)
            else:
                error_messages.append(f"Excel: {error}")

            # Сохранение в Word
            result, error = self.save_to_word_without_dialog(totals_data)
            if result:
                success_messages.append(result)
            else:
                error_messages.append(f"Word: {error}")

        except Exception as e:
            import traceback
            error_messages.append(f"Общая ошибка: {str(e)}\n{traceback.format_exc()}")

        if success_messages:
            self.save_popup.dismiss()
            self.show_success("Файлы сохранены:\n" + "\n".join(success_messages))
        if error_messages:
            self.show_error("Ошибки сохранения:\n" + "\n".join(error_messages))
            # Не закрываем popup, чтобы пользователь видел ошибки

    def save_to_json(self, totals_data=None):
        """Сохранение данных в JSON формате"""
        data = {
            'page_data': self.page_data,
            'section': self.current_section,
            'column_names': self.column_names,
            'export_date': datetime.datetime.now().isoformat()
        }

        if totals_data:
            data['totals'] = totals_data

        filename = f"Перечетная_ведомость_{self.current_section}_{datetime.datetime.now().strftime('%Y-%m-%d_%H-%M')}.json"
        full_path = os.path.join(self.reports_dir, filename)

        try:
            with open(full_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            return f"JSON: {filename}", None
        except Exception as e:
            return None, f"Ошибка сохранения JSON: {str(e)}"

    def save_to_excel_without_dialog(self, totals_data=None):
        """Сохранение в Excel без диалога"""
        filename = f"Перечетная_ведомость_{self.current_section}_{datetime.datetime.now().strftime('%Y-%m-%d_%H-%M')}.xlsx"
        full_path = os.path.join(self.reports_dir, filename)

        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Перечетная ведомость"

            # Заголовок
            ws['A1'] = f'ПЕРЕЧЕТНАЯ ВЕДОМОСТЬ - УЧАСТОК {self.current_section}'
            ws['A1'].font = openpyxl.styles.Font(bold=True, size=14)
            ws.merge_cells('A1:I1')

            # Данные
            all_data = []
            for page in sorted(self.page_data.keys()):
                all_data.extend(self.page_data[page])

            # Заголовки столбцов
            for col_num, header in enumerate(self.column_names, 1):
                cell = ws.cell(row=3, column=col_num, value=header)
                cell.font = openpyxl.styles.Font(bold=True)
                cell.fill = openpyxl.styles.PatternFill(start_color="00FF00", end_color="00FF00", fill_type="solid")

            # Данные
            for row_num, row_data in enumerate(all_data, 4):
                for col_num, cell_value in enumerate(row_data, 1):
                    ws.cell(row=row_num, column=col_num, value=cell_value)

            # Автоподбор ширины столбцов
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width

            # Добавляем лист с итогами, если есть данные
            if totals_data:
                ws_totals = wb.create_sheet("Итоги")
                ws_totals['A1'] = f'ИТОГИ ПО ПЕРЕЧЕТНОЙ ВЕДОМОСТИ - УЧАСТОК {self.current_section}'
                ws_totals['A1'].font = openpyxl.styles.Font(bold=True, size=14)
                ws_totals.merge_cells('A1:D1')

                # Общие итоги
                ws_totals['A3'] = 'Общие показатели:'
                ws_totals['A3'].font = openpyxl.styles.Font(bold=True)
                ws_totals['A4'] = 'Всего деревьев:'
                ws_totals['B4'] = totals_data.get('total_trees', 0)
                ws_totals['A5'] = 'Средний диаметр (см):'
                ws_totals['B5'] = round(totals_data.get('avg_diameter', 0), 1)
                ws_totals['A6'] = 'Средняя высота (м):'
                ws_totals['B6'] = round(totals_data.get('avg_height', 0), 1)

                # Распределение по породам
                species_summary = totals_data.get('species_summary', {})
                if species_summary:
                    ws_totals['A8'] = 'Распределение по породам:'
                    ws_totals['A8'].font = openpyxl.styles.Font(bold=True)

                    row_num = 9
                    for species, data in sorted(species_summary.items()):
                        ws_totals[f'A{row_num}'] = f'Порода: {species}'
                        ws_totals[f'A{row_num}'].font = openpyxl.styles.Font(bold=True)
                        ws_totals[f'B{row_num}'] = f'Количество: {data["count"]}'

                        diameters = data['diameters']
                        heights = data['heights']
                        if diameters:
                            avg_d = sum(diameters) / len(diameters)
                            ws_totals[f'C{row_num}'] = f'Ср. диаметр: {avg_d:.1f} см'
                        if heights:
                            avg_h = sum(heights) / len(heights)
                            ws_totals[f'D{row_num}'] = f'Ср. высота: {avg_h:.1f} м'

                        row_num += 1

                    # Автоподбор ширины для листа итогов
                    for column in ws_totals.columns:
                        max_length = 0
                        column_letter = column[0].column_letter
                        for cell in column:
                            try:
                                if len(str(cell.value)) > max_length:
                                    max_length = len(str(cell.value))
                            except:
                                pass
                        adjusted_width = min(max_length + 2, 50)
                        ws_totals.column_dimensions[column_letter].width = adjusted_width

            wb.save(full_path)
            return f"Excel: {filename}", None
        except Exception as e:
            return None, f"Ошибка сохранения Excel: {str(e)}"

    def save_to_word_without_dialog(self, totals_data=None):
        """Сохранение в Word без диалога"""
        try:
            from docx import Document

            filename = f"Перечетная_ведомость_{self.current_section}_{datetime.datetime.now().strftime('%Y-%m-%d_%H-%M')}.docx"
            full_path = os.path.join(self.reports_dir, filename)

            doc = Document()
            doc.add_heading(f'Перечетная ведомость - Участок {self.current_section}', 0)

            all_data = []
            for page in sorted(self.page_data.keys()):
                all_data.extend(self.page_data[page])

            table = doc.add_table(rows=1, cols=len(self.column_names))
            table.style = 'Table Grid'

            # Заголовки
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(self.column_names):
                hdr_cells[i].text = header

            # Данные
            for row_data in all_data:
                row_cells = table.add_row().cells
                for i, cell_value in enumerate(row_data):
                    row_cells[i].text = str(cell_value) if cell_value else ""

            # Добавляем итоги, если есть данные
            if totals_data:
                doc.add_page_break()
                doc.add_heading('Итоги по перечетной ведомости', 1)

                # Общие итоги
                doc.add_paragraph(f'Всего деревьев: {totals_data.get("total_trees", 0)}')
                doc.add_paragraph(f'Средний диаметр: {totals_data.get("avg_diameter", 0):.1f} см (измерено: {totals_data.get("diameter_count", 0)})')
                doc.add_paragraph(f'Средняя высота: {totals_data.get("avg_height", 0):.1f} м (измерено: {totals_data.get("height_count", 0)})')

                # Распределение по породам
                species_summary = totals_data.get('species_summary', {})
                if species_summary:
                    doc.add_heading('Распределение по породам', 2)

                    # Создаем таблицу для пород
                    species_table = doc.add_table(rows=1, cols=4)
                    species_table.style = 'Table Grid'

                    # Заголовки таблицы пород
                    hdr_cells = species_table.rows[0].cells
                    hdr_cells[0].text = 'Порода'
                    hdr_cells[1].text = 'Количество'
                    hdr_cells[2].text = 'Средний диаметр'
                    hdr_cells[3].text = 'Средняя высота'

                    # Данные по породам
                    for species, data in sorted(species_summary.items()):
                        row_cells = species_table.add_row().cells
                        row_cells[0].text = species
                        row_cells[1].text = str(data['count'])

                        diameters = data['diameters']
                        heights = data['heights']

                        if diameters:
                            avg_d = sum(diameters) / len(diameters)
                            row_cells[2].text = f'{avg_d:.1f} см'
                        else:
                            row_cells[2].text = '-'

                        if heights:
                            avg_h = sum(heights) / len(heights)
                            row_cells[3].text = f'{avg_h:.1f} м'
                        else:
                            row_cells[3].text = '-'

            doc.save(full_path)
            return f"Word: {filename}", None
        except ImportError:
            return None, "Для сохранения в Word установите библиотеку python-docx: pip install python-docx"
        except Exception as e:
            return None, f"Ошибка сохранения Word: {str(e)}"

    def save_to_excel(self, instance):
        filename = self.filename_input.text.strip()
        if not filename:
            self.show_error("Имя файла не может быть пустым!")
            return
        filename = re.sub(r'[\\/*?:"<>|]', "", filename)
        filename = f"{filename}.xlsx" if not filename.endswith(".xlsx") else filename
        full_path = os.path.join(self.reports_dir, filename)

        try:
            # Расчет итогов
            totals_data = self.calculate_totals()

            wb = Workbook()
            ws = wb.active
            ws.title = "Перечетная ведомость"

            # Заголовок
            ws['A1'] = f'ПЕРЕЧЕТНАЯ ВЕДОМОСТЬ - УЧАСТОК {self.current_section}'
            ws['A1'].font = openpyxl.styles.Font(bold=True, size=14)
            ws.merge_cells('A1:I1')

            # Данные
            all_data = []
            for page in sorted(self.page_data.keys()):
                all_data.extend(self.page_data[page])

            # Заголовки столбцов
            for col_num, header in enumerate(self.column_names, 1):
                cell = ws.cell(row=3, column=col_num, value=header)
                cell.font = openpyxl.styles.Font(bold=True)
                cell.fill = openpyxl.styles.PatternFill(start_color="00FF00", end_color="00FF00", fill_type="solid")

            # Данные
            for row_num, row_data in enumerate(all_data, 4):
                for col_num, cell_value in enumerate(row_data, 1):
                    ws.cell(row=row_num, column=col_num, value=cell_value)

            # Автоподбор ширины столбцов
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width

            # Добавляем лист с итогами
            ws_totals = wb.create_sheet("Итоги")
            ws_totals['A1'] = f'ИТОГИ ПО ПЕРЕЧЕТНОЙ ВЕДОМОСТИ - УЧАСТОК {self.current_section}'
            ws_totals['A1'].font = openpyxl.styles.Font(bold=True, size=14)
            ws_totals.merge_cells('A1:D1')

            # Общие итоги
            ws_totals['A3'] = 'Общие показатели:'
            ws_totals['A3'].font = openpyxl.styles.Font(bold=True)
            ws_totals['A4'] = 'Всего деревьев:'
            ws_totals['B4'] = totals_data.get('total_trees', 0)
            ws_totals['A5'] = 'Средний диаметр (см):'
            ws_totals['B5'] = round(totals_data.get('avg_diameter', 0), 1)
            ws_totals['A6'] = 'Средняя высота (м):'
            ws_totals['B6'] = round(totals_data.get('avg_height', 0), 1)

            # Распределение по породам
            species_summary = totals_data.get('species_summary', {})
            if species_summary:
                ws_totals['A8'] = 'Распределение по породам:'
                ws_totals['A8'].font = openpyxl.styles.Font(bold=True)

                row_num = 9
                for species, data in sorted(species_summary.items()):
                    ws_totals[f'A{row_num}'] = f'Порода: {species}'
                    ws_totals[f'A{row_num}'].font = openpyxl.styles.Font(bold=True)
                    ws_totals[f'B{row_num}'] = f'Количество: {data["count"]}'

                    diameters = data['diameters']
                    heights = data['heights']
                    if diameters:
                        avg_d = sum(diameters) / len(diameters)
                        ws_totals[f'C{row_num}'] = f'Ср. диаметр: {avg_d:.1f} см'
                    if heights:
                        avg_h = sum(heights) / len(heights)
                        ws_totals[f'D{row_num}'] = f'Ср. высота: {avg_h:.1f} м'

                    row_num += 1

                # Автоподбор ширины для листа итогов
                for column in ws_totals.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    ws_totals.column_dimensions[column_letter].width = adjusted_width

            wb.save(full_path)
            self.save_popup.dismiss()
            self.show_success(f"Файл сохранен: {filename}")
        except Exception as e:
            self.show_error(f"Ошибка: {str(e)}")

    def load_section(self, instance):
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        cursor.execute('SELECT section_number FROM sections WHERE section_number IS NOT NULL AND section_number != "" ORDER BY id DESC')
        sections = cursor.fetchall()
        conn.close()
        if not sections:
            self.show_error("Нет сохраненных участков!")
            return
        content = FloatLayout()

        close_btn = ModernButton(
            text='X',
            size_hint=(None, None),
            size=(40, 40),
            pos_hint={'right': 0.95, 'top': 0.95},
            bg_color=(1, 0, 0, 1),
            no_shadow=True
        )
        close_btn.bind(on_press=lambda x: self.load_popup.dismiss())

        scroll = ScrollView(size_hint=(1, 0.9), pos_hint={'center_x': 0.5, 'center_y': 0.45})
        layout = GridLayout(cols=1, spacing=5, size_hint_y=None)
        layout.bind(minimum_height=layout.setter('height'))
        for section in sections:
            btn = ModernButton(
                text=section[0],
                size_hint_y=None,
                height=40,
                bg_color=(0, 1, 0, 1),
                color=(0, 0, 0, 1),
                no_shadow=True
            )
            btn.bind(on_release=lambda b, s=section[0]: self.load_saved_section_from_popup(s))
            layout.add_widget(btn)
        scroll.add_widget(layout)

        content.add_widget(close_btn)
        content.add_widget(scroll)
        self.load_popup = Popup(
            title="Выберите участок",
            content=content,
            size_hint=(0.5, 0.6),
            separator_height=0,
            background_color=(0.5, 0.5, 0.5, 1),
            overlay_color=(0, 0, 0, 0.5)
        )
        self.load_popup.open()

    def load_saved_section_from_popup(self, section_number):
        files = glob.glob(os.path.join(self.reports_dir, f"{section_number}_*.xlsx"))
        if files:
            latest_file = max(files, key=os.path.getctime)
            try:
                df = pd.read_excel(latest_file)
                data = df.values.tolist()

                self.current_section = section_number
                self.update_section_label()
                self.page_data.clear()

                for page_num in range(0, len(df), self.rows_per_page):
                    page = page_num // self.rows_per_page
                    page_data = df.iloc[page_num:page_num+self.rows_per_page].values.tolist()
                    self.page_data[page] = page_data

                self.current_page = 0
                self.load_page_data()
                self.update_pagination()
                self.load_popup.dismiss()
                self.show_success("Данные участка загружены!")
            except Exception as e:
                self.show_error(f"Ошибка загрузки: {str(e)}")
        else:
            self.show_error("Файл участка не найден!")

    def load_page_data(self):
        for row in self.inputs:
            for inp in row:
                inp.text = ''
        
        if self.current_page in self.page_data:
            for i, row_data in enumerate(self.page_data[self.current_page]):
                if i >= len(self.inputs):
                    break
                for j, text in enumerate(row_data):
                    if j < len(self.inputs[i]):
                        self.inputs[i][j].text = str(text) if not pd.isna(text) else ''

    def clear_table_data(self, instance=None):
        for row in self.inputs:
            for inp in row:
                inp.text = ''
        self.page_data.clear()
        self.show_success("Данные очищены!")

    def open_excel_file(self, instance):
        if os.path.exists(self.reports_dir):
            os.startfile(self.reports_dir)
        else:
            self.show_error("Папка reports не найдена!")

    def change_page(self, delta):
        new_page = self.current_page + delta
        if 0 <= new_page < self.MAX_PAGES:
            self.current_page = new_page
            self.load_page_data()
            self.update_page_label()

    def update_pagination(self):
        self.total_pages = len(self.page_data) if self.page_data else 1
        self.total_pages = min(self.total_pages, self.MAX_PAGES)
        self.update_page_label()

    def update_page_label(self):
        self.page_label.text = f'Страница {self.current_page+1} из {self.total_pages}'

    def go_back(self, instance):
        App.get_running_app().root.current = 'main'

    def show_error(self, message):
        content = FloatLayout()
        with content.canvas.before:
            Color(0.98, 0.98, 0.98, 1)
            RoundedRectangle(
                pos=(content.x-10, content.y-10),
                size=(content.width+20, content.height+20),
                radius=[50]
            )
        label = Label(
            text=message,
            color=(1, 0, 0, 1),
            font_name='Roboto',
            pos_hint={'center_x': 0.5, 'center_y': 0.5}
        )
        content.add_widget(label)
        Popup(
            title='Ошибка',
            content=content,
            size_hint=(0.6, 0.3),
            background='atlas://data/images/defaulttheme/modalview-background',
            overlay_color=(0, 0, 0, 0.5)
        ).open()

    def show_success(self, message):
        content = FloatLayout()
        label = Label(
            text=message,
            color=(0, 0.5, 0, 1),
            font_name='Roboto',
            pos_hint={'center_x': 0.5, 'center_y': 0.5}
        )
        content.add_widget(label)
        Popup(
            title='Успешно',
            content=content,
            size_hint=(0.6, 0.3),
            background='atlas://data/images/defaulttheme/modalview-background',
            overlay_color=(0, 0, 0, 0.5)
        ).open()
        
    def load_existing_data(self):
        pass

    def edit_column_name(self, col_idx):
        current_name = self.column_names[col_idx]
        content = FloatLayout()

        label = Label(
            text='Введите новое название столбца:',
            font_name='Roboto',
            font_size='18sp',
            color=(0.2, 0.2, 0.2, 1),
            pos_hint={'center_x': 0.5, 'center_y': 0.75},
            size_hint=(None, None),
            size=(300, 30)
        )

        input_field = TextInput(
            text=current_name,
            multiline=False,
            size_hint=(None, None),
            size=(300, 40),
            pos_hint={'center_x': 0.5, 'center_y': 0.5},
            font_name='Roboto'
        )

        btn_box = BoxLayout(
            orientation='horizontal',
            size_hint=(None, None),
            size=(300, 40),
            spacing=10,
            pos_hint={'center_x': 0.5, 'center_y': 0.3}
        )

        save_btn = ModernButton(
            text='Сохранить',
            bg_color=get_color_from_hex('#00FF00'),
            size_hint=(0.5, None),
            height=40,
            no_shadow=True
        )
        exit_btn = ModernButton(
            text='Выйти',
            bg_color=get_color_from_hex('#FF0000'),
            size_hint=(0.5, None),
            height=40,
            no_shadow=True
        )

        save_btn.bind(on_press=lambda x: self.save_column_name(col_idx, input_field.text))
        exit_btn.bind(on_press=lambda x: self.dismiss_column_edit_popup())

        btn_box.add_widget(save_btn)
        btn_box.add_widget(exit_btn)

        content.add_widget(label)
        content.add_widget(input_field)
        content.add_widget(btn_box)

        self.column_edit_popup = Popup(
            title='Изменить название столбца',
            content=content,
            size_hint=(0.6, 0.5),
            background='atlas://data/images/defaulttheme/modalview-background',
            overlay_color=(0, 0, 0, 0.5)
        )
        self.column_edit_popup.open()

    def dismiss_column_edit_popup(self):
        if hasattr(self, 'column_edit_popup'):
            self.column_edit_popup.dismiss()
            del self.column_edit_popup

    def save_column_name(self, col_idx, new_name):
        if new_name.strip():
            self.column_names[col_idx] = new_name.strip()
            self.header_buttons[col_idx].text = new_name.strip()
            self.save_column_config()
            self.column_edit_popup.dismiss()
        else:
            self.show_error("Название не может быть пустым!")

    def calculate_totals(self):
        """Расчет итоговых данных по перечетной ведомости"""
        try:
            # Собираем все данные из всех страниц
            all_data = []
            for page in sorted(self.page_data.keys()):
                all_data.extend(self.page_data[page])

            # Словарь для группировки по породам
            species_summary = {}

            # Определение хвойных и лиственных пород
            coniferous_breeds = ['Сосна', 'Ель', 'Пихта', 'Кедр', 'Лиственница']
            deciduous_breeds = ['Берёза', 'Осина', 'Ольха чёрная', 'Ольха серая', 'Ива', 'Ива кустарниковая']

            def get_breed_type(species):
                for conif in coniferous_breeds:
                    if conif.lower() in species.lower():
                        return 'coniferous'
                for decid in deciduous_breeds:
                    if decid.lower() in species.lower():
                        return 'deciduous'
                return 'unknown'

            total_trees = 0
            total_diameter = 0.0
            total_height = 0.0
            diameter_count = 0
            height_count = 0

            for row in all_data:
                if len(row) >= 6:  # Проверяем, что строка содержит достаточное количество столбцов
                    tree_num = row[0].strip()
                    species = row[1].strip()
                    age = row[3].strip()
                    diameter_str = row[4].strip()
                    height_str = row[5].strip()

                    if tree_num and species:  # Только если есть номер дерева и порода
                        total_trees += 1

                        # Группировка по породам
                        if species not in species_summary:
                            species_summary[species] = {
                                'count': 0,
                                'diameters': [],
                                'heights': [],
                                'ages': [],
                                'type': get_breed_type(species)
                            }

                        species_summary[species]['count'] += 1

                        # Диаметр
                        try:
                            diameter = float(diameter_str)
                            species_summary[species]['diameters'].append(diameter)
                            total_diameter += diameter
                            diameter_count += 1
                        except (ValueError, TypeError):
                            pass

                        # Высота
                        try:
                            height = float(height_str)
                            species_summary[species]['heights'].append(height)
                            total_height += height
                            height_count += 1
                        except (ValueError, TypeError):
                            pass

                        # Возраст
                        if age:
                            species_summary[species]['ages'].append(age)

            # Рассчитываем средние значения
            avg_diameter = total_diameter / diameter_count if diameter_count > 0 else 0
            avg_height = total_height / height_count if height_count > 0 else 0

            # Формируем итоговые данные
            totals_data = {
                'total_trees': total_trees,
                'avg_diameter': avg_diameter,
                'avg_height': avg_height,
                'diameter_count': diameter_count,
                'height_count': height_count,
                'species_summary': species_summary
            }

            return totals_data

        except Exception as e:
            print(f"Error calculating totals: {str(e)}")
            return {}

    def show_total_summary_popup(self, instance=None):
        """Показать popup с итоговыми данными по перечетной ведомости"""
        try:
            totals_data = self.calculate_totals()

            # Создаем popup
            content = BoxLayout(orientation='vertical', spacing=10, padding=10)

            # Заголовок
            title_label = Label(
                text=f'ИТОГИ ПО ПЕРЕЧЕТНОЙ ВЕДОМОСТИ\nУчасток: {self.current_section}',
                font_name='Roboto',
                font_size='20sp',
                bold=True,
                color=(0, 0.5, 0, 1),
                size_hint=(1, None),
                height=60,
                halign='center',
                valign='top'
            )
            content.add_widget(title_label)

            # Блок адрес участка
            address_block = BoxLayout(
                orientation='vertical',
                size_hint=(1, None),
                height=100,
                padding=[10, 10],
                spacing=5
            )
            with address_block.canvas.before:
                Color(rgba=get_color_from_hex('#E8F4FD'))
                address_block.bg = RoundedRectangle(pos=address_block.pos, size=address_block.size, radius=[10])
                address_block.bind(pos=lambda *args: setattr(address_block.bg, 'pos', address_block.pos),
                                 size=lambda *args: setattr(address_block.bg, 'size', address_block.size))

            address_title = Label(
                text='АДРЕС УЧАСТКА',
                font_name='Roboto',
                font_size='16sp',
                bold=True,
                color=(0, 0, 0, 1),
                size_hint=(1, None),
                height=25,
                halign='center'
            )
            address_block.add_widget(address_title)

            # Получаем адресные данные из базы данных
            address_info = "Адрес не указан"
            try:
                conn = sqlite3.connect(self.db_name)
                cursor = conn.cursor()
                cursor.execute('''
                    SELECT quarter, plot, forestry, district_forestry
                    FROM sections
                    WHERE section_number = ?
                    ORDER BY id DESC LIMIT 1
                ''', (self.current_section,))
                address_row = cursor.fetchone()
                if address_row:
                    quarter, plot, forestry, district_forestry = address_row
                    address_parts = []
                    if quarter: address_parts.append(f"Квартал: {quarter}")
                    if plot: address_parts.append(f"Выдел: {plot}")
                    if forestry: address_parts.append(f"Лесничество: {forestry}")
                    if district_forestry: address_parts.append(f"Участковое лесничество: {district_forestry}")
                    if address_parts:
                        address_info = "\n".join(address_parts)
                conn.close()
            except Exception as e:
                print(f"Error loading address: {e}")

            address_label = Label(
                text=address_info,
                font_name='Roboto',
                font_size='14sp',
                color=(0, 0, 0, 1),
                size_hint=(1, None),
                height=60,
                halign='left',
                valign='top'
            )
            address_label.bind(size=lambda *args: setattr(address_label, 'text_size', (address_label.width, None)))
            address_block.add_widget(address_label)
            content.add_widget(address_block)

            # Блок Проект
            project_block = BoxLayout(
                orientation='vertical',
                size_hint=(1, None),
                height=120,
                padding=[10, 10],
                spacing=5
            )
            with project_block.canvas.before:
                Color(rgba=get_color_from_hex('#FFF8E1'))
                project_block.bg = RoundedRectangle(pos=project_block.pos, size=project_block.size, radius=[10])
                project_block.bind(pos=lambda *args: setattr(project_block.bg, 'pos', project_block.pos),
                                 size=lambda *args: setattr(project_block.bg, 'size', project_block.size))

            project_title = Label(
                text='ПРОЕКТ',
                font_name='Roboto',
                font_size='16sp',
                bold=True,
                color=(0, 0, 0, 1),
                size_hint=(1, None),
                height=25,
                halign='center'
            )
            project_block.add_widget(project_title)

            # Получаем данные проекта из меню Действия (таксационные показатели)
            project_info = "Проект не создан"
            try:
                # Попытка получить данные из последнего расчета таксационных показателей
                molodniki_screen = App.get_running_app().root.get_screen('molodniki')
                if hasattr(molodniki_screen, 'current_section') and molodniki_screen.current_section:
                    total_data = molodniki_screen.get_total_data_from_db()
                    if total_data and total_data.get('care_subject'):
                        care_subject = total_data.get('care_subject', '')
                        intensity = total_data.get('intensity', 0)
                        project_info = f"Предмет ухода: {care_subject}\nИнтенсивность рубки: {intensity:.1f}%"
                    else:
                        project_info = "Данные проекта отсутствуют"
                else:
                    project_info = "Молодняки не загружены"
            except Exception as e:
                print(f"Error loading project data: {e}")

            project_label = Label(
                text=project_info,
                font_name='Roboto',
                font_size='14sp',
                color=(0, 0, 0, 1),
                size_hint=(1, None),
                height=80,
                halign='left',
                valign='top'
            )
            project_label.bind(size=lambda *args: setattr(project_label, 'text_size', (project_label.width, None)))
            project_block.add_widget(project_label)
            content.add_widget(project_block)

            # Общие итоги
            summary_text = f"""
Общие показатели:
• Всего деревьев: {totals_data.get('total_trees', 0)}
• Средний диаметр: {totals_data.get('avg_diameter', 0):.1f} см (измерено: {totals_data.get('diameter_count', 0)})
• Средняя высота: {totals_data.get('avg_height', 0):.1f} м (измерено: {totals_data.get('height_count', 0)})
"""

            summary_label = Label(
                text=summary_text,
                font_name='Roboto',
                font_size='16sp',
                color=(0, 0, 0, 1),
                size_hint=(1, None),
                height=100,
                halign='left',
                valign='top'
            )
            summary_label.bind(size=lambda *args: setattr(summary_label, 'text_size', (summary_label.width, None)))
            content.add_widget(summary_label)

            # Блок хвойных пород
            coniferous_summary = totals_data.get('species_summary', {})
            coniferous_data = {species: data for species, data in coniferous_summary.items() if data.get('type') == 'coniferous'}

            if coniferous_data:
                coniferous_block = BoxLayout(
                    orientation='vertical',
                    size_hint=(1, None),
                    height=150,
                    padding=[10, 10],
                    spacing=5
                )
                with coniferous_block.canvas.before:
                    Color(rgba=get_color_from_hex('#F0F8FF'))
                    coniferous_block.bg = RoundedRectangle(pos=coniferous_block.pos, size=coniferous_block.size, radius=[10])
                    coniferous_block.bind(pos=lambda *args: setattr(coniferous_block.bg, 'pos', coniferous_block.pos),
                                         size=lambda *args: setattr(coniferous_block.bg, 'size', coniferous_block.size))

                coniferous_title = Label(
                    text='ХВОЙНЫЕ ПОРОДЫ - СРЕДНИЕ ПОКАЗАТЕЛИ',
                    font_name='Roboto',
                    font_size='16sp',
                    bold=True,
                    color=(0, 0, 0, 1),
                    size_hint=(1, None),
                    height=25,
                    halign='center'
                )
                coniferous_block.add_widget(coniferous_title)

                coniferous_info = []
                for species, data in sorted(coniferous_data.items()):
                    count = data['count']
                    diameters = data['diameters']
                    heights = data['heights']
                    ages = data.get('ages', [])
                    avg_diameter = sum(diameters) / len(diameters) if diameters else 0
                    avg_height = sum(heights) / len(heights) if heights else 0
                    avg_age = sum(int(a) for a in ages if a.isdigit()) / len([a for a in ages if a.isdigit()]) if ages else 0
                    coniferous_info.append(f"{species}: {count} шт, ср. D={avg_diameter:.1f}см, ср. H={avg_height:.1f}м, ср. возраст={avg_age:.1f}лет")

                coniferous_text = "\n".join(coniferous_info)
                coniferous_label = Label(
                    text=coniferous_text,
                    font_name='Roboto',
                    font_size='14sp',
                    color=(0, 0.4, 0, 1),
                    size_hint=(1, None),
                    height=100,
                    halign='left',
                    valign='top'
                )
                coniferous_label.bind(size=lambda *args: setattr(coniferous_label, 'text_size', (coniferous_label.width, None)))
                coniferous_block.add_widget(coniferous_label)
                content.add_widget(coniferous_block)

            # ScrollView для детального отчета по породам
            scroll = ScrollView(size_hint=(1, None), height=400)
            details_layout = GridLayout(cols=1, spacing=10, size_hint_y=None)
            details_layout.bind(minimum_height=details_layout.setter('height'))

            # Заголовок раздела пород
            species_title = Label(
                text='РАСПРЕДЕЛЕНИЕ ПО ПОРОДАМ:',
                font_name='Roboto',
                font_size='18sp',
                bold=True,
                color=(0, 0, 0, 1),
                size_hint=(1, None),
                height=40,
                halign='center'
            )
            details_layout.add_widget(species_title)

            # Детали по каждой породе
            species_summary = totals_data.get('species_summary', {})
            for species, data in sorted(species_summary.items()):
                count = data['count']
                diameters = data['diameters']
                heights = data['heights']
                ages = data['ages']

                avg_diameter = sum(diameters) / len(diameters) if diameters else 0
                avg_height = sum(heights) / len(heights) if heights else 0

                species_text = f"""
Порода: {species}
• Количество деревьев: {count}
• Средний диаметр: {avg_diameter:.1f} см (измерено: {len(diameters)})
• Средняя высота: {avg_height:.1f} м (измерено: {len(heights)})
• Возраста: {', '.join(ages[:10])}{'...' if len(ages) > 10 else ''}  # Показываем первые 10 возрастов
"""

                species_label = Label(
                    text=species_text,
                    font_name='Roboto',
                    font_size='14sp',
                    color=(0, 0, 0.5, 1),
                    size_hint=(1, None),
                    height=120,
                    halign='left',
                    valign='top'
                )
                species_label.bind(size=lambda *args: setattr(species_label, 'text_size', (species_label.width, None)))
                details_layout.add_widget(species_label)

            scroll.add_widget(details_layout)
            content.add_widget(scroll)

            # Кнопка закрытия
            close_btn = ModernButton(
                text='Закрыть',
                bg_color=get_color_from_hex('#808080'),
                size_hint=(1, None),
                height=50
            )
            content.add_widget(close_btn)

            popup = Popup(
                title="Итоги перечетной ведомости",
                content=content,
                size_hint=(0.95, 0.95)
            )

            close_btn.bind(on_press=popup.dismiss)
            popup.open()

        except Exception as e:
            import traceback
            self.show_error(f"Ошибка расчета итогов: {str(e)}\n{traceback.format_exc()}")

class ThemeChooser(Popup):
    def __init__(self, **kwargs):
        super().__init__(title='Выбор темы', size_hint=(0.8, 0.5))
        content = FloatLayout()

        close_btn = ModernButton(
            text='X',
            size_hint=(None, None),
            size=(40, 40),
            pos_hint={'right': 0.95, 'top': 0.95},
            bg_color=(1, 0, 0, 1),
            no_shadow=True
        )
        close_btn.bind(on_press=self.dismiss)

        # Стрелка влево
        left_arrow = ModernButton(
            text='◄',
            size_hint=(None, None),
            size=(50, 50),
            pos_hint={'x': 0.05, 'center_y': 0.5},
            bg_color=(0.5, 0.5, 0.5, 1),
            no_shadow=True
        )
        left_arrow.bind(on_press=self.prev_theme)

        # Стрелка вправо
        right_arrow = ModernButton(
            text='►',
            size_hint=(None, None),
            size=(50, 50),
            pos_hint={'right': 0.95, 'center_y': 0.5},
            bg_color=(0.5, 0.5, 0.5, 1),
            no_shadow=True
        )
        right_arrow.bind(on_press=self.next_theme)

        # Контейнер для иконок тем
        self.themes_container = BoxLayout(
            orientation='horizontal',
            spacing=10,
            size_hint=(0.8, 0.8),
            pos_hint={'center_x': 0.5, 'center_y': 0.5}
        )

        self.update_theme_icons()

        content.add_widget(close_btn)
        content.add_widget(left_arrow)
        content.add_widget(right_arrow)
        content.add_widget(self.themes_container)
        self.content = content

    def update_theme_icons(self):
        self.themes_container.clear_widgets()
        for theme in App.get_running_app().theme_manager.themes:
            btn = Button(size_hint=(None, None), size=(80, 80))
            if theme['type'] == 'color':
                btn.background_color = theme['background']
            else:
                btn.background_normal = theme['background']
            btn.bind(on_release=lambda x, t=theme: self.select_theme(t))
            self.themes_container.add_widget(btn)

    def select_theme(self, theme):
        manager = App.get_running_app().theme_manager
        manager.current_theme_index = manager.themes.index(theme)
        manager.save_config()
        App.get_running_app().reload_theme()
        self.update_theme_icons()
        self.dismiss()

    def prev_theme(self, instance):
        themes = App.get_running_app().theme_manager.themes
        current_idx = App.get_running_app().theme_manager.current_theme_index
        new_idx = (current_idx - 1) % len(themes)
        App.get_running_app().theme_manager.current_theme_index = new_idx
        App.get_running_app().theme_manager.save_config()
        App.get_running_app().reload_theme()
        self.update_theme_icons()

    def next_theme(self, instance):
        themes = App.get_running_app().theme_manager.themes
        current_idx = App.get_running_app().theme_manager.current_theme_index
        new_idx = (current_idx + 1) % len(themes)
        App.get_running_app().theme_manager.current_theme_index = new_idx
        App.get_running_app().theme_manager.save_config()
        App.get_running_app().reload_theme()
        self.update_theme_icons()

    def add_theme(self, instance):
        Tk().withdraw()
        file_path = filedialog.askopenfilename(
            filetypes=[("Image files", "*.jpg *.jpeg *.png")]
        )
        if file_path:
            App.get_running_app().theme_manager.add_theme(file_path)
            self.update_theme_icons()


class ForestryApp(App):
    theme_manager = ThemeManager()
    
    def build(self):
        Config.set('graphics', 'multisamples', '4')
        self.apply_theme()
        sm = ScreenManager()
        sm.add_widget(MainMenu(name='main'))
        sm.add_widget(TableScreen(name='table'))
        sm.add_widget(ExtendedMolodnikiTableScreen(name='molodniki'))
        return sm
    
    def apply_theme(self):
        theme = self.theme_manager.current_theme
        if theme['type'] == 'image':
            Window.clearcolor = (1, 1, 1, 1)
        else:
            Window.clearcolor = theme['background']
    
    def reload_theme(self):
        self.apply_theme()
        for screen in self.root.screens:
            if hasattr(screen, 'setup_ui'):
                screen.setup_ui()
            if hasattr(screen, '_update_background'):
                screen._update_background(self.theme_manager.current_theme)

if __name__ == '__main__':
    ForestryApp().run()
