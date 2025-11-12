# Расширенная таблица молодняков по новой структуре
# Структура: 6 основных столбцов + динамические подстолбцы для пород

from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.button import Button
from kivy.uix.screenmanager import ScreenManager, Screen
from kivy.uix.scrollview import ScrollView
from kivy.uix.gridlayout import GridLayout
from kivy.uix.textinput import TextInput
from kivy.uix.popup import Popup
from kivy.uix.label import Label
from kivy.uix.floatlayout import FloatLayout
from kivy.uix.widget import Widget
from kivy.properties import (NumericProperty, BooleanProperty,
                          ObjectProperty, ListProperty, StringProperty)
from kivy.core.window import Window
from kivy.graphics import Color, Rectangle, Line, RoundedRectangle
from kivy.clock import Clock
from kivy.animation import Animation
from kivy.core.text import LabelBase
from kivy.utils import get_color_from_hex
from kivy.config import Config
from kivy.core.image import Image as CoreImage
import sqlite3
import pandas as pd
import os
import datetime
import re
import json
import shutil
import glob
import openpyxl
from openpyxl import load_workbook, Workbook
from tkinter import Tk, filedialog

LabelBase.register(name='Roboto',
                 fn_regular='fonts/Roboto-Medium.ttf',
                 fn_bold='fonts/Roboto-Bold.ttf')

class ModernButton(Button):
    bg_color = ListProperty([1, 1, 1, 1])
    no_shadow = BooleanProperty(False)

    def __init__(self, **kwargs):
        self.no_shadow = kwargs.pop('no_shadow', False)
        super().__init__(**kwargs)
        self.background_color = (0, 0, 0, 0)
        self.font_name = 'Roboto'
        self.font_size = '16sp'
        self.bold = False
        self.size_hint = (None, None)
        self.height = 50
        self.padding = (20, 5)

        with self.canvas.before:
            if not self.no_shadow:
                Color(rgba=(0, 0, 0, 0.1))
                self.shadow = RoundedRectangle(
                    pos=(self.x+3, self.y-3),
                    size=self.size,
                    radius=[10]
                )
            self.bg_color_instruction = Color(rgba=self.bg_color)
            self.background = RoundedRectangle(
                pos=self.pos,
                size=self.size,
                radius=[10]
            )

        self.bind(pos=self.update_graphics, size=self.update_graphics)
        self.bind(text=self.update_width)

    def update_width(self, instance, value):
        self.width = self.texture_size[0] + 60

    def update_graphics(self, *args):
        self.background.pos = self.pos
        if not self.no_shadow:
            self.shadow.pos = (self.x+3, self.y-3)
            self.shadow.size = self.size
        self.background.size = self.size

    def on_touch_down(self, touch):
        if self.collide_point(*touch.pos):
            Animation(rgba=[c * 0.9 for c in self.bg_color], d=0.1).start(self.bg_color_instruction)
        return super().on_touch_down(touch)

    def on_touch_up(self, touch):
        Animation(rgba=self.bg_color, d=0.2).start(self.bg_color_instruction)
        return super().on_touch_up(touch)

class AutoCompleteTextInput(TextInput):
    next_widget = ObjectProperty(None)
    prev_widget = ObjectProperty(None)
    row_index = NumericProperty(0)
    col_index = NumericProperty(0)
    suggestions = ListProperty([])

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.bind(text=self.show_suggestions)
        self.popup = None

    def keyboard_on_key_down(self, window, keycode, text, modifiers):
        key = keycode[1]
        if key == 'down':
            self.focus_next('down')
        elif key == 'up':
            self.focus_previous('up')
        elif key == 'right':
            self.focus_next('right')
        elif key == 'left':
            self.focus_previous('left')
        else:
            super().keyboard_on_key_down(window, keycode, text, modifiers)
        return True

    def show_suggestions(self, instance, value):
        if not value or len(value) < 3:
            return

        conn = sqlite3.connect('forest_data.db')
        cursor = conn.cursor()
        cursor.execute('''
            SELECT value FROM molodniki_suggestions
            WHERE column_index = ? AND value LIKE ?
            ORDER BY LENGTH(value) ASC, value ASC
            LIMIT 1
        ''', (self.col_index, f'{value}%'))
        results = cursor.fetchall()
        conn.close()

        if results:
            self.text = results[0][0]

    def get_table_screen(self):
        return App.get_running_app().root.get_screen('molodniki')

    def focus_next(self, direction):
        table_screen = self.get_table_screen()
        if direction == 'right' and self.next_widget:
            self.next_widget.focus = True
        elif direction == 'down':
            next_row = self.row_index + 1
            if next_row < len(table_screen.inputs):
                table_screen.inputs[next_row][self.col_index].focus = True

    def focus_previous(self, direction):
        table_screen = self.get_table_screen()
        if direction == 'left' and self.prev_widget:
            self.prev_widget.focus = True
        elif direction == 'up':
            prev_row = self.row_index - 1
            if prev_row >= 0:
                table_screen.inputs[prev_row][self.col_index].focus = True

class Joypad(FloatLayout):
    def __init__(self, table_screen, **kwargs):
        super().__init__(**kwargs)
        self.table_screen = table_screen
        self.size_hint = (None, None)
        self.size = (140, 140)
        self.pos_hint = {'right': 0.95, 'y': 0.02}

        with self.canvas.before:
            Color(0.2, 0.2, 0.2, 1)
            self.bg_rect = RoundedRectangle(
                pos=self.pos,
                size=self.size,
                radius=[70]
            )

        arrows = [
            ('▲', (0.5, 0.7), 'up', (60, 40)),
            ('▼', (0.5, 0.3), 'down', (60, 40)),
            ('◄', (0.3, 0.5), 'left', (40, 60)),
            ('►', (0.7, 0.5), 'right', (40, 60))
        ]

        for symbol, pos, direction, size in arrows:
            btn = ModernButton(
                text=symbol,
                size_hint=(None, None),
                size=size,
                pos_hint={'center_x': pos[0], 'center_y': pos[1]},
                bg_color=(0.1, 0.1, 0.1, 1),
                color=(1, 1, 1, 1),
                font_size='20sp',
                bold=True
            )
            btn.bind(on_press=lambda x, d=direction: self.move_focus(d))
            self.add_widget(btn)

        self.bind(pos=self.update_bg, size=self.update_bg)

    def update_bg(self, *args):
        self.bg_rect.pos = self.pos
        self.bg_rect.size = self.size

    def move_focus(self, direction):
        current = self.table_screen.focused_cell
        if not current: return
        row, col = current

        if direction == 'up': row = max(0, row-1)
        elif direction == 'down': row = min(len(self.table_screen.inputs)-1, row+1)
        elif direction == 'left': col = max(0, col-1)
        elif direction == 'right': col = min(5, col+1)  # 6 столбцов (0-5)

        self.table_screen.focused_cell = [row, col]
        inp = self.table_screen.inputs[row][col]
        inp.focus = True
        inp.cursor = (len(inp.text), 0)
        Clock.schedule_once(lambda dt: self._update_cursor(inp), 0.01)

    def _update_cursor(self, inp):
        inp.focus = True
        inp.cursor = (len(inp.text), 0)
        inp.text = inp.text

class ExtendedMolodnikiTableScreen(Screen):
    current_page = NumericProperty(0)
    total_pages = NumericProperty(1)
    unsaved_changes = BooleanProperty(False)
    focused_cell = ListProperty([0, 0])
    edit_mode = BooleanProperty(False)
    current_section = StringProperty("")
    current_quarter = StringProperty("")
    current_plot = StringProperty("")
    current_forestry = StringProperty("")
    current_radius = StringProperty("5.64")
    MAX_PAGES = 200

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.theme_manager = App.get_running_app().theme_manager
        self.reports_dir = "reports"
        os.makedirs(self.reports_dir, exist_ok=True)
        self.db_name = 'forest_data.db'
        self.rows_per_page = 30
        self.page_data = {}
        self.setup_database()
        self.create_ui()
        self.load_existing_data()
        Window.bind(on_key_down=self.key_action)

    def key_action(self, window, key, scancode, codepoint, modifier):
        if key == 115 and 'ctrl' in modifier:
            self.save_current_page()

    def setup_database(self):
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()

        # Создаем таблицу для хранения данных молодняков
        cursor.execute('''CREATE TABLE IF NOT EXISTS molodniki_data (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        page_number INTEGER,
                        row_index INTEGER,
                        nn INTEGER,
                        gps_point TEXT,
                        predmet_uhoda TEXT,
                        radius REAL DEFAULT 5.64,
                        primechanie TEXT,
                        section_name TEXT,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')

        # Создаем индексы для быстрого поиска
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_molodniki_data_page ON molodniki_data (page_number, row_index)')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_molodniki_data_section ON molodniki_data (section_name)')

        # Создаем таблицу для хранения пород (множественные породы на одну запись)
        cursor.execute('''CREATE TABLE IF NOT EXISTS molodniki_breeds (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        molodniki_data_id INTEGER,
                        breed_name TEXT,
                        breed_type TEXT, -- 'coniferous' или 'deciduous'
                        do_05 INTEGER DEFAULT 0,
                        _05_15 INTEGER DEFAULT 0,
                        bolee_15 INTEGER DEFAULT 0,
                        density INTEGER DEFAULT 0,
                        height REAL DEFAULT 0.0,
                        age INTEGER DEFAULT 0,
                        composition_coefficient REAL DEFAULT 0.0,
                        FOREIGN KEY(molodniki_data_id) REFERENCES molodniki_data(id) ON DELETE CASCADE)''')

        # Создаем индекс для поиска данных пород
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_molodniki_breeds ON molodniki_breeds (molodniki_data_id)')

        # Создаем таблицу для хранения итогов по страницам
        cursor.execute('''CREATE TABLE IF NOT EXISTS molodniki_totals (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        page_number INTEGER,
                        section_name TEXT,
                        total_composition TEXT,
                        total_area REAL DEFAULT 0.0,
                        avg_age REAL DEFAULT 0.0,
                        avg_density REAL DEFAULT 0.0,
                        avg_height REAL DEFAULT 0.0,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')

        cursor.execute('CREATE INDEX IF NOT EXISTS idx_molodniki_totals_page ON molodniki_totals (page_number, section_name)')

        # Создаем таблицу для хранения данных пород (JSON)
        cursor.execute('''CREATE TABLE IF NOT EXISTS molodniki_suggestions (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        column_index INTEGER,
                        value TEXT,
                        UNIQUE(column_index, value))''')

        cursor.execute('CREATE INDEX IF NOT EXISTS idx_molodniki_suggestions ON molodniki_suggestions (column_index, value)')

        conn.commit()
        conn.close()

    def create_ui(self):
        main_layout = BoxLayout(orientation='horizontal', padding=10, spacing=10)

        with self.canvas.before:
            self.bg_color = Color(1, 1, 1, 1)
            self.bg_rect = Rectangle(pos=self.pos, size=self.size)
            self.bind(pos=self._update_bg, size=self._update_bg)

        self._update_background(self.theme_manager.current_theme)

        # Табличная часть (левая панель) - уменьшаем для видимости кнопок
        table_panel = BoxLayout(orientation='vertical', size_hint_x=0.75)

        # Заголовок участка и адресная строка
        header_layout = BoxLayout(orientation='vertical', size_hint=(1, None), height=80, spacing=5)

        # Адресная строка с кнопками
        address_layout = BoxLayout(orientation='horizontal', size_hint=(1, None), height=35, spacing=5)

        # Кнопки Квартал, Выдел, Лесничество, Радиус
        quarter_btn = ModernButton(
            text='Квартал',
            bg_color=get_color_from_hex('#87CEEB'),
            size_hint=(None, None),
            size=(100, 35),
            font_size='14sp'
        )
        quarter_btn.bind(on_press=self.show_quarter_popup)

        plot_btn = ModernButton(
            text='Выдел',
            bg_color=get_color_from_hex('#87CEEB'),
            size_hint=(None, None),
            size=(80, 35),
            font_size='14sp'
        )
        plot_btn.bind(on_press=self.show_plot_popup)

        forestry_btn = ModernButton(
            text='Лесничество',
            bg_color=get_color_from_hex('#87CEEB'),
            size_hint=(None, None),
            size=(120, 35),
            font_size='14sp'
        )
        forestry_btn.bind(on_press=self.show_forestry_popup)

        radius_btn = ModernButton(
            text='Радиус',
            bg_color=get_color_from_hex('#FF8C00'),
            size_hint=(None, None),
            size=(90, 35),
            font_size='14sp'
        )
        radius_btn.bind(on_press=self.show_radius_popup)

        address_layout.add_widget(quarter_btn)
        address_layout.add_widget(plot_btn)
        address_layout.add_widget(forestry_btn)
        address_layout.add_widget(radius_btn)

        # Адресная строка (текстовое поле для отображения адреса)
        self.address_label = Label(
            text=f"Адрес: {self.current_quarter} кв. {self.current_plot} выд. {self.current_forestry}",
            font_name='Roboto',
            size_hint=(1, None),
            height=35,
            color=self._get_text_color(),
            halign='left',
            valign='middle'
        )
        self.address_label.bind(size=self.address_label.setter('text_size'))
        address_layout.add_widget(self.address_label)

        header_layout.add_widget(address_layout)

        # Заголовок участка
        self.section_label = Label(
            text=f"Молодняки - Участок: {self.current_section}",
            font_name='Roboto',
            size_hint=(1, None),
            height=30,
            color=self._get_text_color()
        )
        header_layout.add_widget(self.section_label)

        table_panel.add_widget(header_layout)

        # Пагинация
        pagination = BoxLayout(size_hint_y=None, height=40, spacing=5)
        self.page_label = Label(
            text=f'Страница {self.current_page+1} из {self.total_pages}',
            size_hint_x=0.4,
            font_name='Roboto',
            color=self._get_text_color()
        )
        prev_btn = ModernButton(
            text='← Предыдущая',
            size_hint_x=0.3,
            bg_color=get_color_from_hex('#00FF00'),
            color=self._get_text_color()
        )
        prev_btn.bind(on_press=lambda x: self.change_page(-1))
        next_btn = ModernButton(
            text='Следующая →',
            size_hint_x=0.3,
            bg_color=get_color_from_hex('#00FF00'),
            color=self._get_text_color()
        )
        next_btn.bind(on_press=lambda x: self.change_page(1))
        pagination.add_widget(prev_btn)
        pagination.add_widget(self.page_label)
        pagination.add_widget(next_btn)
        table_panel.add_widget(pagination)

        # Основная таблица (6 столбцов: основные данные)
        scroll = ScrollView(do_scroll_x=True, do_scroll_y=True, bar_width=10)
        self.table = GridLayout(cols=6, size_hint=(None, None), spacing=1)
        self.table.bind(minimum_height=self.table.setter('height'),
                       minimum_width=self.table.setter('width'))

        # Заголовки столбцов (6 столбцов: основные данные)
        headers = [
            '№ППР', 'GPS точка', 'Предмет ухода', 'Порода', 'Примечания', 'Тип Леса'
        ]

        self.header_bgs = []
        for i, header in enumerate(headers):
            lbl = Label(
                text=header,
                size_hint_y=None,
                height=30,
                font_name='Roboto',
                bold=True,
                halign='center',
                size_hint_x=None,
                width=120,  # Все столбцы одинаковой ширины
                color=(0, 0, 0, 1)  # Черный цвет для заголовков
            )
            with lbl.canvas.before:
                Color(rgba=get_color_from_hex('#00FF00'))
                bg = Rectangle(pos=lbl.pos, size=lbl.size)
                self.header_bgs.append(bg)
            lbl.bind(pos=lambda i,v, b=bg: setattr(b, 'pos', i.pos), size=lambda i,v, b=bg: setattr(b, 'size', i.size))
            self.table.add_widget(lbl)

        # Создаем строки таблицы (6 столбцов: основные данные)
        self.inputs = []
        for row_idx in range(self.rows_per_page):
            row = []
            for col_idx in range(6):
                inp = AutoCompleteTextInput(multiline=False, size_hint_y=None, height=30)
                inp.row_index = row_idx
                inp.col_index = col_idx
                inp.bind(focus=self.update_focus)
                inp.font_name = 'Roboto'
                inp.size_hint_x = None
                inp.width = 120  # Все столбцы одинаковой ширины

                # Настройка фильтров ввода для числовых полей
                if col_idx == 0:  # №ППР
                    inp.input_filter = 'int'
                    inp.bind(focus=self.auto_fill_nn)
                elif col_idx == 3:  # Порода - открываем popup выбора типа
                    inp.bind(focus=self.show_breed_popup)

                row.append(inp)
                self.table.add_widget(inp)
            self.inputs.append(row)

        # Добавляем строку "ИТОГО" (6 столбцов)
        self.total_labels = []
        for col_idx in range(6):
            if col_idx == 0:
                lbl = Label(text="ИТОГО:", font_name='Roboto', bold=True,
                           size_hint_y=None, height=30, color=(0, 0, 0, 1),
                           size_hint_x=None, width=120)
            else:
                lbl = Label(text="0", font_name='Roboto', bold=True,
                           size_hint_y=None, height=30, color=(0, 0, 0, 1),
                           size_hint_x=None, width=120)
                lbl.col_index = col_idx
                self.total_labels.append(lbl)
            self.table.add_widget(lbl)

        scroll.add_widget(self.table)
        table_panel.add_widget(scroll)
        main_layout.add_widget(table_panel)

        # Правая панель управления (увеличиваем для видимости кнопок)
        control_panel = BoxLayout(
            orientation='vertical',
            size_hint_x=0.25,
            spacing=10,
            padding=[0, 10, 0, 0]
        )

        controls = GridLayout(
            cols=1,
            size_hint_y=None,
            height=400,
            spacing=8,
            pos_hint={'top': 1}
        )

        button_handlers = {
            'Радиус': self.show_radius_popup,
            'Сохранить Excel': self.show_save_dialog,
            'Сохранить Word': self.save_to_word,
            'Сохранить страницу': self.save_current_page,
            'Загрузить': self.load_section,
            'Открыть папку': self.open_excel_file,
            'Редакт. режим': self.toggle_edit_mode,
            'Очистить': self.clear_table_data,
            'В меню': self.go_back
        }

        button_colors = {
            'Радиус': '#FF8C00',
            'Сохранить Excel': '#00FF00',
            'Сохранить Word': '#4169E1',
            'Сохранить страницу': '#00FFFF',
            'Загрузить': '#006400',
            'Открыть папку': '#0000FF',
            'Редакт. режим': '#00BFFF',
            'Очистить': '#800000',
            'В меню': '#FF0000'
        }

        for text, color in button_colors.items():
            btn = ModernButton(
                text=text,
                bg_color=get_color_from_hex(color),
                size_hint=(None, None),
                size=(280, 50),
                pos_hint={'center_x': 0.5}
            )
            btn.bind(on_press=button_handlers[text])
            controls.add_widget(btn)

        control_panel.add_widget(controls)

        # Джойстик - центрируем внизу
        joypad_container = BoxLayout(
            size_hint=(1, None),
            height=150,
            padding=[0, 20, 0, 0]
        )

        self.joypad = Joypad(self)
        self.joypad.pos_hint = {'center_x': 0.5, 'center_y': 0.5}
        joypad_container.add_widget(self.joypad)

        spacer = BoxLayout(size_hint_y=1)
        control_panel.add_widget(spacer)
        control_panel.add_widget(joypad_container)

        main_layout.add_widget(control_panel)
        self.add_widget(main_layout)

    def _update_bg(self, instance, value):
        self.bg_rect.pos = self.pos
        self.bg_rect.size = self.size

    def _update_background(self, theme):
        if theme['type'] == 'image':
            try:
                self.bg_color.rgba = (1, 1, 1, 1)
                self.bg_rect.texture = CoreImage(theme['background']).texture
            except Exception as e:
                print(f"Error loading background image: {str(e)}")
        else:
            self.bg_color.rgba = theme['background']
            self.bg_rect.texture = None

    def _get_text_color(self):
        theme = self.theme_manager.current_theme
        if theme['type'] == 'image':
            return get_color_from_hex('#FFFFFF')
        else:
            return get_color_from_hex(theme['text_color'])

    def update_section_label(self):
        self.section_label.text = f"Молодняки - Участок: {self.current_section}"

    def toggle_edit_mode(self, instance):
        self.edit_mode = not self.edit_mode
        instance.bg_color = get_color_from_hex('#FFA500' if self.edit_mode else '#00BFFF')

    def update_focus(self, instance, value):
        if value:
            self.focused_cell = [instance.row_index, instance.col_index]

    def move_focus(self, direction):
        current = self.focused_cell
        if not current: return
        row, col = current

        if direction == 'up': row = max(0, row-1)
        elif direction == 'down': row = min(len(self.inputs)-1, row+1)
        elif direction == 'left': col = max(0, col-1)
        elif direction == 'right': col = min(5, col+1)  # 6 столбцов (0-5)

        self.focused_cell = [row, col]
        inp = self.inputs[row][col]
        inp.focus = True
        inp.cursor = (len(inp.text), 0)

    def auto_fill_nn(self, instance, value):
        if self.edit_mode: return
        if value and instance.focus:
            try:
                current_number = int(instance.text)
            except ValueError:
                current_number = 0
            for row_idx, row in enumerate(self.inputs):
                if row_idx > instance.row_index:
                    try:
                        prev_num = int(self.inputs[row_idx-1][0].text)
                        row[0].text = str(prev_num + 1)
                    except (ValueError, IndexError):
                        pass

    def update_plot_total(self, instance, value):
        """Обновляем итог по площадке при изменении данных"""
        row_idx = instance.row_index
        if row_idx >= len(self.plot_totals):
            return

        # Собираем данные пород из столбца "Порода"
        breeds_text = self.inputs[row_idx][3].text  # Столбец "Порода"
        breeds_data = self.parse_breeds_data(breeds_text)

        if not breeds_data:
            self.plot_totals[row_idx].text = ""
            return

        # Вычисляем итоговые значения для площадки
        total_density = 0
        total_height = 0.0
        total_age = 0
        breed_count = 0
        breed_names = []

        for breed_info in breeds_data:
            breed_count += 1
            breed_name = breed_info.get('name', 'Неизвестная')
            breed_names.append(breed_name)

            # Для хвойных пород густота = сумма (до 0.5м + 0.5-1.5м + >1.5м)
            if breed_info.get('type') == 'coniferous':
                coniferous_density = (breed_info.get('do_05', 0) +
                                    breed_info.get('05_15', 0) +
                                    breed_info.get('bolee_15', 0))
                if coniferous_density > 0:
                    total_density += coniferous_density
            elif 'density' in breed_info and breed_info['density']:
                total_density += breed_info['density']

            if 'height' in breed_info and breed_info['height']:
                total_height += breed_info['height']
            if 'age' in breed_info and breed_info['age']:
                total_age += breed_info['age']

        # Формируем текст итогов с списком пород и общей густотой
        summary_parts = []
        if breed_names:
            summary_parts.append(f"Породы: {', '.join(breed_names)}")
        if total_density > 0:
            summary_parts.append(f"Общая густота: {total_density}")
        if total_height > 0:
            avg_height = total_height / breed_count if breed_count > 0 else 0
            summary_parts.append(f"Средняя высота: {avg_height:.1f}м")
        if total_age > 0:
            avg_age = total_age / breed_count if breed_count > 0 else 0
            summary_parts.append(f"Средний возраст: {avg_age:.1f} лет")

        self.plot_totals[row_idx].text = "\n".join(summary_parts) if summary_parts else ""

    def calculate_total_breeds(self, instance, value):
        """Автоматический расчет 'Всего тыс.шт/га' при изменении пород"""
        row_idx = instance.row_index
        row = self.inputs[row_idx]

        try:
            total = 0
            for col_idx in range(10):  # Столбцы пород A-J
                if row[col_idx].text:
                    total += float(row[col_idx].text)

            row[10].text = f"{total:.1f}" if total > 0 else ""
        except (ValueError, IndexError):
            pass

        # Обновляем итоги
        self.update_totals()

    def show_breed_popup(self, instance, value):
        """Показать popup для выбора типа породы"""
        if not value: return

        content = BoxLayout(orientation='vertical', spacing=10, padding=10)

        # Кнопки выбора типа породы
        type_layout = BoxLayout(orientation='horizontal', spacing=10)
        coniferous_btn = ModernButton(
            text='Хвойные',
            bg_color=get_color_from_hex('#228B22'),
            size_hint=(0.5, None),
            height=50
        )
        deciduous_btn = ModernButton(
            text='Лиственные',
            bg_color=get_color_from_hex('#32CD32'),
            size_hint=(0.5, None),
            height=50
        )
        type_layout.add_widget(coniferous_btn)
        type_layout.add_widget(deciduous_btn)
        content.add_widget(type_layout)

        # Кнопка отмены
        cancel_btn = ModernButton(
            text='Отмена',
            bg_color=get_color_from_hex('#FF6347'),
            size_hint=(0.5, 1),
            height=50
        )
        content.add_widget(cancel_btn)

        popup = Popup(
            title="Выберите тип породы",
            content=content,
            size_hint=(0.8, 0.5)
        )

        def select_coniferous(btn):
            self.show_breed_selection_popup(instance, 'coniferous')
            popup.dismiss()

        def select_deciduous(btn):
            self.show_breed_selection_popup(instance, 'deciduous')
            popup.dismiss()

        coniferous_btn.bind(on_press=select_coniferous)
        deciduous_btn.bind(on_press=select_deciduous)
        cancel_btn.bind(on_press=popup.dismiss)

        popup.open()

    def show_breed_selection_popup(self, instance, breed_type):
        """Показать popup для выбора конкретной породы из словаря"""
        content = BoxLayout(orientation='vertical', spacing=10, padding=10)

        # Заголовок
        title_label = Label(
            text=f"Выберите {'хвойную' if breed_type == 'coniferous' else 'лиственную'} породу",
            font_name='Roboto',
            bold=True,
            size_hint=(1, None),
            height=30
        )
        content.add_widget(title_label)

        # Списки пород
        if breed_type == 'coniferous':
            breeds = [
                'Сосна семенная', 'Сосна культуры', 'Сосна подрост',
                'Ель семенная', 'Ель культуры', 'Ель подрост',
                'Пихта семенная', 'Пихта культуры', 'Пихта подрост',
                'Кедр семенной', 'Кедр культуры', 'Кедр подрост',
                'Лиственница семенная', 'Лиственница культуры', 'Лиственница подрост'
            ]
        else:
            breeds = [
                'Берёза', 'Осина', 'Ольха чёрная', 'Ольха серая',
                'Ива', 'Ива кустарниковая'
            ]

        # ScrollView для списка пород
        scroll = ScrollView(size_hint=(1, None), height=300)
        breeds_layout = GridLayout(cols=1, spacing=5, size_hint_y=None)
        breeds_layout.bind(minimum_height=breeds_layout.setter('height'))

        for breed in breeds:
            btn = ModernButton(
                text=breed,
                bg_color=get_color_from_hex('#87CEEB'),
                size_hint=(1, None),
                height=50,
                font_size='14sp'
            )
            btn.bind(on_press=lambda x, b=breed: self.select_breed(instance, breed_type, b))
            breeds_layout.add_widget(btn)

        scroll.add_widget(breeds_layout)
        content.add_widget(scroll)

        # Кнопка "Другая порода"
        other_btn = ModernButton(
            text='Другая порода',
            bg_color=get_color_from_hex('#DDA0DD'),
            size_hint=(1, None),
            height=50
        )
        other_btn.bind(on_press=lambda x: self.select_breed(instance, breed_type, 'other'))
        content.add_widget(other_btn)

        # Кнопка отмены
        cancel_btn = ModernButton(
            text='Отмена',
            bg_color=get_color_from_hex('#FF6347'),
            size_hint=(1, None),
            height=50
        )
        content.add_widget(cancel_btn)

        popup = Popup(
            title=f"Выбор {'хвойной' if breed_type == 'coniferous' else 'лиственной'} породы",
            content=content,
            size_hint=(0.85, 0.85)
        )

        cancel_btn.bind(on_press=popup.dismiss)
        popup.open()

    def select_breed(self, instance, breed_type, selected_breed):
        """Обработка выбора породы"""
        if selected_breed == 'other':
            # Показываем popup для ввода названия другой породы
            self.show_custom_breed_popup(instance, breed_type)
        else:
            # Показываем popup с параметрами породы, передавая название выбранной породы
            self.show_breed_details_popup(instance, breed_type, selected_breed)

    def show_breed_details_popup(self, instance, breed_type, selected_breed=None):
        """Показать popup для управления множественными породами"""
        content = BoxLayout(orientation='vertical', spacing=10, padding=10)

        # Заголовок
        title_label = Label(
            text=f"Управление породами - {selected_breed}",
            font_name='Roboto',
            bold=True,
            size_hint=(1, None),
            height=30
        )
        content.add_widget(title_label)

        # Список существующих пород в этой строке
        # Если instance.text содержит JSON, парсим его, иначе получаем данные из базы
        existing_breeds = self.parse_breeds_data(instance.text)
        if not existing_breeds:
            # Пытаемся получить данные из сохраненных данных страницы
            row_idx = instance.row_index
            if self.current_page in self.page_data and row_idx < len(self.page_data[self.current_page]):
                saved_text = self.page_data[self.current_page][row_idx][3]  # Столбец "Порода"
                existing_breeds = self.parse_breeds_data(saved_text)

        # Отображаем список существующих пород
        if existing_breeds:
            breeds_list_label = Label(
                text="Уже добавленные породы:",
                font_name='Roboto',
                bold=True,
                size_hint=(1, None),
                height=30,
                color=(0.3, 0.3, 0.3, 1)
            )
            content.add_widget(breeds_list_label)

            # ScrollView для списка пород
            breeds_scroll = ScrollView(size_hint=(1, None), height=80)
            breeds_list_layout = BoxLayout(orientation='vertical', spacing=5, size_hint_y=None)
            breeds_list_layout.bind(minimum_height=breeds_list_layout.setter('height'))

            for i, breed_info in enumerate(existing_breeds):
                breed_name = breed_info.get('name', 'Неизвестная')
                breed_type = breed_info.get('type', 'unknown')
                params = []
                if 'density' in breed_info and breed_info['density']:
                    params.append(f"Густота: {breed_info['density']}")
                if 'height' in breed_info and breed_info['height']:
                    params.append(f"Высота: {breed_info['height']}м")
                if 'age' in breed_info and breed_info['age']:
                    params.append(f"Возраст: {breed_info['age']} лет")
                if breed_type == 'coniferous':
                    if 'do_05' in breed_info and breed_info['do_05']:
                        params.append(f"До 0.5м: {breed_info['do_05']}")
                    if '05_15' in breed_info and breed_info['05_15']:
                        params.append(f"0.5-1.5м: {breed_info['05_15']}")
                    if 'bolee_15' in breed_info and breed_info['bolee_15']:
                        params.append(f">1.5м: {breed_info['bolee_15']}")

                breed_text = f"{i+1}. {breed_name}: {'; '.join(params)}" if params else f"{i+1}. {breed_name}"
                breed_label = Label(
                    text=breed_text,
                    font_name='Roboto',
                    size_hint=(1, None),
                    height=25,
                    color=(0.2, 0.2, 0.2, 1),
                    text_size=(None, None),
                    halign='left',
                    valign='top'
                )
                breed_label.bind(size=lambda *args: setattr(breed_label, 'text_size', (breed_label.width, None)))
                breeds_list_layout.add_widget(breed_label)

            breeds_scroll.add_widget(breeds_list_layout)
            content.add_widget(breeds_scroll)
        else:
            breeds_list_label = Label(
                text="Породы ещё не добавлены",
                font_name='Roboto',
                size_hint=(1, None),
                height=30,
                color=(0.5, 0.5, 0.5, 1)
            )
            content.add_widget(breeds_list_label)

        # Поля ввода для параметров породы
        fields_layout = GridLayout(cols=2, spacing=5, size_hint=(1, None), height=200)

        if breed_type == 'coniferous':
            # Хвойные: до 0.5м, 0.5-1.5м, >1.5м, высота, густота, возраст
            fields = [
                ('До 0.5м:', 'do_05'),
                ('0.5-1.5м:', '05_15'),
                ('>1.5м:', 'bolee_15'),
                ('Высота (м):', 'height'),
                ('Густота:', 'density'),
                ('Возраст (лет):', 'age')
            ]
        else:
            # Лиственные: густота, высота, возраст
            fields = [
                ('Густота:', 'density'),
                ('Высота (м):', 'height'),
                ('Возраст (лет):', 'age')
            ]

        self.breed_inputs = {}
        for label_text, field_key in fields:
            lbl = Label(text=label_text, font_name='Roboto', size_hint=(None, None), size=(100, 30))
            inp = TextInput(multiline=False, size_hint=(None, None), size=(100, 30))
            if field_key in ['density', 'age']:
                inp.input_filter = 'int'
            elif field_key == 'height':
                inp.input_filter = 'float'
            elif field_key in ['do_05', '05_15', 'bolee_15']:
                inp.input_filter = 'int'
                # Для хвойных пород автоматически рассчитываем густоту
                if breed_type == 'coniferous':
                    inp.bind(text=self.update_coniferous_density)
            fields_layout.add_widget(lbl)
            fields_layout.add_widget(inp)
            self.breed_inputs[field_key] = inp

        content.add_widget(fields_layout)

        # Кнопки управления
        btn_layout = BoxLayout(orientation='horizontal', spacing=5, size_hint=(1, None), height=50)
        add_btn = ModernButton(
            text='Добавить породу',
            bg_color=get_color_from_hex('#00FF00'),
            size_hint=(0.25, 1),
            height=50
        )
        save_btn = ModernButton(
            text='Сохранить',
            bg_color=get_color_from_hex('#32CD32'),
            size_hint=(0.25, 1),
            height=50
        )
        view_btn = ModernButton(
            text='Просмотр',
            bg_color=get_color_from_hex('#87CEEB'),
            size_hint=(0.25, 1),
            height=50
        )
        cancel_btn = ModernButton(
            text='Отмена',
            bg_color=get_color_from_hex('#FF6347'),
            size_hint=(0.25, 1),
            height=50
        )
        btn_layout.add_widget(add_btn)
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(view_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title=f"Параметры породы: {selected_breed}",
            content=content,
            size_hint=(0.9, 0.95)
        )

        def add_breed(btn):
            # Собираем данные породы
            breed_data = {
                'name': selected_breed,
                'type': breed_type
            }

            for key, inp in self.breed_inputs.items():
                if inp.text.strip():
                    try:
                        if key in ['density', 'age']:
                            breed_data[key] = int(inp.text)
                        elif key == 'height':
                            breed_data[key] = float(inp.text)
                        else:
                            breed_data[key] = float(inp.text)
                    except ValueError:
                        breed_data[key] = inp.text

            # Добавляем к существующим породам
            existing_breeds = self.parse_breeds_data(instance.text)
            existing_breeds.append(breed_data)

            # Сохраняем как JSON массив
            instance.text = json.dumps(existing_breeds, ensure_ascii=False, indent=2)

            # Обновляем итог по площадке
            self.update_plot_total(instance, instance.text)

            # Очищаем поля для ввода следующей породы
            for inp in self.breed_inputs.values():
                inp.text = ''

            # После добавления породы возвращаемся к выбору типа породы для следующей
            self.show_breed_popup(instance, True)

            self.show_success(f"Порода '{selected_breed}' добавлена! Выберите тип следующей породы.")

        def save_breeds(btn):
            # Сохраняем все данные по площадке и закрываем popup
            existing_breeds = self.parse_breeds_data(instance.text)
            # Если пород нет, создаем пустой массив
            if not existing_breeds:
                existing_breeds = []
            instance.text = json.dumps(existing_breeds, ensure_ascii=False, indent=2)
            # Обновляем итог по площадке с сохраненными данными
            self.update_plot_total(instance, instance.text)
            self.show_success("Все данные по площадке сохранены!")
            popup.dismiss()

        def view_breeds(btn):
            # Показываем список всех пород в этой строке
            popup.dismiss()
            self.show_breeds_list_popup(instance)

        add_btn.bind(on_press=add_breed)
        save_btn.bind(on_press=save_breeds)
        view_btn.bind(on_press=view_breeds)
        cancel_btn.bind(on_press=popup.dismiss)

        popup.open()

    def show_breeds_list_popup(self, instance):
        """Показать popup со списком всех пород в этой строке"""
        content = BoxLayout(orientation='vertical', spacing=10, padding=10)

        # Заголовок
        title_label = Label(
            text="Список пород в этой строке",
            font_name='Roboto',
            bold=True,
            size_hint=(1, None),
            height=30
        )
        content.add_widget(title_label)

        # Получаем список пород
        breeds_data = self.parse_breeds_data(instance.text)

        if not breeds_data:
            no_breeds_label = Label(
                text="Породы не найдены",
                font_name='Roboto',
                size_hint=(1, None),
                height=50,
                color=(0.5, 0.5, 0.5, 1)
            )
            content.add_widget(no_breeds_label)
        else:
            # ScrollView для списка пород
            scroll = ScrollView(size_hint=(1, None), height=300)
            breeds_layout = GridLayout(cols=1, spacing=5, size_hint_y=None)
            breeds_layout.bind(minimum_height=breeds_layout.setter('height'))

            for i, breed_info in enumerate(breeds_data):
                # Создаем карточку породы
                breed_card = BoxLayout(
                    orientation='vertical',
                    size_hint=(1, None),
                    height=120,
                    padding=5
                )
                breed_card.canvas.before.clear()
                with breed_card.canvas.before:
                    Color(rgba=get_color_from_hex('#E8F4FD'))
                    Rectangle(pos=breed_card.pos, size=breed_card.size)
                    Color(rgba=get_color_from_hex('#B0BEC5'))
                    Line(rectangle=(breed_card.x, breed_card.y, breed_card.width, breed_card.height), width=1)

                # Название породы
                name_label = Label(
                    text=f"{i+1}. {breed_info.get('name', 'Неизвестная порода')}",
                    font_name='Roboto',
                    bold=True,
                    size_hint=(1, None),
                    height=25,
                    color=(0, 0, 0, 1)
                )
                breed_card.add_widget(name_label)

                # Параметры породы
                params_text = []
                if 'density' in breed_info and breed_info['density']:
                    params_text.append(f"Густота: {breed_info['density']}")
                if 'height' in breed_info and breed_info['height']:
                    params_text.append(f"Высота: {breed_info['height']}м")
                if 'age' in breed_info and breed_info['age']:
                    params_text.append(f"Возраст: {breed_info['age']} лет")
                if 'do_05' in breed_info and breed_info['do_05']:
                    params_text.append(f"До 0.5м: {breed_info['do_05']}")
                if '05_15' in breed_info and breed_info['05_15']:
                    params_text.append(f"0.5-1.5м: {breed_info['05_15']}")
                if 'bolee_15' in breed_info and breed_info['bolee_15']:
                    params_text.append(f">1.5м: {breed_info['bolee_15']}")

                params_label = Label(
                    text="; ".join(params_text) if params_text else "Нет параметров",
                    font_name='Roboto',
                    size_hint=(1, None),
                    height=40,
                    color=(0.3, 0.3, 0.3, 1),
                    text_size=(None, None),
                    halign='left',
                    valign='top'
                )
                params_label.bind(size=lambda *args: setattr(params_label, 'text_size', (params_label.width, None)))
                breed_card.add_widget(params_label)

                # Кнопки управления
                btn_layout = BoxLayout(orientation='horizontal', spacing=5, size_hint=(1, None), height=30)
                edit_btn = ModernButton(
                    text='Изменить',
                    bg_color=get_color_from_hex('#87CEEB'),
                    size_hint=(0.5, 1),
                    font_size='12sp'
                )
                delete_btn = ModernButton(
                    text='Удалить',
                    bg_color=get_color_from_hex('#FF6347'),
                    size_hint=(0.5, 1),
                    font_size='12sp'
                )
                btn_layout.add_widget(edit_btn)
                delete_btn.bind(on_press=lambda x, idx=i: self.delete_breed_from_list(instance, idx))
                btn_layout.add_widget(delete_btn)
                breed_card.add_widget(btn_layout)

                # Привязываем события
                def edit_breed(btn, idx=i):
                    self.edit_breed_in_list(instance, idx)

                edit_btn.bind(on_press=edit_breed)

                breeds_layout.add_widget(breed_card)

            scroll.add_widget(breeds_layout)
            content.add_widget(scroll)

        # Кнопка закрытия
        close_btn = ModernButton(
            text='Закрыть',
            bg_color=get_color_from_hex('#808080'),
            size_hint=(1, None),
            height=50
        )
        content.add_widget(close_btn)

        popup = Popup(
            title="Управление породами",
            content=content,
            size_hint=(0.85, 0.9)
        )

        close_btn.bind(on_press=popup.dismiss)
        popup.open()

    def edit_breed_in_list(self, instance, breed_index):
        """Редактировать породу в списке"""
        breeds_data = self.parse_breeds_data(instance.text)
        if 0 <= breed_index < len(breeds_data):
            breed_info = breeds_data[breed_index]
            # Здесь можно открыть popup для редактирования
            self.show_edit_breed_popup(instance, breed_index, breed_info)

    def delete_breed_from_list(self, instance, breed_index):
        """Удалить породу из списка"""
        breeds_data = self.parse_breeds_data(instance.text)
        if 0 <= breed_index < len(breeds_data):
            breed_name = breeds_data[breed_index].get('name', 'Неизвестная порода')
            # Удаляем породу
            breeds_data.pop(breed_index)
            # Сохраняем обновленный список
            instance.text = json.dumps(breeds_data, ensure_ascii=False, indent=2) if breeds_data else ''
            self.update_totals()
            self.show_success(f"Порода '{breed_name}' удалена!")

    def show_edit_breed_popup(self, instance, breed_index, breed_info):
        """Показать popup для редактирования породы"""
        content = BoxLayout(orientation='vertical', spacing=10, padding=10)

        # Заголовок
        title_label = Label(
            text=f"Редактирование породы: {breed_info.get('name', '')}",
            font_name='Roboto',
            bold=True,
            size_hint=(1, None),
            height=30
        )
        content.add_widget(title_label)

        # Поля для редактирования
        fields_layout = GridLayout(cols=2, spacing=5, size_hint=(1, None), height=200)

        breed_type = breed_info.get('type', 'deciduous')
        if breed_type == 'coniferous':
            fields = [
                ('До 0.5м:', 'do_05'),
                ('0.5-1.5м:', '05_15'),
                ('>1.5м:', 'bolee_15'),
                ('Высота (м):', 'height'),
                ('Густота:', 'density'),
                ('Возраст (лет):', 'age')
            ]
        else:
            fields = [
                ('Густота:', 'density'),
                ('Высота (м):', 'height'),
                ('Возраст (лет):', 'age')
            ]

        self.edit_inputs = {}
        for label_text, field_key in fields:
            lbl = Label(text=label_text, font_name='Roboto', size_hint=(None, None), size=(100, 30))
            inp = TextInput(
                multiline=False,
                size_hint=(None, None),
                size=(100, 30),
                text=str(breed_info.get(field_key, ''))
            )
            if field_key in ['density', 'age']:
                inp.input_filter = 'int'
            elif field_key == 'height':
                inp.input_filter = 'float'
            fields_layout.add_widget(lbl)
            fields_layout.add_widget(inp)
            self.edit_inputs[field_key] = inp

        content.add_widget(fields_layout)

        # Кнопки
        btn_layout = BoxLayout(orientation='horizontal', spacing=10, size_hint=(1, None), height=50)
        save_btn = ModernButton(
            text='Сохранить',
            bg_color=get_color_from_hex('#00FF00'),
            size_hint=(0.5, 1),
            height=50
        )
        cancel_btn = ModernButton(
            text='Отмена',
            bg_color=get_color_from_hex('#FF6347'),
            size_hint=(0.5, 1),
            height=50
        )
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title=f"Редактирование породы",
            content=content,
            size_hint=(0.9, 0.8)
        )

        def save_edit(btn):
            # Собираем обновленные данные
            for key, inp in self.edit_inputs.items():
                if inp.text.strip():
                    try:
                        if key in ['density', 'age']:
                            breed_info[key] = int(inp.text)
                        elif key == 'height':
                            breed_info[key] = float(inp.text)
                        else:
                            breed_info[key] = float(inp.text)
                    except ValueError:
                        breed_info[key] = inp.text
                else:
                    breed_info[key] = 0 if key in ['density', 'age', 'do_05', '05_15', 'bolee_15'] else 0.0

            # Обновляем породу в списке
            breeds_data = self.parse_breeds_data(instance.text)
            if 0 <= breed_index < len(breeds_data):
                breeds_data[breed_index] = breed_info
                instance.text = json.dumps(breeds_data, ensure_ascii=False, indent=2)
                self.update_totals()
                self.show_success("Порода обновлена!")
                popup.dismiss()

        save_btn.bind(on_press=save_edit)
        cancel_btn.bind(on_press=popup.dismiss)

        popup.open()

    def show_custom_breed_popup(self, instance, breed_type):
        """Показать popup для ввода названия другой породы"""
        content = BoxLayout(orientation='vertical', spacing=10, padding=10)

        # Заголовок
        title_label = Label(
            text="Введите название другой породы",
            font_name='Roboto',
            bold=True,
            size_hint=(1, None),
            height=30
        )
        content.add_widget(title_label)

        # Поле ввода названия породы
        self.custom_breed_input = TextInput(
            hint_text="Название породы",
            multiline=False,
            size_hint=(1, None),
            height=40,
            font_name='Roboto'
        )
        content.add_widget(self.custom_breed_input)

        # Кнопки
        btn_layout = BoxLayout(orientation='horizontal', spacing=10, size_hint=(1, None), height=50)
        save_btn = ModernButton(
            text='Сохранить',
            bg_color=get_color_from_hex('#00FF00'),
            size_hint=(0.5, 1),
            height=50
        )
        cancel_btn = ModernButton(
            text='Отмена',
            bg_color=get_color_from_hex('#FF6347'),
            size_hint=(0.5, 1),
            height=50
        )
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title=f"Ввод {'хвойной' if breed_type == 'coniferous' else 'лиственной'} породы",
            content=content,
            size_hint=(0.8, 0.6)
        )

        def save_custom_breed(btn):
            breed_name = self.custom_breed_input.text.strip()
            if breed_name:
                instance.text = breed_name
                self.show_breed_details_popup(instance, breed_type, breed_name)
                popup.dismiss()
            else:
                self.show_error("Название породы не может быть пустым!")

        save_btn.bind(on_press=save_custom_breed)
        cancel_btn.bind(on_press=popup.dismiss)

        popup.open()

    def update_coniferous_density(self, instance, value):
        """Автоматический расчет густоты для хвойных пород"""
        if 'density' in self.breed_inputs:
            density_input = self.breed_inputs['density']
            try:
                do_05 = int(self.breed_inputs.get('do_05', TextInput(text='0')).text or '0')
                _05_15 = int(self.breed_inputs.get('05_15', TextInput(text='0')).text or '0')
                bolee_15 = int(self.breed_inputs.get('bolee_15', TextInput(text='0')).text or '0')

                total_density = do_05 + _05_15 + bolee_15
                density_input.text = str(total_density) if total_density > 0 else ''
            except (ValueError, AttributeError):
                pass

    def update_address_label(self):
        """Обновить текст адресной строки"""
        address_parts = []
        if self.current_quarter:
            address_parts.append(f"{self.current_quarter} кв.")
        if self.current_plot:
            address_parts.append(f"{self.current_plot} выд.")
        if self.current_forestry:
            address_parts.append(self.current_forestry)

        address_text = "Адрес: " + " ".join(address_parts) if address_parts else "Адрес: не указан"
        self.address_label.text = address_text

    def load_existing_data(self):
        """Загружаем существующие данные из базы данных"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()

        try:
            # Получаем все уникальные номера страниц для текущего участка
            cursor.execute('''
                SELECT DISTINCT page_number FROM molodniki_data
                WHERE section_name = ?
                ORDER BY page_number
            ''', (self.current_section,))

            page_numbers = [row[0] for row in cursor.fetchall()]

            if page_numbers:
                # Загружаем данные для каждой страницы
                for page_num in page_numbers:
                    cursor.execute('''
                        SELECT row_index, nn, gps_point, predmet_uhoda, poroda, primechanie, radius
                        FROM molodniki_data
                        WHERE page_number = ? AND section_name = ?
                        ORDER BY row_index
                    ''', (page_num, self.current_section))

                    page_data = []
                    rows_data = cursor.fetchall()

                    # Создаем пустые строки для всех возможных позиций
                    for row_idx in range(self.rows_per_page):
                        row_data = ['', '', '', '', '', '']  # 6 столбцов
                        page_data.append(row_data)

                    # Заполняем данными из базы
                    for row_data in rows_data:
                        row_idx, nn, gps_point, predmet_uhoda, poroda, primechanie, radius = row_data
                        if row_idx < len(page_data):
                            page_data[row_idx] = [
                                str(nn) if nn is not None else '',
                                str(gps_point) if gps_point is not None else '',
                                str(predmet_uhoda) if predmet_uhoda is not None else '',
                                str(poroda) if poroda is not None else '',
                                str(primechanie) if primechanie is not None else '',
                                str(radius) if radius is not None else '',
                                ''  # Итого по площадке (пока пустой)
                            ]

                    self.page_data[page_num] = page_data

                # Устанавливаем текущую страницу на первую доступную
                self.current_page = min(page_numbers)
                self.load_page_data()
                self.update_pagination()

        except Exception as e:
            print(f"Error loading existing data: {e}")
        finally:
            conn.close()

    def load_page_data(self):
        for row in self.inputs:
            for inp in row:
                inp.text = ''

        if self.current_page in self.page_data:
            for i, row_data in enumerate(self.page_data[self.current_page]):
                if i >= len(self.inputs):
                    break
                for j, text in enumerate(row_data):
                    if j < len(self.inputs[i]):
                        self.inputs[i][j].text = str(text) if not pd.isna(text) else ''

        self.update_totals()

    def clear_table_data(self, instance=None):
        for row in self.inputs:
            for inp in row:
                inp.text = ''
        self.page_data.clear()
        self.update_totals()
        self.show_success("Данные очищены!")

    def open_excel_file(self, instance):
        if os.path.exists(self.reports_dir):
            os.startfile(self.reports_dir)
        else:
            self.show_error("Папка reports не найдена!")

    def change_page(self, delta):
        new_page = self.current_page + delta
        if 0 <= new_page < self.MAX_PAGES:
            self.current_page = new_page
            self.load_page_data()
            self.update_page_label()

    def update_pagination(self):
        self.total_pages = len(self.page_data) if self.page_data else 1
        self.total_pages = min(self.total_pages, self.MAX_PAGES)
        self.update_page_label()

    def update_page_label(self):
        self.page_label.text = f'Страница {self.current_page+1} из {self.total_pages}'

    def go_back(self, instance):
        App.get_running_app().root.current = 'main'

    def show_error(self, message):
        Popup(
            title='Ошибка',
            content=Label(text=message, color=(1, 0, 0, 1)),
            size_hint=(0.6, 0.3)
        ).open()

    def show_success(self, message):
        Popup(
            title='Успешно',
            content=Label(text=message, color=(0, 0.5, 0, 1)),
            size_hint=(0.6, 0.3)
        ).open()

    def show_quarter_popup(self, instance):
        """Показать popup для ввода квартала"""
        content = BoxLayout(orientation='vertical', spacing=10, padding=10)

        # Заголовок
        title_label = Label(
            text="Введите номер квартала",
            font_name='Roboto',
            bold=True,
            size_hint=(1, None),
            height=30
        )
        content.add_widget(title_label)

        # Поле ввода квартала
        self.quarter_input = TextInput(
            hint_text="Номер квартала",
            multiline=False,
            size_hint=(1, None),
            height=40,
            font_name='Roboto',
            input_filter='int',
            text=self.current_quarter
        )
        content.add_widget(self.quarter_input)

        # Кнопки
        btn_layout = BoxLayout(orientation='horizontal', spacing=10, size_hint=(1, None), height=40)
        save_btn = ModernButton(
            text='Сохранить',
            bg_color=get_color_from_hex('#00FF00'),
            size_hint=(0.5, 1),
            height=50
        )
        cancel_btn = ModernButton(
            text='Отмена',
            bg_color=get_color_from_hex('#FF6347'),
            size_hint=(0.5, 1),
            height=50
        )
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="Настройка квартала",
            content=content,
            size_hint=(0.6, 0.5)
        )

        def save_quarter(btn):
            quarter = self.quarter_input.text.strip()
            if quarter:
                self.current_quarter = quarter
                self.update_address_label()
                self.show_success(f"Квартал установлен: {quarter}")
                popup.dismiss()
            else:
                self.show_error("Номер квартала не может быть пустым!")

        save_btn.bind(on_press=save_quarter)
        cancel_btn.bind(on_press=popup.dismiss)

        popup.open()

    def show_plot_popup(self, instance):
        """Показать popup для ввода выдела"""
        content = BoxLayout(orientation='vertical', spacing=10, padding=10)

        # Заголовок
        title_label = Label(
            text="Введите номер выдела",
            font_name='Roboto',
            bold=True,
            size_hint=(1, None),
            height=30
        )
        content.add_widget(title_label)

        # Поле ввода выдела
        self.plot_input = TextInput(
            hint_text="Номер выдела",
            multiline=False,
            size_hint=(1, None),
            height=40,
            font_name='Roboto',
            input_filter='int',
            text=self.current_plot
        )
        content.add_widget(self.plot_input)

        # Кнопки
        btn_layout = BoxLayout(orientation='horizontal', spacing=10, size_hint=(1, None), height=50)
        save_btn = ModernButton(
            text='Сохранить',
            bg_color=get_color_from_hex('#00FF00'),
            size_hint=(0.5, 1),
            height=50
        )
        cancel_btn = ModernButton(
            text='Отмена',
            bg_color=get_color_from_hex('#FF6347'),
            size_hint=(0.5, 1),
            height=50
        )
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="Настройка выдела",
            content=content,
            size_hint=(0.6, 0.5)
        )

        def save_plot(btn):
            plot = self.plot_input.text.strip()
            if plot:
                self.current_plot = plot
                self.update_address_label()
                self.show_success(f"Выдел установлен: {plot}")
                popup.dismiss()
            else:
                self.show_error("Номер выдела не может быть пустым!")

        save_btn.bind(on_press=save_plot)
        cancel_btn.bind(on_press=popup.dismiss)

        popup.open()

    def show_forestry_popup(self, instance):
        """Показать popup для ввода лесничества"""
        content = BoxLayout(orientation='vertical', spacing=10, padding=10)

        # Заголовок
        title_label = Label(
            text="Введите название лесничества",
            font_name='Roboto',
            bold=True,
            size_hint=(1, None),
            height=30
        )
        content.add_widget(title_label)

        # Поле ввода лесничества
        self.forestry_input = TextInput(
            hint_text="Название лесничества",
            multiline=False,
            size_hint=(1, None),
            height=40,
            font_name='Roboto',
            text=self.current_forestry
        )
        content.add_widget(self.forestry_input)

        # Кнопки
        btn_layout = BoxLayout(orientation='horizontal', spacing=10, size_hint=(1, None), height=50)
        save_btn = ModernButton(
            text='Сохранить',
            bg_color=get_color_from_hex('#00FF00'),
            size_hint=(0.5, 1),
            height=50
        )
        cancel_btn = ModernButton(
            text='Отмена',
            bg_color=get_color_from_hex('#FF6347'),
            size_hint=(0.5, 1),
            height=50
        )
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="Настройка лесничества",
            content=content,
            size_hint=(0.6, 0.5)
        )

        def save_forestry(btn):
            forestry = self.forestry_input.text.strip()
            if forestry:
                self.current_forestry = forestry
                self.update_address_label()
                self.show_success(f"Лесничество установлено: {forestry}")
                popup.dismiss()
            else:
                self.show_error("Название лесничества не может быть пустым!")

        save_btn.bind(on_press=save_forestry)
        cancel_btn.bind(on_press=popup.dismiss)

        popup.open()

    def update_totals(self):
        """Обновление строки итогов с поддержкой множественных пород"""
        breed_composition = {}  # Для расчета коэффициента состава
        total_stats = {'density': [], 'height': [], 'age': []}

        # Собираем данные по всем породам для расчета коэффициента состава
        for row in self.inputs:
            predmet_text = row[2].text  # Столбец "Предмет ухода"
            if predmet_text:
                # Парсим коэффициент состава (например: "3С1Е2Б")
                composition = self.parse_composition(predmet_text)
                for breed, count in composition.items():
                    if breed not in breed_composition:
                        breed_composition[breed] = []
                    breed_composition[breed].append(count)

            # Собираем статистику по породам для других расчетов
            breeds_text = row[3].text  # Столбец "Порода"
            if breeds_text:
                breeds_data = self.parse_breeds_data(breeds_text)
                for breed_info in breeds_data:
                    # Для хвойных пород густота = сумма (до 0.5м + 0.5-1.5м + >1.5м)
                    if breed_info.get('type') == 'coniferous':
                        coniferous_density = (breed_info.get('do_05', 0) +
                                            breed_info.get('05_15', 0) +
                                            breed_info.get('bolee_15', 0))
                        if coniferous_density > 0:
                            total_stats['density'].append(coniferous_density)
                    elif 'density' in breed_info and breed_info['density']:
                        total_stats['density'].append(breed_info['density'])

                    if 'height' in breed_info and breed_info['height']:
                        total_stats['height'].append(breed_info['height'])
                    if 'age' in breed_info and breed_info['age']:
                        total_stats['age'].append(breed_info['age'])

        # Вычисляем средний коэффициент состава
        avg_composition = {}
        for breed, counts in breed_composition.items():
            if counts:
                avg_composition[breed] = sum(counts) / len(counts)

        # Формируем текстовое представление среднего состава
        composition_text = ""
        for breed in sorted(avg_composition.keys()):
            count = avg_composition[breed]
            if count > 0:
                composition_text += f"{int(count)}{breed}"

        # Вычисляем средние значения для других параметров
        avg_density = sum(total_stats['density']) / len(total_stats['density']) if total_stats['density'] else 0
        avg_height = sum(total_stats['height']) / len(total_stats['height']) if total_stats['height'] else 0
        avg_age = sum(total_stats['age']) / len(total_stats['age']) if total_stats['age'] else 0

        # Обновляем labels итогов
        for lbl in self.total_labels:
            if hasattr(lbl, 'col_index'):
                col_idx = lbl.col_index
                if col_idx == 1:  # GPS точка - не применяем расчет
                    lbl.text = ""
                elif col_idx == 2:  # Предмет ухода - коэффициент состава
                    lbl.text = composition_text if composition_text else ""
                elif col_idx == 3:  # Порода - средний возраст
                    lbl.text = f"{avg_age:.1f}" if avg_age > 0 else "0"
                elif col_idx == 4:  # Примечания - средняя густота
                    lbl.text = f"{avg_density:.1f}" if avg_density > 0 else "0"
                elif col_idx == 5:  # Тип Леса - средняя высота
                    lbl.text = f"{avg_height:.1f}" if avg_height > 0 else "0"

    def parse_composition(self, text):
        """Парсит текстовое представление состава пород (например: '3С1Е2Б')"""
        composition = {}
        # Регулярное выражение для поиска паттернов вида "число + буква"
        import re
        matches = re.findall(r'(\d+)([А-ЯA-Z])', text.upper())
        for count, breed in matches:
            try:
                composition[breed] = int(count)
            except ValueError:
                pass
        return composition

    def parse_breeds_data(self, breeds_text):
        """Парсит данные пород из текстового поля (JSON или простой текст)"""
        if not breeds_text:
            return []

        # Пытаемся распарсить как JSON (новый формат)
        try:
            if breeds_text.startswith('['):  # Массив пород
                return json.loads(breeds_text)
            elif breeds_text.startswith('{'):  # Одна порода
                return [json.loads(breeds_text)]
        except (json.JSONDecodeError, TypeError):
            pass

        # Если не JSON, возвращаем пустой список (старый формат обрабатывается отдельно)
        return []

    def calculate_page_totals(self):
        """Вычисляет итоговые значения для текущей страницы"""
        totals = {
            'composition': '',
            'total_area': 0.0,
            'avg_age': 0.0,
            'avg_density': 0.0,
            'avg_height': 0.0
        }

        breed_composition = {}
        total_stats = {'density': [], 'height': [], 'age': []}
        total_area = 0.0

        # Собираем данные по всем строкам на странице
        for row in self.inputs:
            predmet_text = row[2].text  # Столбец "Предмет ухода"
            if predmet_text:
                composition = self.parse_composition(predmet_text)
                for breed, count in composition.items():
                    if breed not in breed_composition:
                        breed_composition[breed] = []
                    breed_composition[breed].append(count)

            # Получаем радиус для расчета площади
            radius = 5.64  # Значение по умолчанию
            try:
                if row[5].text:  # Столбец "Тип Леса" содержит радиус
                    radius = float(row[5].text)
            except (ValueError, IndexError):
                pass

            area = 3.14159 * (radius ** 2)  # π * r²
            total_area += area

            # Собираем статистику по породам из столбца "Порода"
            breeds_text = row[3].text  # Столбец "Порода"
            if breeds_text:
                breeds_data = self.parse_breeds_data(breeds_text)
                for breed_info in breeds_data:
                    # Для хвойных пород густота = сумма (до 0.5м + 0.5-1.5м + >1.5м)
                    if breed_info.get('type') == 'coniferous':
                        coniferous_density = (breed_info.get('do_05', 0) +
                                            breed_info.get('05_15', 0) +
                                            breed_info.get('bolee_15', 0))
                        if coniferous_density > 0:
                            total_stats['density'].append(coniferous_density)
                    elif 'density' in breed_info and breed_info['density']:
                        total_stats['density'].append(breed_info['density'])

                    if 'height' in breed_info and breed_info['height']:
                        total_stats['height'].append(breed_info['height'])
                    if 'age' in breed_info and breed_info['age']:
                        total_stats['age'].append(breed_info['age'])

        # Вычисляем средний коэффициент состава
        avg_composition = {}
        for breed, counts in breed_composition.items():
            if counts:
                avg_composition[breed] = sum(counts) / len(counts)

        composition_text = ""
        for breed in sorted(avg_composition.keys()):
            count = avg_composition[breed]
            if count > 0:
                composition_text += f"{int(count)}{breed}"

        totals['composition'] = composition_text
        totals['total_area'] = total_area
        totals['avg_density'] = sum(total_stats['density']) / len(total_stats['density']) if total_stats['density'] else 0.0
        totals['avg_height'] = sum(total_stats['height']) / len(total_stats['height']) if total_stats['height'] else 0.0
        totals['avg_age'] = sum(total_stats['age']) / len(total_stats['age']) if total_stats['age'] else 0.0

        return totals

    def save_current_page(self, instance=None):
        """Сохраняем текущую страницу в базу данных с поддержкой множественных пород"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()

        try:
            # Удаляем старые данные для этой страницы и участка
            cursor.execute('''
                DELETE FROM molodniki_data
                WHERE page_number = ? AND section_name = ?
            ''', (self.current_page, self.current_section))

            # Сохраняем данные каждой строки
            for row_idx, row in enumerate(self.inputs):
                row_data = [inp.text.strip() for inp in row]
                if any(row_data[:5]):  # Сохраняем если есть данные в основных полях
                    # Получаем радиус из поля или используем значение по умолчанию
                    radius = 5.64  # Значение по умолчанию
                    try:
                        if row_data[5]:  # Если в поле "Тип Леса" указан радиус
                            radius = float(row_data[5])
                    except (ValueError, IndexError):
                        pass

                    cursor.execute('''
                        INSERT INTO molodniki_data
                        (page_number, row_index, nn, gps_point, predmet_uhoda, radius, primechanie, section_name)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    ''', (
                        self.current_page,
                        row_idx,
                        row_data[0] or None,  # №ППР
                        row_data[1] or None,  # GPS точка
                        row_data[2] or None,  # Предмет ухода
                        radius,               # Радиус
                        row_data[4] or None,  # Примечания
                        self.current_section
                    ))

                    molodniki_data_id = cursor.lastrowid

                    # Обрабатываем породы из столбца "Порода" (row_data[3])
                    if row_data[3]:
                        breeds_data = self.parse_breeds_data(row_data[3])
                        for breed_info in breeds_data:
                            # Вычисляем коэффициент состава: (густота * π * радиус²) / 10000
                            composition_coeff = 0.0
                            if breed_info.get('density') and radius:
                                area = 3.14159 * (radius ** 2)  # π * r²
                                composition_coeff = (breed_info['density'] * area) / 10000

                            cursor.execute('''
                                INSERT INTO molodniki_breeds
                                (molodniki_data_id, breed_name, breed_type, do_05, _05_15, bolee_15,
                                 density, height, age, composition_coefficient)
                                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                            ''', (
                                molodniki_data_id,
                                breed_info.get('name', ''),
                                breed_info.get('type', 'deciduous'),
                                breed_info.get('do_05', 0),
                                breed_info.get('05_15', 0),
                                breed_info.get('bolee_15', 0),
                                breed_info.get('density', 0),
                                breed_info.get('height', 0.0),
                                breed_info.get('age', 0),
                                composition_coeff
                            ))

            # Сохраняем итоги страницы
            totals = self.calculate_page_totals()
            cursor.execute('''
                INSERT OR REPLACE INTO molodniki_totals
                (page_number, section_name, total_composition, total_area, avg_age, avg_density, avg_height)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (
                self.current_page,
                self.current_section,
                totals['composition'],
                totals['total_area'],
                totals['avg_age'],
                totals['avg_density'],
                totals['avg_height']
            ))

            conn.commit()
            self.show_success("Страница сохранена в базу данных!")

        except Exception as e:
            conn.rollback()
            self.show_error(f"Ошибка сохранения: {str(e)}")
        finally:
            conn.close()

        # Также сохраняем в памяти для быстрого доступа
        page_data = []
        for row in self.inputs:
            page_data.append([inp.text for inp in row])
        self.page_data[self.current_page] = page_data
        self.update_totals()

    def show_save_dialog(self, instance=None):
        content = BoxLayout(orientation='vertical', spacing=10, padding=10)

        # Заголовок
        title_label = Label(
            text="Введите имя файла",
            font_name='Roboto',
            bold=True,
            size_hint=(1, None),
            height=30
        )
        content.add_widget(title_label)

        self.filename_input = TextInput(
            hint_text="Имя файла",
            multiline=False,
            size_hint=(1, None),
            height=40,
            font_name='Roboto'
        )
        default_name = f"Молодняки_расширенный_{self.current_section}_{datetime.datetime.now().strftime('%Y-%m-%d_%H-%M')}"
        self.filename_input.text = default_name
        content.add_widget(self.filename_input)

        # Кнопки
        btn_layout = BoxLayout(orientation='horizontal', spacing=10, size_hint=(1, None), height=50)
        save_btn = ModernButton(
            text='Сохранить',
            bg_color=get_color_from_hex('#00FF00'),
            size_hint=(0.5, 1),
            height=50
        )
        cancel_btn = ModernButton(
            text='Отмена',
            bg_color=get_color_from_hex('#FF6347'),
            size_hint=(0.5, 1),
            height=50
        )
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        self.save_popup = Popup(
            title="Сохранение отчета Excel",
            content=content,
            size_hint=(0.7, 0.5)
        )
        save_btn.bind(on_press=self.save_to_excel)
        cancel_btn.bind(on_press=self.save_popup.dismiss)
        self.save_popup.open()

    def save_to_excel(self, instance):
        filename = self.filename_input.text.strip()
        if not filename:
            self.show_error("Имя файла не может быть пустым!")
            return
        filename = re.sub(r'[\\/*?:"<>|]', "", filename)
        filename = f"{filename}.xlsx" if not filename.endswith(".xlsx") else filename
        full_path = os.path.join(self.reports_dir, filename)

        try:
            # Создаем новый workbook
            wb = Workbook()
            ws = wb.active
            ws.title = "Молодняки"

            # Добавляем адресную строку в начало
            address_parts = []
            if self.current_quarter:
                address_parts.append(f"Квартал: {self.current_quarter}")
            if self.current_plot:
                address_parts.append(f"Выдел: {self.current_plot}")
            if self.current_forestry:
                address_parts.append(f"Лесничество: {self.current_forestry}")
            if self.current_radius:
                address_parts.append(f"Радиус: {self.current_radius} м")

            address_text = " | ".join(address_parts) if address_parts else "Адрес не указан"
            ws['A1'] = f"Адрес: {address_text}"
            ws['A1'].font = openpyxl.styles.Font(bold=True, size=12)

            # Добавляем пустую строку
            ws.append([])

            # Заголовки столбцов
            headers = [
                '№ППР', 'GPS точка', 'Предмет ухода', 'Порода', 'Примечания', 'Тип Леса', 'Итого по площадке'
            ]
            for col_num, header in enumerate(headers, 1):
                cell = ws.cell(row=3, column=col_num, value=header)
                cell.font = openpyxl.styles.Font(bold=True)
                cell.fill = openpyxl.styles.PatternFill(start_color="00FF00", end_color="00FF00", fill_type="solid")

            # Собираем все данные из всех страниц
            all_data = []
            for page in sorted(self.page_data.keys()):
                all_data.extend(self.page_data[page])

            # Обрабатываем данные пород и добавляем в таблицу
            current_row = 4  # Начинаем с 4-й строки (после заголовков)
            for row in all_data:
                processed_row = []
                for i, cell_value in enumerate(row):
                    if i == 3 and cell_value:  # Столбец "Порода"
                        try:
                            # Если это JSON массив с данными пород, преобразуем в читаемый формат
                            breeds_data = json.loads(cell_value)
                            if isinstance(breeds_data, list):
                                # Множественные породы - создаем отдельные строки для каждой породы
                                for breed_idx, breed_info in enumerate(breeds_data):
                                    if isinstance(breed_info, dict):
                                        breed_row = ['', '', '', '', '', '']  # Копируем основные данные (6 столбцов)
                                        if breed_idx == 0:  # Первая порода в строке
                                            breed_row[0] = row[0]  # №ППР
                                            breed_row[1] = row[1]  # GPS точка
                                            breed_row[2] = row[2]  # Предмет ухода
                                            breed_row[4] = row[4]  # Примечания
                                            breed_row[5] = row[5]  # Тип Леса

                                        # Формируем строку породы
                                        breed_name = breed_info.get('name', 'Неизвестная')
                                        params = []
                                        if 'density' in breed_info and breed_info['density']:
                                            params.append(f"Густота: {breed_info['density']}")
                                        if 'height' in breed_info and breed_info['height']:
                                            params.append(f"Высота: {breed_info['height']}м")
                                        if 'age' in breed_info and breed_info['age']:
                                            params.append(f"Возраст: {breed_info['age']} лет")
                                        if 'do_05' in breed_info and breed_info['do_05']:
                                            params.append(f"До 0.5м: {breed_info['do_05']}")
                                        if '05_15' in breed_info and breed_info['05_15']:
                                            params.append(f"0.5-1.5м: {breed_info['05_15']}")
                                        if 'bolee_15' in breed_info and breed_info['bolee_15']:
                                            params.append(f">1.5м: {breed_info['bolee_15']}")

                                        breed_row[3] = f"{breed_name}: {'; '.join(params)}"
                                        ws.append(breed_row)
                                        current_row += 1
                            else:
                                # Одна порода (старый формат)
                                breed_info = []
                                for key, value in breeds_data.items():
                                    if key == 'do_05':
                                        breed_info.append(f"До 0.5м: {value}")
                                    elif key == '05_15':
                                        breed_info.append(f"0.5-1.5м: {value}")
                                    elif key == 'bolee_15':
                                        breed_info.append(f">1.5м: {value}")
                                    elif key == 'height':
                                        breed_info.append(f"Высота: {value}м")
                                    elif key == 'density':
                                        breed_info.append(f"Густота: {value}")
                                    elif key == 'age':
                                        breed_info.append(f"Возраст: {value} лет")
                                processed_row.append("; ".join(breed_info))
                        except (json.JSONDecodeError, TypeError):
                            # Если это просто название породы
                            processed_row.append(cell_value)
                    else:
                        processed_row.append(cell_value)

                # Если не было множественных пород, добавляем обычную строку
                if processed_row:
                    ws.append(processed_row)
                    current_row += 1

            # Автоподбор ширины столбцов
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)  # Максимальная ширина 50
                ws.column_dimensions[column_letter].width = adjusted_width

            wb.save(full_path)
            self.save_popup.dismiss()
            self.show_success(f"Файл сохранен: {filename}")
        except Exception as e:
            self.show_error(f"Ошибка: {str(e)}")

    def save_to_word(self, instance):
        try:
            from docx import Document

            filename = f"Молодняки_расширенный_{self.current_section}_{datetime.datetime.now().strftime('%Y-%m-%d_%H-%M')}.docx"
            full_path = os.path.join(self.reports_dir, filename)

            doc = Document()
            doc.add_heading(f'Расширенный отчет по молоднякам - Участок {self.current_section}', 0)

            # Собираем все данные из всех страниц
            all_data = []
            for page in sorted(self.page_data.keys()):
                all_data.extend(self.page_data[page])

            # Создаем таблицу (6 столбцов)
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'

            # Заголовки
            headers = [
                '№ППР', 'GPS точка', 'Предмет ухода', 'Порода', 'Примечания', 'Тип Леса', 'Итого по площадке'
            ]
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header

            # Обрабатываем и добавляем данные
            for row_data in all_data:
                if any(row_data):  # Только непустые строки
                    row_cells = table.add_row().cells
                    for i, cell_value in enumerate(row_data):
                        if i < 6:  # 6 столбцов
                            if i == 3 and cell_value:  # Столбец "Порода"
                                try:
                                    # Если это JSON массив с данными пород, преобразуем в читаемый формат
                                    breeds_data = json.loads(cell_value)
                                    if isinstance(breeds_data, list):
                                        # Множественные породы
                                        breed_summaries = []
                                        for breed_info in breeds_data:
                                            if isinstance(breed_info, dict):
                                                breed_name = breed_info.get('name', 'Неизвестная')
                                                params = []
                                                if 'density' in breed_info and breed_info['density']:
                                                    params.append(f"Густота: {breed_info['density']}")
                                                if 'height' in breed_info and breed_info['height']:
                                                    params.append(f"Высота: {breed_info['height']}м")
                                                if 'age' in breed_info and breed_info['age']:
                                                    params.append(f"Возраст: {breed_info['age']} лет")
                                                if 'do_05' in breed_info and breed_info['do_05']:
                                                    params.append(f"До 0.5м: {breed_info['do_05']}")
                                                if '05_15' in breed_info and breed_info['05_15']:
                                                    params.append(f"0.5-1.5м: {breed_info['05_15']}")
                                                if 'bolee_15' in breed_info and breed_info['bolee_15']:
                                                    params.append(f">1.5м: {breed_info['bolee_15']}")
                                                breed_summaries.append(f"{breed_name}: {'; '.join(params)}")
                                        row_cells[i].text = " | ".join(breed_summaries)
                                    else:
                                        # Одна порода (старый формат)
                                        breed_info = []
                                        for key, value in breeds_data.items():
                                            if key == 'do_05':
                                                breed_info.append(f"До 0.5м: {value}")
                                            elif key == '05_15':
                                                breed_info.append(f"0.5-1.5м: {value}")
