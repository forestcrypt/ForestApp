#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Скрипт для чтения цветовой маркировки в шаблоне"""

from docx import Document
from docx.shared import RGBColor
import json

def get_text_color(run):
    """Получает цвет текста"""
    if run.font.color and run.font.color.rgb:
        return run.font.color.rgb
    return None

def is_red(rgb):
    """Проверяет, красный ли цвет"""
    if rgb:
        return rgb[0] > 150 and rgb[1] < 100 and rgb[2] < 100
    return False

def is_purple(rgb):
    """Проверяет, фиолетовый ли цвет"""
    if rgb:
        return rgb[0] > 150 and rgb[1] < 100 and rgb[2] > 150
    return False

def analyze_document(doc_path):
    """Анализирует документ и возвращает информацию о цветах"""
    doc = Document(doc_path)
    
    result = {
        'paragraphs': [],
        'tables': []
    }
    
    # Анализируем параграфы
    for i, para in enumerate(doc.paragraphs):
        para_info = {
            'index': i,
            'text': para.text.strip(),
            'runs': []
        }
        
        for run in para.runs:
            text = run.text.strip()
            if text:
                rgb = get_text_color(run)
                color_type = 'black'
                if is_red(rgb):
                    color_type = 'red'
                elif is_purple(rgb):
                    color_type = 'purple'
                
                para_info['runs'].append({
                    'text': text,
                    'color': color_type,
                    'rgb': str(rgb) if rgb else None
                })
        
        if para_info['text']:
            result['paragraphs'].append(para_info)
    
    # Анализируем таблицы
    for i, table in enumerate(doc.tables):
        table_info = {
            'index': i,
            'rows': len(table.rows),
            'columns': len(table.columns),
            'cells': []
        }
        
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_info = {
                    'row': row_idx,
                    'col': col_idx,
                    'text': cell.text.strip(),
                    'runs': []
                }
                
                for run in cell.paragraphs[0].runs if cell.paragraphs else []:
                    text = run.text.strip()
                    if text:
                        rgb = get_text_color(run)
                        color_type = 'black'
                        if is_red(rgb):
                            color_type = 'red'
                        elif is_purple(rgb):
                            color_type = 'purple'
                        
                        cell_info['runs'].append({
                            'text': text,
                            'color': color_type,
                            'rgb': str(rgb) if rgb else None
                        })
                
                if cell_info['text']:
                    table_info['cells'].append(cell_info)
        
        result['tables'].append(table_info)
    
    return result

if __name__ == '__main__':
    import sys
    import io
    # Устанавливаем UTF-8 для вывода
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    
    doc_path = 'reports/Шаблон инт.docx'
    print(f"Analysis: {doc_path}")
    print("=" * 80)
    
    result = analyze_document(doc_path)
    
    print("\n=== PARAGRAPHS WITH COLOR MARKING ===")
    for para in result['paragraphs'][:40]:
        text_preview = para['text'][:80].replace('\n', ' ')
        print(f"\n[{para['index']}] {text_preview}")
        for run in para['runs']:
            if run['color'] != 'black':
                text_safe = run['text'].replace('\n', ' ')
                print(f"    - '{text_safe}' - {run['color'].upper()} ({run['rgb']})")
    
    print("\n=== TABLES ===")
    for table in result['tables']:
        print(f"\nTable {table['index']}: {table['rows']} rows, {table['columns']} cols")
        for cell in table['cells'][:20]:
            text_safe = cell['text'][:50].replace('\n', ' ')
            print(f"  [{cell['row']},{cell['col']}]: {text_safe}")
            for run in cell['runs']:
                if run['color'] != 'black':
                    text_safe = run['text'].replace('\n', ' ')
                    print(f"      - '{text_safe}' - {run['color'].upper()}")
    
    # Сохраняем результат в JSON
    with open('reports/шаблон_анализ.json', 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    
    print(f"\n=== Result saved to reports/шаблон_анализ.json ===")
