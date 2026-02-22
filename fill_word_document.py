#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для подстановки данных из приложения "Молодняки" в Word-документ
"Проект рубок ухода в молодняках"
"""

import os
import sys
import json
import argparse
from docx import Document
import sqlite3
import datetime

# Добавляем текущую директорию в путь для импорта модулей
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

class WordDocumentFiller:
    def __init__(self, db_name='forest_data.db', data_file=None, address_data=None, total_data=None):
        self.db_name = db_name
        self.document_path = 'reports/Шаблон проект.docx'
        self.data_file = data_file

        # Данные адресной строки (можно настроить)
        self.address_data = address_data or {
            'quarter': '1',  # Квартал
            'plot': '15',    # Выдел
            'section': 'Володозерское',  # Участковое лесничество
            'forestry': 'Сегежское лесничество',  # Лесничество
            'target_purpose': 'Эксплуатационные леса',  # Целевое назначение
            'plot_area': '25.5',  # Площадь участка (га)
            'forest_type': 'Сосняк черничный'  # Тип леса
        }

        # Данные из Итого
        self.total_data = total_data or {}

    def load_data_from_json(self, file_path):
        """Загружаем данные из JSON файла"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            # Извлекаем данные адресной строки и итоговые данные
            self.address_data = data.get('address_data', self.address_data)
            self.total_data = data.get('total_data', self.total_data)

            print("Данные успешно загружены из JSON файла")
            return True

        except Exception as e:
            print(f"Ошибка загрузки данных из JSON файла: {e}")
            return False

    def load_data_from_db(self):
        """Загружаем данные из базы данных приложения"""
        try:
            conn = sqlite3.connect(self.db_name)
            cursor = conn.cursor()

            # Получаем последние данные из таблицы molodniki_totals
            cursor.execute('''
                SELECT * FROM molodniki_totals
                ORDER BY created_at DESC
                LIMIT 1
            ''')

            totals_row = cursor.fetchone()
            if totals_row:
                self.total_data = {
                    'page_number': totals_row[1],
                    'section_name': totals_row[2],
                    'total_composition': totals_row[3],
                    'avg_age': totals_row[5],
                    'avg_density': totals_row[6],
                    'avg_height': totals_row[7]
                }

            # Рассчитываем количество площадок
            cursor.execute('''
                SELECT COUNT(DISTINCT id) FROM molodniki_data
            ''')
            total_plots = cursor.fetchone()[0]
            self.total_data['total_plots'] = total_plots

            # Получаем данные пород
            cursor.execute('''
                SELECT breed_name, breed_type, density, height, age, do_05, _05_15, bolee_15
                FROM molodniki_breeds
                ORDER BY id DESC
                LIMIT 20
            ''')

            breeds = cursor.fetchall()
            self.total_data['breeds'] = []
            for breed in breeds:
                self.total_data['breeds'].append({
                    'name': breed[0],
                    'type': breed[1],
                    'density': breed[2],
                    'height': breed[3],
                    'age': breed[4],
                    'do_05': breed[5],
                    '_05_15': breed[6],
                    'bolee_15': breed[7]
                })

            # Прописываем данные адреса в итого для доступа к ним
            self.total_data['address_data'] = self.address_data

            conn.close()
            print("Данные успешно загружены из базы данных")
            return True

        except Exception as e:
            print(f"Ошибка загрузки данных из БД: {e}")
            return False

    def calculate_totals(self):
        """Используем данные Итого напрямую без дополнительных расчетов"""
        # Данные уже загружены из таблицы molodniki_totals в load_data_from_db()
        # Никаких дополнительных расчетов не требуется

        # Убеждаемся, что все необходимые поля присутствуют
        if not self.total_data.get('composition'):
            self.total_data['composition'] = self.total_data.get('total_composition', 'Не определен')

        if not self.total_data.get('care_subject'):
            self.total_data['care_subject'] = 'Не определен'  # Предмет ухода из Итого

        if not self.total_data.get('intensity'):
            self.total_data['intensity'] = '25%'  # Интенсивность из Итого или по умолчанию

    def calculate_forest_type(self):
        """Определяем тип леса как среднее по площадкам из данных таблицы"""
        try:
            # Получаем данные адреса из итого
            address_data = self.total_data.get('address_data', self.address_data)

            # Получаем данные из базы данных о типах леса по площадкам
            conn = sqlite3.connect(self.db_name)
            cursor = conn.cursor()

            # Проверяем, существует ли столбец tip_lesa
            cursor.execute("PRAGMA table_info(molodniki_data)")
            columns = [col[1] for col in cursor.fetchall()]

            if 'tip_lesa' not in columns:
                print("Столбец tip_lesa не найден, используем значение по умолчанию")
                conn.close()
                return address_data.get('forest_type', 'Бр В2')

            # Получаем все типы леса из таблицы molodniki_data
            cursor.execute('''
                SELECT tip_lesa, COUNT(*) as count
                FROM molodniki_data
                WHERE tip_lesa IS NOT NULL AND tip_lesa != ''
                GROUP BY tip_lesa
                ORDER BY count DESC
            ''')

            forest_types = cursor.fetchall()
            conn.close()

            if forest_types:
                # Возвращаем наиболее часто встречающийся тип леса
                predominant_type = forest_types[0][0]
                return predominant_type
            else:
                return address_data.get('forest_type', 'Бр В2')

        except Exception as e:
            print(f"Ошибка при определении типа леса: {e}")
            address_data = self.total_data.get('address_data', self.address_data)
            return address_data.get('forest_type', 'Бр В2')

    def get_breed_letter(self, breed_name):
        """Получаем букву для породы в формуле состава"""
        letters = {
            'Сосна': 'С',
            'Ель': 'Е',
            'Пихта': 'П',
            'Кедр': 'К',
            'Лиственница': 'Л',
            'Берёза': 'Б',
            'Осина': 'Ос',
            'Ольха': 'О'
        }

        for name, letter in letters.items():
            if name.lower() in breed_name.lower():
                return letter

        return breed_name[0].upper()

    def get_test_data(self):
        """Тестовые данные для демонстрации"""
        return {
            'composition': '8С2БДр',
            'care_subject': '300шт/гаС + 50шт/гаБ',
            'avg_age': '25',
            'avg_density': '350',
            'avg_height': '12.5',
            'intensity': '25%',
            'breeds': [
                {'name': 'Сосна', 'type': 'coniferous', 'density': 300, 'height': 15, 'age': 30, 'do_05': 50, '_05_15': 150, 'bolee_15': 100},
                {'name': 'Берёза', 'type': 'deciduous', 'density': 50, 'height': 10, 'age': 20}
            ]
        }

    def fill_document(self):
        """Заполняем документ данными"""
        if not os.path.exists(self.document_path):
            print(f"Файл {self.document_path} не найден!")
            return False

        try:
            doc = Document(self.document_path)

            # Получаем данные из дополнительных функций
            care_queue = self.total_data.get('care_queue', 'первая')
            characteristics = self.total_data.get('characteristics', {})
            care_date = self.total_data.get('care_date', 'сент 2025 года')
            technology = self.total_data.get('technology', 'Равномерное изреживание молодняка. Срубленные деревья необходимо приземлить на месте. По пространственному размещению по площади лесного участка вырубаемых и оставляемых деревьев должна применяться равномерная рубка, в том числе при групповом или куртином размещении деревьев целевых древесных пород. Отбор деревьев производиться так, чтобы обеспечить равномерность размещения по площади оставляемых на выращивание деревьев целевых пород.')
            forest_purpose = self.total_data.get('forest_purpose', 'Эксплуатационные леса')

            # Формируем мероприятие с очередью
            activity_name = self.total_data.get('activity_name', 'осветление')
            care_activity_text = f"{activity_name}, {care_queue} очередь"

            # Формируем характеристики молодняков
            best_text = characteristics.get('best', 'здоровая, хорошо укорененная сосна, с хорошо сформированной кроной')
            auxiliary_text = characteristics.get('auxiliary', 'деревья всех пород обеспечивающие сохранение целостности и устойчивости насаждения')
            undesirable_text = characteristics.get('undesirable', 'деревья мешающие росту и формированию крон отобранных лучших и вспомогательных деревьев; деревья неудовлетворительного состояния')

            # Функция для округления чисел до одной запятой
            def format_number(value, default='Н/Д'):
                if value is None or value == 'Н/Д':
                    return default
                try:
                    num = float(value)
                    return f"{num:.1f}"
                except (ValueError, TypeError):
                    return str(value)

            # Получаем данные адреса из итого (если не хранятся отдельно)
            address_data = self.total_data.get('address_data', self.address_data)

            # Определяем тип леса как среднее по площадкам
            forest_type = self.calculate_forest_type()

            # Рассчитываем параметры площадок
            current_radius = float(address_data.get('radius', 1.78))
            plot_area_m2 = 3.14159 * current_radius ** 2

            # Словарь замен для различных плейсхолдеров в документе
            replacements = {
                # Вид рубки - из Детали (очередь рубки - тип мероприятий)
                '(Детали- очередь рубки-тип мероприятий)': care_activity_text,

                # Лесничество и участковое лесничество - из Адрес
                '(Адрес- лесничество)': address_data.get('forestry', ''),
                '(Адрес-участковое лесничество)': address_data.get('district_forestry', ''),

                # Целевое назначение лесов - из Детали
                '(Детали-назначение лесов)': forest_purpose,

                # Квартал, Выдел, Площадь - из Адрес
                'Квартал (Адрес)': address_data.get('quarter', ''),
                'Выдел (Адрес)': address_data.get('plot', ''),
                'площадь (Адрес-площадь)': address_data.get('plot_area', ''),

                # Тип леса - из Данных по площадкам
                '(Данных по площадкам)': address_data.get('forest_type', forest_type),

                # Очередь рубки - из Функции
                '(Функции)': care_queue,

                # Количество площадок и радиус - из Итого и Адрес
                '(Итого и Адрес-радиус)': f"{self.total_data.get('total_plots', 0)} шт. {plot_area_m2:.0f}м²(R-{current_radius:.2f}м)",

                # Коэффициент состава - из Итого (оставляем для совместимости, но не заменяем "Итого" в таблице)
                # '5С5БДр': self.total_data.get('composition', 'Не определен'),

                # Возраст по породам - из таблицы Итого
                'Возраст по породам': format_number(self.total_data.get('avg_age')),
                'Диаметр по породам': 'Н/Д',  # Пока не определено
                'Высота по породам': format_number(self.total_data.get('avg_height')),
                'Кол-во деревьев по породам тыс. шт/га': format_number(self.total_data.get('avg_density')),
                'Сомкнутость (полнота С)': 'Н/Д',  # Пока не определено
                'Подрост: Состав, возраст Высота Кол-во': 'Н/Д',  # Пока не определено

                # Характеристика деревьев - из Детали (Характеристика молодняков)
                '(Детали - Характеристика молодняков)': f"Лучшие: {best_text}\nВспомогательные: {auxiliary_text}\nНежелательные: {undesirable_text}",

                # Дата рубки - из Детали
                '(Детали-Дата рубки)': care_date,

                # Интенсивность - из Итого
                '(Итого)': self.total_data.get('intensity', '25%'),

                # Технология ухода - из Детали
                '(Детали-Технология ухода)': technology,

                # Планируемые затраты - пустое поле
                'Пусто': '',
            }

            # Дополнительные замены для различных вариантов плейсхолдеров
            additional_replacements = {
                # Альтернативные плейсхолдеры для лесничества
                'в __ _Сегежском лесничестве': f"в {address_data.get('forestry', '')} лесничестве",
                'Сегежском участковом лесничестве': f"{address_data.get('district_forestry', '')} участковом лесничестве",

                # Альтернативные плейсхолдеры для выдела и площади
                'квартал 225 выдел 33': f"квартал {address_data.get('quarter', '')} выдел {address_data.get('plot', '')}",
                'площадь общ-24 га': f"площадь общ-{address_data.get('plot_area', '')} га",
                'площадь экс-24 га': f"площадь экс-{address_data.get('plot_area', '')} га",

                # Альтернативные плейсхолдеры для типа леса
                'Бр В2': address_data.get('forest_type', forest_type),

                # Альтернативные плейсхолдеры для коэффициента состава
                '5С5БДр': self.total_data.get('composition', 'Не определен'),

                # Альтернативные плейсхолдеры для возраста, высоты, густоты
                '10.666666666666666': format_number(self.total_data.get('avg_age')),
                '2.3333333333333335': format_number(self.total_data.get('avg_height')),
                '3683.680201929093': format_number(self.total_data.get('avg_density')),

                # Альтернативные плейсхолдеры для предмета ухода
                'Подставляем данные из Итого': self.total_data.get('care_subject', 'Не определен'),

                # Альтернативные плейсхолдеры для интенсивности
                '25%': self.total_data.get('intensity', '25%'),

                # Дополнительные плейсхолдеры для адресной строки
                'Сегежское лесничество': address_data.get('forestry', ''),
                'Володозерское': address_data.get('section', ''),
                'участковое лесничество': f"{address_data.get('district_forestry', '')} участковое лесничество",

                # Дополнительные плейсхолдеры для итогов
                'Предмет ухода': self.total_data.get('care_subject', 'Не определен'),
                'Интенсивность': self.total_data.get('intensity', '25%'),

                # Дополнительные плейсхолдеры для дополнительных функций
                'очередь ухода': care_queue,
                'характеристика молодняков': f"Лучшие: {best_text}\nВспомогательные: {auxiliary_text}\nНежелательные: {undesirable_text}",
                'дата ухода': care_date,
                'технология ухода': technology,

                # Плейсхолдеры для характеристик молодняков по отдельности
                'Лучшие:': f"Лучшие: {best_text}",
                'Вспомогательные:': f"Вспомогательные: {auxiliary_text}",
                'Нежелательные:': f"Нежелательные: {undesirable_text}",
            }

            # Объединяем словари замен
            replacements.update(additional_replacements)

            # Сначала обрабатываем таблицы (специальная обработка характеристик)
            for table in doc.tables:
                for row_idx, row in enumerate(table.rows):
                    # Специальная обработка для таблицы характеристик насаждения
                    # Ищем таблицу с заголовками характеристик
                    if row_idx == 0 and len(row.cells) >= 8:  # Заголовок таблицы
                        headers = []
                        for cell in row.cells:
                            try:
                                text = str(cell.text).strip() if cell.text is not None else ""
                                headers.append(text)
                            except (AttributeError, TypeError):
                                headers.append("")

                        if len(headers) > 0 and 'Выдел' in headers[0] and 'Состав' in ''.join(headers):
                            # Это таблица характеристик насаждения
                            self.fill_characteristics_table(table, headers)
                            break  # Прерываем цикл после обработки таблицы

            # Проходим по всем параграфам и заменяем текст
            for paragraph in doc.paragraphs:
                for old_text, new_text in replacements.items():
                    if old_text in paragraph.text:
                        paragraph.text = paragraph.text.replace(old_text, str(new_text))

            # Также проверяем таблицы (общие замены)
            for table in doc.tables:
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for paragraph in cell.paragraphs:
                            for old_text, new_text in replacements.items():
                                if old_text in paragraph.text:
                                    paragraph.text = paragraph.text.replace(old_text, str(new_text))

            # Сохраняем документ
            timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = self.document_path.replace('.docx', f'_заполненный_{timestamp}.docx')
            doc.save(output_path)

            print(f"Документ успешно заполнен и сохранен как: {output_path}")
            return True

        except Exception as e:
            print(f"Ошибка при заполнении документа: {e}")
            return False

    def fill_characteristics_table(self, table, headers):
        """Заполняет таблицу характеристик насаждения"""
        try:
            # Определяем индексы колонок
            col_indices = {}
            for i, header in enumerate(headers):
                header_lower = header.lower()
                if 'выдел' in header_lower and 'площадь' in header_lower:
                    col_indices['plot_area'] = i
                elif 'состав древостоя' in header_lower:
                    col_indices['composition'] = i
                elif 'возраст' in header_lower and 'породам' in header_lower:
                    col_indices['age_by_breeds'] = i
                elif 'диаметр' in header_lower and 'породам' in header_lower:
                    col_indices['diameter_by_breeds'] = i
                elif 'высота' in header_lower and 'породам' in header_lower:
                    col_indices['height_by_breeds'] = i
                elif 'кол-во деревьев' in header_lower or 'тыс. шт/га' in header_lower:
                    col_indices['density_by_breeds'] = i
                elif 'сомкнутость' in header_lower or 'полнота' in header_lower:
                    col_indices['crown_density'] = i
                elif 'подрост' in header_lower:
                    col_indices['undergrowth'] = i

            # Функция для форматирования чисел
            def format_number(value, default=''):
                if value is None or value == 'Н/Д':
                    return default
                try:
                    num = float(value)
                    return f"{num:.1f}"
                except (ValueError, TypeError):
                    return str(value)

            # Данные для заполнения
            address_data = self.total_data.get('address_data', self.address_data)
            plot_info = f"{address_data.get('quarter', '')} {address_data.get('plot', '')}, {address_data.get('plot_area', '')}"
            composition = self.total_data.get('composition', 'Не определен')

            # Заполняем строки таблицы
            # Строка 3 (индекс 2) - исходное состояние, Строка 4 (индекс 3) - проектируемое состояние
            for row_idx in [2, 3]:  # Только эти две строки содержат данные для заполнения
                if row_idx >= len(table.rows):
                    continue

                row = table.rows[row_idx]
                row_type = 'исх' if row_idx == 2 else 'проект'

                # Заполняем данные в зависимости от типа колонки
                # Колонки 0-1: Выдел/Площадь
                if 0 < len(row.cells):
                    row.cells[0].text = plot_info  # Выдел
                if 1 < len(row.cells):
                    row.cells[1].text = plot_info  # Площадь

                # Колонки 2-3: Состав древостоя
                if 2 < len(row.cells):
                    row.cells[2].text = composition  # исходный
                if 3 < len(row.cells):
                    row.cells[3].text = composition  # проектируемый

                # Колонки 4-5: Возраст по породам
                age_value = format_number(self.total_data.get('avg_age'))
                if 4 < len(row.cells):
                    row.cells[4].text = age_value  # исходный
                if 5 < len(row.cells):
                    row.cells[5].text = age_value  # проектируемый

                # Колонки 6-7: Диаметр по породам
                if 6 < len(row.cells):
                    row.cells[6].text = 'Н/Д'  # исходный
                if 7 < len(row.cells):
                    row.cells[7].text = 'Н/Д'  # проектируемый

                # Колонки 8-9: Высота по породам
                height_value = format_number(self.total_data.get('avg_height'))
                if 8 < len(row.cells):
                    row.cells[8].text = height_value  # исходный
                if 9 < len(row.cells):
                    row.cells[9].text = height_value  # проектируемый

                # Колонки 10-11: Кол-во деревьев по породам тыс. шт/га
                density_value = format_number(self.total_data.get('avg_density'))
                if row_type == 'проект':
                    intensity = self.total_data.get('intensity', '25%')
                    try:
                        intensity_value = float(str(intensity).strip('%')) / 100
                        original_density = self.total_data.get('avg_density', 0)
                        projected_density = original_density * (1 - intensity_value)
                        density_value = format_number(projected_density)
                    except (ValueError, TypeError):
                        pass

                if 10 < len(row.cells):
                    row.cells[10].text = density_value  # исходный или проектируемый
                if 11 < len(row.cells):
                    row.cells[11].text = density_value  # исходный или проектируемый

                # Колонки 12-13: Сомкнутость (полнота С)
                if 12 < len(row.cells):
                    row.cells[12].text = 'Н/Д'  # исходный
                if 13 < len(row.cells):
                    row.cells[13].text = 'Н/Д'  # проектируемый

                # Колонки 14-15: Подрост
                if 14 < len(row.cells):
                    row.cells[14].text = 'Н/Д'  # исходный
                if 15 < len(row.cells):
                    row.cells[15].text = 'Н/Д'  # проектируемый

        except Exception as e:
            print(f"Ошибка при заполнении таблицы характеристик: {e}")

    def run(self):
        """Основной метод выполнения"""
        print("Начинаем заполнение Word-документа...")

        # Если указан файл данных, загружаем из него
        if self.data_file:
            if not self.load_data_from_json(self.data_file):
                print(f"Ошибка: Не удалось загрузить данные из файла {self.data_file}.")
                return False
        else:
            # Загружаем данные из БД (для обратной совместимости)
            if not self.load_data_from_db():
                print("Ошибка: Не удалось загрузить данные из базы данных. Проверьте наличие данных в таблице molodniki_totals.")
                return False

        # Рассчитываем итоги (используем данные Итого напрямую)
        self.calculate_totals()

        # Заполняем документ
        success = self.fill_document()

        if success:
            print("Задача выполнена успешно!")
        else:
            print("Произошла ошибка при выполнении задачи.")

        return success

def fill_document_from_json(json_file_path):
    """
    Функция для заполнения Word-документа данными из JSON файла.
    Эта функция создает экземпляр WordDocumentFiller и запускает процесс заполнения.
    
    Args:
        json_file_path (str): Путь к JSON файлу с данными
        
    Returns:
        bool: True если заполнение прошло успешно, False в случае ошибки
    """
    try:
        filler = WordDocumentFiller(data_file=json_file_path)
        return filler.run()
    except Exception as e:
        print(f"Ошибка при заполнении документа из JSON: {e}")
        return False

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Заполнение Word-документа данными из приложения Молодняки')
    parser.add_argument('--data-file', type=str, help='Путь к JSON файлу с данными (address_data и total_data)')

    args = parser.parse_args()

    filler = WordDocumentFiller(data_file=args.data_file)
    filler.run()
