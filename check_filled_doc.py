#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для проверки заполненного документа
"""

from docx import Document

def check_filled_document():
    """Проверяем заполненный документ"""
    try:
        doc = Document('reports/Шаблон проект_заполненный_20260217_224739.docx')

        print("=== ПРОВЕРКА ЗАПОЛНЕННОГО ДОКУМЕНТА ===")
        print("Параграфы с данными:")
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text and len(text) > 5 and not text.startswith('ПРОЕКТ') and not text.startswith('Проект'):
                print(f'{i+1}: "{text[:120]}..."')

        print("\nТаблица характеристик (первые 5 строк):")
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                if row_idx < 5:  # Показываем первые 5 строк
                    cells_text = [cell.text.strip() for cell in row.cells]
                    print(f'  Строка {row_idx+1}: {cells_text}')

    except Exception as e:
        print(f"Ошибка: {e}")

if __name__ == '__main__':
    check_filled_document()
