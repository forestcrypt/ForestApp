#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Создание идеального шаблона проекта ухода с правильными плейсхолдерами
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def create_ideal_template():
    """Создаёт идеальный шаблон проекта ухода"""
    
    # Создаём новый документ
    doc = Document()
    
    # Настройка полей страницы
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    
    # === ЗАГОЛОВОК ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('ПРОЕКТ РУБОК УХОДА')
    title_run.bold = True
    title_run.font.size = Pt(16)
    
    # === Вид рубки ===
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Вид рубки: ').bold = True
    p.add_run('(Детали-очередь рубки-тип мероприятий)')
    
    # === Адресная строка ===
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('в ')
    p.add_run('(Адрес-лесничество)').italic = True
    p.add_run(' лесничестве, ')
    p.add_run('(Адрес-участковое лесничество)').italic = True
    p.add_run(' участковом лесничестве')
    
    # === Целевое назначение ===
    p = doc.add_paragraph()
    p.add_run('целевое назначение лесов - ').bold = True
    p.add_run('(Детали-назначение лесов)')
    
    # === Квартал, выдел, площадь ===
    p = doc.add_paragraph()
    p.add_run('квартал ').bold = True
    p.add_run('(Адрес-квартал)').italic = True
    p.add_run(', выдел ').bold = True
    p.add_run('(Адрес-выдел)').italic = True
    p.add_run(', площадь ').bold = True
    p.add_run('(Адрес-площадь)').italic = True
    p.add_run(' га')
    
    # === Тип леса ===
    p = doc.add_paragraph()
    p.add_run('Тип леса: ').bold = True
    p.add_run('(Данных по площадкам-тип леса)').italic = True
    
    doc.add_paragraph()
    
    # === 1. Потребность насаждения ===
    p = doc.add_paragraph()
    run = p.add_run('1. Потребность насаждения в проведении рубки ухода: ')
    run.bold = True
    p.add_run('(Функции-очередь рубки)').italic = True
    
    # === 2. Количество площадок ===
    p = doc.add_paragraph()
    run = p.add_run('2. Проектируемое количество и размеры учетных площадей: ')
    run.bold = True
    p.add_run('(Итого и Адрес-радиус)').italic = True
    
    # === 3. Таблица характеристик ===
    p = doc.add_paragraph()
    run = p.add_run('3. Характеристика насаждения: исходная (до рубки) - проектируемая (после рубки):')
    run.bold = True
    
    # Создаём таблицу
    table = doc.add_table(rows=5, cols=17)
    table.style = 'Table Grid'
    
    # Заголовки таблицы
    headers = [
        'Квартал\nВыдел\nПлощадь, га',
        'Состав древостоя\nисх',
        'Состав древостоя\nпроект',
        'Возраст по породам\nисх',
        'Возраст по породам\nпроект',
        'Диаметр по породам\nисх',
        'Диаметр по породам\nпроект',
        'Высота по породам\nисх',
        'Высота по породам\nпроект',
        'Кол-во деревьев по породам\nтыс. шт/га исх',
        'Кол-во деревьев по породам\nтыс. шт/га проект',
        'Сомкнутость (полнота С)\nисх',
        'Сомкнутость (полнота С)\nпроект',
        'Подрост\nСостав исх',
        'Подрост\nСостав проект',
        'Подрост\nВозраст исх',
        'Подрост\nВозраст проект'
    ]
    
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Строка данных (будет заполнена)
    data_row = table.rows[1]
    data_row.cells[0].text = '(Адрес-квартал) (Адрес-выдел) (Адрес-площадь)'
    data_row.cells[1].text = '(Итого-состав исх)'
    data_row.cells[2].text = '(Итого-состав проект)'
    data_row.cells[3].text = '(Итого-возраст)'
    data_row.cells[4].text = '(Итого-возраст)'
    data_row.cells[5].text = '(Итого-диаметр)'
    data_row.cells[6].text = '(Итого-диаметр)'
    data_row.cells[7].text = '(Итого-высота)'
    data_row.cells[8].text = '(Итого-высота)'
    data_row.cells[9].text = '(Итого-густота)'
    data_row.cells[10].text = '(Итого-густота)'
    data_row.cells[11].text = 'Н/Д'
    data_row.cells[12].text = 'Н/Д'
    data_row.cells[13].text = 'Н/Д'
    data_row.cells[14].text = 'Н/Д'
    data_row.cells[15].text = 'Н/Д'
    data_row.cells[16].text = 'Н/Д'
    
    # === 4. Характеристика деревьев ===
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('4. Характеристика деревьев по классам хозяйственно-биологической классификации:')
    run.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run('Лучшие: ')
    run.bold = True
    p.add_run('(Детали-характеристика-лучшие)').italic = True
    
    p = doc.add_paragraph()
    run = p.add_run('Вспомогательные: ')
    run.bold = True
    p.add_run('(Детали-характеристика-вспомогательные)').italic = True
    
    p = doc.add_paragraph()
    run = p.add_run('Нежелательные: ')
    run.bold = True
    p.add_run('(Детали-характеристика-нежелательные)').italic = True
    
    # === 5. Время проведения ===
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('5. Планируемое время проведения рубки ухода: ')
    run.bold = True
    p.add_run('(Детали-дата рубки)').italic = True
    
    # === 6. Интенсивность ===
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('6. Интенсивность рубки: ')
    run.bold = True
    p.add_run('(Итого-интенсивность)').italic = True
    
    # === 7. Технология ===
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('7. Технология рубки: ')
    run.bold = True
    p.add_run('(Детали-технология ухода)').italic = True
    
    # === 8. Планируемые затраты ===
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('8. Планируемые затраты на проведение рубки ухода (на 1 га): ')
    run.bold = True
    p.add_run('___не заполнять___').italic = True
    
    # === 9. Сортиментный состав ===
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('9. Сортиментный состав вырубаемой части древостоя: ')
    run.bold = True
    p.add_run('___не заполнять___').italic = True
    
    doc.add_paragraph()
    
    # === Подписи ===
    p = doc.add_paragraph()
    run = p.add_run('Проект составил: ')
    run.bold = True
    p.add_run('___________________ / _________________ /')
    
    p = doc.add_paragraph()
    run = p.add_run('Проект согласовал: ')
    run.bold = True
    p.add_run('___________________ / _________________ /')
    
    # Сохраняем документ
    output_path = 'reports/Шаблон проект_идеальный.docx'
    doc.save(output_path)
    
    print(f"[OK] Идеальный шаблон создан: {output_path}")
    print("\n=== Плейсхолдеры в шаблоне ===")
    print("Адресные данные:")
    print("  - (Адрес-лесничество)")
    print("  - (Адрес-участковое лесничество)")
    print("  - (Адрес-квартал)")
    print("  - (Адрес-выдел)")
    print("  - (Адрес-площадь)")
    print("\nДетали ухода:")
    print("  - (Детали-очередь рубки-тип мероприятий)")
    print("  - (Детали-назначение лесов)")
    print("  - (Детали-дата рубки)")
    print("  - (Детали-технология ухода)")
    print("  - (Детали-характеристика-лучшие)")
    print("  - (Детали-характеристика-вспомогательные)")
    print("  - (Детали-характеристика-нежелательные)")
    print("\nДанные из Итого:")
    print("  - (Итого-состав исх)")
    print("  - (Итого-состав проект)")
    print("  - (Итого-возраст)")
    print("  - (Итого-диаметр)")
    print("  - (Итого-высота)")
    print("  - (Итого-густота)")
    print("  - (Итого-интенсивность)")
    print("\nДругие:")
    print("  - (Функции-очередь рубки)")
    print("  - (Итого и Адрес-радиус)")
    print("  - (Данных по площадкам-тип леса)")
    
    return output_path

if __name__ == '__main__':
    create_ideal_template()
