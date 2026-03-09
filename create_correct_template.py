#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Создание правильного шаблона проекта ухода по образцу
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

def create_correct_template():
    """Создаёт правильный шаблон проекта ухода"""
    
    doc = Document()
    
    # Настройка полей
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    
    # === ПОДПИСИ ===
    # Согласовано
    p = doc.add_paragraph()
    p.add_run('Согласовано:\n').bold = True
    p.add_run('Начальник ГКУ РК\n')
    p.add_run('«Муезерское центральное лесничество»\n')
    p.add_run('Зипунников В.Н.\n\n')
    p.add_run('________________________\n')
    p.add_run('"____" _____________20__г.')
    
    # Утверждаю
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('Утверждаю:\n\n').bold = True
    p.add_run('________________________\n')
    p.add_run('"____" _____________20__г.')
    
    doc.add_paragraph()
    
    # === ЗАГОЛОВОК ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('ПРОЕКТ РУБОК УХОДА')
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph()
    
    # === Вид рубки ===
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run('(Детали-очередь рубки-тип мероприятий)')
    run.italic = True
    
    # === Адресная строка ===
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.add_run('в ')
    run = p.add_run('(Адрес-лесничество)')
    run.italic = True
    p.add_run(' лесничестве,\n')
    run = p.add_run('(Адрес-участковое лесничество)')
    run.italic = True
    p.add_run(' участковом лесничестве,\n')
    p.add_run('целевое назначение лесов ')
    run = p.add_run('(Детали-назначение лесов)')
    run.italic = True
    p.add_run(',\n')
    p.add_run('квартал ')
    run = p.add_run('(Адрес-квартал)')
    run.italic = True
    p.add_run(' выдел ')
    run = p.add_run('(Адрес-выдел)')
    run.italic = True
    p.add_run(' площадь ')
    run = p.add_run('(Адрес-площадь)')
    run.italic = True
    p.add_run(' га.')
    
    # === Тип леса ===
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.add_run('Тип (группа типов) леса и тип лесорастительных условий ')
    run = p.add_run('(Данных по площадкам-тип леса)')
    run.italic = True
    
    doc.add_paragraph()
    
    # === 1. Потребность ===
    p = doc.add_paragraph()
    run = p.add_run('1. Потребность насаждения в проведении рубки ухода (в первую, вторую третью, четвертую очередь-наличие других насаждений, требующих ухода в первую, вторую или третью очередь:')
    run.bold = True
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(5)
    run = p.add_run('(Функции-очередь рубки)')
    run.italic = True
    
    # === 2. Количество площадок ===
    p = doc.add_paragraph()
    run = p.add_run('2. Проектируемое количество и размеры учетных площадей:')
    run.bold = True
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(5)
    run = p.add_run('(Итого и Адрес-радиус)')
    run.italic = True
    
    # === 3. Таблица характеристик ===
    p = doc.add_paragraph()
    run = p.add_run('3. Характеристика насаждения: исходная (до рубки) - проектируемая (после рубки):')
    run.bold = True
    
    # Создаём таблицу
    table = doc.add_table(rows=5, cols=15)
    table.style = 'Table Grid'
    
    # Заголовки (первая строка)
    headers_row1 = [
        'Выдел', 'Площадь', 'Состав\nдревостоя', 'Состав\nдревостоя',
        'Возраст\nпо\nпородам', 'Возраст\nпо\nпородам',
        'Диаметр\nпо\nпородам', 'Диаметр\nпо\nпородам',
        'Высота\nпо\nпородам', 'Высота\nпо\nпородам',
        'Кол-во\nДеревьев по породам\nтыс. шт/га', 'Кол-во\nДеревьев по породам\nтыс. шт/га',
        'Сомкнутость\n(полнота С)\nпо породам', 'Сомкнутость\n(полнота С)\nпо породам',
        'Подрост:\nСостав, возраст\nВысота\nКол-во'
    ]
    
    for i, header in enumerate(headers_row1):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Вторая строка заголовков (исх/проект)
    headers_row2 = [
        '', '', 'исх', 'проект',
        'исх', 'проект',
        'исх', 'проект',
        'исх', 'проект',
        'исх', 'проект',
        'исх', 'проект',
        'исх/проект'
    ]
    
    for i, header in enumerate(headers_row2):
        cell = table.rows[1].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Строка данных (будет заполнена)
    data_row = table.rows[2]
    data_row.cells[0].text = '(Адрес-выдел)'
    data_row.cells[1].text = '(Адрес-площадь)'
    data_row.cells[2].text = '(Итого-состав исх)'
    data_row.cells[3].text = '(Итого-состав проект)'
    data_row.cells[4].text = '(Итого-возраст)'
    data_row.cells[5].text = '(Итого-возраст)'
    data_row.cells[6].text = '(Итого-диаметр)'
    data_row.cells[7].text = '(Итого-диаметр)'
    data_row.cells[8].text = '(Итого-высота)'
    data_row.cells[9].text = '(Итого-высота)'
    data_row.cells[10].text = '(Итого-густота)'
    data_row.cells[11].text = '(Итого-густота)'
    data_row.cells[12].text = '*'
    data_row.cells[13].text = '*'
    data_row.cells[14].text = '*'
    
    doc.add_paragraph()
    
    # === 4. Характеристика деревьев ===
    p = doc.add_paragraph()
    run = p.add_run('4. Характеристика деревьев по классам хозяйственно-биологической классификации с учетом выделения главных и второстепенных пород (для молодняков обязательно даются параметры перспективных-лучших и нежелательных деревьев по высоте)')
    run.bold = True
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run('Лучшие: ')
    run.bold = True
    run = p.add_run('(Детали-характеристика-лучшие)')
    run.italic = True
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run('Вспомогательные: ')
    run.bold = True
    run = p.add_run('(Детали-характеристика-вспомогательные)')
    run.italic = True
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run('Нежелательные (подлежащие вырубке): ')
    run.bold = True
    run = p.add_run('(Детали-характеристика-нежелательные)')
    run.italic = True
    
    # === 5. Время проведения ===
    p = doc.add_paragraph()
    run = p.add_run('5. Планируемое время проведения рубки ухода (месяцы, год)')
    run.bold = True
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(5)
    run = p.add_run('(Детали-дата рубки)')
    run.italic = True
    
    # === 6. Интенсивность ===
    p = doc.add_paragraph()
    run = p.add_run('6. Интенсивность рубки')
    run.bold = True
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(5)
    run = p.add_run('(Итого-интенсивность)')
    run.italic = True
    
    # === 7. Технология ===
    p = doc.add_paragraph()
    run = p.add_run('7. Проектируемая технология ухода')
    run.bold = True
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run('(Детали-технология ухода)')
    run.italic = True
    
    doc.add_paragraph()
    
    # === 8-11. Пустые разделы ===
    p = doc.add_paragraph()
    run = p.add_run('8. Планируемые затраты на проведение рубок ухода (на 1 га площади участка)')
    run.bold = True
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('9. Сортиментный состав вырубаемой части древостоя (по предварительной оценке)')
    run.bold = True
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('10. Ожидаемый доход от реализации заготовленной древесины')
    run.bold = True
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('11. Прибыль на 1 га участка леса, пройденного уходом')
    run.bold = True
    doc.add_paragraph()
    
    # === Подписи ===
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('Проект согласовал:\n\n')
    p.add_run('________________________ / ________________________ /')
    
    # Сохраняем
    output_path = 'reports/Шаблон проект_правильный.docx'
    doc.save(output_path)
    
    print(f"[OK] Правильный шаблон создан: {output_path}")
    print("\n=== Плейсхолдеры ===")
    print("Адрес: (Адрес-лесничество), (Адрес-участковое лесничество), (Адрес-квартал), (Адрес-выдел), (Адрес-площадь)")
    print("Детали: (Детали-очередь рубки-тип мероприятий), (Детали-назначение лесов), (Детали-дата рубки), (Детали-технология ухода)")
    print("       (Детали-характеристика-лучшие), (Детали-характеристика-вспомогательные), (Детали-характеристика-нежелательные)")
    print("Итого: (Итого-состав исх), (Итого-состав проект), (Итого-возраст), (Итого-диаметр), (Итого-высота), (Итого-густота), (Итого-интенсивность)")
    print("Другие: (Функции-очередь рубки), (Итого и Адрес-радиус), (Данных по площадкам-тип леса)")
    
    return output_path

if __name__ == '__main__':
    create_correct_template()
