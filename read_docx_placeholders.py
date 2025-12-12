#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для чтения плейсхолдеров из Word-документа
"""

import os
import sys
from docx import Document

def read_placeholders(doc_path):
    """Читаем документ и находим все плейсхолдеры"""
    if not os.path.exists(doc_path):
        print(f"Файл {doc_path} не найден!")
        return

    try:
        doc = Document(doc_path)
        placeholders = []
        all_text = []

        # Проходим по всем параграфам
        for paragraph in doc.paragraphs:
            text = paragraph.text
            all_text.append(f"Параграф: {text}")
            if 'Подставляем данные' in text:
                placeholders.append(f"Параграф: {text}")

        # Проходим по таблицам
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        all_text.append(f"Таблица: {text}")
                        if 'Подставляем данные' in text:
                            placeholders.append(f"Таблица: {text}")

        print("Найденные плейсхолдеры:")
        for placeholder in placeholders:
            print(f"  {placeholder}")

        if not placeholders:
            print("\nПлейсхолдеры не найдены. Показываем первые 20 текстовых блоков:")
            for i, text in enumerate(all_text[:20]):
                print(f"  {i+1}: {text[:100]}...")

        return placeholders

    except Exception as e:
        print(f"Ошибка при чтении документа: {e}")
        return []

if __name__ == '__main__':
    import sys
    if len(sys.argv) > 1:
        doc_path = sys.argv[1]
    else:
        doc_path = 'reports/проект 224-56-38га Волдозерское.docx'
    read_placeholders(doc_path)
