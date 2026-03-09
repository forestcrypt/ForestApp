#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Создание своего шаблона проекта ухода на основе образца
Все данные из меню Итого заполняются - по каждой породе и общие
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

def create_custom_template():
    """Создаёт свой шаблон проекта ухода"""
    
    doc = Document()
    
    # Настройка полей
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(1.5)
    
    # === ПОДПИСИ ===
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    
    # Согласовано
    cell1 = table.cell(0, 0)
    p = cell1.paragraphs[0]
    p.add_run('Согласовано:\n').bold = True
    p.add_run('Начальник ГКУ РК\n')
    p.add_run('«Муезерское центральное лесничество»\n')
    p.add_run('Зипунников В.Н.\n\n')
    p.add_run('________________________\n')
    p.add_run('"____" _____________20__г.')
    
    # Утверждаю
    cell2 = table.cell(0, 1)
    p = cell2.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('Утверждаю:\n\n').bold = True
    p.add_run('________________________\n')
    p.add_run('"____" _____________20__г.')
    
    doc.add_paragraph()
    
    # === ЗАГОЛОВОК ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('ПРОЕКТ РУБОК УХОДА')
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph()
    
    # === Вид рубки ===
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run('{care_activity_text}')
    run.italic = True
    
    # === Адресная строка ===
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.add_run('в ')
    run = p.add_run('{forestry}')
    run.italic = True
    p.add_run(' лесничестве,\n')
    run = p.add_run('{district_forestry}')
    run.italic = True
    p.add_run(' участковом лесничестве,\n')
    p.add_run('целевое назначение лесов ')
    run = p.add_run('{forest_purpose}')
    run.italic = True
    p.add_run(',\n')
    p.add_run('квартал ')
    run = p.add_run('{quarter}')
    run.italic = True
    p.add_run(' выдел ')
    run = p.add_run('{plot}')
    run.italic = True
    p.add_run(' площадь ')
    run = p.add_run('{plot_area}')
    run.italic = True
    p.add_run(' га.')
    
    # === Тип леса ===
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.add_run('Тип (группа типов) леса и тип лесорастительных условий ')
    run = p.add_run('{forest_type}')
    run.italic = True
    
    doc.add_paragraph()
    
    # === 1. Потребность ===
    p = doc.add_paragraph()
    run = p.add_run('1. Потребность насаждения в проведении рубки ухода:')
    run.bold = True
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(5)
    run = p.add_run('{care_queue}')
    run.italic = True
    
    # === 2. Количество площадок ===
    p = doc.add_paragraph()
    run = p.add_run('2. Проектируемое количество и размеры учетных площадей:')
    run.bold = True
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(5)
    run = p.add_run('{radius_info}')
    run.italic = True
    
    # === 3. Таблица характеристик по породам ===
    p = doc.add_paragraph()
    run = p.add_run('3. Характеристика насаждения по породам: исходная (до рубки) - проектируемая (после рубки):')
    run.bold = True
    
    # Таблица по породам
    table = doc.add_table(rows=2, cols=11)
    table.style = 'Table Grid'
    
    # Заголовки
    headers = [
        'Порода',
        'Состав\nисх',
        'Состав\nпроект',
        'Возраст\nисх',
        'Возраст\nпроект',
        'Диаметр\nисх',
        'Диаметр\nпроект',
        'Высота\nисх',
        'Высота\nпроект',
        'Густота\nтыс. шт/га\nисх',
        'Густота\nтыс. шт/га\nпроект'
    ]
    
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Строка данных для каждой породы (будет заполнена циклом)
    data_row = table.rows[1]
    data_row.cells[0].text = '{breed_name}'
    data_row.cells[1].text = '{breed_composition_isx}'
    data_row.cells[2].text = '{breed_composition_project}'
    data_row.cells[3].text = '{breed_age_isx}'
    data_row.cells[4].text = '{breed_age_project}'
    data_row.cells[5].text = '{breed_diameter_isx}'
    data_row.cells[6].text = '{breed_diameter_project}'
    data_row.cells[7].text = '{breed_height_isx}'
    data_row.cells[8].text = '{breed_height_project}'
    data_row.cells[9].text = '{breed_density_isx}'
    data_row.cells[10].text = '{breed_density_project}'
    
    doc.add_paragraph()
    
    # === 4. Общая характеристика по всем породам ===
    p = doc.add_paragraph()
    run = p.add_run('4. Общая характеристика насаждения по всем породам:')
    run.bold = True
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.add_run('Состав древостоя: исх - ')
    run = p.add_run('{total_composition_isx}')
    run.italic = True
    p.add_run(', проект - ')
    run = p.add_run('{total_composition_project}')
    run.italic = True
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.add_run('Средний возраст: исх - ')
    run = p.add_run('{total_age_isx}')
    run.italic = True
    p.add_run(', проект - ')
    run = p.add_run('{total_age_project}')
    run.italic = True
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.add_run('Средняя высота: исх - ')
    run = p.add_run('{total_height_isx}')
    run.italic = True
    p.add_run(', проект - ')
    run = p.add_run('{total_height_project}')
    run.italic = True
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.add_run('Средний диаметр: исх - ')
    run = p.add_run('{total_diameter_isx}')
    run.italic = True
    p.add_run(', проект - ')
    run = p.add_run('{total_diameter_project}')
    run.italic = True
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.add_run('Густота: исх - ')
    run = p.add_run('{total_density_isx}')
    run.italic = True
    p.add_run(', проект - ')
    run = p.add_run('{total_density_project}')
    run.italic = True
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.add_run('Интенсивность рубки: ')
    run = p.add_run('{intensity}')
    run.bold = True
    
    doc.add_paragraph()
    
    # === 5. Характеристика деревьев ===
    p = doc.add_paragraph()
    run = p.add_run('5. Характеристика деревьев по классам хозяйственно-биологической классификации:')
    run.bold = True
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.add_run('Лучшие: ').bold = True
    run = p.add_run('{best}')
    run.italic = True
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.add_run('Вспомогательные: ').bold = True
    run = p.add_run('{auxiliary}')
    run.italic = True
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.add_run('Нежелательные (подлежащие вырубке): ').bold = True
    run = p.add_run('{undesirable}')
    run.italic = True
    
    # === 6. Предмет ухода ===
    p = doc.add_paragraph()
    run = p.add_run('6. Предмет ухода:')
    run.bold = True
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run('{care_subject}')
    run.italic = True
    
    # === 7. Время проведения ===
    p = doc.add_paragraph()
    run = p.add_run('7. Планируемое время проведения рубки ухода:')
    run.bold = True
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(5)
    run = p.add_run('{care_date}')
    run.italic = True
    
    # === 8. Технология ===
    p = doc.add_paragraph()
    run = p.add_run('8. Проектируемая технология ухода:')
    run.bold = True
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run('{technology}')
    run.italic = True
    
    doc.add_paragraph()
    
    # === 9-12. Пустые разделы ===
    p = doc.add_paragraph()
    run = p.add_run('9. Планируемые затраты на проведение рубок ухода (на 1 га):')
    run.bold = True
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('10. Сортиментный состав вырубаемой части древостоя:')
    run.bold = True
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('11. Ожидаемый доход от реализации заготовленной древесины:')
    run.bold = True
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('12. Прибыль на 1 га участка леса, пройденного уходом:')
    run.bold = True
    doc.add_paragraph()
    
    # === Подписи ===
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('Проект согласовал:\n\n')
    p.add_run('________________________ / ________________________ /')
    
    # Сохраняем
    output_path = 'reports/Шаблон проект_наш.docx'
    doc.save(output_path)
    
    print(f"[OK] Наш шаблон создан: {output_path}")
    print("\n=== Плейсхолдеры ===")
    print("\nАдресные:")
    print("  {forestry}, {district_forestry}, {quarter}, {plot}, {plot_area}")
    print("\nДетали:")
    print("  {care_activity_text}, {care_queue}, {care_date}, {technology}")
    print("  {forest_purpose}, {best}, {auxiliary}, {undesirable}, {care_subject}")
    print("\nПо породам (для каждой породы отдельно):")
    print("  {breed_name}, {breed_composition_isx}, {breed_composition_project}")
    print("  {breed_age_isx}, {breed_age_project}, {breed_diameter_isx}, {breed_diameter_project}")
    print("  {breed_height_isx}, {breed_height_project}, {breed_density_isx}, {breed_density_project}")
    print("\nОбщие по всем породам:")
    print("  {total_composition_isx}, {total_composition_project}")
    print("  {total_age_isx}, {total_age_project}, {total_height_isx}, {total_height_project}")
    print("  {total_diameter_isx}, {total_diameter_project}")
    print("  {total_density_isx}, {total_density_project}, {intensity}")
    print("\nДругие:")
    print("  {radius_info}, {forest_type}")
    
    return output_path

if __name__ == '__main__':
    create_custom_template()
