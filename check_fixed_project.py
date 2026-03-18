#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Чтение НОВОГО заполненного проекта для проверки исправлений
"""

import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
import glob
import os

# Находим последний заполненный файл
files = glob.glob('reports/Шаблон проект_наш_заполненный_*.docx')
if not files:
    print("Не найдено заполненных файлов!")
    sys.exit(1)

latest_file = max(files, key=os.path.getctime)
print(f"Читаем файл: {latest_file}")

doc = Document(latest_file)

print("=" * 60)
print("ПРОВЕРКА ИСПРАВЛЕНИЙ")
print("=" * 60)

# Читаем параграфы
print("\n=== ПАРАГРАФЫ ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        # Выводим только важные параграфы
        if any(keyword in text for keyword in ['осветление', 'Брусничник', 'Состав', 'Высота', 'Диаметр', 'Густота', 'Интенсивность', 'Лучшие', 'Вспомогательные', 'Нежелательные', 'Предмет ухода']):
            print(f"\nПараграф {i}:")
            print(f"  {text}")

# Читаем таблицу с породами
print("\n=== ТАБЛИЦА С ПОРОДАМИ ===")
for i, table in enumerate(doc.tables):
    if len(table.columns) == 11:
        print(f"\nТаблица {i+1} (породы):")
        
        for j, row in enumerate(table.rows):
            row_text = []
            for k, cell in enumerate(row.cells):
                cell_text = cell.text.strip().replace('\n', ' ')
                if cell_text and k < 11:  # Только первые 11 колонок
                    row_text.append(f"[{k}] {cell_text}")
            
            if row_text:
                print(f"  Строка {j}: {' | '.join(row_text)}")

print("\n" + "=" * 60)
print("ПРОВЕРКА ЗАВЕРШЕНА")
print("=" * 60)
