# Расширенная таблица молодняков по новой структуре
        # Структура: 6 основных столбцов + динамические подстолбцы для пород

from kivy.app import App
from kivy.metrics import dp
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.button import Button
from kivy.uix.screenmanager import Screen
from kivy.uix.scrollview import ScrollView
from kivy.uix.gridlayout import GridLayout
from kivy.uix.textinput import TextInput
from kivy.uix.popup import Popup
from kivy.uix.label import Label
from kivy.uix.floatlayout import FloatLayout
from kivy.uix.checkbox import CheckBox
from kivy.properties import (NumericProperty, BooleanProperty,
                          ObjectProperty, ListProperty, StringProperty)
from kivy.core.window import Window
from kivy.graphics import Color, Rectangle, Line, RoundedRectangle
from kivy.clock import Clock
from kivy.animation import Animation
from kivy.core.text import LabelBase
from kivy.utils import get_color_from_hex
from kivy.core.image import Image as CoreImage
import sqlite3
import pandas as pd
import os
import datetime
import re
import json
import sys
import openpyxl
from openpyxl import Workbook
from tkinter import Tk, filedialog

from kivymd.uix.boxlayout import MDBoxLayout
from kivymd.uix.gridlayout import MDGridLayout
from kivymd.uix.card import MDCard
from kivymd.uix.button import MDButton, MDButtonText, MDIconButton
from kivymd.uix.label import MDLabel, MDIcon
from kivymd.uix.scrollview import MDScrollView
from kivymd.uix.appbar import MDTopAppBar, MDTopAppBarLeadingButtonContainer, MDTopAppBarTrailingButtonContainer, MDTopAppBarTitle, MDActionTopAppBarButton
from kivymd.uix.textfield import MDTextField

from ui_styles import Colors, Spacing, Fonts

LabelBase.register(name='Roboto',
                 fn_regular='fonts/Roboto-Medium.ttf',
                 fn_bold='fonts/Roboto-Bold.ttf')

class ModernButton(Button):
    bg_color = ListProperty([1, 1, 1, 1])
    no_shadow = BooleanProperty(False)
    auto_width = BooleanProperty(True)

    def __init__(self, **kwargs):
        self.no_shadow = kwargs.pop('no_shadow', False)
        self.auto_width = kwargs.pop('auto_width', True)
        super().__init__(**kwargs)
        self.background_color = (0, 0, 0, 0)
        self.font_name = 'Roboto'
        self.font_size = '16sp'
        self.bold = False
        self.size_hint = (None, None)
        self.height = 50
        self.padding = (20, 5)

        with self.canvas.before:
            if not self.no_shadow:
                Color(rgba=(0, 0, 0, 0.1))
                self.shadow = RoundedRectangle(
                    pos=(self.x+3, self.y-3),
                    size=self.size,
                    radius=[10]
                )
            self.bg_color_instruction = Color(rgba=self.bg_color)
            self.background = RoundedRectangle(
                pos=self.pos,
                size=self.size,
                radius=[10]
            )

        self.bind(pos=self.update_graphics, size=self.update_graphics)
        if self.auto_width:
            self.bind(text=self.update_width)
        
        # Обновляем text_size при изменении размера для поддержки halign/valign
        self.bind(size=self._update_text_size)

    def _update_text_size(self, *args):
        """Обновляем text_size при изменении размера кнопки"""
        self.text_size = (self.width - self.padding[0] * 2, self.height)

    def update_width(self, instance, value):
        self.width = self.texture_size[0] + 60

    def update_graphics(self, *args):
        self.background.pos = self.pos
        if not self.no_shadow:
            self.shadow.pos = (self.x+3, self.y-3)
            self.shadow.size = self.size
        self.background.size = self.size

    def on_touch_down(self, touch):
        if self.collide_point(*touch.pos):
            Animation(rgba=[c * 0.9 for c in self.bg_color], d=0.1).start(self.bg_color_instruction)
        return super().on_touch_down(touch)

    def on_touch_up(self, touch):
        Animation(rgba=self.bg_color, d=0.2).start(self.bg_color_instruction)
        return super().on_touch_up(touch)

class AutoCompleteTextInput(TextInput):
    next_widget = ObjectProperty(None)
    prev_widget = ObjectProperty(None)
    row_index = NumericProperty(0)
    col_index = NumericProperty(0)
    suggestions = ListProperty([])

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.bind(text=self.show_suggestions)
        self.popup = None

    def keyboard_on_key_down(self, window, keycode, text, modifiers):
        key = keycode[1]
        if key == 'down':
            self.focus_next('down')
        elif key == 'up':
            self.focus_previous('up')
        elif key == 'right':
            self.focus_next('right')
        elif key == 'left':
            self.focus_previous('left')
        else:
            super().keyboard_on_key_down(window, keycode, text, modifiers)
        return True

    def show_suggestions(self, instance, value):
        if not value or len(value) < 3:
            return

        conn = sqlite3.connect('forest_data.db')
        cursor = conn.cursor()
        cursor.execute('''
            SELECT value FROM molodniki_suggestions
            WHERE column_index = ? AND value LIKE ?
            ORDER BY LENGTH(value) ASC, value ASC
            LIMIT 1
        ''', (self.col_index, f'{value}%'))
        results = cursor.fetchall()
        conn.close()

        if results:
            self.text = results[0][0]

    def get_table_screen(self):
        return App.get_running_app().root.get_screen('molodniki')

    def focus_next(self, direction):
        table_screen = self.get_table_screen()
        if direction == 'right' and self.next_widget:
            self.next_widget.focus = True
        elif direction == 'down':
            if not hasattr(table_screen, 'inputs'):
                return
            next_row = self.row_index + 1
            if next_row < len(table_screen.inputs):
                table_screen.inputs[next_row][self.col_index].focus = True

    def focus_previous(self, direction):
        table_screen = self.get_table_screen()
        if direction == 'left' and self.prev_widget:
            self.prev_widget.focus = True
        elif direction == 'up':
            if not hasattr(table_screen, 'inputs'):
                return
            prev_row = self.row_index - 1
            if prev_row >= 0:
                table_screen.inputs[prev_row][self.col_index].focus = True

class Joypad(FloatLayout):
    def __init__(self, table_screen, **kwargs):
        super().__init__(**kwargs)
        self.table_screen = table_screen
        self.size_hint = (None, None)
        self.size = (140, 140)
        self.pos_hint = {'right': 0.95, 'y': 0.02}

        with self.canvas.before:
            Color(0.2, 0.2, 0.2, 1)
            self.bg_rect = RoundedRectangle(
                pos=self.pos,
                size=self.size,
                radius=[70]
            )

        arrows = [
            ('▲', (0.5, 0.7), 'up', (60, 40)),
            ('▼', (0.5, 0.3), 'down', (60, 40)),
            ('◄', (0.3, 0.5), 'left', (40, 60)),
            ('►', (0.7, 0.5), 'right', (40, 60))
        ]

        for symbol, pos, direction, size in arrows:
            btn = ModernButton(
                text=symbol,
                size_hint=(None, None),
                size=size,
                pos_hint={'center_x': pos[0], 'center_y': pos[1]},
                bg_color=(0.1, 0.1, 0.1, 1),
                color=(1, 1, 1, 1),
                font_size='20sp',
                bold=True
            )
            btn.bind(on_press=lambda x, d=direction: self.move_focus(d))
            self.add_widget(btn)

        self.bind(pos=self.update_bg, size=self.update_bg)

    def update_bg(self, *args):
        self.bg_rect.pos = self.pos
        self.bg_rect.size = self.size

    def move_focus(self, direction):
        current = self.table_screen.focused_cell
        if not current: return
        row, col = current

        if direction == 'up': row = max(0, row-1)
        elif direction == 'down': row = min(len(self.table_screen.inputs)-1, row+1)
        elif direction == 'left': col = max(0, col-1)
        elif direction == 'right': col = min(5, col+1)  # 6 столбцов (0-5)

        self.table_screen.focused_cell = [row, col]
        inp = self.table_screen.inputs[row][col]
        inp.focus = True
        inp.cursor = (len(inp.text), 0)
        Clock.schedule_once(lambda dt: self._update_cursor(inp), 0.01)

    def _update_cursor(self, inp):
        inp.focus = True
        inp.cursor = (len(inp.text), 0)
        inp.text = inp.text

class MolodnikiTreeDataInputPopup(Popup):
    def __init__(self, table_screen, row_index, **kwargs):
        super().__init__(
            title="",
            size_hint=(0.85, None),
            height=dp(520),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3],
            **kwargs
        )
        self.table_screen = table_screen
        self.row_index = row_index
        self.fields = [
            ('GPS точка', 1),
            ('Предмет ухода', 2),
            ('Порода', 3),
            ('Примечания', 4),
            ('Тип Леса', 5)
        ]
        self.data = {}
        self.create_ui()

    def create_ui(self):
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)

        title_label = MDLabel(
            text=f"Площадка №{self.row_index + 1}",
            font_style='Headline', role='medium',
            theme_text_color='Custom', text_color=Colors.GREEN,
            size_hint_y=None, height=dp(40)
        )
        content.add_widget(title_label)

        scroll = ScrollView(size_hint=(1, None), height=dp(320))
        scroll_content = MDGridLayout(cols=1, spacing=Spacing.MD, adaptive_height=True,
                                      size_hint_y=None, padding=[0, 0])

        self.input_fields = []
        for field_name, col_index in self.fields:
            field_label = MDLabel(
                text=field_name,
                font_style='Label', role='medium',
                theme_text_color='Custom', text_color=[1,1,1,1],
                size_hint_y=None, height=dp(24)
            )
            input_field = AutoCompleteTextInput(
                multiline=False,
                size_hint_y=None,
                height=dp(44),
                background_color=[0.15,0.15,0.15,1],
                foreground_color=[1,1,1,1],
                cursor_color=[0.3,0.8,0.3,1],
                col_index=col_index,
                font_name='Roboto',
                padding=[dp(8), dp(10)]
            )
            if col_index == 3:
                input_field.bind(focus=self.show_breed_popup)
            self.input_fields.append(input_field)
            scroll_content.add_widget(field_label)
            scroll_content.add_widget(input_field)

        if self.table_screen.current_page in self.table_screen.page_data and self.row_index < len(self.table_screen.page_data[self.table_screen.current_page]):
            row_data = self.table_screen.page_data[self.table_screen.current_page][self.row_index]
            for i, (field_name, col_index) in enumerate(self.fields):
                if col_index < len(row_data) and row_data[col_index]:
                    self.input_fields[i].text = str(row_data[col_index])

        scroll.add_widget(scroll_content)
        content.add_widget(scroll)

        btn_layout = MDBoxLayout(orientation='horizontal', spacing=Spacing.MD, adaptive_height=True)
        save_btn = MDButton(style='filled', md_bg_color=Colors.BTN_SUCCESS,
                            size_hint=(0.5, None), height=dp(48))
        save_btn.add_widget(MDButtonText(text='Сохранить'))
        save_btn.bind(on_release=self.save_data)
        cancel_btn = MDButton(style='outlined', size_hint=(0.5, None), height=dp(48),
                              line_color=Colors.DANGER)
        cancel_btn.add_widget(MDButtonText(text='Отмена',
                              theme_text_color='Custom', text_color=Colors.DANGER))
        cancel_btn.bind(on_release=self.dismiss)
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        self.content = content
        self.open()

    def show_breed_popup(self, instance, value):
        """Показать popup для выбора типа породы"""
        if not value: return

        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)

        title_label = MDLabel(
            text='Выберите тип породы',
            font_style='Headline', role='medium',
            theme_text_color='Custom', text_color=Colors.GREEN,
            size_hint_y=None, height=dp(44)
        )
        content.add_widget(title_label)

        type_layout = MDBoxLayout(orientation='horizontal', spacing=Spacing.MD,
                                  adaptive_height=True)
        coniferous_btn = MDButton(style='filled', md_bg_color=Colors.BTN_SUCCESS,
                                  size_hint=(0.5, None), height=dp(56))
        coniferous_btn.add_widget(MDButtonText(text='Хвойные'))
        deciduous_btn = MDButton(style='filled', md_bg_color=Colors.GREEN,
                                 size_hint=(0.5, None), height=dp(56))
        deciduous_btn.add_widget(MDButtonText(text='Лиственные'))
        type_layout.add_widget(coniferous_btn)
        type_layout.add_widget(deciduous_btn)
        content.add_widget(type_layout)

        cancel_btn = MDButton(style='outlined', size_hint=(1, None), height=dp(48),
                              line_color=Colors.DANGER)
        cancel_btn.add_widget(MDButtonText(text='Отмена',
                              theme_text_color='Custom', text_color=Colors.DANGER))
        content.add_widget(cancel_btn)

        popup = Popup(
            title="", content=content, size_hint=(0.7, None), height=dp(220),
            separator_height=0, background_color=[0,0,0,0.3], overlay_color=[0,0,0,0.3]
        )

        def select_coniferous(btn):
            self.show_breed_selection_popup(instance, 'coniferous')
            popup.dismiss()

        def select_deciduous(btn):
            self.show_breed_selection_popup(instance, 'deciduous')
            popup.dismiss()

        coniferous_btn.bind(on_release=select_coniferous)
        deciduous_btn.bind(on_release=select_deciduous)
        cancel_btn.bind(on_release=popup.dismiss)

        popup.open()

    def show_breed_selection_popup(self, instance, breed_type):
        """Показать popup для выбора конкретной породы из словаря"""
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)

        title_label = MDLabel(
            text=f"Выберите {'хвойную' if breed_type == 'coniferous' else 'лиственную'} породу",
            font_style='Headline', role='medium',
            theme_text_color='Custom', text_color=Colors.GREEN,
            size_hint_y=None, height=dp(40)
        )
        content.add_widget(title_label)

        if breed_type == 'coniferous':
            breeds = ['Сосна', 'Ель', 'Лиственница', 'Кедр', 'Пихта']
        else:
            breeds = ['Берёза', 'Осина', 'Ольха чёрная', 'Ольха серая', 'Ива']

        custom_breeds = self.load_custom_breeds(breed_type)
        all_breeds = breeds + custom_breeds

        scroll = ScrollView(size_hint=(1, None), height=dp(300))
        breeds_layout = MDGridLayout(cols=1, spacing=Spacing.SM, adaptive_height=True,
                                     size_hint_y=None, padding=[0, 0])

        for breed in all_breeds:
            btn = MDButton(style='filled', md_bg_color=get_color_from_hex('#37474F'),
                           size_hint=(1, None), height=dp(48))
            btn.add_widget(MDButtonText(text=breed, theme_text_color='Custom', text_color=[1,1,1,1]))
            btn.bind(on_release=lambda x, b=breed: self.select_breed(instance, breed_type, b))
            breeds_layout.add_widget(btn)

        scroll.add_widget(breeds_layout)
        content.add_widget(scroll)

        buttons_layout = MDBoxLayout(orientation='vertical', spacing=Spacing.SM,
                                     adaptive_height=True)

        other_btn = MDButton(style='filled', md_bg_color=Colors.BTN_PURPLE,
                             size_hint=(1, None), height=dp(48))
        other_btn.add_widget(MDButtonText(text='Новая'))
        other_btn.bind(on_release=lambda x: self.select_breed(instance, breed_type, 'other'))
        buttons_layout.add_widget(other_btn)

        clear_btn = MDButton(style='filled', md_bg_color=Colors.BTN_WARNING,
                             size_hint=(1, None), height=dp(48))
        clear_btn.add_widget(MDButtonText(text='Очистить'))
        clear_btn.bind(on_release=lambda x: App.get_running_app().root.get_screen('molodniki').show_clear_breeds_popup(breed_type))
        buttons_layout.add_widget(clear_btn)

        cancel_btn = MDButton(style='outlined', size_hint=(1, None), height=dp(48),
                              line_color=Colors.DANGER)
        cancel_btn.add_widget(MDButtonText(text='Отмена',
                              theme_text_color='Custom', text_color=Colors.DANGER))
        buttons_layout.add_widget(cancel_btn)
        content.add_widget(buttons_layout)

        popup = Popup(
            title="", content=content, size_hint=(0.85, 0.85),
            separator_height=0, background_color=[0,0,0,0.3], overlay_color=[0,0,0,0.3]
        )

        cancel_btn.bind(on_release=popup.dismiss)
        popup.open()

    def update_plot_breeds_display(self, plot_breeds_list, existing_breeds_for_plot):
        """Обновить отображение списка пород на площадке"""
        # Очищаем текущий список
        plot_breeds_list.clear_widgets()

        if existing_breeds_for_plot:
            for i, breed_info in enumerate(existing_breeds_for_plot):
                if not isinstance(breed_info, dict):
                    continue
                breed_name = breed_info.get('name', 'Неизвестная')

                # Рассчитываем густоту (для хвойных - сумма градаций)
                density = breed_info.get('density', 0)
                if breed_info.get('type') == 'coniferous':
                    conif_density = (breed_info.get('do_05', 0) + breed_info.get('05_15', 0) + breed_info.get('bolee_15', 0))
                    if conif_density > 0:
                        density = conif_density

                height = breed_info.get('height', 0)
                age = breed_info.get('age', 0)
                diameter = breed_info.get('diameter', 0)

                # Формируем строку с параметрами породы
                params_parts = []
                if density:
                    params_parts.append(f"Густота: {density}")
                if height:
                    params_parts.append(f"Высота: {height}м")
                if diameter:
                    params_parts.append(f"Диаметр: {diameter}см")
                if age:
                    params_parts.append(f"Возраст: {age}л")

                params_text = ", ".join(params_parts) if params_parts else "Нет данных"

                # Для хвойных пород добавляем градации по высоте
                gradation_text = ""
                if breed_info.get('type') == 'coniferous':
                    gradations = []
                    if breed_info.get('do_05', 0) > 0:
                        gradations.append(f"До 0.5м: {breed_info['do_05']}")
                    if breed_info.get('05_15', 0) > 0:
                        gradations.append(f"0.5-1.5м: {breed_info['05_15']}")
                    if breed_info.get('bolee_15', 0) > 0:
                        gradations.append(f">1.5м: {breed_info['bolee_15']}")
                    if gradations:
                        gradation_text = " | " + ", ".join(gradations)

                breed_item = Label(
                    text=f'{i+1}. {breed_name} - {params_text}{gradation_text}',
                    font_name='Roboto',
                    font_size='12sp',
                    color=(0.2, 0.2, 0.2, 1),
                    size_hint=(1, None),
                    height=35 if gradation_text else 30,
                    halign='left',
                    valign='middle'
                )
                breed_item.bind(size=lambda *args: setattr(breed_item, 'text_size', (breed_item.width, None)))
                plot_breeds_list.add_widget(breed_item)
        else:
            no_breeds_label = Label(
                text='Породы еще не добавлены',
                font_name='Roboto',
                font_size='12sp',
                color=(0.5, 0.5, 0.5, 1),
                size_hint=(1, None),
                height=30,
                halign='left'
            )
            no_breeds_label.bind(size=lambda *args: setattr(no_breeds_label, 'text_size', (no_breeds_label.width, None)))
            plot_breeds_list.add_widget(no_breeds_label)

    def select_breed(self, instance, breed_type, selected_breed):
        """Обработка выбора породы"""
        if selected_breed == 'other':
            # Показываем popup для ввода названия другой породы
            self.show_custom_breed_popup(instance, breed_type)
        else:
            # Показываем popup с параметрами породы, передавая название выбранной породы
            self.show_breed_details_popup(instance, breed_type, selected_breed)

    def show_breed_details_popup(self, instance, breed_type, selected_breed=None):
        """Показать popup для управления множественными породами"""
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD, md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)

        # Заголовок
        title_label = MDLabel(
            text=f"Управление породами - {selected_breed}",
            font_style='Headline',
            role='medium',
            theme_text_color='Custom',
            text_color=Colors.GREEN,
            size_hint=(1, None),
            height=50
        )
        content.add_widget(title_label)

        # Бокс для отображения сохраненных пород по площадке (в начале popup)
        plot_breeds_box = MDBoxLayout(
            orientation='vertical',
            size_hint=(1, None),
            height=200,
            padding=[10, 10],
            spacing=Spacing.SM,
            md_bg_color=get_color_from_hex('#1E3A2E')
        )

        # Получаем номер площадки из row_index
        plot_number = self.row_index + 1

        plot_breeds_title = MDLabel(
            text=f'Площадка №{plot_number} - Сохраненные породы:',
            font_style='Title',
            role='small',
            theme_text_color='Custom',
            text_color=[1,1,1,1],
            size_hint=(1, None),
            height=30
        )
        plot_breeds_box.add_widget(plot_breeds_title)

        # ScrollView для списка пород на площадке
        plot_breeds_scroll = ScrollView(size_hint=(1, None), height=150)
        plot_breeds_list = MDGridLayout(cols=1, spacing=Spacing.SM, adaptive_height=True, size_hint_y=None, padding=[0, 0])
        plot_breeds_list.bind(minimum_height=plot_breeds_list.setter('height'))
        
        # Получаем существующие породы для этой площадки из instance.text (поле ввода породы)
        existing_breeds_for_plot = []
        if hasattr(instance, 'text') and instance.text:
            try:
                parsed = json.loads(instance.text) if isinstance(instance.text, str) else []
                if isinstance(parsed, list):
                    existing_breeds_for_plot = parsed
            except (json.JSONDecodeError, TypeError) as e:
                print(f"DEBUG: Error parsing breeds from instance.text: {e}")
                pass

        # Также проверяем page_data как резервный источник
        if not existing_breeds_for_plot:
            if self.table_screen.current_page in self.table_screen.page_data:
                page_data = self.table_screen.page_data[self.table_screen.current_page]
                if self.row_index < len(page_data) and page_data[self.row_index][3]:
                    try:
                        parsed = json.loads(page_data[self.row_index][3]) if isinstance(page_data[self.row_index][3], str) else []
                        if isinstance(parsed, list):
                            existing_breeds_for_plot = parsed
                    except (json.JSONDecodeError, TypeError) as e:
                        print(f"DEBUG: Error parsing breeds from page_data: {e}")
                        pass

        # Используем новый метод для отображения пород
        self.update_plot_breeds_display(plot_breeds_list, existing_breeds_for_plot)
        
        plot_breeds_scroll.add_widget(plot_breeds_list)
        plot_breeds_box.add_widget(plot_breeds_scroll)
        content.add_widget(plot_breeds_box)

        # Поля ввода для параметров породы с прокруткой
        scroll_fields = ScrollView(size_hint=(1, None), height=200)
        fields_layout = MDGridLayout(cols=2, spacing=5, size_hint_y=None, md_bg_color=Colors.DARK_SURFACE)
        fields_layout.bind(minimum_height=fields_layout.setter('height'))

        if breed_type == 'coniferous':
            fields = [
                ('До 0.5м:', 'do_05'),
                ('0.5-1.5м:', '05_15'),
                ('>1.5м:', 'bolee_15'),
                ('Высота (м):', 'height'),
                ('Диаметр (см):', 'diameter'),
                ('Густота:', 'density'),
                ('Возраст (лет):', 'age')
            ]
        else:
            fields = [
                ('Густота:', 'density'),
                ('Высота (м):', 'height'),
                ('Диаметр (см):', 'diameter'),
                ('Возраст (лет):', 'age')
            ]

        self.breed_inputs = {}
        for label_text, field_key in fields:
            lbl = MDLabel(text=label_text, font_style='Title', role='small', theme_text_color='Custom', text_color=[1,1,1,1], size_hint=(None, None), size=(120, 40))
            inp = MDTextField(
                mode="outlined",
                size_hint=(None, None),
                size=(120, 40),
                line_color_focus=Colors.GREEN
            )
            if field_key in ['density', 'age']:
                inp.input_filter = 'int'
            elif field_key == 'height':
                inp.input_filter = 'float'
            elif field_key in ['do_05', '05_15', 'bolee_15']:
                inp.input_filter = 'int'
                if breed_type == 'coniferous':
                    inp.bind(text=self.update_coniferous_density)
            fields_layout.add_widget(lbl)
            fields_layout.add_widget(inp)
            self.breed_inputs[field_key] = inp

        scroll_fields.add_widget(fields_layout)
        content.add_widget(scroll_fields)

        # Кнопки управления - четыре отдельные кнопки
        btn_layout = MDBoxLayout(orientation='horizontal', spacing=Spacing.MD, size_hint=(1, None), height=70, md_bg_color=Colors.DARK_SURFACE)

        add_btn = MDButton(style='filled', size_hint=(0.25, None), height=dp(70))
        add_btn.add_widget(MDButtonText(text='Добавить', theme_text_color='Custom', text_color=Colors.GREEN))

        save_btn = MDButton(style='filled', size_hint=(0.25, None), height=dp(70))
        save_btn.add_widget(MDButtonText(text='Сохранить', theme_text_color='Custom', text_color=Colors.GREEN))

        delete_btn = MDButton(style='filled', size_hint=(0.25, None), height=dp(70))
        delete_btn.add_widget(MDButtonText(text='Удалить', theme_text_color='Custom', text_color=[1,0.3,0.3,1]))

        exit_btn = MDButton(style='filled', size_hint=(0.25, None), height=dp(70))
        exit_btn.add_widget(MDButtonText(text='Выйти', theme_text_color='Custom', text_color=[0.7,0.7,0.7,1]))

        btn_layout.add_widget(add_btn)
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(delete_btn)
        btn_layout.add_widget(exit_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            size_hint=(0.85, 0.85),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3]
        )

        def add_breed(btn):
            """Добавить породу - открывает меню выбора типа породы"""
            # Сначала сохраняем текущие данные если они есть
            has_data = any(inp.text.strip() for inp in self.breed_inputs.values())
            if has_data:
                # Сохраняем текущую породу
                breed_data = {
                    'name': selected_breed,
                    'type': breed_type
                }

                for key, inp in self.breed_inputs.items():
                    if inp.text.strip():
                        try:
                            if key in ['density', 'age']:
                                breed_data[key] = int(inp.text)
                            elif key == 'height':
                                breed_data[key] = float(inp.text)
                            else:
                                breed_data[key] = float(inp.text)
                        except ValueError:
                            breed_data[key] = 0 if key in ['density', 'age', 'do_05', '05_15', 'bolee_15'] else 0.0

                # Рассчитать возраст для хвойных пород, если не введен
                if breed_type == 'coniferous':
                    if 'age' not in breed_data or breed_data['age'] == 0:
                        do_05 = breed_data.get('do_05', 0)
                        _05_15 = breed_data.get('05_15', 0)
                        bolee_15 = breed_data.get('bolee_15', 0)
                        height = breed_data.get('height', 0)

                        if bolee_15 > 0:
                            breed_data['age'] = 20
                        elif _05_15 > 0:
                            breed_data['age'] = 10
                        elif do_05 > 0:
                            breed_data['age'] = 5
                        elif height > 0:
                            if height < 0.5:
                                breed_data['age'] = 3
                            elif height < 1.5:
                                breed_data['age'] = 10
                            else:
                                breed_data['age'] = 20
                        else:
                            breed_data['age'] = 10

                existing_breeds = self.table_screen.parse_breeds_data(instance.text)
                existing_breeds.append(breed_data)
                instance.text = json.dumps(existing_breeds, ensure_ascii=False, indent=2)

                # Update page_data
                if self.table_screen.current_page not in self.table_screen.page_data:
                    self.table_screen.page_data[self.table_screen.current_page] = []
                while len(self.table_screen.page_data[self.table_screen.current_page]) <= self.row_index:
                    self.table_screen.page_data[self.table_screen.current_page].append(['', '', '', '', '', ''])
                self.table_screen.page_data[self.table_screen.current_page][self.row_index][3] = instance.text

                self.table_screen.update_plot_total(instance, instance.text)

                # Очищаем поля ввода
                for inp in self.breed_inputs.values():
                    inp.text = ''

                # Обновляем отображение с градациями по высоте
                self.update_plot_breeds_display(plot_breeds_list, existing_breeds)

                # Сохраняем в базу данных
                self.table_screen.save_current_page()

                # Показываем сообщение об успехе
                self.table_screen.show_success(f"Порода добавлена! Всего пород: {len(existing_breeds)}")

            # Закрываем текущий popup с параметрами породы
            popup.dismiss()
            
            # Открываем popup выбора типа породы для добавления следующей породы
            from kivy.clock import Clock
            Clock.schedule_once(lambda dt: self.show_breed_popup(instance, True), 0.2)

        def save_data(btn):
            """Сохранить данные - отображает породы в блоке и сохраняет в БД"""
            # Если есть данные в полях ввода, добавляем их
            has_data = any(inp.text.strip() for inp in self.breed_inputs.values())
            if has_data:
                breed_data = {
                    'name': selected_breed,
                    'type': breed_type
                }

                for key, inp in self.breed_inputs.items():
                    if inp.text.strip():
                        try:
                            if key in ['density', 'age']:
                                breed_data[key] = int(inp.text)
                            elif key == 'height':
                                breed_data[key] = float(inp.text)
                            else:
                                breed_data[key] = float(inp.text)
                        except ValueError:
                            breed_data[key] = 0 if key in ['density', 'age', 'do_05', '05_15', 'bolee_15'] else 0.0

                # Рассчитать возраст для хвойных пород, если не введен
                if breed_type == 'coniferous':
                    if 'age' not in breed_data or breed_data['age'] == 0:
                        do_05 = breed_data.get('do_05', 0)
                        _05_15 = breed_data.get('05_15', 0)
                        bolee_15 = breed_data.get('bolee_15', 0)
                        height = breed_data.get('height', 0)

                        if bolee_15 > 0:
                            breed_data['age'] = 20
                        elif _05_15 > 0:
                            breed_data['age'] = 10
                        elif do_05 > 0:
                            breed_data['age'] = 5
                        elif height > 0:
                            if height < 0.5:
                                breed_data['age'] = 3
                            elif height < 1.5:
                                breed_data['age'] = 10
                            else:
                                breed_data['age'] = 20
                        else:
                            breed_data['age'] = 10

                existing_breeds = self.table_screen.parse_breeds_data(instance.text)
                existing_breeds.append(breed_data)
                instance.text = json.dumps(existing_breeds, ensure_ascii=False, indent=2)

                # Update page_data
                if self.table_screen.current_page not in self.table_screen.page_data:
                    self.table_screen.page_data[self.table_screen.current_page] = []
                while len(self.table_screen.page_data[self.table_screen.current_page]) <= self.row_index:
                    self.table_screen.page_data[self.table_screen.current_page].append(['', '', '', '', '', ''])
                self.table_screen.page_data[self.table_screen.current_page][self.row_index][3] = instance.text

                self.table_screen.update_plot_total(instance, instance.text)

                # Очищаем поля ввода
                for inp in self.breed_inputs.values():
                    inp.text = ''

            # Получаем актуальный список пород для обновления отображения
            existing_breeds = self.table_screen.parse_breeds_data(instance.text)

            # Обновляем отображение сохраненных пород с градациями по высоте
            self.update_plot_breeds_display(plot_breeds_list, existing_breeds)

            # Сохраняем в базу данных
            self.table_screen.save_current_page()

            # Показываем сообщение об успехе
            self.table_screen.show_success(f"Данные сохранены! Всего пород: {len(existing_breeds)}")

        def exit_popup(btn):
            """Выйти - автоматически сохраняет все данные в БД"""
            # Если есть несохраненные данные в полях ввода, сохраняем их
            has_data = any(inp.text.strip() for inp in self.breed_inputs.values())
            if has_data:
                breed_data = {
                    'name': selected_breed,
                    'type': breed_type
                }

                for key, inp in self.breed_inputs.items():
                    if inp.text.strip():
                        try:
                            if key in ['density', 'age']:
                                breed_data[key] = int(inp.text)
                            elif key == 'height':
                                breed_data[key] = float(inp.text)
                            else:
                                breed_data[key] = float(inp.text)
                        except ValueError:
                            breed_data[key] = 0 if key in ['density', 'age', 'do_05', '05_15', 'bolee_15'] else 0.0

                # Рассчитать возраст для хвойных пород
                if breed_type == 'coniferous':
                    if 'age' not in breed_data or breed_data['age'] == 0:
                        do_05 = breed_data.get('do_05', 0)
                        _05_15 = breed_data.get('05_15', 0)
                        bolee_15 = breed_data.get('bolee_15', 0)
                        height = breed_data.get('height', 0)

                        if bolee_15 > 0:
                            breed_data['age'] = 20
                        elif _05_15 > 0:
                            breed_data['age'] = 10
                        elif do_05 > 0:
                            breed_data['age'] = 5
                        elif height > 0:
                            if height < 0.5:
                                breed_data['age'] = 3
                            elif height < 1.5:
                                breed_data['age'] = 10
                            else:
                                breed_data['age'] = 20
                        else:
                            breed_data['age'] = 10

                existing_breeds = self.table_screen.parse_breeds_data(instance.text)
                existing_breeds.append(breed_data)
                instance.text = json.dumps(existing_breeds, ensure_ascii=False, indent=2)

                # Update page_data
                if self.table_screen.current_page not in self.table_screen.page_data:
                    self.table_screen.page_data[self.table_screen.current_page] = []
                while len(self.table_screen.page_data[self.table_screen.current_page]) <= self.row_index:
                    self.table_screen.page_data[self.table_screen.current_page].append(['', '', '', '', '', ''])
                self.table_screen.page_data[self.table_screen.current_page][self.row_index][3] = instance.text

                self.table_screen.update_plot_total(instance, instance.text)

            # Получаем актуальный список пород для обновления отображения
            existing_breeds = self.table_screen.parse_breeds_data(instance.text)

            # Обновляем отображение пород с градациями
            self.update_plot_breeds_display(plot_breeds_list, existing_breeds)

            # Всегда сохраняем в базу данных при выходе
            self.table_screen.save_current_page()

            self.table_screen.show_success(f"Данные автоматически сохранены в БД! Всего пород: {len(existing_breeds)}")

            popup.dismiss()

        add_btn.bind(on_release=add_breed)
        save_btn.bind(on_release=save_data)
        delete_btn.bind(on_release=lambda x: self.show_delete_breed_popup(instance, plot_breeds_list))
        exit_btn.bind(on_release=exit_popup)

        popup.open()

    def show_delete_breed_popup(self, instance, plot_breeds_list):
        """Показать popup для удаления пород"""
        # Получаем список пород
        existing_breeds = self.table_screen.parse_breeds_data(instance.text)

        if not existing_breeds:
            self.table_screen.show_error("Нет пород для удаления!")
            return

        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD, md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)

        # Заголовок
        title_label = MDLabel(
            text="Выберите породы для удаления:",
            font_style='Headline',
            role='medium',
            theme_text_color='Custom',
            text_color=Colors.GREEN,
            size_hint=(1, None),
            height=50
        )
        content.add_widget(title_label)

        # ScrollView для списка пород
        scroll = ScrollView(size_hint=(1, None), height=350)
        breeds_layout = MDGridLayout(cols=1, spacing=Spacing.SM, adaptive_height=True, size_hint_y=None, padding=[0, 0])

        # Чекбоксы для выбора пород
        self.breed_checkboxes = {}
        for i, breed_info in enumerate(existing_breeds):
            breed_name = breed_info.get('name', 'Неизвестная')

            # Создаём строку с чекбоксом и названием породы
            breed_row = MDBoxLayout(orientation='horizontal', spacing=Spacing.MD, size_hint=(1, None), height=50)

            # Чекбокс (используем Button как чекбокс)
            from kivy.uix.checkbox import CheckBox
            checkbox = CheckBox(size_hint=(None, None), size=(50, 50), active=False)
            self.breed_checkboxes[i] = checkbox

            # Название породы
            breed_label = MDLabel(
                text=f"{breed_name}",
                font_style='Title',
                role='small',
                theme_text_color='Custom',
                text_color=[1,1,1,1],
                size_hint=(1, None),
                height=50
            )

            breed_row.add_widget(checkbox)
            breed_row.add_widget(breed_label)
            breeds_layout.add_widget(breed_row)

        scroll.add_widget(breeds_layout)
        content.add_widget(scroll)

        # Кнопки управления
        btn_layout = MDBoxLayout(orientation='horizontal', spacing=Spacing.MD, size_hint=(1, None), height=70, md_bg_color=Colors.DARK_SURFACE)

        confirm_btn = MDButton(style='filled', size_hint=(0.34, None), height=dp(70))
        confirm_btn.add_widget(MDButtonText(text='Удалить выбранные', theme_text_color='Custom', text_color=[1,0.3,0.3,1]))

        clear_all_btn = MDButton(style='filled', size_hint=(0.33, None), height=dp(70))
        clear_all_btn.add_widget(MDButtonText(text='Очистить все', theme_text_color='Custom', text_color=[0.7,0.7,0.7,1]))

        cancel_btn = MDButton(style='filled', size_hint=(0.33, None), height=dp(70))
        cancel_btn.add_widget(MDButtonText(text='Отмена', theme_text_color='Custom', text_color=[0.7,0.7,0.7,1]))

        btn_layout.add_widget(confirm_btn)
        btn_layout.add_widget(clear_all_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            size_hint=(0.85, 0.85),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3]
        )

        def confirm_delete(btn):
            # Получаем индексы пород для удаления
            indices_to_delete = sorted(
                [idx for idx, cb in self.breed_checkboxes.items() if cb.active],
                reverse=True
            )

            if not indices_to_delete:
                self.table_screen.show_error("Выберите хотя бы одну породу для удаления!")
                return

            # Удаляем породы
            for idx in indices_to_delete:
                if 0 <= idx < len(existing_breeds):
                    existing_breeds.pop(idx)

            # Обновляем данные
            instance.text = json.dumps(existing_breeds, ensure_ascii=False, indent=2) if existing_breeds else ''

            # Update page_data
            if self.table_screen.current_page in self.table_screen.page_data:
                page_data = self.table_screen.page_data[self.table_screen.current_page]
                if self.row_index < len(page_data):
                    page_data[self.row_index][3] = instance.text

            self.table_screen.update_plot_total(instance, instance.text)

            # Сохраняем в базу данных
            self.table_screen.save_current_page()

            # Обновляем отображение пород в основном popup
            self.update_plot_breeds_display(plot_breeds_list, existing_breeds)

            self.table_screen.show_success(f"Удалено пород: {len(indices_to_delete)}")
            popup.dismiss()

        def clear_all(btn):
            # Очищаем все породы
            instance.text = ''

            # Update page_data
            if self.table_screen.current_page in self.table_screen.page_data:
                page_data = self.table_screen.page_data[self.table_screen.current_page]
                if self.row_index < len(page_data):
                    page_data[self.row_index][3] = ''

            self.table_screen.update_plot_total(instance, '')

            # Сохраняем в базу данных
            self.table_screen.save_current_page()

            # Обновляем отображение пород в основном popup
            self.update_plot_breeds_display(plot_breeds_list, [])

            self.table_screen.show_success("Все породы очищены!")
            popup.dismiss()

        confirm_btn.bind(on_release=confirm_delete)
        clear_all_btn.bind(on_release=clear_all)
        cancel_btn.bind(on_release=popup.dismiss)

        popup.open()

    def show_breed_choice_popup(self, instance, selected_breed):
        """Показать popup с выбором после добавления первой породы"""
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD, md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)

        # Заголовок
        title_label = MDLabel(
            text=f"Порода '{selected_breed}' добавлена!",
            font_style='Headline',
            role='medium',
            theme_text_color='Custom',
            text_color=Colors.GREEN,
            size_hint=(1, None),
            height=40
        )
        content.add_widget(title_label)

        # Информация о номере породы
        info_label = MDLabel(
            text="Выберите действие:",
            font_style='Title',
            role='small',
            theme_text_color='Custom',
            text_color=[1,1,1,1],
            size_hint=(1, None),
            height=30
        )
        content.add_widget(info_label)

        btn_layout = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, size_hint=(1, None), height=150)
        add_more_btn = MDButton(style='filled', md_bg_color=get_color_from_hex('#00FF00'), size_hint=(1, None), height=dp(65))
        add_more_btn.add_widget(MDButtonText(text='Добавить еще породу'))
        save_exit_btn = MDButton(style='filled', md_bg_color=get_color_from_hex('#32CD32'), size_hint=(1, None), height=dp(65))
        save_exit_btn.add_widget(MDButtonText(text='Сохранить и выйти'))
        btn_layout.add_widget(add_more_btn)
        btn_layout.add_widget(save_exit_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            size_hint=(0.85, 0.85),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3]
        )

        def add_more_breed(btn):
            popup.dismiss()
            self.show_breed_popup(instance, True)

        def save_and_exit(btn):
            popup.dismiss()
            self.table_screen.show_success("Данные по площадке сохранены!")

        add_more_btn.bind(on_release=add_more_breed)
        save_exit_btn.bind(on_release=save_and_exit)

        popup.open()

    def show_custom_breed_popup(self, instance, breed_type):
        """Показать popup для ввода названия другой породы"""
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD, md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)

        # Заголовок
        title_label = MDLabel(
            text="Введите название другой породы",
            font_style='Headline',
            role='medium',
            theme_text_color='Custom',
            text_color=Colors.GREEN,
            size_hint=(1, None),
            height=50
        )
        content.add_widget(title_label)

        self.custom_breed_input = MDTextField(
            hint_text="Название породы",
            mode="outlined",
            size_hint=(1, None),
            height=50,
            line_color_focus=Colors.GREEN
        )
        content.add_widget(self.custom_breed_input)

        btn_layout = MDBoxLayout(orientation='horizontal', spacing=Spacing.MD, size_hint=(1, None), height=70, md_bg_color=Colors.DARK_SURFACE)
        save_btn = MDButton(style='filled', size_hint=(0.5, None), height=dp(70))
        save_btn.add_widget(MDButtonText(text='Сохранить', theme_text_color='Custom', text_color=Colors.GREEN))
        cancel_btn = MDButton(style='filled', size_hint=(0.5, None), height=dp(70))
        cancel_btn.add_widget(MDButtonText(text='Отмена', theme_text_color='Custom', text_color=[1,0.3,0.3,1]))
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            size_hint=(0.85, 0.85),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3]
        )

        def save_custom_breed(btn):
            breed_name = self.custom_breed_input.text.strip()
            if breed_name:
                # Проверяем, не является ли порода запрещенной
                forbidden_breeds = ['семенная', 'культуры', 'подрост']
                if any(forbidden.lower() in breed_name.lower() for forbidden in forbidden_breeds):
                    self.table_screen.show_error("Эта порода не разрешена для использования!")
                    return

                # Сохраняем новую породу в базу данных
                self.save_custom_breed_to_db(breed_name, breed_type)

                # Возвращаемся в меню выбора типа породы с новым названием
                popup.dismiss()
                from kivy.clock import Clock
                Clock.schedule_once(lambda dt: self.show_breed_selection_popup(instance, breed_type), 0.15)
            else:
                self.table_screen.show_error("Название породы не может быть пустым!")

        save_btn.bind(on_release=save_custom_breed)
        cancel_btn.bind(on_release=popup.dismiss)

        popup.open()

    def update_coniferous_density(self, instance, value):
        """Автоматический расчет густоты для хвойных пород"""
        if 'density' in self.breed_inputs:
            density_input = self.breed_inputs['density']
            try:
                do_05 = int(self.breed_inputs.get('do_05', TextInput(text='0')).text or '0')
                _05_15 = int(self.breed_inputs.get('05_15', TextInput(text='0')).text or '0')
                bolee_15 = int(self.breed_inputs.get('bolee_15', TextInput(text='0')).text or '0')

                total_density = do_05 + _05_15 + bolee_15
                density_input.text = str(total_density) if total_density > 0 else ''
            except (ValueError, AttributeError):
                pass

    def save_custom_breed_to_db(self, breed_name, breed_type):
        """Сохранить новую породу в базу данных"""
        conn = sqlite3.connect('forest_data.db')
        cursor = conn.cursor()
        try:
            cursor.execute('''
                INSERT INTO custom_breeds (breed_name, breed_type)
                VALUES (?, ?)
            ''', (breed_name, breed_type))
            conn.commit()
            return True
        except sqlite3.IntegrityError:
            # Порода уже существует
            return False
        finally:
            conn.close()

    def load_custom_breeds(self, breed_type):
        """Загрузить пользовательские породы из базы данных"""
        conn = sqlite3.connect('forest_data.db')
        cursor = conn.cursor()
        cursor.execute('''
            SELECT breed_name FROM custom_breeds
            WHERE breed_type = ?
            ORDER BY breed_name
        ''', (breed_type,))
        results = [row[0] for row in cursor.fetchall()]
        conn.close()
        return results

    def save_data(self, instance):
        # Save to page_data
        for i, (field_name, col_index) in enumerate(self.fields):
            value = self.input_fields[i].text.strip()
            if value:
                if self.table_screen.current_page not in self.table_screen.page_data:
                    self.table_screen.page_data[self.table_screen.current_page] = []
                while len(self.table_screen.page_data[self.table_screen.current_page]) <= self.row_index:
                    self.table_screen.page_data[self.table_screen.current_page].append(['', '', '', '', '', ''])
                self.table_screen.page_data[self.table_screen.current_page][self.row_index][col_index] = value

        # Save to database
        self.table_screen.save_current_page()

        # Show success
        self.table_screen.show_success("Данные площадки молодняков сохранены!")
        self.dismiss()

class ExtendedMolodnikiTableScreen(Screen):
    current_page = NumericProperty(0)
    total_pages = NumericProperty(1)
    unsaved_changes = BooleanProperty(False)
    focused_cell = ListProperty([0, 0])
    edit_mode = BooleanProperty(False)
    current_section = StringProperty("")
    current_quarter = StringProperty("")
    current_plot = StringProperty("")
    current_forestry = StringProperty("")
    current_radius = StringProperty("5.64")
    current_plot_area_ha = StringProperty("")
    plot_area_input = StringProperty("")
    current_address_display_text = StringProperty("")
    MAX_PAGES = 200

    # Данные проекта ухода
    care_queue = StringProperty("")
    characteristics = StringProperty("")
    care_date = StringProperty("")
    technology = StringProperty("")
    forest_purpose = StringProperty("")

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        print("DEBUG: ExtendedMolodnikiTableScreen __init__ started")
        self.theme_manager = App.get_running_app().theme_manager
        self.reports_dir = "reports"
        os.makedirs(self.reports_dir, exist_ok=True)
        self.db_name = 'forest_data.db'
        self.rows_per_page = 30
        self.page_data = {}
        self.setup_database()
        self.create_ui()
        self.load_existing_data()
        Window.bind(on_key_down=self.key_action)

        # Инициализация данных проекта
        self.project_data = {
            'address': {
                'quarter': '',
                'plot': '',
                'forestry': '',
                'district_forestry': '',
                'radius': '5.64',
                'plot_area': ''
            },
            'details': {
                'care_queue': '',
                'characteristics': '',
                'care_date': '',
                'technology': '',
                'forest_purpose': ''
            },
            'document_name': 'Проект'
        }

        # Убираем вызов update_section_label, так как section_label больше не существует

    def key_action(self, window, key, scancode, codepoint, modifier):
        if key == 115 and 'ctrl' in modifier:
            self.save_current_page()

    def setup_database(self):
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()

        # Создаем таблицу для хранения данных молодняков
        cursor.execute('''CREATE TABLE IF NOT EXISTS molodniki_data (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        page_number INTEGER,
                        row_index INTEGER,
                        nn INTEGER,
                        gps_point TEXT,
                        predmet_uhoda TEXT,
                        radius REAL DEFAULT 5.64,
                        primechanie TEXT,
                        section_name TEXT,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')

        # Создаем индексы для быстрого поиска
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_molodniki_data_page ON molodniki_data (page_number, row_index)')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_molodniki_data_section ON molodniki_data (section_name)')

        # Создаем таблицу для хранения пород (множественные породы на одну запись)
        cursor.execute('''CREATE TABLE IF NOT EXISTS molodniki_breeds (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        molodniki_data_id INTEGER,
                        breed_name TEXT,
                        breed_type TEXT, -- 'coniferous' или 'deciduous'
                        do_05 INTEGER DEFAULT 0,
                        _05_15 INTEGER DEFAULT 0,
                        bolee_15 INTEGER DEFAULT 0,
                        density INTEGER DEFAULT 0,
                        height REAL DEFAULT 0.0,
                        diameter REAL DEFAULT 0.0,
                        age INTEGER DEFAULT 0,
                        composition_coefficient REAL DEFAULT 0.0,
                        FOREIGN KEY(molodniki_data_id) REFERENCES molodniki_data(id) ON DELETE CASCADE)''')

        # Создаем индекс для поиска данных пород
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_molodniki_breeds ON molodniki_breeds (molodniki_data_id)')

        # Добавляем недостающие столбцы, если они отсутствуют
        try:
            cursor.execute('ALTER TABLE molodniki_breeds ADD COLUMN diameter REAL DEFAULT 0.0')
        except sqlite3.OperationalError:
            pass  # Столбец уже существует

        try:
            cursor.execute('ALTER TABLE molodniki_breeds ADD COLUMN composition_coefficient REAL DEFAULT 0.0')
        except sqlite3.OperationalError:
            pass  # Столбец уже существует

        # Добавляем колонку poroda если её нет (для хранения JSON пород)
        try:
            cursor.execute("ALTER TABLE molodniki_data ADD COLUMN poroda TEXT DEFAULT ''")
        except sqlite3.OperationalError:
            pass  # Столбец уже существует

        # Добавляем колонку radius если её нет
        try:
            cursor.execute("ALTER TABLE molodniki_data ADD COLUMN radius REAL DEFAULT 5.64")
        except sqlite3.OperationalError:
            pass  # Столбец уже существует

        # Создаем таблицу для хранения итогов по страницам
        cursor.execute('''CREATE TABLE IF NOT EXISTS molodniki_totals (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        page_number INTEGER,
                        section_name TEXT,
                        total_composition TEXT,
                        total_area REAL DEFAULT 0.0,
                        avg_age REAL DEFAULT 0.0,
                        avg_density REAL DEFAULT 0.0,
                        avg_height REAL DEFAULT 0.0,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')

        cursor.execute('CREATE INDEX IF NOT EXISTS idx_molodniki_totals_page ON molodniki_totals (page_number, section_name)')

        # Создаем таблицу для хранения настроек участка
        cursor.execute('''CREATE TABLE IF NOT EXISTS molodniki_settings (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        section_name TEXT UNIQUE,
                        radius REAL DEFAULT 5.64,
                        plot_area REAL DEFAULT 0.0,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')

        cursor.execute('CREATE INDEX IF NOT EXISTS idx_molodniki_settings_section ON molodniki_settings (section_name)')

        # Создаем таблицу для хранения данных пород (JSON)
        cursor.execute('''CREATE TABLE IF NOT EXISTS molodniki_suggestions (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        column_index INTEGER,
                        value TEXT,
                        UNIQUE(column_index, value))''')

        cursor.execute('CREATE INDEX IF NOT EXISTS idx_molodniki_suggestions ON molodniki_suggestions (column_index, value)')

        # Создаем таблицу для хранения пользовательских пород
        cursor.execute('''CREATE TABLE IF NOT EXISTS custom_breeds (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        breed_name TEXT UNIQUE,
                        breed_type TEXT, -- 'coniferous' или 'deciduous'
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')

        cursor.execute('CREATE INDEX IF NOT EXISTS idx_custom_breeds ON custom_breeds (breed_type)')

        conn.commit()
        conn.close()

    def save_custom_breed_to_db(self, breed_name, breed_type):
        """Сохранить новую породу в базу данных"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        try:
            cursor.execute('''
                INSERT INTO custom_breeds (breed_name, breed_type)
                VALUES (?, ?)
            ''', (breed_name, breed_type))
            conn.commit()
            return True
        except sqlite3.IntegrityError:
            # Порода уже существует
            return False
        finally:
            conn.close()

    def load_custom_breeds(self, breed_type):
        """Загрузить пользовательские породы из базы данных"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        cursor.execute('''
            SELECT breed_name FROM custom_breeds
            WHERE breed_type = ?
            ORDER BY breed_name
        ''', (breed_type,))
        results = [row[0] for row in cursor.fetchall()]
        conn.close()
        return results

    def show_clear_breeds_popup(self, breed_type):
        """Показать popup для очистки пользовательских пород с выбором через галочки"""
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD, md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)

        # Заголовок
        title_label = MDLabel(
            text=f"Очистка {'хвойных' if breed_type == 'coniferous' else 'лиственных'} пород",
            font_style='Headline',
            role='medium',
            theme_text_color='Custom',
            text_color=Colors.GREEN,
            size_hint=(1, None),
            height=50
        )
        content.add_widget(title_label)

        # Описание
        desc_label = MDLabel(
            text="Отметьте породы для удаления:",
            font_style='Title',
            role='small',
            theme_text_color='Custom',
            text_color=[1,1,1,1],
            size_hint=(1, None),
            height=35
        )
        content.add_widget(desc_label)

        # Загружаем пользовательские породы
        custom_breeds = self.load_custom_breeds(breed_type)

        if not custom_breeds:
            no_breeds_label = MDLabel(
                text="Нет пользовательских пород для удаления",
                font_style='Title',
                role='small',
                theme_text_color='Custom',
                text_color=[0.5,0.5,0.5,1],
                size_hint=(1, None),
                height=40
            )
            content.add_widget(no_breeds_label)
        else:
            # ScrollView для списка пород
            scroll = ScrollView(size_hint=(1, None), height=300)
            breeds_layout = MDGridLayout(cols=1, spacing=Spacing.SM, adaptive_height=True, size_hint_y=None, padding=[0, 0])

            # Хранилище для чекбоксов
            self.breed_checkboxes = {}

            for breed in custom_breeds:
                # Создаем горизонтальный layout для чекбокса и названия
                breed_row = MDBoxLayout(orientation='horizontal', size_hint_y=None, height=50, spacing=Spacing.MD)

                # Чекбокс (используем Button как чекбокс)
                from kivy.uix.checkbox import CheckBox
                checkbox = CheckBox(active=False, size_hint=(None, None), size=(50, 50))
                self.breed_checkboxes[breed] = checkbox

                # Название породы
                breed_label = MDLabel(
                    text=breed,
                    font_style='Title',
                    role='small',
                    theme_text_color='Custom',
                    text_color=[1,1,1,1],
                    size_hint=(1, None),
                    height=50
                )

                breed_row.add_widget(checkbox)
                breed_row.add_widget(breed_label)
                breeds_layout.add_widget(breed_row)

            scroll.add_widget(breeds_layout)
            content.add_widget(scroll)

        # Кнопки управления
        btn_layout = MDBoxLayout(orientation='horizontal', spacing=Spacing.MD, size_hint=(1, None), height=70)

        # Кнопка "Выбрать все"
        select_all_btn = MDButton(style='filled', md_bg_color=get_color_from_hex('#4169E1'), size_hint=(0.33, None), height=dp(70))
        select_all_btn.add_widget(MDButtonText(text='Выбрать все'))

        # Кнопка "Удалить"
        delete_btn = MDButton(style='filled', md_bg_color=get_color_from_hex('#FF0000'), size_hint=(0.33, None), height=dp(70))
        delete_btn.add_widget(MDButtonText(text='Удалить'))

        # Кнопка "Отмена"
        cancel_btn = MDButton(style='filled', md_bg_color=get_color_from_hex('#808080'), size_hint=(0.33, None), height=dp(70))
        cancel_btn.add_widget(MDButtonText(text='Отмена'))

        btn_layout.add_widget(select_all_btn)
        btn_layout.add_widget(delete_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            size_hint=(0.85, 0.85),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3]
        )

        def select_all(instance):
            """Выбрать все породы"""
            for checkbox in self.breed_checkboxes.values():
                checkbox.active = True

        def delete_selected(instance):
            """Удалить выбранные породы"""
            breeds_to_delete = [breed for breed, checkbox in self.breed_checkboxes.items() if checkbox.active]
            
            if not breeds_to_delete:
                # Показываем сообщение что ничего не выбрано
                error_content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD, md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)
                error_content.add_widget(MDLabel(
                    text="Не выбрано ни одной породы для удаления",
                    font_style='Title',
                    role='small',
                    theme_text_color='Custom',
                    text_color=[1,1,1,1]
                ))
                error_popup = Popup(
                    title="",
                    content=error_content,
                    size_hint=(0.85, 0.85),
                    separator_height=0,
                    background_color=[0,0,0,0.3],
                    overlay_color=[0,0,0,0.3],
                    auto_dismiss=True
                )
                error_popup.open()
                return

            # Подтверждение удаления
            confirm_content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD, md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)
            
            # Заголовок
            confirm_title = MDLabel(
                text="Подтверждение удаления",
                font_style='Headline',
                role='medium',
                theme_text_color='Custom',
                text_color=Colors.GREEN,
                size_hint=(1, None),
                height=40
            )
            confirm_content.add_widget(confirm_title)
            
            confirm_label = MDLabel(
                text=f"Вы уверены, что хотите удалить {len(breeds_to_delete)} пород(ы)?",
                font_style='Title',
                role='small',
                theme_text_color='Custom',
                text_color=[1,1,1,1],
                size_hint=(1, None),
                height=50
            )
            confirm_content.add_widget(confirm_label)

            confirm_btn_layout = MDBoxLayout(orientation='horizontal', spacing=Spacing.MD, size_hint=(1, None), height=70)
            confirm_yes_btn = MDButton(style='filled', md_bg_color=get_color_from_hex('#FF6347'), size_hint=(0.5, None), height=dp(70))
            confirm_yes_btn.add_widget(MDButtonText(text='Да, удалить'))
            confirm_no_btn = MDButton(style='filled', md_bg_color=get_color_from_hex('#32CD32'), size_hint=(0.5, None), height=dp(70))
            confirm_no_btn.add_widget(MDButtonText(text='Отмена'))
            confirm_btn_layout.add_widget(confirm_yes_btn)
            confirm_btn_layout.add_widget(confirm_no_btn)
            confirm_content.add_widget(confirm_btn_layout)

            confirm_popup = Popup(
                title="",
                content=confirm_content,
                size_hint=(0.85, 0.85),
                separator_height=0,
                background_color=[0,0,0,0.3],
                overlay_color=[0,0,0,0.3]
            )

            def do_delete(instance):
                """Выполнить удаление"""
                conn = sqlite3.connect(self.db_name)
                cursor = conn.cursor()
                for breed in breeds_to_delete:
                    cursor.execute('''
                        DELETE FROM custom_breeds
                        WHERE breed_name = ? AND breed_type = ?
                    ''', (breed, breed_type))
                conn.commit()
                conn.close()
                
                confirm_popup.dismiss()
                popup.dismiss()
                
                # Показываем сообщение об успехе
                success_content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD, md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)
                success_content.add_widget(MDLabel(
                    text=f"Удалено {len(breeds_to_delete)} пород(ы)",
                    font_style='Title',
                    role='small',
                    theme_text_color='Custom',
                    text_color=Colors.GREEN
                ))
                success_popup = Popup(
                    title="",
                    content=success_content,
                    size_hint=(0.85, 0.85),
                    separator_height=0,
                    background_color=[0,0,0,0.3],
                    overlay_color=[0,0,0,0.3],
                    auto_dismiss=True
                )
                success_popup.open()

            confirm_yes_btn.bind(on_release=do_delete)
            confirm_no_btn.bind(on_release=confirm_popup.dismiss)
            confirm_popup.open()

        select_all_btn.bind(on_release=select_all)
        delete_btn.bind(on_release=delete_selected)
        cancel_btn.bind(on_release=popup.dismiss)

        popup.open()

    def create_ui(self):
        main_layout = MDBoxLayout(orientation='vertical', md_bg_color=[0.12, 0.12, 0.12, 1])

        # Верхняя панель
        toolbar = MDTopAppBar(
            type='small', elevation=2,
            md_bg_color=Colors.SECONDARY,
        )
        leading = MDTopAppBarLeadingButtonContainer()
        leading.add_widget(MDActionTopAppBarButton(
            icon='arrow-left', on_release=self.go_back))
        toolbar.add_widget(leading)
        self._toolbar_title = MDTopAppBarTitle(text=f'Участок №{self.current_section}')
        toolbar.add_widget(self._toolbar_title)
        trailing = MDTopAppBarTrailingButtonContainer()
        trailing.add_widget(MDActionTopAppBarButton(
            icon='refresh', on_release=self.clear_table_data))
        toolbar.add_widget(trailing)
        main_layout.add_widget(toolbar)

        # Контент
        scroll = MDScrollView()
        content = MDBoxLayout(
            orientation='vertical', spacing=Spacing.LG,
            padding=[Spacing.MD, Spacing.MD],
            adaptive_height=True,
        )

        # Карточка с адресной информацией
        address_card = MDCard(
            orientation='vertical', padding=Spacing.MD,
            size_hint_y=None, height=dp(56),
            radius=[Spacing.RADIUS_LG], elevation=1,
            md_bg_color=[0.18, 0.18, 0.18, 0.95],
        )
        self.address_label = MDLabel(
            text=f"Участок №{self.current_section}  |  {self.current_quarter} {self.current_plot} {self.current_forestry}",
            theme_text_color='Custom', text_color=[1,1,1,1],
            font_style='Title', role='small',
            adaptive_height=True, valign='middle',
        )
        address_card.add_widget(self.address_label)
        content.add_widget(address_card)

        # Сетка кнопок в 2 колонки
        buttons = [
            ('map-marker', 'Адрес', Colors.BTN_INFO, self.show_address_popup),
            ('file-outline', 'Файл', Colors.BTN_WARNING, self.show_file_popup),
            ('clipboard-text-outline', 'Детали', Colors.BTN_PURPLE, self.show_additional_functions_popup),
            ('calculator', 'Итого', Colors.BTN_SUCCESS, self.show_total_summary_popup),
            ('file-document-edit', 'Проект', Colors.BTN_TEAL, self.generate_care_project),
            ('exit-to-app', 'Меню', Colors.BTN_DANGER, self.go_back),
        ]

        grid = MDBoxLayout(
            orientation='vertical', spacing=Spacing.MD,
            size_hint_y=None, adaptive_height=True,
        )
        for i in range(0, len(buttons), 2):
            row = MDBoxLayout(
                orientation='horizontal', spacing=Spacing.MD,
                size_hint_y=None, height=dp(72),
            )
            for j in range(2):
                if i + j < len(buttons):
                    icon, text, color, cb = buttons[i + j]
                    card = MDCard(
                        orientation='horizontal', padding=[Spacing.MD, Spacing.SM],
                        size_hint=(0.5, 1), radius=[Spacing.RADIUS_MD],
                        elevation=1, md_bg_color=[0.18, 0.18, 0.18, 0.95],
                        ripple_behavior=True,
                        on_release=lambda x, c=cb: c(x),
                    )
                    icon_btn = MDIconButton(
                        icon=icon,
                        theme_icon_color='Custom', icon_color=color,
                        font_size=dp(24),
                    )
                    card.add_widget(icon_btn)
                    card.add_widget(MDLabel(
                        text=text, font_size='14sp',
                        theme_text_color='Custom', text_color=[1,1,1,1],
                        adaptive_height=True, valign='middle',
                    ))
                    row.add_widget(card)
            grid.add_widget(row)
        content.add_widget(grid)

        scroll.add_widget(content)
        main_layout.add_widget(scroll)
        self.add_widget(main_layout)

    def _update_bg(self, instance, value):
        self.bg_rect.pos = self.pos
        self.bg_rect.size = self.size

    def _update_background(self, theme):
        if theme['type'] == 'image':
            try:
                self.bg_color.rgba = (1, 1, 1, 1)
                self.bg_rect.texture = CoreImage(theme['background']).texture
            except Exception as e:
                print(f"Error loading background image: {str(e)}")
        else:
            self.bg_color.rgba = theme['background']
            self.bg_rect.texture = None

    def _get_text_color(self):
        theme = self.theme_manager.current_theme
        if theme['type'] == 'image':
            return get_color_from_hex('#FFFFFF')
        else:
            return get_color_from_hex(theme['text_color'])

    def update_section_label(self):
        if hasattr(self, '_toolbar_title') and self._toolbar_title:
            self._toolbar_title.text = f'Участок №{self.current_section}'
        if hasattr(self, 'address_label') and self.address_label:
            self.address_label.text = f"Участок №{self.current_section}"

    def toggle_edit_mode(self, instance):
        self.edit_mode = not self.edit_mode
        instance.bg_color = get_color_from_hex('#FFA500' if self.edit_mode else '#00BFFF')

    def update_focus(self, instance, value):
        if value:
            self.focused_cell = [instance.row_index, instance.col_index]

    def move_focus(self, direction):
        current = self.focused_cell
        if not current: return
        row, col = current

        if direction == 'up': row = max(0, row-1)
        elif direction == 'down': row = min(len(self.inputs)-1, row+1)
        elif direction == 'left': col = max(0, col-1)
        elif direction == 'right': col = min(5, col+1)  # 6 столбцов (0-5)

        self.focused_cell = [row, col]
        inp = self.inputs[row][col]
        inp.focus = True
        inp.cursor = (len(inp.text), 0)

    def show_tree_popup(self, instance, value):
        """Показать popup для ввода данных площадки молодняков"""
        if value and instance.text.strip():
            if not self.edit_mode:
                # In normal mode, only show popup if other columns are empty
                if not any(inp.text.strip() for inp in self.inputs[instance.row_index][1:]):
                    MolodnikiTreeDataInputPopup(self, instance.row_index).open()
            else:
                # In edit mode, always show popup for editing existing data
                MolodnikiTreeDataInputPopup(self, instance.row_index).open()

    def auto_fill_nn(self, instance, value):
        if self.edit_mode: return
        if value and instance.focus:
            try:
                current_number = int(instance.text)
            except ValueError:
                current_number = 0
            for row_idx, row in enumerate(self.inputs):
                if row_idx > instance.row_index:
                    try:
                        prev_num = int(self.inputs[row_idx-1][0].text)
                        row[0].text = str(prev_num + 1)
                    except (ValueError, IndexError):
                        pass

    def show_breed_popup(self, instance, value):
        """Показать popup для выбора типа породы"""
        if not value: return

        content = MDBoxLayout(orientation='vertical', spacing=10, padding=10, md_bg_color=Colors.DARK_SURFACE)

        # Кнопки выбора типа породы
        type_layout = MDBoxLayout(orientation='horizontal', spacing=10)
        coniferous_btn = MDButton(style='filled', size_hint=(0.5, None), height=50)
        coniferous_btn.add_widget(MDButtonText(text='Хвойные', theme_text_color='Custom', text_color=Colors.GREEN))
        deciduous_btn = MDButton(style='filled', size_hint=(0.5, None), height=50)
        deciduous_btn.add_widget(MDButtonText(text='Лиственные', theme_text_color='Custom', text_color=Colors.GREEN))
        type_layout.add_widget(coniferous_btn)
        type_layout.add_widget(deciduous_btn)
        content.add_widget(type_layout)

        # Кнопка отмены
        cancel_btn = MDButton(style='filled', size_hint=(0.5, 1), height=50)
        cancel_btn.add_widget(MDButtonText(text='Отмена', theme_text_color='Custom', text_color=[1,0.3,0.3,1]))
        content.add_widget(cancel_btn)

        popup = Popup(
            title="",
            content=content,
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3],
            size_hint=(0.8, 0.5)
        )

        def select_coniferous(btn):
            self.show_breed_selection_popup(instance, 'coniferous')
            popup.dismiss()

        def select_deciduous(btn):
            self.show_breed_selection_popup(instance, 'deciduous')
            popup.dismiss()

        coniferous_btn.bind(on_release=select_coniferous)
        deciduous_btn.bind(on_release=select_deciduous)
        cancel_btn.bind(on_release=popup.dismiss)

        popup.open()

    def show_breed_selection_popup(self, instance, breed_type):
        """Показать popup для выбора конкретной породы из словаря"""
        content = MDBoxLayout(orientation='vertical', spacing=10, padding=10, md_bg_color=Colors.DARK_SURFACE)

        # Заголовок
        title_label = MDLabel(
            text=f"Выберите {'хвойную' if breed_type == 'coniferous' else 'лиственную'} породу",
            bold=True,
            size_hint=(1, None),
            height=30,
            theme_text_color='Custom',
            text_color=Colors.TEXT_ON_DARK
        )
        content.add_widget(title_label)

        # Списки пород
        if breed_type == 'coniferous':
            breeds = [
                'Сосна', 'Ель', 'Лиственница', 'Кедр', 'Пихта'
            ]
        else:
            breeds = [
                'Берёза', 'Осина', 'Ольха чёрная', 'Ольха серая',
                'Ива'
            ]

        # Загружаем пользовательские породы из базы данных
        custom_breeds = self.load_custom_breeds(breed_type)
        # Добавляем пользовательские породы к стандартным
        all_breeds = breeds + custom_breeds

        # ScrollView для списка пород
        scroll = ScrollView(size_hint=(1, None), height=300)
        breeds_layout = MDGridLayout(cols=1, spacing=5, size_hint_y=None, md_bg_color=Colors.DARK_SURFACE)
        breeds_layout.bind(minimum_height=breeds_layout.setter('height'))

        for breed in all_breeds:
            btn = MDButton(style='filled', size_hint=(1, None), height=50)
            btn.add_widget(MDButtonText(text=breed, theme_text_color='Custom', text_color=[0.7,0.7,0.7,1]))
            btn.bind(on_release=lambda x, b=breed: self.select_breed(instance, breed_type, b))
            breeds_layout.add_widget(btn)

        scroll.add_widget(breeds_layout)
        content.add_widget(scroll)

        # Кнопка "Новая"
        other_btn = MDButton(style='filled', size_hint=(1, None), height=50)
        other_btn.add_widget(MDButtonText(text='Новая', theme_text_color='Custom', text_color=Colors.GREEN))
        other_btn.bind(on_release=lambda x: self.select_breed(instance, breed_type, 'other'))
        content.add_widget(other_btn)

        # Кнопка "Очистить"
        clear_btn = MDButton(style='filled', size_hint=(1, None), height=50)
        clear_btn.add_widget(MDButtonText(text='Очистить', theme_text_color='Custom', text_color=[0.7,0.7,0.7,1]))
        clear_btn.bind(on_release=lambda x: self.show_clear_breeds_popup(breed_type))
        content.add_widget(clear_btn)

        # Кнопка отмены
        cancel_btn = MDButton(style='filled', size_hint=(1, None), height=50)
        cancel_btn.add_widget(MDButtonText(text='Отмена', theme_text_color='Custom', text_color=[1,0.3,0.3,1]))
        content.add_widget(cancel_btn)

        popup = Popup(
            title="",
            content=content,
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3],
            size_hint=(0.85, 0.85)
        )

        cancel_btn.bind(on_release=popup.dismiss)
        popup.open()

    def select_breed(self, instance, breed_type, selected_breed):
        """Обработка выбора породы"""
        if selected_breed == 'other':
            # Показываем popup для ввода названия другой породы
            self.show_custom_breed_popup(instance, breed_type)
        else:
            # Показываем popup с параметрами породы, передавая название выбранной породы
            self.show_breed_details_popup(instance, breed_type, selected_breed)

    def show_breed_details_popup(self, instance, breed_type, selected_breed=None):
        """Показать popup для ввода параметров породы (единый поток сохранения)"""
        content = MDBoxLayout(orientation='vertical', spacing=10, padding=10, md_bg_color=Colors.DARK_SURFACE)

        # Получаем номер площадки из row_index
        plot_number = instance.row_index + 1 if hasattr(instance, 'row_index') else 1

        # Бокс для отображения сохраненных пород по площадке (в начале popup)
        plot_breeds_box = MDBoxLayout(
            orientation='vertical',
            size_hint=(1, None),
            height=150,
            padding=[10, 10],
            spacing=5,
            md_bg_color=Colors.CARD_BG
        )
        
        plot_breeds_title = MDLabel(
            text=f'Площадка №{plot_number} - Сохраненные породы:',
            font_size='14sp',
            bold=True,
            theme_text_color='Custom',
            text_color=Colors.GREEN,
            size_hint=(1, None),
            height=25,
            halign='left'
        )
        plot_breeds_title.bind(size=lambda *args: setattr(plot_breeds_title, 'text_size', (plot_breeds_title.width, None)))
        plot_breeds_box.add_widget(plot_breeds_title)
        
        # ScrollView для списка пород на площадке
        plot_breeds_scroll = ScrollView(size_hint=(1, None), height=110)
        plot_breeds_list = MDGridLayout(cols=1, spacing=3, size_hint_y=None, md_bg_color=Colors.CARD_BG)
        plot_breeds_list.bind(minimum_height=plot_breeds_list.setter('height'))
        
        # Получаем существующие породы для этой площадки
        existing_breeds_for_plot = []
        if hasattr(instance, 'text') and instance.text:
            try:
                existing_breeds_for_plot = json.loads(instance.text) if isinstance(instance.text, str) else []
            except (json.JSONDecodeError, TypeError):
                pass
        
        if existing_breeds_for_plot:
            for i, breed_info in enumerate(existing_breeds_for_plot):
                breed_name = breed_info.get('name', 'Неизвестная')
                breed_type_display = "Хвойная" if breed_info.get('type') == 'coniferous' else "Лиственная"
                density = breed_info.get('density', 0)
                height = breed_info.get('height', 0)
                
                # Для хвойных показываем сумму градаций
                if breed_info.get('type') == 'coniferous':
                    conif_density = (breed_info.get('do_05', 0) + breed_info.get('05_15', 0) + breed_info.get('bolee_15', 0))
                    if conif_density > 0:
                        density = conif_density
                
                breed_item = MDLabel(
                    text=f'{i+1}. {breed_name} ({breed_type_display}) - Густота: {density}, Высота: {height}м',
                    font_size='12sp',
                    theme_text_color='Custom',
                    text_color=[0.7,0.7,0.7,1],
                    size_hint=(1, None),
                    height=25,
                    halign='left'
                )
                breed_item.bind(size=lambda *args: setattr(breed_item, 'text_size', (breed_item.width, None)))
                plot_breeds_list.add_widget(breed_item)
        else:
            no_breeds_label = MDLabel(
                text='Породы еще не добавлены',
                font_size='12sp',
                theme_text_color='Custom',
                text_color=[0.5,0.5,0.5,1],
                size_hint=(1, None),
                height=25,
                halign='left'
            )
            no_breeds_label.bind(size=lambda *args: setattr(no_breeds_label, 'text_size', (no_breeds_label.width, None)))
            plot_breeds_list.add_widget(no_breeds_label)
        
        plot_breeds_scroll.add_widget(plot_breeds_list)
        plot_breeds_box.add_widget(plot_breeds_scroll)
        content.add_widget(plot_breeds_box)

        # Заголовок с названием породы
        title_label = MDLabel(
            text=f"Добавление породы: {selected_breed}",
            font_size='18sp',
            bold=True,
            size_hint=(1, None),
            height=40,
            theme_text_color='Custom',
            text_color=Colors.GREEN
        )
        content.add_widget(title_label)

        # Информация о типе породы
        type_text = "Хвойная порода" if breed_type == 'coniferous' else "Лиственная порода"
        type_label = MDLabel(
            text=f"Тип: {type_text}",
            font_size='14sp',
            size_hint=(1, None),
            height=25,
            theme_text_color='Custom',
            text_color=[0.7,0.7,0.7,1]
        )
        content.add_widget(type_label)

        # ScrollView для полей ввода параметров породы
        fields_scroll = ScrollView(size_hint=(1, None), height=200)
        fields_layout = MDGridLayout(cols=2, spacing=10, size_hint_y=None, md_bg_color=Colors.DARK_SURFACE)
        fields_layout.bind(minimum_height=fields_layout.setter('height'))

        if breed_type == 'coniferous':
            fields = [
                ('До 0.5м (шт):', 'do_05'),
                ('0.5-1.5м (шт):', '05_15'),
                ('>1.5м (шт):', 'bolee_15'),
                ('Высота (м):', 'height'),
                ('Диаметр (см):', 'diameter'),  # ✅ ДОБАВЛЕНО
                ('Густота (авто):', 'density'),
                ('Возраст (лет):', 'age')
            ]
        else:
            fields = [
                ('Густота (шт):', 'density'),
                ('Высота (м):', 'height'),
                ('Диаметр (см):', 'diameter'),  # ✅ ДОБАВЛЕНО
                ('Возраст (лет):', 'age')
            ]

        self.breed_inputs = {}
        for label_text, field_key in fields:
            lbl = MDLabel(text=label_text, size_hint=(None, None), size=(120, 40), halign='left', valign='middle', theme_text_color='Custom', text_color=Colors.TEXT_ON_DARK)
            lbl.bind(size=lambda *args: setattr(lbl, 'text_size', (lbl.width, None)))
            inp = MDTextField(
                mode="outlined",
                size_hint=(None, None),
                size=(120, 40),
                line_color_focus=Colors.GREEN
            )
            if field_key in ['density', 'age']:
                inp.input_filter = 'int'
            elif field_key == 'height':
                inp.input_filter = 'float'
            elif field_key in ['do_05', '05_15', 'bolee_15']:
                inp.input_filter = 'int'
                if breed_type == 'coniferous':
                    inp.bind(text=self.update_coniferous_density)
            # Делаем поле густоты только для чтения для хвойных (рассчитывается автоматически)
            if field_key == 'density' and breed_type == 'coniferous':
                inp.disabled = True
            fields_layout.add_widget(lbl)
            fields_layout.add_widget(inp)
            self.breed_inputs[field_key] = inp

        # Заполняем поля данными из существующей породы, если она есть
        if existing_breeds_for_plot:
            for existing_breed in existing_breeds_for_plot:
                if existing_breed.get('name') == selected_breed:
                    # Заполняем поля данными
                    for key, inp in self.breed_inputs.items():
                        if key in existing_breed:
                            inp.text = str(existing_breed[key])
                    break

        fields_scroll.add_widget(fields_layout)
        content.add_widget(fields_scroll)

        # Подсказка для хвойных пород
        if breed_type == 'coniferous':
            hint_label = MDLabel(
                text="* Густота для хвойных рассчитывается автоматически как сумма градаций высот",
                font_size='12sp',
                size_hint=(1, None),
                height=30,
                theme_text_color='Custom',
                text_color=[0.5,0.5,0.5,1]
            )
            content.add_widget(hint_label)

        # Кнопки управления - добавление, удаление, сохранить и выход
        btn_layout = MDBoxLayout(orientation='horizontal', spacing=10, size_hint=(1, None), height=50, md_bg_color=Colors.DARK_SURFACE)

        # Главная кнопка - "Добавить"
        save_add_btn = MDButton(style='filled', size_hint=(0.25, 1), height=50)
        save_add_btn.add_widget(MDButtonText(text='Добавить', theme_text_color='Custom', text_color=Colors.GREEN, bold=True))

        # Кнопка "Сохранить"
        save_btn = MDButton(style='filled', size_hint=(0.25, 1), height=50)
        save_btn.add_widget(MDButtonText(text='Сохранить', theme_text_color='Custom', text_color=Colors.GREEN))

        # Кнопка удаления
        delete_btn = MDButton(style='filled', size_hint=(0.25, 1), height=50)
        delete_btn.add_widget(MDButtonText(text='Удалить', theme_text_color='Custom', text_color=[1,0.3,0.3,1]))

        # Кнопка выхода
        exit_btn = MDButton(style='filled', size_hint=(0.25, 1), height=50)
        exit_btn.add_widget(MDButtonText(text='Выход', theme_text_color='Custom', text_color=[0.7,0.7,0.7,1]))

        btn_layout.add_widget(save_add_btn)
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(delete_btn)
        btn_layout.add_widget(exit_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3],
            size_hint=(0.9, 0.95)
        )

        def save_and_add(btn):
            """Единая функция: добавить породу и сохранить данные"""
            # Собираем данные из полей ввода
            breed_data = {
                'name': selected_breed,
                'type': breed_type
            }

            has_data = False
            for key, inp in self.breed_inputs.items():
                if inp.text.strip():
                    has_data = True
                    try:
                        if key in ['density', 'age']:
                            breed_data[key] = int(inp.text)
                        elif key in ['height', 'diameter']:  # ✅ ДОБАВЛЕНО: diameter как float
                            breed_data[key] = float(inp.text)
                        else:
                            breed_data[key] = float(inp.text)
                    except ValueError:
                        breed_data[key] = 0

            # Для хвойных рассчитываем возраст, если не введен
            if breed_type == 'coniferous':
                if 'age' not in breed_data or breed_data['age'] == 0:
                    do_05 = breed_data.get('do_05', 0)
                    _05_15 = breed_data.get('05_15', 0)
                    bolee_15 = breed_data.get('bolee_15', 0)

                    if bolee_15 > 0:
                        breed_data['age'] = 20
                    elif _05_15 > 0:
                        breed_data['age'] = 10
                    elif do_05 > 0:
                        breed_data['age'] = 5
                    else:
                        breed_data['age'] = 10  # default age

            # Добавляем породу к существующим
            existing_breeds = self.parse_breeds_data(instance.text)
            existing_breeds.append(breed_data)
            instance.text = json.dumps(existing_breeds, ensure_ascii=False, indent=2)

            # Обновляем page_data
            if self.current_page not in self.page_data:
                self.page_data[self.current_page] = []
            if hasattr(instance, 'row_index'):
                row_idx = instance.row_index
                while len(self.page_data[self.current_page]) <= row_idx:
                    self.page_data[self.current_page].append(['', '', '', '', '', ''])
                self.page_data[self.current_page][row_idx][3] = instance.text

            self.update_plot_total(instance, instance.text)

            # Показываем результат и предлагаем добавить еще породу
            popup.dismiss()

            # Показываем выбор следующего действия
            self.show_after_add_popup(instance, selected_breed, len(existing_breeds))

        def save_only(btn):
            """Сохранить данные без добавления новой породы"""
            # Собираем данные из полей ввода
            breed_data = {
                'name': selected_breed,
                'type': breed_type
            }

            has_data = False
            for key, inp in self.breed_inputs.items():
                if inp.text.strip():
                    has_data = True
                    try:
                        if key in ['density', 'age']:
                            breed_data[key] = int(inp.text)
                        elif key in ['height', 'diameter']:
                            breed_data[key] = float(inp.text)
                        else:
                            breed_data[key] = float(inp.text)
                    except ValueError:
                        breed_data[key] = 0

            # Для хвойных рассчитываем возраст, если не введен
            if breed_type == 'coniferous':
                if 'age' not in breed_data or breed_data['age'] == 0:
                    do_05 = breed_data.get('do_05', 0)
                    _05_15 = breed_data.get('05_15', 0)
                    bolee_15 = breed_data.get('bolee_15', 0)

                    if bolee_15 > 0:
                        breed_data['age'] = 20
                    elif _05_15 > 0:
                        breed_data['age'] = 10
                    elif do_05 > 0:
                        breed_data['age'] = 5
                    else:
                        breed_data['age'] = 10

            # Обновляем существующую породу или добавляем новую
            existing_breeds = self.parse_breeds_data(instance.text)
            
            # Ищем существующую породу с таким именем
            found = False
            for i, b in enumerate(existing_breeds):
                if b.get('name') == selected_breed:
                    existing_breeds[i] = breed_data
                    found = True
                    break
            
            if not found:
                existing_breeds.append(breed_data)
            
            instance.text = json.dumps(existing_breeds, ensure_ascii=False, indent=2) if existing_breeds else ''

            # Обновляем page_data
            if self.current_page not in self.page_data:
                self.page_data[self.current_page] = []
            if hasattr(instance, 'row_index'):
                row_idx = instance.row_index
                while len(self.page_data[self.current_page]) <= row_idx:
                    self.page_data[self.current_page].append(['', '', '', '', '', ''])
                self.page_data[self.current_page][row_idx][3] = instance.text

            self.update_plot_total(instance, instance.text)
            self.show_success(f"Порода '{selected_breed}' сохранена!")

        def exit_to_plot_menu(btn):
            """Выход в меню заполнения площадок"""
            # Сохраняем данные перед выходом
            save_only(btn)
            popup.dismiss()

        save_add_btn.bind(on_release=save_and_add)
        save_btn.bind(on_release=save_only)
        delete_btn.bind(on_release=lambda x: self.show_delete_breed_popup(instance, plot_breeds_list))
        exit_btn.bind(on_release=exit_to_plot_menu)

        popup.open()

    def show_delete_breed_popup(self, instance, plot_breeds_list):
        """Показать popup для удаления пород"""
        # Получаем список пород
        existing_breeds = self.parse_breeds_data(instance.text)
        
        if not existing_breeds:
            self.show_error("Нет пород для удаления!")
            return
        
        content = MDBoxLayout(orientation='vertical', spacing=10, padding=10, md_bg_color=Colors.DARK_SURFACE)
        
        title_label = MDLabel(
            text="Выберите породы для удаления:",
            bold=True,
            size_hint=(1, None),
            height=40,
            theme_text_color='Custom',
            text_color=[1,0.3,0.3,1]
        )
        content.add_widget(title_label)
        
        # ScrollView для списка пород
        scroll = ScrollView(size_hint=(1, None), height=300)
        breeds_layout = MDGridLayout(cols=1, spacing=5, size_hint_y=None, md_bg_color=Colors.DARK_SURFACE)
        breeds_layout.bind(minimum_height=breeds_layout.setter('height'))
        
        # Чекбоксы для выбора пород
        self.breed_checkboxes = {}
        for i, breed_info in enumerate(existing_breeds):
            breed_name = breed_info.get('name', 'Неизвестная')
            
            # Создаём строку с чекбоксом и названием породы
            breed_row = MDBoxLayout(orientation='horizontal', spacing=10, size_hint=(1, None), height=40, md_bg_color=Colors.DARK_SURFACE)
            
            checkbox = CheckBox(size_hint=(None, None), size=(40, 40), active=False)
            self.breed_checkboxes[i] = checkbox
            
            breed_label = MDLabel(
                text=f"{breed_name}",
                size_hint=(1, None),
                height=40,
                halign='left',
                theme_text_color='Custom',
                text_color=Colors.TEXT_ON_DARK
            )
            
            breed_row.add_widget(checkbox)
            breed_row.add_widget(breed_label)
            breeds_layout.add_widget(breed_row)
        
        scroll.add_widget(breeds_layout)
        content.add_widget(scroll)

        # Кнопки управления
        btn_layout = MDBoxLayout(orientation='horizontal', spacing=10, size_hint=(1, None), height=50, md_bg_color=Colors.DARK_SURFACE)

        confirm_btn = MDButton(style='filled', size_hint=(0.34, 1), height=50)
        confirm_btn.add_widget(MDButtonText(text='Удалить выбранные', theme_text_color='Custom', text_color=[1,0.3,0.3,1]))

        clear_all_btn = MDButton(style='filled', size_hint=(0.33, 1), height=50)
        clear_all_btn.add_widget(MDButtonText(text='Очистить все', theme_text_color='Custom', text_color=[0.7,0.7,0.7,1]))

        cancel_btn = MDButton(style='filled', size_hint=(0.33, 1), height=50)
        cancel_btn.add_widget(MDButtonText(text='Отмена', theme_text_color='Custom', text_color=[0.7,0.7,0.7,1]))

        btn_layout.add_widget(confirm_btn)
        btn_layout.add_widget(clear_all_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3],
            size_hint=(0.8, 0.7)
        )

        def confirm_delete(btn):
            # Получаем индексы пород для удаления
            indices_to_delete = sorted(
                [idx for idx, cb in self.breed_checkboxes.items() if cb.active],
                reverse=True
            )

            if not indices_to_delete:
                self.show_error("Выберите хотя бы одну породу для удаления!")
                return

            # Удаляем породы
            for idx in indices_to_delete:
                if 0 <= idx < len(existing_breeds):
                    existing_breeds.pop(idx)

            # Обновляем данные
            instance.text = json.dumps(existing_breeds, ensure_ascii=False, indent=2) if existing_breeds else ''

            # Update page_data
            if self.current_page in self.page_data:
                page_data = self.page_data[self.current_page]
                if hasattr(instance, 'row_index') and instance.row_index < len(page_data):
                    page_data[instance.row_index][3] = instance.text

            self.update_plot_total(instance, instance.text)

            # Сохраняем в базу данных
            self.save_current_page()

            # Обновляем отображение пород в основном popup
            self.update_plot_breeds_display(plot_breeds_list, existing_breeds)

            self.show_success(f"Удалено пород: {len(indices_to_delete)}")
            popup.dismiss()

        def clear_all(btn):
            # Очищаем все породы
            instance.text = ''

            # Update page_data
            if self.current_page in self.page_data:
                page_data = self.page_data[self.current_page]
                if hasattr(instance, 'row_index') and instance.row_index < len(page_data):
                    page_data[instance.row_index][3] = ''

            self.update_plot_total(instance, '')

            # Сохраняем в базу данных
            self.save_current_page()

            # Обновляем отображение пород в основном popup
            self.update_plot_breeds_display(plot_breeds_list, [])

            self.show_success("Все породы очищены!")
            popup.dismiss()

        confirm_btn.bind(on_release=confirm_delete)
        clear_all_btn.bind(on_release=clear_all)
        cancel_btn.bind(on_release=popup.dismiss)

        popup.open()

    def show_after_add_popup(self, instance, added_breed, total_breeds_count):
        """Показать popup после добавления породы с выбором следующего действия"""
        content = MDBoxLayout(orientation='vertical', spacing=15, padding=15, md_bg_color=Colors.DARK_SURFACE)

        # Заголовок с информацией
        title_label = MDLabel(
            text=f"✓ Порода '{added_breed}' добавлена!\n\nВсего пород на площадке: {total_breeds_count}",
            font_size='18sp',
            bold=True,
            theme_text_color='Custom',
            text_color=Colors.GREEN,
            size_hint=(1, None),
            height=80,
            halign='center'
        )
        content.add_widget(title_label)

        # Кнопки выбора
        btn_layout = MDBoxLayout(orientation='vertical', spacing=10, size_hint=(1, None), height=150, md_bg_color=Colors.DARK_SURFACE)
        
        add_more_btn = MDButton(style='filled', size_hint=(1, None), height=50)
        add_more_btn.add_widget(MDButtonText(text='+ Добавить ещё породу', theme_text_color='Custom', text_color=Colors.GREEN))
        
        finish_btn = MDButton(style='filled', size_hint=(1, None), height=50)
        finish_btn.add_widget(MDButtonText(text='✓ Завершить редактирование', theme_text_color='Custom', text_color=Colors.GREEN, bold=True))
        
        btn_layout.add_widget(add_more_btn)
        btn_layout.add_widget(finish_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3],
            size_hint=(0.8, 0.55)
        )

        def add_more(btn):
            popup.dismiss()
            # Возвращаемся к выбору типа породы
            self.show_breed_popup(instance, True)

        def finish(btn):
            popup.dismiss()
            self.show_success(f"Данные площадки сохранены! Всего пород: {total_breeds_count}")

        add_more_btn.bind(on_release=add_more)
        finish_btn.bind(on_release=finish)

        popup.open()

    def show_breeds_list_popup(self, instance):
        """Показать popup со списком всех пород в этой строке"""
        content = MDBoxLayout(orientation='vertical', spacing=10, padding=10, md_bg_color=Colors.DARK_SURFACE)

        title_label = MDLabel(
            text="Список пород в этой строке",
            bold=True,
            size_hint=(1, None),
            height=30,
            theme_text_color='Custom',
            text_color=Colors.TEXT_ON_DARK
        )
        content.add_widget(title_label)

        breeds_data = self.parse_breeds_data(instance.text)

        if not breeds_data:
            no_breeds_label = MDLabel(
                text="Породы не найдены",
                size_hint=(1, None),
                height=50,
                theme_text_color='Custom',
                text_color=[0.5,0.5,0.5,1]
            )
            content.add_widget(no_breeds_label)
        else:
            scroll = ScrollView(size_hint=(1, None), height=300)
            breeds_layout = MDGridLayout(cols=1, spacing=5, size_hint_y=None, md_bg_color=Colors.DARK_SURFACE)
            breeds_layout.bind(minimum_height=breeds_layout.setter('height'))

            for i, breed_info in enumerate(breeds_data):
                breed_card = MDBoxLayout(
                    orientation='vertical',
                    size_hint=(1, None),
                    height=120,
                    padding=5,
                    md_bg_color=Colors.CARD_BG
                )

                name_label = MDLabel(
                    text=f"{i+1}. {breed_info.get('name', 'Неизвестная порода')}",
                    bold=True,
                    size_hint=(1, None),
                    height=25,
                    theme_text_color='Custom',
                    text_color=Colors.TEXT_ON_DARK
                )
                breed_card.add_widget(name_label)

                params_text = []
                breed_type = breed_info.get('type', 'deciduous')

                # В зависимости от типа породы показываем разные параметры
                if breed_type == 'coniferous':
                    # Для хвойных показываем градации + общую густоту и другие параметры
                    if 'do_05' in breed_info and breed_info['do_05']:
                        params_text.append(f"До 0.5м: {breed_info['do_05']}")
                    if '05_15' in breed_info and breed_info['05_15']:
                        params_text.append(f"0.5-1.5м: {breed_info['05_15']}")
                    if 'bolee_15' in breed_info and breed_info['bolee_15']:
                        params_text.append(f">1.5м: {breed_info['bolee_15']}")
                    if 'density' in breed_info and breed_info['density']:
                        params_text.append(f"Общая густота: {breed_info['density']}")
                    if 'height' in breed_info and breed_info['height']:
                        params_text.append(f"Высота: {breed_info['height']}м")
                    if 'age' in breed_info and breed_info['age']:
                        params_text.append(f"Возраст: {breed_info['age']} лет")
                else:
                    # Для лиственных показываем только основную информацию (без градаций)
                    if 'density' in breed_info and breed_info['density']:
                        params_text.append(f"Густота: {breed_info['density']}")
                    if 'height' in breed_info and breed_info['height']:
                        params_text.append(f"Высота: {breed_info['height']}м")
                    if 'diameter' in breed_info and breed_info['diameter']:  # ✅ ДОБАВЛЕНО
                        params_text.append(f"Диаметр: {breed_info['diameter']}см")
                    if 'age' in breed_info and breed_info['age']:
                        params_text.append(f"Возраст: {breed_info['age']} лет")

                params_label = MDLabel(
                    text="; ".join(params_text) if params_text else "Нет параметров",
                    size_hint=(1, None),
                    height=40,
                    theme_text_color='Custom',
                    text_color=[0.7,0.7,0.7,1],
                    halign='left',
                    valign='top'
                )
                params_label.bind(size=lambda *args: setattr(params_label, 'text_size', (params_label.width, None)))
                breed_card.add_widget(params_label)

                btn_layout = MDBoxLayout(orientation='horizontal', spacing=5, size_hint=(1, None), height=30, md_bg_color=Colors.CARD_BG)
                edit_btn = MDButton(style='filled', size_hint=(0.5, 1))
                edit_btn.add_widget(MDButtonText(text='Изменить', theme_text_color='Custom', text_color=[0.7,0.7,0.7,1]))
                delete_btn = MDButton(style='filled', size_hint=(0.5, 1))
                delete_btn.add_widget(MDButtonText(text='Удалить', theme_text_color='Custom', text_color=[1,0.3,0.3,1]))
                btn_layout.add_widget(edit_btn)
                delete_btn.bind(on_release=lambda x, idx=i: self.delete_breed_from_list(instance, idx))
                btn_layout.add_widget(delete_btn)
                breed_card.add_widget(btn_layout)

                def edit_breed(btn, idx=i):
                    self.edit_breed_in_list(instance, idx)

                edit_btn.bind(on_release=edit_breed)

                breeds_layout.add_widget(breed_card)

            scroll.add_widget(breeds_layout)
            content.add_widget(scroll)

        close_btn = MDButton(style='filled', size_hint=(1, None), height=50)
        close_btn.add_widget(MDButtonText(text='Закрыть', theme_text_color='Custom', text_color=[0.7,0.7,0.7,1]))
        content.add_widget(close_btn)

        popup = Popup(
            title="",
            content=content,
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3],
            size_hint=(0.85, 0.9)
        )

        close_btn.bind(on_release=popup.dismiss)
        popup.open()

    def edit_breed_in_list(self, instance, breed_index):
        """Редактировать породу в списке"""
        breeds_data = self.parse_breeds_data(instance.text)
        if 0 <= breed_index < len(breeds_data):
            breed_info = breeds_data[breed_index]
            self.show_edit_breed_popup(instance, breed_index, breed_info)

    def delete_breed_from_list(self, instance, breed_index):
        """Удалить породу из списка"""
        breeds_data = self.parse_breeds_data(instance.text)
        if 0 <= breed_index < len(breeds_data):
            breed_name = breeds_data[breed_index].get('name', 'Неизвестная порода')
            breeds_data.pop(breed_index)
            instance.text = json.dumps(breeds_data, ensure_ascii=False, indent=2) if breeds_data else ''
        self.update_totals()
        self.show_success("Порода удалена!")
        if hasattr(self, 'popup') and self.popup:
            self.popup.dismiss()

    def save_totals_to_excel(self, breeds_data, current_radius, plot_area_ha, plot_count, total_plot_area_ha):
        """Сохранить итоговые данные в Excel на новом листе"""
        timestamp = datetime.datetime.now().strftime('%M%S')  # Только минуты и секунды
        document_name = self.project_data.get('document_name', 'Проект')
        # Очень короткое имя файла на основе названия проекта (макс 10 символов)
        short_name = document_name.replace(' ', '').replace('/', '_').replace('.', '')[:10]
        filename = f"{short_name}_{self.current_section}_{timestamp}.xlsx"
        full_path = os.path.join(self.reports_dir, filename)

        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Итоги"

            # Заголовок
            ws['A1'] = f'ИТОГИ ПО УЧАСТКУ МОЛОДНЯКОВ - {self.current_section}'
            ws['A1'].font = openpyxl.styles.Font(bold=True, size=14)
            ws.merge_cells('A1:E1')

            # Информация о радиусе
            ws['A3'] = f'Радиус участка: {current_radius:.2f} м'
            ws['A4'] = f'1 дерево = {10000 / (3.14159 * (current_radius ** 2)):.0f} тыс.шт./га'

            # Коэффициент состава
            ws['A6'] = 'КОЭФФИЦИЕНТ СОСТАВА НАСАЖДЕНИЯ'
            ws['A6'].font = openpyxl.styles.Font(bold=True, size=12)

            # Расчет коэффициента состава
            total_densities = {}
            total_density_all = 0  # Общая густота всех пород (сумма средних)
            
            for breed_name, data in breeds_data.items():
                if data['plots']:
                    # Рассчитываем СРЕДНЮЮ густоту породы (сумма плотностей / кол-во площадок)
                    # Плотности уже в шт/га, делим на 1000 для перевода в тыс.шт/га
                    avg_density = sum(p.get('density', 0) for p in data['plots']) / len(data['plots'])
                    avg_density_thousands = avg_density / 1000  # переводим в тыс.шт/га
                    total_density_all += avg_density_thousands
                    if avg_density_thousands > 0:
                        total_densities[breed_name] = avg_density_thousands

            if total_densities:
                # Метод наибольшего остатка для распределения коэффициентов
                # Сначала рассчитываем точные коэффициенты
                exact_coeffs = []
                for breed_name, density in sorted(total_densities.items(), key=lambda x: x[1], reverse=True):
                    exact_coeff = (density / total_density_all * 10) if total_density_all > 0 else 1
                    exact_coeffs.append(exact_coeff)
                
                # Округляем вниз
                coeffs_floor = [int(coeff) for coeff in exact_coeffs]
                # Считаем остатки
                remainders = [(exact_coeffs[i] - coeffs_floor[i], i) for i in range(len(exact_coeffs))]
                # Сортируем по убыванию остатков
                remainders.sort(reverse=True)
                
                # Распределяем единицы начиная с наибольших остатков
                coeffs = coeffs_floor[:]
                total = sum(coeffs)
                i = 0
                while total < 10 and i < len(remainders):
                    idx = remainders[i][1]
                    coeffs[idx] += 1
                    total += 1
                    i += 1
                
                # Формируем формулу состава
                sorted_breeds = sorted(total_densities.items(), key=lambda x: x[1], reverse=True)
                composition_parts = []
                for i, (breed_name, _) in enumerate(sorted_breeds):
                    if i < len(coeffs):
                        breed_letter = self.get_breed_letter(breed_name)
                        composition_parts.append(f"{coeffs[i]}{breed_letter}")

                composition_text = ''.join(composition_parts) + "Др"
                ws['A7'] = f"Формула состава: {composition_text}"

            # Хвойные породы
            row = 9
            ws[f'A{row}'] = 'ХВОЙНЫЕ ПОРОДЫ - ВЫСОТА ПО ГРАДАЦИЯМ'
            ws[f'A{row}'].font = openpyxl.styles.Font(bold=True, size=12)

            has_coniferous = False
            for breed_name, data in sorted(breeds_data.items()):
                if data['type'] == 'coniferous' and data['plots']:
                    has_coniferous = True
                    row += 1
                    zones = data.get('coniferous_zones', {})
                    # Густота по градациям = общее кол-во деревьев в градации / общая площадь всех площадок в га
                    plot_area_ha = 3.14159 * (float(current_radius) if current_radius else 1.78) ** 2 / 10000
                    total_area_ha = plot_area_ha * plot_count
                    total_do_05 = sum(p.get('do_05', 0) for p in data['plots'])
                    total_05_15 = sum(p.get('05_15', 0) for p in data['plots'])
                    total_bolee_15 = sum(p.get('bolee_15', 0) for p in data['plots'])
                    avg_do_05 = total_do_05 / total_area_ha if total_area_ha > 0 else 0
                    avg_05_15 = total_05_15 / total_area_ha if total_area_ha > 0 else 0
                    avg_bolee_15 = total_bolee_15 / total_area_ha if total_area_ha > 0 else 0
                    # Высота = средняя только на тех площадках, где есть порода
                    avg_height_total = sum(p['height'] for p in data['plots'] if p['height'] > 0) / len([p for p in data['plots'] if p['height'] > 0]) if any(p['height'] > 0 for p in data['plots']) else 0

                    ws[f'A{row}'] = f"{breed_name}:"
                    ws[f'B{row}'] = f"до 0.5м: {avg_do_05:.1f} шт/га"
                    row += 1
                    ws[f'B{row}'] = f"0.5-1.5м: {avg_05_15:.1f} шт/га"
                    row += 1
                    ws[f'B{row}'] = f">1.5м: {avg_bolee_15:.1f} шт/га"
                    row += 1
                    ws[f'B{row}'] = f"средняя высота породы: {avg_height_total:.1f}м"
                    row += 1

            # Лиственные породы
            if has_coniferous:
                row += 1
            ws[f'A{row}'] = 'ЛИСТВЕННЫЕ ПОРОДЫ - СРЕДНИЕ ПОКАЗАТЕЛИ'
            ws[f'A{row}'].font = openpyxl.styles.Font(bold=True, size=12)

            has_deciduous = False
            for breed_name, data in sorted(breeds_data.items()):
                if data['type'] == 'deciduous' and data['plots']:
                    has_deciduous = True
                    row += 1
                    # Густота = общее количество деревьев / общая площадь всех площадок в га
                    total_trees = sum(p.get('density_raw', 0) for p in data['plots'])
                    plot_area_ha = 3.14159 * (float(current_radius) if current_radius else 1.78) ** 2 / 10000
                    total_area_ha = plot_area_ha * plot_count
                    avg_density = total_trees / total_area_ha if total_area_ha > 0 else 0
                    # Высота = средняя только на тех площадках, где есть высота
                    avg_heights = [p['height'] for p in data['plots'] if p['height'] > 0]
                    avg_height = sum(avg_heights) / len(avg_heights) if avg_heights else 0
                    # Возраст = средний только на тех площадках, где есть возраст
                    avg_ages = [p['age'] for p in data['plots'] if p['age'] > 0]
                    avg_age = sum(avg_ages) / len(avg_ages) if avg_ages else 0

                    ws[f'A{row}'] = f"{breed_name}:"
                    ws[f'B{row}'] = f"Средняя густота: {avg_density:.1f} шт/га"
                    row += 1
                    ws[f'B{row}'] = f"Средняя высота: {avg_height:.1f}м"
                    row += 1
                    ws[f'B{row}'] = f"Средний возраст: {avg_age:.1f} лет"
                    row += 1

            # Автоподбор ширины столбцов
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width

            wb.save(full_path)
            self.show_success(f"Итоги сохранены в Excel: {filename}")
        except Exception as e:
            self.show_error(f"Ошибка сохранения итогов в Excel: {str(e)}")

    def save_totals_to_word(self, breeds_data, current_radius, plot_area_ha, plot_count, total_plot_area_ha):
        """Сохранить итоговые данные в Word"""
        try:
            from docx import Document
            from docx.shared import Inches

            timestamp = datetime.datetime.now().strftime('%M%S')  # Только минуты и секунды
            document_name = self.project_data.get('document_name', 'Проект')
            # Очень короткое имя файла на основе названия проекта (макс 10 символов)
            short_name = document_name.replace(' ', '').replace('/', '_').replace('.', '')[:10]
            filename = f"{short_name}_{self.current_section}_{timestamp}.docx"
            full_path = os.path.join(self.reports_dir, filename)

            doc = Document()
            doc.add_heading(f'ИТОГИ ПО УЧАСТКУ МОЛОДНЯКОВ - {self.current_section}', 0)

            # Информация о радиусе
            doc.add_paragraph(f'Радиус участка: {current_radius:.2f} м')
            doc.add_paragraph(f'1 дерево = {10000 / (3.14159 * (current_radius ** 2)):.0f} тыс.шт./га')

            # Коэффициент состава
            doc.add_heading('КОЭФФИЦИЕНТ СОСТАВА НАСАЖДЕНИЯ', level=2)

            total_densities = {}
            total_density_all = 0  # Общая густота всех пород (сумма средних)
            
            for breed_name, data in breeds_data.items():
                if data['plots']:
                    # Рассчитываем СРЕДНЮЮ густоту породы (сумма плотностей / кол-во площадок)
                    # Плотности уже в шт/га, делим на 1000 для перевода в тыс.шт/га
                    avg_density = sum(p.get('density', 0) for p in data['plots']) / len(data['plots'])
                    avg_density_thousands = avg_density / 1000  # переводим в тыс.шт/га
                    total_density_all += avg_density_thousands
                    if avg_density_thousands > 0:
                        total_densities[breed_name] = avg_density_thousands

            if total_densities:
                # Метод наибольшего остатка для распределения коэффициентов
                # Сначала рассчитываем точные коэффициенты
                exact_coeffs = []
                for breed_name, density in sorted(total_densities.items(), key=lambda x: x[1], reverse=True):
                    exact_coeff = (density / total_density_all * 10) if total_density_all > 0 else 1
                    exact_coeffs.append(exact_coeff)
                
                # Округляем вниз
                coeffs_floor = [int(coeff) for coeff in exact_coeffs]
                # Считаем остатки
                remainders = [(exact_coeffs[i] - coeffs_floor[i], i) for i in range(len(exact_coeffs))]
                # Сортируем по убыванию остатков
                remainders.sort(reverse=True)
                
                # Распределяем единицы начиная с наибольших остатков
                coeffs = coeffs_floor[:]
                total = sum(coeffs)
                i = 0
                while total < 10 and i < len(remainders):
                    idx = remainders[i][1]
                    coeffs[idx] += 1
                    total += 1
                    i += 1
                
                # Формируем формулу состава
                sorted_breeds = sorted(total_densities.items(), key=lambda x: x[1], reverse=True)
                composition_parts = []
                for i, (breed_name, _) in enumerate(sorted_breeds):
                    if i < len(coeffs):
                        breed_letter = self.get_breed_letter(breed_name)
                        composition_parts.append(f"{coeffs[i]}{breed_letter}")

                composition_text = ''.join(composition_parts) + "Др"
                doc.add_paragraph(f"Формула состава: {composition_text}")

            # Хвойные породы
            doc.add_heading('ХВОЙНЫЕ ПОРОДЫ - ВЫСОТА ПО ГРАДАЦИЯМ', level=2)

            has_coniferous = False
            for breed_name, data in sorted(breeds_data.items()):
                if data['type'] == 'coniferous' and data['plots']:
                    has_coniferous = True
                    zones = data.get('coniferous_zones', {})
                    # Густота по градациям = общее кол-во деревьев в градации / общая площадь всех площадок в га
                    plot_area_ha = 3.14159 * (float(current_radius) if current_radius else 1.78) ** 2 / 10000
                    total_area_ha = plot_area_ha * plot_count
                    total_do_05 = sum(p.get('do_05', 0) for p in data['plots'])
                    total_05_15 = sum(p.get('05_15', 0) for p in data['plots'])
                    total_bolee_15 = sum(p.get('bolee_15', 0) for p in data['plots'])
                    avg_do_05 = total_do_05 / total_area_ha if total_area_ha > 0 else 0
                    avg_05_15 = total_05_15 / total_area_ha if total_area_ha > 0 else 0
                    avg_bolee_15 = total_bolee_15 / total_area_ha if total_area_ha > 0 else 0
                    # Высота = средняя только на тех площадках, где есть порода
                    avg_height_total = sum(p['height'] for p in data['plots'] if p['height'] > 0) / len([p for p in data['plots'] if p['height'] > 0]) if any(p['height'] > 0 for p in data['plots']) else 0

                    p = doc.add_paragraph()
                    p.add_run(f"{breed_name}:").bold = True
                    doc.add_paragraph(f"• до 0.5м: {avg_do_05:.1f} шт/га")
                    doc.add_paragraph(f"• 0.5-1.5м: {avg_05_15:.1f} шт/га")
                    doc.add_paragraph(f"• >1.5м: {avg_bolee_15:.1f} шт/га")
                    doc.add_paragraph(f"• средняя высота породы: {avg_height_total:.1f}м")

            # Лиственные породы
            doc.add_heading('ЛИСТВЕННЫЕ ПОРОДЫ - СРЕДНИЕ ПОКАЗАТЕЛИ', level=2)

            has_deciduous = False
            for breed_name, data in sorted(breeds_data.items()):
                if data['type'] == 'deciduous' and data['plots']:
                    has_deciduous = True
                    # Густота = общее количество деревьев / общая площадь всех площадок в га
                    total_trees = sum(p.get('density_raw', 0) for p in data['plots'])
                    plot_area_ha = 3.14159 * (float(current_radius) if current_radius else 1.78) ** 2 / 10000
                    total_area_ha = plot_area_ha * plot_count
                    avg_density = total_trees / total_area_ha if total_area_ha > 0 else 0
                    # Высота = средняя только на тех площадках, где есть высота
                    avg_heights = [p['height'] for p in data['plots'] if p['height'] > 0]
                    avg_height = sum(avg_heights) / len(avg_heights) if avg_heights else 0
                    # Возраст = средний только на тех площадках, где есть возраст
                    avg_ages = [p['age'] for p in data['plots'] if p['age'] > 0]
                    avg_age = sum(avg_ages) / len(avg_ages) if avg_ages else 0

                    p = doc.add_paragraph()
                    p.add_run(f"{breed_name}:").bold = True
                    doc.add_paragraph(f"• Средняя густота: {avg_density:.1f} шт/га")
                    doc.add_paragraph(f"• Средняя высота: {avg_height:.1f}м")
                    doc.add_paragraph(f"• Средний возраст: {avg_age:.1f} лет")

            doc.save(full_path)
            self.show_success(f"Итоги сохранены в Word: {filename}")
        except ImportError:
            self.show_error("Для сохранения в Word установите библиотеку python-docx: pip install python-docx")
        except Exception as e:
            self.show_error(f"Ошибка сохранения итогов в Word: {str(e)}")

    def show_plot_area_input_popup(self, instance):
        """Показать popup для ввода площади участка в гектарах"""
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)

        title_label = MDLabel(
            text="Введите площадь обследуемого участка",
            font_style='Title', role='medium',
            theme_text_color='Custom', text_color=Colors.GREEN,
            adaptive_height=True
        )
        content.add_widget(title_label)

        self.plot_area_input_field = MDTextField(
            hint_text="Площадь участка (га)",
            mode='outlined',
            size_hint_y=None, height=dp(52),
            input_filter='float',
            text=self.plot_area_input if hasattr(self, 'plot_area_input') and self.plot_area_input else ''
        )
        self.plot_area_input_field.bind(text=self.update_plot_area_display)
        content.add_widget(self.plot_area_input_field)

        info_label = MDLabel(
            text="Укажите площадь обследуемого участка в гектарах.\n"
                 "Это значение используется для расчета площади перечета\n"
                 "по всем площадкам и отображается в итоговых отчетах.",
            theme_text_color='Custom', text_color=[0.8,0.8,0.8,1],
            adaptive_height=True
        )
        content.add_widget(info_label)

        btn_layout = MDBoxLayout(orientation='horizontal', spacing=Spacing.MD, adaptive_height=True)
        save_btn = MDButton(style='filled', md_bg_color=Colors.BTN_SUCCESS,
                            size_hint=(0.5, None), height=dp(48))
        save_btn.add_widget(MDButtonText(text='Сохранить'))
        cancel_btn = MDButton(style='outlined',
                              size_hint=(0.5, None), height=dp(48),
                              line_color=Colors.DANGER)
        cancel_btn.add_widget(MDButtonText(text='Отмена',
                              theme_text_color='Custom', text_color=Colors.DANGER))
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            size_hint=(0.85, None),
            height=dp(300),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3]
        )

        def save_plot_area(btn):
            try:
                plot_area = float(self.plot_area_input_field.text.strip())
                if plot_area <= 0:
                    self.show_error("Площадь участка должна быть положительным числом!")
                    return
                self.plot_area_input = str(plot_area)
                self.project_data['address']['plot_area'] = str(plot_area)
                self.show_success(f"Площадь участка {plot_area} га сохранена")
                popup.dismiss()
                self.update_address_popup_display()
            except ValueError:
                self.show_error("Введите корректное числовое значение площади!")

        save_btn.bind(on_release=save_plot_area)
        cancel_btn.bind(on_release=popup.dismiss)

        popup.open()

    def show_plot_area_ha_popup(self, instance):
        """Показать popup с информацией о площади участка в гектарах"""
        try:
            current_radius = float(self.current_radius) if self.current_radius else 5.64
            plot_area_m2 = 3.14159 * (current_radius ** 2)
            plot_area_ha = plot_area_m2 / 10000

            # Расчет площади перечета по всем площадкам
            total_plot_area_ha = 0.0
            plot_count = 0

            for page_num, page_rows in self.page_data.items():
                for row in page_rows:
                    if len(row) >= 4 and row[3]:  # Есть данные о породах
                        try:
                            breeds_data = json.loads(row[3]) if isinstance(row[3], str) else []
                            if breeds_data:
                                plot_count += 1
                                total_plot_area_ha += plot_area_ha
                        except (json.JSONDecodeError, TypeError):
                            continue

            content = MDBoxLayout(orientation='vertical', spacing=10, padding=10, md_bg_color=Colors.DARK_SURFACE)

            title_label = MDLabel(
                text="Площадь участка в гектарах",
                font_size='18sp',
                bold=True,
                theme_text_color='Custom',
                text_color=Colors.GREEN,
                size_hint=(1, None),
                height=40
            )
            content.add_widget(title_label)

            info_text = f"""
Одиночная площадка:
Радиус: {current_radius:.2f} м
Площадь: {plot_area_ha:.4f} га

Всего площадок: {plot_count}
Совокупная площадь перечета: {total_plot_area_ha:.4f} га

Расчет совокупной площади:
{plot_count} площадок × {plot_area_ha:.4f} га = {total_plot_area_ha:.4f} га

Пример расчета густоты на гектар:
Если на площадке 10 деревьев, то густота = 10 / {plot_area_ha:.4f} ≈ {10/plot_area_ha:.1f} шт/га
"""

            info_label = MDLabel(
                text=info_text,
                font_size='14sp',
                theme_text_color='Custom',
                text_color=Colors.TEXT_ON_DARK,
                size_hint=(1, None),
                height=250,
                halign='left',
                valign='top'
            )
            info_label.bind(size=lambda *args: setattr(info_label, 'text_size', (info_label.width, None)))
            content.add_widget(info_label)

            close_btn = MDButton(style='filled', size_hint=(1, None), height=50)
            close_btn.add_widget(MDButtonText(text='Закрыть', theme_text_color='Custom', text_color=[0.7,0.7,0.7,1]))
            content.add_widget(close_btn)

            popup = Popup(
                title="",
                content=content,
                separator_height=0,
                background_color=[0,0,0,0.3],
                overlay_color=[0,0,0,0.3],
                size_hint=(0.8, 0.8)
            )

            close_btn.bind(on_release=popup.dismiss)
            popup.open()

        except Exception as e:
            self.show_error(f"Ошибка расчета площади: {str(e)}")

    def show_plot_area_combined_popup(self, instance):
        """Показать объединенное popup для работы с площадью участка"""
        content = MDBoxLayout(orientation='vertical', spacing=20, padding=20, md_bg_color=Colors.DARK_SURFACE)

        try:
            current_radius = float(self.current_radius) if self.current_radius else 5.64
            plot_area_m2 = 3.14159 * (current_radius ** 2)
            plot_area_ha = plot_area_m2 / 10000

            # Расчет площади перечета по всем площадкам
            total_plot_area_ha = 0.0
            plot_count = 0

            for page_num, page_rows in self.page_data.items():
                for row in page_rows:
                    if len(row) >= 4 and row[3]:  # Есть данные о породах
                        try:
                            breeds_data = json.loads(row[3]) if isinstance(row[3], str) else []
                            if breeds_data:
                                plot_count += 1
                                total_plot_area_ha += plot_area_ha
                        except (json.JSONDecodeError, TypeError):
                            continue

            title_label = MDLabel(
                text="Площадь участка",
                font_size='20sp',
                bold=True,
                theme_text_color='Custom',
                text_color=Colors.GREEN,
                size_hint=(1, None),
                height=50,
                halign='center'
            )
            content.add_widget(title_label)

            # Раздел ввода площади участка
            input_section = MDBoxLayout(orientation='vertical', spacing=10, size_hint=(1, None), height=120, md_bg_color=Colors.DARK_SURFACE)

            input_title = MDLabel(
                text="Ввод площади участка",
                font_size='16sp',
                bold=True,
                size_hint=(1, None),
                height=30,
                halign='center',
                theme_text_color='Custom',
                text_color=Colors.TEXT_ON_DARK
            )
            input_section.add_widget(input_title)

            plot_area_input_field = MDTextField(
                hint_text="Площадь участка (га)",
                mode="outlined",
                size_hint=(1, None),
                height=50,
                line_color_focus=Colors.GREEN,
                text=str(self._get_current_plot_area_input()) if hasattr(self, 'plot_area_input') and self.plot_area_input else ''
            )
            input_section.add_widget(plot_area_input_field)

            content.add_widget(input_section)

            # Раздел информации о площади
            info_label = MDLabel(
                text="Информация о площади участка:",
                font_size='16sp',
                bold=True,
                size_hint=(1, None),
                height=30,
                halign='center',
                theme_text_color='Custom',
                text_color=Colors.TEXT_ON_DARK
            )
            content.add_widget(info_label)

            info_text = ScrollView(size_hint=(1, None), height=250)
            info_layout = MDBoxLayout(orientation='vertical', spacing=5, padding=10, size_hint_y=None, md_bg_color=Colors.DARK_SURFACE)
            info_layout.bind(minimum_height=info_layout.setter('height'))

            info_data = MDLabel(
                text=f"""Одиночная площадка:
Радиус: {current_radius:.2f} м
Площадь: {plot_area_ha:.4f} га

Всего площадок: {plot_count}
Совокупная площадь перечета: {total_plot_area_ha:.4f} га

Расчет совокупной площади:
{plot_count} площадок × {plot_area_ha:.4f} га = {total_plot_area_ha:.4f} га

Пример расчета густоты на гектар:
Если на площадке 10 деревьев, то густота = 10 / {plot_area_ha:.4f} ≈ {10/plot_area_ha:.1f} шт/га""",
                font_size='14sp',
                theme_text_color='Custom',
                text_color=Colors.TEXT_ON_DARK,
                size_hint=(1, None),
                height=200,
                halign='left',
                valign='top'
            )
            info_data.bind(size=lambda *args: setattr(info_data, 'text_size', (info_data.width, None)))
            info_layout.add_widget(info_data)
            info_text.add_widget(info_layout)

            content.add_widget(info_text)

            # Кнопки управления (объединение сохранения и обновления в одну кнопку)
            btn_layout = MDBoxLayout(orientation='horizontal', spacing=10, size_hint=(1, None), height=60, md_bg_color=Colors.DARK_SURFACE)

            combined_btn = MDButton(style='filled', size_hint=(0.7, 1), height=60)
            combined_btn.add_widget(MDButtonText(text='Сохранить и обновить', theme_text_color='Custom', text_color=Colors.GREEN))

            close_btn = MDButton(style='filled', size_hint=(0.3, 1), height=60)
            close_btn.add_widget(MDButtonText(text='Закрыть', theme_text_color='Custom', text_color=[1,0.3,0.3,1]))

            btn_layout.add_widget(combined_btn)
            btn_layout.add_widget(close_btn)

            content.add_widget(btn_layout)

            popup = Popup(
                title="",
                content=content,
                separator_height=0,
                background_color=[0,0,0,0.3],
                overlay_color=[0,0,0,0.3],
                size_hint=(0.8, 0.9)
            )

            def save_and_refresh(btn):
                # Сначала сохраняем площадь участка
                try:
                    plot_area = float(plot_area_input_field.text.strip())
                    if plot_area <= 0:
                        self.show_error("Площадь участка должна быть положительным числом!")
                        return

                    self.plot_area_input = str(plot_area)
                    self.show_success(f"Площадь участка {plot_area} га сохранена")
                except ValueError:
                    self.show_error("Введите корректное числовое значение площади!")
                    return

                # Затем обновляем информацию о площади
                popup.dismiss()
                self.show_plot_area_combined_popup(instance)

            combined_btn.bind(on_release=save_and_refresh)
            close_btn.bind(on_release=popup.dismiss)

            popup.open()

        except Exception as e:
            self.show_error(f"Ошибка расчета площади: {str(e)}")

    def _get_current_plot_area_input(self):
        """Получить текущее значение площади участка"""
        # If stored in instance variable
        if hasattr(self, 'plot_area_input') and self.plot_area_input:
            return self.plot_area_input
        return ''

    def update_plot_total(self, instance, value):
        """Обновляем итог по площадке при изменении данных"""
        row_idx = instance.row_index

        breeds_text = self.inputs[row_idx][3].text
        breeds_data = self.parse_breeds_data(breeds_text)

        if not breeds_data:
            return

        total_density = 0
        total_height = 0.0
        total_age = 0
        breed_count = 0
        breed_names = []

        for breed_info in breeds_data:
            breed_count += 1
            breed_name = breed_info.get('name', 'Неизвестная')
            breed_names.append(breed_name)

            if breed_info.get('type') == 'coniferous':
                coniferous_density = (breed_info.get('do_05', 0) +
                                    breed_info.get('05_15', 0) +
                                    breed_info.get('bolee_15', 0))
                if coniferous_density > 0:
                    total_density += coniferous_density
            elif 'density' in breed_info and breed_info['density']:
                total_density += breed_info['density']

            if 'height' in breed_info and breed_info['height']:
                total_height += breed_info['height']
            if 'age' in breed_info and breed_info['age']:
                total_age += breed_info['age']

        # Обновляем общие итоги
        self.update_totals()

    def show_care_queue_popup(self, instance):
        """Показать popup для выбора мероприятий рубки"""
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)

        title_label = MDLabel(
            text="Выберите мероприятие:",
            font_style='Title', role='medium',
            theme_text_color='Custom', text_color=Colors.GREEN,
            size_hint_y=None, height=dp(40)
        )
        content.add_widget(title_label)

        self.activity_input = MDTextField(
            hint_text="Название мероприятия (необязательно)",
            mode='outlined',
            size_hint_y=None, height=dp(52)
        )
        content.add_widget(self.activity_input)

        queue_label = MDLabel(
            text="Очередь:",
            font_style='Label', role='medium',
            theme_text_color='Custom', text_color=[1,1,1,1],
            size_hint_y=None, height=dp(30)
        )
        content.add_widget(queue_label)

        self.queue_checkboxes = {}
        queue_grid = MDGridLayout(cols=3, spacing=Spacing.MD, adaptive_height=True,
                                  size_hint_y=None, padding=[0, 0])
        for q in ['первая', 'вторая', 'третья']:
            cell = MDBoxLayout(orientation='horizontal', spacing=Spacing.XS,
                               size_hint_y=None, height=dp(40))
            cb = CheckBox(size_hint=(None, 1), width=dp(36), color=[0.3, 0.8, 0.3, 1],
                          active=(q == 'первая'))
            lbl = MDLabel(text=q, theme_text_color='Custom', text_color=[1,1,1,1],
                          size_hint=(1, 1), valign='middle')
            cell.add_widget(cb)
            cell.add_widget(lbl)
            queue_grid.add_widget(cell)
            self.queue_checkboxes[q] = cb
        content.add_widget(queue_grid)

        self.activity_checkboxes = {}
        activities = ['осветление', 'прочистка']

        activities_label = MDLabel(
            text="Типы мероприятий:",
            font_style='Label', role='medium',
            theme_text_color='Custom', text_color=[1,1,1,1],
            size_hint_y=None, height=dp(30)
        )
        content.add_widget(activities_label)

        act_grid = MDGridLayout(cols=2, spacing=Spacing.MD, adaptive_height=True,
                                size_hint_y=None, padding=[0, 0])
        for activity in activities:
            cell = MDBoxLayout(orientation='horizontal', spacing=Spacing.XS,
                               size_hint_y=None, height=dp(40))
            checkbox = CheckBox(size_hint=(None, 1), width=dp(36), color=[0.3, 0.8, 0.3, 1])
            label = MDLabel(
                text=activity,
                theme_text_color='Custom', text_color=[1,1,1,1],
                size_hint=(1, 1), valign='middle'
            )
            cell.add_widget(checkbox)
            cell.add_widget(label)
            act_grid.add_widget(cell)
            self.activity_checkboxes[activity] = checkbox
        content.add_widget(act_grid)

        btn_layout = MDBoxLayout(orientation='horizontal', spacing=Spacing.MD,
                                 adaptive_height=True)
        save_btn = MDButton(style='filled', md_bg_color=Colors.BTN_SUCCESS,
                            size_hint=(0.5, None), height=dp(48))
        save_btn.add_widget(MDButtonText(text='Сохранить'))
        cancel_btn = MDButton(style='outlined',
                              size_hint=(0.5, None), height=dp(48),
                              line_color=Colors.DANGER)
        cancel_btn.add_widget(MDButtonText(text='Отмена',
                              theme_text_color='Custom', text_color=Colors.DANGER))
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            size_hint=(0.85, None),
            height=dp(440),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3]
        )

        def save_activity(btn):
            activity_text = self.activity_input.text.strip()
            selected_queues = [q for q, cb in self.queue_checkboxes.items() if cb.active]
            selected_activities = [activity for activity, checkbox in self.activity_checkboxes.items() if checkbox.active]

            if not activity_text and not selected_activities:
                self.show_error("Введите мероприятие или выберите тип мероприятия!")
                return

            if not selected_queues:
                self.show_error("Выберите очередь!")
                return

            selected_queue = ', '.join(selected_queues)
            self.care_queue = selected_queue
            if activity_text:
                self.care_queue += f" - {activity_text}"
            if selected_activities:
                self.care_queue += f" ({', '.join(selected_activities)})"

            self.project_data['details']['care_queue'] = self.care_queue

            result_parts = []
            if activity_text:
                result_parts.append(f"Мероприятие: {activity_text}")
            result_parts.append(f"Очередь: {selected_queue}")
            if selected_activities:
                result_parts.append(f"Типы: {', '.join(selected_activities)}")

            self.show_success(f"Мероприятие сохранено: {'; '.join(result_parts)}")
            if hasattr(self, 'current_details_info'):
                self.update_details_display()
            popup.dismiss()

        save_btn.bind(on_release=save_activity)
        cancel_btn.bind(on_release=popup.dismiss)

        popup.open()

    def show_characteristics_popup(self, instance):
        """Показать popup для характеристики молодняков"""
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)

        title_label = MDLabel(
            text="Характеристика молодняков:",
            font_style='Title', role='medium',
            theme_text_color='Custom', text_color=Colors.GREEN,
            size_hint_y=None, height=dp(40)
        )
        content.add_widget(title_label)

        self.characteristics_inputs = {}
        characteristics = ['Лучшие', 'Вспомогательные', 'Нежелательные']

        existing_characteristics = {}
        if self.project_data['details'].get('characteristics'):
            try:
                existing_data = self.project_data['details']['characteristics']
                for line in existing_data.split('\n'):
                    if ':' in line:
                        key, value = line.split(':', 1)
                        key = key.strip()
                        value = value.strip()
                        if key in characteristics:
                            existing_characteristics[key] = value
            except:
                pass

        for char in characteristics:
            char_label = MDLabel(
                text=f"{char}:",
                font_style='Label', role='medium',
                theme_text_color='Custom', text_color=[1,1,1,1],
                size_hint_y=None, height=dp(25)
            )
            content.add_widget(char_label)
            char_input = MDTextField(
                hint_text="Введите название породы",
                mode='outlined',
                text=existing_characteristics.get(char, ''),
                multiline=True,
                size_hint_y=None, height=dp(60)
            )
            content.add_widget(char_input)
            self.characteristics_inputs[char] = char_input

        btn_layout = MDBoxLayout(orientation='horizontal', spacing=Spacing.MD, adaptive_height=True)
        save_btn = MDButton(style='filled', md_bg_color=Colors.BTN_SUCCESS,
                            size_hint=(0.5, None), height=dp(48))
        save_btn.add_widget(MDButtonText(text='Сохранить'))
        cancel_btn = MDButton(style='outlined',
                              size_hint=(0.5, None), height=dp(48),
                              line_color=Colors.DANGER)
        cancel_btn.add_widget(MDButtonText(text='Отмена',
                              theme_text_color='Custom', text_color=Colors.DANGER))
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            size_hint=(0.85, None),
            height=dp(420),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3]
        )

        def save_characteristics(btn):
            filled_characteristics = {}
            for char, input_field in self.characteristics_inputs.items():
                value = input_field.text.strip()
                if value:
                    filled_characteristics[char] = value

            if filled_characteristics:
                characteristics_text = "\n".join([f"{k}: {v}" for k, v in filled_characteristics.items()])
                self.project_data['details']['characteristics'] = characteristics_text
                self.characteristics = characteristics_text
                self.show_success(f"Характеристики сохранены:\n{characteristics_text}")
                self.update_details_display()
            else:
                self.show_error("Заполните хотя бы одну характеристику!")
                return
            popup.dismiss()

        save_btn.bind(on_release=save_characteristics)
        cancel_btn.bind(on_release=popup.dismiss)

        popup.open()

    def show_date_popup(self, instance):
        """Показать popup для ввода даты рубки"""
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)

        title_label = MDLabel(
            text="Введите дату рубки:",
            font_style='Title', role='medium',
            theme_text_color='Custom', text_color=Colors.GREEN,
            size_hint_y=None, height=dp(40)
        )
        content.add_widget(title_label)

        self.date_input = MDTextField(
            hint_text="ДД.ММ.ГГГГ",
            mode='outlined',
            size_hint_y=None, height=dp(52)
        )
        content.add_widget(self.date_input)

        info_label = MDLabel(
            text="Формат: ДД.ММ.ГГГГ\nНапример: 15.06.2025",
            font_style='Body', role='small',
            theme_text_color='Custom', text_color=[0.8,0.8,0.8,1],
            size_hint_y=None, height=dp(50)
        )
        content.add_widget(info_label)

        btn_layout = MDBoxLayout(orientation='horizontal', spacing=Spacing.MD, adaptive_height=True)
        save_btn = MDButton(style='filled', md_bg_color=Colors.BTN_SUCCESS,
                            size_hint=(0.5, None), height=dp(48))
        save_btn.add_widget(MDButtonText(text='Сохранить'))
        cancel_btn = MDButton(style='outlined',
                              size_hint=(0.5, None), height=dp(48),
                              line_color=Colors.DANGER)
        cancel_btn.add_widget(MDButtonText(text='Отмена',
                              theme_text_color='Custom', text_color=Colors.DANGER))
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            size_hint=(0.85, None),
            height=dp(280),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3]
        )

        def save_date(btn):
            date_text = self.date_input.text.strip()
            if date_text:
                import re
                if re.match(r'^\d{2}\.\d{2}\.\d{4}$', date_text):
                    self.care_date = date_text
                    self.project_data['details']['care_date'] = date_text
                    self.show_success(f"Дата рубки сохранена: {date_text}")
                    self.update_details_display()
                else:
                    self.show_error("Неверный формат даты! Используйте ДД.ММ.ГГГГ")
                    return
            else:
                self.show_error("Введите дату рубки!")
                return
            popup.dismiss()

        save_btn.bind(on_release=save_date)
        cancel_btn.bind(on_release=popup.dismiss)

        popup.open()

    def show_technology_popup(self, instance):
        """Показать popup для ввода технологии ухода"""
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)

        title_label = MDLabel(
            text="Введите технологию ухода:",
            font_style='Title', role='medium',
            theme_text_color='Custom', text_color=Colors.GREEN,
            size_hint_y=None, height=dp(40)
        )
        content.add_widget(title_label)

        self.technology_input = MDTextField(
            hint_text="Опишите технологию ухода",
            mode='outlined',
            multiline=True,
            size_hint_y=None, height=dp(100)
        )
        content.add_widget(self.technology_input)

        btn_layout = MDBoxLayout(orientation='horizontal', spacing=Spacing.MD, adaptive_height=True)
        save_btn = MDButton(style='filled', md_bg_color=Colors.BTN_SUCCESS,
                            size_hint=(0.5, None), height=dp(48))
        save_btn.add_widget(MDButtonText(text='Сохранить'))
        cancel_btn = MDButton(style='outlined',
                              size_hint=(0.5, None), height=dp(48),
                              line_color=Colors.DANGER)
        cancel_btn.add_widget(MDButtonText(text='Отмена',
                              theme_text_color='Custom', text_color=Colors.DANGER))
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            size_hint=(0.85, None),
            height=dp(320),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3]
        )

        def save_technology(btn):
            technology_text = self.technology_input.text.strip()
            if technology_text:
                self.technology = technology_text
                self.project_data['details']['technology'] = technology_text
                self.show_success(f"Технология ухода сохранена: {technology_text[:50]}...")
                self.update_details_display()
            else:
                self.show_error("Введите технологию ухода!")
                return
            popup.dismiss()

        save_btn.bind(on_release=save_technology)
        cancel_btn.bind(on_release=popup.dismiss)

        popup.open()

    def show_forest_purpose_popup(self, instance):
        """Показать popup для выбора назначения лесов"""
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)

        title_label = MDLabel(
            text="Выберите назначение лесов:",
            font_style='Title', role='medium',
            theme_text_color='Custom', text_color=Colors.GREEN,
            size_hint_y=None, height=dp(40)
        )
        content.add_widget(title_label)

        forest_purposes = [
            ('Эксплуатационные', 'Эксплуатационные леса'),
            ('Защитные', 'Защитные леса'),
            ('Резервные', 'Резервные леса')
        ]

        self.forest_purpose_checkboxes = {}
        self.selected_forest_purpose = None

        purpose_grid = MDGridLayout(cols=1, spacing=Spacing.MD, adaptive_height=True,
                                    size_hint_y=None, padding=[0, 0])
        for short_name, full_name in forest_purposes:
            cell = MDBoxLayout(orientation='horizontal', spacing=Spacing.SM,
                               size_hint_y=None, height=dp(40))
            checkbox = CheckBox(size_hint=(None, 1), width=dp(36), color=[0.3, 0.8, 0.3, 1])
            label = MDLabel(
                text=f"{short_name} ({full_name})",
                theme_text_color='Custom', text_color=[1,1,1,1],
                size_hint=(1, 1), valign='middle'
            )
            cell.add_widget(checkbox)
            cell.add_widget(label)
            purpose_grid.add_widget(cell)
            self.forest_purpose_checkboxes[full_name] = checkbox
        content.add_widget(purpose_grid)

        btn_layout = MDBoxLayout(orientation='horizontal', spacing=Spacing.MD, adaptive_height=True)
        save_btn = MDButton(style='filled', md_bg_color=Colors.BTN_SUCCESS,
                            size_hint=(0.5, None), height=dp(48))
        save_btn.add_widget(MDButtonText(text='Сохранить'))
        cancel_btn = MDButton(style='outlined',
                              size_hint=(0.5, None), height=dp(48),
                              line_color=Colors.DANGER)
        cancel_btn.add_widget(MDButtonText(text='Отмена',
                              theme_text_color='Custom', text_color=Colors.DANGER))
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            size_hint=(0.85, None),
            height=dp(360),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3]
        )

        def save_forest_purpose(btn):
            selected_purpose = None
            for purpose, checkbox in self.forest_purpose_checkboxes.items():
                if checkbox.active:
                    selected_purpose = purpose
                    break

            if selected_purpose:
                self.forest_purpose = selected_purpose
                self.project_data['details']['forest_purpose'] = selected_purpose
                self.show_success(f"Назначение лесов установлено: {selected_purpose}")
                self.update_details_display()
                popup.dismiss()
            else:
                self.show_error("Выберите назначение лесов!")

        save_btn.bind(on_release=save_forest_purpose)
        cancel_btn.bind(on_release=popup.dismiss)
        self.forest_purpose_popup = popup
        popup.open()

    def select_forest_purpose(self, purpose):
        """Обработка выбора назначения лесов"""
        self.selected_forest_purpose = purpose
        self.show_success(f"Назначение лесов установлено: {purpose}")
        if hasattr(self, 'forest_purpose_popup'):
            self.forest_purpose_popup.dismiss()

    def show_additional_functions_popup(self, instance):
        """Показать popup с дополнительными функциями"""
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE)

        title_label = MDLabel(
            text="Детали проекта",
            font_style='Headline', role='medium',
            theme_text_color='Custom', text_color=Colors.GREEN,
            size_hint_y=None, height=dp(50)
        )
        content.add_widget(title_label)

        scroll = ScrollView(size_hint=(1, 1))
        scroll_content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD,
                                     adaptive_height=True, size_hint_y=None)
        scroll_content.bind(minimum_height=scroll_content.setter('height'))

        buttons_layout = MDGridLayout(cols=2, spacing=Spacing.MD,
                                      adaptive_height=True, size_hint_y=None,
                                      padding=[0, 0])

        def make_detail_card(text, icon_name, callback):
            card = MDCard(style='elevated', size_hint=(1, None), height=dp(70),
                          md_bg_color=Colors.CARD_BG, ripple_behavior=True,
                          on_release=callback, focus_behavior=True)
            card.add_widget(MDBoxLayout(
                MDIcon(icon=icon_name, theme_text_color='Custom', text_color=Colors.GREEN,
                       size_hint=(None, 1), width=dp(40)),
                MDLabel(text=text, font_style='Label', role='medium',
                        theme_text_color='Custom', text_color=[1,1,1,1],
                        halign='left', adaptive_height=True),
                orientation='horizontal', spacing=Spacing.SM, padding=[Spacing.SM, 0],
                adaptive_height=True
            ))
            return card

        buttons_layout.add_widget(make_detail_card('Вид рубки', 'content-cut', self.show_care_queue_popup))
        buttons_layout.add_widget(make_detail_card('ХарактерМ', 'texture', self.show_characteristics_popup))
        buttons_layout.add_widget(make_detail_card('Дата рубки', 'calendar', self.show_date_popup))
        buttons_layout.add_widget(make_detail_card('Технология\nухода', 'wrench', self.show_technology_popup))
        buttons_layout.add_widget(make_detail_card('Назначение\nлесов', 'pine-tree', self.show_forest_purpose_popup))

        scroll_content.add_widget(buttons_layout)

        care_queue_val = self.project_data['details'].get('care_queue', '') or self.care_queue or 'Не указана'
        characteristics_val = self.project_data['details'].get('characteristics', '') or self.characteristics or 'Не указана'
        care_date_val = self.project_data['details'].get('care_date', '') or self.care_date or 'Не указана'
        technology_val = self.project_data['details'].get('technology', '') or self.technology or 'Не указана'
        forest_purpose_val = self.project_data['details'].get('forest_purpose', '') or self.forest_purpose or 'Не указано'

        current_card = MDCard(style='elevated', md_bg_color=Colors.CARD_BG,
                              adaptive_height=True, padding=Spacing.MD, spacing=Spacing.SM)

        current_title = MDLabel(
            text="Текущие значения проекта:",
            font_style='Title', role='small',
            theme_text_color='Custom', text_color=Colors.GREEN,
            size_hint_y=None, height=dp(30)
        )
        current_card.add_widget(current_title)

        self.current_details_info = MDLabel(
            text=f"Очередь рубки: {care_queue_val}\n"
                 f"Характеристика молодняков: {characteristics_val}\n"
                 f"Дата рубки: {care_date_val}\n"
                 f"Технология ухода: {technology_val}\n"
                 f"Назначение лесов: {forest_purpose_val}",
            font_style='Body', role='medium',
            theme_text_color='Custom', text_color=[1,1,1,1],
            adaptive_height=True
        )
        current_card.add_widget(self.current_details_info)
        scroll_content.add_widget(current_card)

        scroll.add_widget(scroll_content)
        content.add_widget(scroll)

        cancel_btn = MDButton(style='outlined', size_hint=(1, None), height=dp(48),
                              line_color=Colors.DANGER)
        cancel_btn.add_widget(MDButtonText(text='Закрыть',
                              theme_text_color='Custom', text_color=Colors.DANGER))
        content.add_widget(cancel_btn)

        popup = Popup(
            title="",
            content=content,
            size_hint=(0.9, 0.9),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3]
        )

        cancel_btn.bind(on_release=popup.dismiss)
        popup.open()

    def get_breed_letter(self, breed_name):
        """Получение первой буквы для коэффициента состава породы"""
        breed_letters = {
            'Сосна': 'С',
            'Ель': 'Е',
            'Пихта': 'П',
            'Кедр': 'К',
            'Лиственница': 'Л',
            'Берёза': 'Б',
            'Осина': 'Ос',
            'Ольха чёрная': 'ОЧ',
            'Ольха серая': 'ОС',
            'Ива': 'И',
            'Ива кустарниковая': 'ИК'
        }

        for full_name, letter in breed_letters.items():
            if full_name.lower() in breed_name.lower():
                return letter

        # Возвращаем первую букву имени породы, если не найдено
        return breed_name[0].upper() if breed_name else 'Н'

    def show_error(self, message):
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)
        content.add_widget(MDIcon(icon='alert-circle', theme_text_color='Custom', text_color=Colors.DANGER,
                                   size_hint=(None, None), size=(dp(48), dp(48)), halign='center'))
        content.add_widget(MDLabel(text=message, theme_text_color='Custom', text_color=[1,0.3,0.3,1],
                                    halign='center', adaptive_height=True))
        Popup(
            title="",
            content=content,
            size_hint=(0.7, None),
            height=dp(180),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3]
        ).open()

    def show_success(self, message):
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)
        content.add_widget(MDIcon(icon='check-circle', theme_text_color='Custom', text_color=Colors.GREEN,
                                   size_hint=(None, None), size=(dp(48), dp(48)), halign='center'))
        content.add_widget(MDLabel(text=message, theme_text_color='Custom', text_color=[0.3,0.8,0.3,1],
                                    halign='center', adaptive_height=True))
        Popup(
            title="",
            content=content,
            size_hint=(0.7, None),
            height=dp(180),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3]
        ).open()

    def show_quarter_popup(self, instance):
        """Показать popup для ввода квартала"""
        content = MDBoxLayout(orientation='vertical', spacing=10, padding=10, md_bg_color=Colors.DARK_SURFACE)

        title_label = MDLabel(
            text="Введите номер квартала",
            bold=True,
            size_hint=(1, None),
            height=30,
            theme_text_color='Custom',
            text_color=Colors.TEXT_ON_DARK
        )
        content.add_widget(title_label)

        self.quarter_input = MDTextField(
            hint_text="Номер квартала",
            mode="outlined",
            size_hint=(1, None),
            height=40,
            line_color_focus=Colors.GREEN,
            text=self.current_quarter
        )
        content.add_widget(self.quarter_input)

        btn_layout = MDBoxLayout(orientation='horizontal', spacing=10, size_hint=(1, None), height=40, md_bg_color=Colors.DARK_SURFACE)
        save_btn = MDButton(style='filled', size_hint=(0.5, 1), height=50)
        save_btn.add_widget(MDButtonText(text='Сохранить', theme_text_color='Custom', text_color=Colors.GREEN))
        cancel_btn = MDButton(style='filled', size_hint=(0.5, 1), height=50)
        cancel_btn.add_widget(MDButtonText(text='Отмена', theme_text_color='Custom', text_color=[1,0.3,0.3,1]))
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3],
            size_hint=(0.6, 0.5)
        )

        def save_quarter(btn):
            quarter = self.quarter_input.text.strip()
            if quarter:
                self.current_quarter = quarter
                self.project_data['address']['quarter'] = quarter
                self.update_address_label()
                self.update_address_popup_display()
                self.show_success(f"Квартал установлен: {quarter}")
                popup.dismiss()
                self.update_address_popup_display()
            else:
                self.show_error("Номер квартала не может быть пустым!")

        save_btn.bind(on_release=save_quarter)
        cancel_btn.bind(on_release=popup.dismiss)

        popup.open()

    def show_plot_popup(self, instance):
        """Показать popup для ввода выдела"""
        content = MDBoxLayout(orientation='vertical', spacing=10, padding=10, md_bg_color=Colors.DARK_SURFACE)

        title_label = MDLabel(
            text="Введите номер выдела",
            bold=True,
            size_hint=(1, None),
            height=30,
            theme_text_color='Custom',
            text_color=Colors.TEXT_ON_DARK
        )
        content.add_widget(title_label)

        self.plot_input = MDTextField(
            hint_text="Номер выдела",
            mode="outlined",
            size_hint=(1, None),
            height=40,
            line_color_focus=Colors.GREEN,
            text=self.current_plot
        )
        self.plot_input.bind(text=self.update_plot_display)
        content.add_widget(self.plot_input)

        btn_layout = MDBoxLayout(orientation='horizontal', spacing=10, size_hint=(1, None), height=50, md_bg_color=Colors.DARK_SURFACE)
        save_btn = MDButton(style='filled', size_hint=(0.5, 1), height=50)
        save_btn.add_widget(MDButtonText(text='Сохранить', theme_text_color='Custom', text_color=Colors.GREEN))
        cancel_btn = MDButton(style='filled', size_hint=(0.5, 1), height=50)
        cancel_btn.add_widget(MDButtonText(text='Отмена', theme_text_color='Custom', text_color=[1,0.3,0.3,1]))
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3],
            size_hint=(0.6, 0.5)
        )

        def save_plot(btn):
            plot = self.plot_input.text.strip()
            if plot:
                self.current_plot = plot
                self.project_data['address']['plot'] = plot
                self.update_address_label()
                self.show_success(f"Выдел установлен: {plot}")
                popup.dismiss()
                self.update_address_popup_display()
            else:
                self.show_error("Номер выдела не может быть пустым!")

        save_btn.bind(on_release=save_plot)
        cancel_btn.bind(on_release=popup.dismiss)

        popup.open()

    def show_forestry_popup(self, instance):
        """Показать popup для ввода лесничества"""
        content = MDBoxLayout(orientation='vertical', spacing=10, padding=10, md_bg_color=Colors.DARK_SURFACE)

        title_label = MDLabel(
            text="Введите название лесничества",
            bold=True,
            size_hint=(1, None),
            height=30,
            theme_text_color='Custom',
            text_color=Colors.GREEN
        )
        content.add_widget(title_label)

        # Поле для лесничества
        forestry_label = MDLabel(
            text="Лесничество:",
            size_hint=(1, None),
            height=25,
            theme_text_color='Custom',
            text_color=Colors.TEXT_ON_DARK
        )
        content.add_widget(forestry_label)

        self.forestry_input = MDTextField(
            hint_text="Название лесничества",
            mode="outlined",
            size_hint=(1, None),
            height=40,
            line_color_focus=Colors.GREEN,
            text=self.current_forestry
        )
        self.forestry_input.bind(text=lambda instance, value: self.update_forestry_display(value))
        content.add_widget(self.forestry_input)

        # Поле для участкового лесничества
        district_forestry_label = MDLabel(
            text="Участковое лесничество:",
            size_hint=(1, None),
            height=25,
            theme_text_color='Custom',
            text_color=Colors.TEXT_ON_DARK
        )
        content.add_widget(district_forestry_label)

        self.district_forestry_input = MDTextField(
            hint_text="Название участкового лесничества",
            mode="outlined",
            size_hint=(1, None),
            height=40,
            line_color_focus=Colors.GREEN,
            text=getattr(self, 'current_district_forestry', '')
        )
        self.district_forestry_input.bind(text=lambda instance, value: self.update_district_forestry_display(value))
        content.add_widget(self.district_forestry_input)

        btn_layout = MDBoxLayout(orientation='horizontal', spacing=10, size_hint=(1, None), height=50, md_bg_color=Colors.DARK_SURFACE)
        save_btn = MDButton(style='filled', size_hint=(0.5, 1), height=50)
        save_btn.add_widget(MDButtonText(text='Сохранить', theme_text_color='Custom', text_color=Colors.GREEN))
        cancel_btn = MDButton(style='filled', size_hint=(0.5, 1), height=50)
        cancel_btn.add_widget(MDButtonText(text='Отмена', theme_text_color='Custom', text_color=[1,0.3,0.3,1]))
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3],
            size_hint=(0.6, 0.7)
        )

        def save_forestry(btn):
            forestry = self.forestry_input.text.strip()
            district_forestry = self.district_forestry_input.text.strip()
            if forestry:
                self.current_forestry = forestry
                self.current_district_forestry = district_forestry
                self.project_data['address']['forestry'] = forestry
                self.project_data['address']['district_forestry'] = district_forestry
                self.update_address_label()
                self.show_success(f"Лесничество установлено: {forestry}" + (f", участковое: {district_forestry}" if district_forestry else ""))
                popup.dismiss()
                self.update_address_popup_display()
            else:
                self.show_error("Название лесничества не может быть пустым!")

        save_btn.bind(on_release=save_forestry)
        cancel_btn.bind(on_release=popup.dismiss)

        popup.open()

    def parse_care_subject_density(self, care_text):
        """Парсит предмет ухода и возвращает оставляемую густоту на гектар"""
        if not care_text:
            return 0

        care_text = care_text.strip().upper()

        # Регулярное выражение для поиска чисел и букв
        # Примеры: "3С", "2Б1С", "1Е0.5С" и т.д.
        matches = re.findall(r'(\d+(?:\.\d+)?)([А-ЯA-Z]+)', care_text)

        if not matches:
            return 0

        total_density = 0
        for number_str, breed_code in matches:
            try:
                density = float(number_str)
                total_density += density
            except ValueError:
                continue

        # Предмет ухода показывает сколько деревьев оставить на гектар
        # Например, "3С" значит оставить 3000 сосен на гектар
        return total_density * 1000  # Умножаем на 1000, так как числа обычно означают тысячи деревьев

    def parse_care_subject_by_breeds(self, care_text):
        """Парсит предмет ухода и возвращает словарь {порода: густота в тыс. шт/га}"""
        if not care_text:
            return {}

        care_text = care_text.strip().upper()

        matches = re.findall(r'(\d+(?:\.\d+)?)([А-ЯA-Z]+)', care_text)

        if not matches:
            return {}

        breed_densities = {}
        for number_str, breed_code in matches:
            try:
                density = float(number_str)
                if breed_code not in breed_densities:
                    breed_densities[breed_code] = 0
                breed_densities[breed_code] += density
            except ValueError:
                continue

        return breed_densities

    def _show_open_project_popup(self, output_file, stdout_text):
        """Спросить пользователя: открыть проект в Word или Excel"""
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)
        content.add_widget(MDLabel(
            text="Проект ухода успешно создан!",
            font_style='Headline', role='medium',
            theme_text_color='Custom', text_color=Colors.GREEN,
            size_hint_y=None, height=dp(44), halign='center'))
        content.add_widget(MDLabel(
            text="Открыть проект?",
            theme_text_color='Custom', text_color=[1,1,1,1],
            size_hint_y=None, height=dp(30), halign='center'))

        btn_layout = MDBoxLayout(orientation='horizontal', spacing=Spacing.MD,
                                 adaptive_height=True, pos_hint={'center_x': 0.5})

        popup = Popup(title="", content=content, size_hint=(0.7, None), height=dp(220),
                      separator_height=0, background_color=[0,0,0,0.3], overlay_color=[0,0,0,0.3])

        word_btn = MDButton(style='filled', md_bg_color=Colors.BTN_SUCCESS,
                            size_hint=(0.3, None), height=dp(48))
        word_btn.add_widget(MDButtonText(text='Word'))
        word_btn.bind(on_release=lambda x: self._open_project_file(output_file, 'word', popup))

        excel_btn = MDButton(style='filled', md_bg_color=Colors.BTN_WARNING,
                             size_hint=(0.3, None), height=dp(48))
        excel_btn.add_widget(MDButtonText(text='Excel'))
        excel_btn.bind(on_release=lambda x: self._open_project_file(output_file, 'excel', popup))

        no_btn = MDButton(style='outlined', size_hint=(0.3, None), height=dp(48),
                          line_color=Colors.DANGER)
        no_btn.add_widget(MDButtonText(text='Нет', theme_text_color='Custom', text_color=Colors.DANGER))
        no_btn.bind(on_release=lambda x: popup.dismiss())

        btn_layout.add_widget(word_btn)
        btn_layout.add_widget(excel_btn)
        btn_layout.add_widget(no_btn)
        content.add_widget(btn_layout)

        popup.open()

    def _open_project_file(self, output_file, file_type, popup):
        """Открыть сгенерированный файл проекта"""
        try:
            popup.dismiss()
            if file_type == 'word' and output_file and os.path.exists(output_file):
                os.startfile(output_file)
            elif file_type == 'excel':
                self.save_to_excel_without_dialog()
        except Exception as e:
            self.show_error(f"Ошибка при открытии файла: {str(e)}")

    def generate_care_project(self, instance):
        """Генерирует проект ухода в Word документе с использованием данных из меню Итого"""
        try:
            # Сохраняем текущую страницу перед генерацией отчета
            if not self.save_current_page():
                self.show_error("Не удалось сохранить текущую страницу!")
                return

            # Извлекаем Тип Леса из данных участка (столбец 5 в page_data)
            forest_type = 'Смешанный лес'  # По умолчанию
            for page_num, page_rows in self.page_data.items():
                for row in page_rows:
                    if len(row) >= 6 and row[5]:  # Столбец 5 - Тип Леса
                        forest_type = str(row[5]).strip()
                        break
                if forest_type != 'Смешанный лес':
                    break

            # Собираем данные из адресной строки и project_data
            address_data = {
                'quarter': str(self.project_data['address'].get('quarter', '')),
                'plot': str(self.project_data['address'].get('plot', '')),
                'section': str(self.current_section or ''),
                'forestry': str(self.project_data['address'].get('forestry', '')),
                'district_forestry': str(self.project_data['address'].get('district_forestry', '')),
                'plot_area': str(self.project_data['address'].get('plot_area', '')),
                'radius': str(self.project_data['address'].get('radius', '5.64')),
                'target_purpose': 'Эксплуатационные леса',
                'forest_type': forest_type  # Извлекаем из данных участка
            }

            # Добавляем данные деталей проекта
            details_data = {
                'care_queue': str(self.project_data['details'].get('care_queue', '')),
                'characteristics': str(self.project_data['details'].get('characteristics', '')),
                'care_date': str(self.project_data['details'].get('care_date', '')),
                'technology': str(self.project_data['details'].get('technology', '')),
                'forest_purpose': str(self.project_data['details'].get('forest_purpose', ''))
            }

            # Получаем итоговые данные из меню Итого (рассчитанные данные)
            total_data = self.get_total_data_from_db()

            # Добавляем данные деталей в total_data для передачи в fill_word_document.py
            total_data.update({
                'care_queue': details_data['care_queue'],
                'characteristics': details_data['characteristics'],
                'care_date': details_data['care_date'],
                'technology': details_data['technology'],
                'forest_purpose': details_data['forest_purpose'],
                'address_data': address_data
            })

            # Создаем временный JSON файл с данными для скрипта
            import tempfile
            import subprocess

            temp_data = {
                'address_data': address_data,
                'total_data': total_data
            }

            # Сохраняем данные во временный файл
            with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False, encoding='utf-8') as f:
                json.dump(temp_data, f, ensure_ascii=False, indent=2)
                temp_file = f.name
            
            # ОТЛАДОЧНЫЙ ВЫВОД
            print("=" * 60)
            print("ОТЛАДКА: Данные для проекта ухода")
            print("=" * 60)
            print(f"Адрес: Квартал {address_data.get('quarter')}, Выдел {address_data.get('plot')}")
            print(f"Интенсивность: {total_data.get('intensity')}")
            print(f"Средний диаметр: {total_data.get('avg_diameter')}")
            print(f"Породы: {len(total_data.get('breeds', []))}")
            for breed in total_data.get('breeds', []):
                print(f"  - {breed.get('name')}: d={breed.get('diameter')}, h={breed.get('height')}, density={breed.get('density')}")
            print("=" * 60)

            # Вызываем скрипт fill_our_template.py с параметром
            script_path = os.path.join(os.path.dirname(__file__), 'fill_our_template.py')
            result = subprocess.run([
                sys.executable, script_path, '--data-file', temp_file
            ], capture_output=True, text=False)

            # Удаляем временный файл
            try:
                os.unlink(temp_file)
            except:
                pass

            # Декодируем вывод с обработкой ошибок кодировки
            def decode_output(output_bytes):
                try:
                    if output_bytes is None:
                        return ""
                    if not output_bytes:
                        return ""
                    return output_bytes.decode('utf-8')
                except UnicodeDecodeError:
                    try:
                        return output_bytes.decode('cp1251')
                    except UnicodeDecodeError:
                        return output_bytes.decode('utf-8', errors='replace')
                except Exception:
                    return ""

            if result.returncode == 0:
                stdout_text = decode_output(result.stdout)
                # Извлекаем путь к файлу из вывода
                output_file = ''
                for line in stdout_text.split('\n'):
                    if '[OK]' in line:
                        output_file = line.split(':', 1)[-1].strip() if ':' in line else ''
                        break
                self._show_open_project_popup(output_file, stdout_text)
            else:
                stderr_text = decode_output(result.stderr)
                self.show_error(f"Ошибка при создании проекта ухода:\n{stderr_text}")

        except Exception as e:
            self.show_error(f"Ошибка при генерации проекта ухода: {str(e)}")

    def get_total_data_from_db(self):
        """Получает итоговые данные из рассчитанных данных меню Итого"""
        try:
            # Используем рассчитанные данные из меню Итого вместо данных из БД
            # Получаем данные аналогично методу show_total_summary_popup

            default_radius = float(self.current_radius) if self.current_radius else 5.64
            plot_area_ha = 3.14159 * (default_radius ** 2) / 10000

            # Словарь для сбора данных по породам
            breeds_data = {}

            # Обрабатываем все страницы
            for page_num, page_rows in self.page_data.items():
                for row in page_rows:
                    if len(row) < 4:
                        continue

                    # Столбец "Порода" в row[3]
                    breeds_text = row[3]
                    if not breeds_text:
                        continue

                    try:
                        breeds_list = json.loads(breeds_text) if isinstance(breeds_text, str) else []
                    except json.JSONDecodeError:
                        continue

                    for breed_info in breeds_list:
                        if not isinstance(breed_info, dict):
                            continue

                        breed_name = breed_info.get('name', '').strip()
                        if not breed_name:
                            continue

                        breed_type = breed_info.get('type', 'deciduous')
                        density = 0
                        height = None
                        age = None

                        # Расчет густоты и высоты в зависимости от типа породы
                        if breed_type == 'coniferous':
                            do_05 = breed_info.get('do_05', 0)
                            _05_15 = breed_info.get('05_15', 0)
                            bolee_15 = breed_info.get('bolee_15', 0)
                            density = (do_05 + _05_15 + bolee_15) / plot_area_ha if plot_area_ha > 0 else 0

                            # Для хвойных пород определяем высоту по градациям или среднюю
                            if any([do_05, _05_15, bolee_15]):
                                # Высота определяется по градациям
                                if bolee_15 > 0:
                                    height = 2.0  # >1.5m
                                elif _05_15 > 0:
                                    height = 1.0  # 0.5-1.5m
                                elif do_05 > 0:
                                    height = 0.3  # до 0.5m
                                else:
                                    height = 0.0
                            else:
                                height = breed_info.get('height', 0) or 0
                        else:
                            # Для лиственных пород - обычная плотность и средняя высота
                            density_value = breed_info.get('density', 0)
                            density = density_value / plot_area_ha if plot_area_ha > 0 else 0
                            height = breed_info.get('height', 0) or 0

                        age = breed_info.get('age', 0) or 0
                        diameter = breed_info.get('diameter', 0) or 0

                        # Сбор данных по породе
                        if breed_name not in breeds_data:
                            breeds_data[breed_name] = {
                                'type': breed_type,
                                'plots': [],
                                'coniferous_zones': {'do_05': 0, '05_15': 0, 'bolee_15': 0} if breed_type == 'coniferous' else None,
                                'diameters': []
                            }

                        # Добавляем данные
                        plot_data = {
                            'density': density,
                            'height': height,
                            'age': age,
                            'diameter': diameter  # ✅ ДОБАВЛЕНО: сохраняем диаметр в plot_data
                        }

                        if breed_type == 'coniferous':
                            plot_data.update({
                                'do_05_density': do_05 / plot_area_ha if plot_area_ha > 0 else 0,
                                '05_15_density': _05_15 / plot_area_ha if plot_area_ha > 0 else 0,
                                'bolee_15_density': bolee_15 / plot_area_ha if plot_area_ha > 0 else 0
                            })

                        breeds_data[breed_name]['plots'].append(plot_data)
                        breeds_data[breed_name]['diameters'].append(diameter)

                        if breed_type == 'coniferous':
                            breeds_data[breed_name]['coniferous_zones']['do_05'] += plot_data['do_05_density']
                            breeds_data[breed_name]['coniferous_zones']['05_15'] += plot_data['05_15_density']
                            breeds_data[breed_name]['coniferous_zones']['bolee_15'] += plot_data['bolee_15_density']

            # Расчет коэффициента состава на основе СРЕДНЕЙ густоты пород
            total_densities = {}
            total_density_all = 0  # Общая густота всех пород (сумма средних)
            
            for breed_name, data in breeds_data.items():
                if data['plots']:
                    # Рассчитываем СРЕДНЮЮ густоту породы (сумма плотностей / кол-во площадок)
                    if data['plots'][0].get('type') == 'coniferous':
                        # Для хвойных суммируем густоту по градациям и делим на кол-во площадок
                        total_density = sum(
                            (p.get('do_05_density', 0) + p.get('05_15_density', 0) + p.get('bolee_15_density', 0))
                            for p in data['plots']
                        ) / len(data['plots'])
                    else:
                        # Для лиственных обычная средняя густота
                        total_density = sum(p.get('density', 0) for p in data['plots']) / len(data['plots'])
                    
                    total_density_all += total_density
                    if total_density > 0:
                        total_densities[breed_name] = total_density

            # Расчет коэффициентов состава
            composition_text = ""
            if total_densities:
                composition_parts = []

                # Сортируем по убыванию плотности
                for breed_name, density in sorted(total_densities.items(), key=lambda x: x[1], reverse=True):
                    breed_letter = self.get_breed_letter(breed_name)
                    # Просто добавляем букву, коэффициент будет рассчитан ниже
                    composition_parts.append(f"0{breed_letter}")

                # Корректировка чтобы сумма равнялась 10 (метод наибольшего остатка)
                # Сначала рассчитываем точные коэффициенты
                exact_coeffs = []
                for breed_name, density in sorted(total_densities.items(), key=lambda x: x[1], reverse=True):
                    exact_coeff = (density / total_density_all * 10) if total_density_all > 0 else 1
                    exact_coeffs.append(exact_coeff)
                
                # Округляем вниз
                coeffs_floor = [int(coeff) for coeff in exact_coeffs]
                # Считаем остатки
                remainders = [(exact_coeffs[i] - coeffs_floor[i], i) for i in range(len(exact_coeffs))]
                # Сортируем по убыванию остатков
                remainders.sort(reverse=True)
                
                # Распределяем единицы начиная с наибольших остатков
                coeffs = coeffs_floor[:]
                total = sum(coeffs)
                i = 0
                while total < 10 and i < len(remainders):
                    idx = remainders[i][1]
                    coeffs[idx] += 1
                    total += 1
                    i += 1
                
                # Обновляем composition_parts
                sorted_breeds = sorted(total_densities.items(), key=lambda x: x[1], reverse=True)
                composition_parts = []
                for i, (breed_name, _) in enumerate(sorted_breeds):
                    if i < len(coeffs):
                        breed_letter = self.get_breed_letter(breed_name)
                        composition_parts.append(f"{coeffs[i]}{breed_letter}")

                composition_text = ''.join(composition_parts) + "Др"

            # Расчет предмета ухода и интенсивности
            care_data = []
            total_density_all_plots = 0
            total_remaining_density = 0
            plot_count_with_care = 0

            for page_num, page_rows in self.page_data.items():
                for row in page_rows:
                    if len(row) >= 4 and row[3]:
                        plot_density = 0
                        breeds_text = row[3]
                        if breeds_text:
                            try:
                                breeds_list = json.loads(breeds_text) if isinstance(breeds_text, str) else []
                                for breed_info in breeds_list:
                                    if isinstance(breed_info, dict):
                                        if breed_info.get('type') == 'coniferous':
                                            do_05 = breed_info.get('do_05', 0)
                                            _05_15 = breed_info.get('05_15', 0)
                                            bolee_15 = breed_info.get('bolee_15', 0)
                                            plot_density += (do_05 + _05_15 + bolee_15) / plot_area_ha if plot_area_ha > 0 else 0
                                        else:
                                            density = breed_info.get('density', 0)
                                            plot_density += density / plot_area_ha if plot_area_ha > 0 else 0
                            except (json.JSONDecodeError, TypeError):
                                pass

                        if plot_density > 0:
                            total_density_all_plots += plot_density

                    if len(row) >= 4 and row[2]:
                        care_text = row[2].strip()
                        if care_text:
                            plot_density = 0
                            breeds_text = row[3]
                            if breeds_text:
                                try:
                                    breeds_list = json.loads(breeds_text) if isinstance(breeds_text, str) else []
                                    for breed_info in breeds_list:
                                        if isinstance(breed_info, dict):
                                            if breed_info.get('type') == 'coniferous':
                                                do_05 = breed_info.get('do_05', 0)
                                                _05_15 = breed_info.get('05_15', 0)
                                                bolee_15 = breed_info.get('bolee_15', 0)
                                                plot_density += (do_05 + _05_15 + bolee_15) / plot_area_ha if plot_area_ha > 0 else 0
                                            else:
                                                density = breed_info.get('density', 0)
                                                plot_density += density / plot_area_ha if plot_area_ha > 0 else 0
                                except (json.JSONDecodeError, TypeError):
                                    pass

                            if plot_density > 0:
                                remaining_density = self.parse_care_subject_density(care_text)
                                if remaining_density > 0:
                                    care_data.append({
                                        'care_text': care_text,
                                        'plot_density': plot_density,
                                        'remaining_density': remaining_density
                                    })
                                    total_remaining_density += remaining_density
                                    plot_count_with_care += 1

            # Расчет среднего предмета ухода
            care_subject = ""
            intensity = 25.0  # По умолчанию

            if care_data:
                care_breed_totals = {}
                care_plot_count = 0

                for item in care_data:
                    care_text = item['care_text']
                    breed_densities = self.parse_care_subject_by_breeds(care_text)
                    for breed, density in breed_densities.items():
                        if breed not in care_breed_totals:
                            care_breed_totals[breed] = 0
                        care_breed_totals[breed] += density
                    care_plot_count += 1

                if care_breed_totals and care_plot_count > 0:
                    avg_care_parts = []
                    for breed, total_density in sorted(care_breed_totals.items()):
                        avg_density = total_density / care_plot_count
                        avg_care_parts.append(f"{avg_density * 1000:.0f}шт/га{breed}")
                    care_subject = ''.join(avg_care_parts)

                    # Расчет интенсивности
                    if plot_count_with_care > 0:
                        avg_remaining_density = total_remaining_density / plot_count_with_care
                        # Используем среднюю густоту по площадкам для расчёта интенсивности
                        num_plots = len([row for page in self.page_data.values() for row in page if any(cell for cell in row[:3] if cell)])
                        avg_overall_density_for_intensity = total_density_all_plots / num_plots if num_plots > 0 else 0

                        if avg_overall_density_for_intensity > 0:
                            intensity = ((avg_overall_density_for_intensity - avg_remaining_density) / avg_overall_density_for_intensity) * 100
                            print(f"[DEBUG] Интенсивность: {avg_overall_density_for_intensity:.1f} - {avg_remaining_density:.0f} / {avg_overall_density_for_intensity:.1f} = {intensity:.1f}%")

            # Расчет средних значений по участку (ОБЩАЯ густота, высота, диаметр, возраст)
            # ВАЖНО: Считаем по ПЛОЩАДКАМ, а не по породам!
            # 1. Для каждой площадки: суммируем густоту по всем породам
            # 2. Для каждой площадки: средние высота/диаметр/возраст по породам
            # 3. Среднее по всем площадкам
            
            # Сначала собираем данные по площадкам из исходных строк
            plot_data_list = []  # Список данных по каждой площадке
            
            for page_num, page_rows in self.page_data.items():
                for row_idx, row in enumerate(page_rows):
                    if len(row) < 4:
                        continue
                    
                    breeds_text = row[3]
                    if not breeds_text:
                        continue
                    
                    # Данные по этой площадке
                    plot_total_density = 0
                    plot_height_sum = 0
                    plot_height_count = 0
                    plot_diameter_sum = 0
                    plot_diameter_count = 0
                    plot_age_sum = 0
                    plot_age_count = 0
                    
                    try:
                        breeds_list = json.loads(breeds_text) if isinstance(breeds_text, str) else []
                        
                        for breed_info in breeds_list:
                            if not isinstance(breed_info, dict):
                                continue
                            
                            breed_type = breed_info.get('type', 'deciduous')
                            
                            # Расчёт густоты
                            if breed_type == 'coniferous':
                                do_05 = breed_info.get('do_05', 0)
                                _05_15 = breed_info.get('05_15', 0)
                                bolee_15 = breed_info.get('bolee_15', 0)
                                total_trees = do_05 + _05_15 + bolee_15
                                density = total_trees / plot_area_ha if plot_area_ha > 0 else 0
                                
                                # Высота для хвойных по градациям
                                if bolee_15 > 0:
                                    height = 2.0
                                elif _05_15 > 0:
                                    height = 1.0
                                elif do_05 > 0:
                                    height = 0.3
                                else:
                                    height = breed_info.get('height', 0) or 0
                            else:
                                density_value = breed_info.get('density', 0)
                                density = density_value / plot_area_ha if plot_area_ha > 0 else 0
                                height = breed_info.get('height', 0) or 0
                            
                            diameter = breed_info.get('diameter', 0) or 0
                            age = breed_info.get('age', 0) or 0
                            
                            # Суммируем по площадке
                            plot_total_density += density
                            if height > 0:
                                plot_height_sum += height
                                plot_height_count += 1
                            if diameter > 0:
                                plot_diameter_sum += diameter
                                plot_diameter_count += 1
                            if age > 0:
                                plot_age_sum += age
                                plot_age_count += 1
                        
                        # Сохраняем данные площадки
                        plot_data_list.append({
                            'density': plot_total_density,
                            'height': plot_height_sum / plot_height_count if plot_height_count > 0 else 0,
                            'diameter': plot_diameter_sum / plot_diameter_count if plot_diameter_count > 0 else 0,
                            'age': plot_age_sum / plot_age_count if plot_age_count > 0 else 0
                        })
                        
                    except (json.JSONDecodeError, TypeError) as e:
                        print(f"[WARNING] Ошибка обработки строки: {e}")
                        continue
            
            # Рассчитываем средние по ВСЕМ ПЛОЩАДКАМ
            num_plots = len(plot_data_list)
            
            if num_plots > 0:
                avg_overall_density = sum(p['density'] for p in plot_data_list) / num_plots
                avg_overall_height = sum(p['height'] for p in plot_data_list) / num_plots
                avg_overall_diameter = sum(p['diameter'] for p in plot_data_list) / num_plots
                avg_overall_age = sum(p['age'] for p in plot_data_list) / num_plots
            else:
                avg_overall_density = 0
                avg_overall_height = 0
                avg_overall_diameter = 0
                avg_overall_age = 0

            # Формируем итоговые данные
            total_data = {
                'page_number': self.current_page,
                'section_name': self.current_section or '',
                'total_composition': composition_text,
                'avg_age': avg_overall_age,
                'avg_density': avg_overall_density,
                'avg_height': avg_overall_height,
                'avg_diameter': avg_overall_diameter,  # РАССЧИТЫВАЕМ!
                'total_plots': len([row for page in self.page_data.values() for row in page if any(cell for cell in row[:3] if cell)]),
                'composition': composition_text,
                'care_subject': care_subject,
                'intensity': intensity if intensity > 0 else 25,  # Если не рассчитана, по умолчанию 25%
                'breeds': []
            }

            # Добавляем данные по породам
            for breed_name, data in breeds_data.items():
                if data['plots']:
                    avg_density = sum(p['density'] for p in data['plots']) / len(data['plots'])
                    avg_height = sum(p['height'] for p in data['plots'] if p['height'] > 0) / len([p for p in data['plots'] if p['height'] > 0]) if any(p['height'] > 0 for p in data['plots']) else 0
                    avg_age = sum(p['age'] for p in data['plots'] if p['age'] > 0) / len([p for p in data['plots'] if p['age'] > 0]) if any(p['age'] > 0 for p in data['plots']) else 0
                    avg_diameter = sum(p.get('diameter', 0) for p in data['plots'] if p.get('diameter', 0) > 0) / len([p for p in data['plots'] if p.get('diameter', 0) > 0]) if any(p.get('diameter', 0) > 0 for p in data['plots']) else 0

                    breed_data = {
                        'name': breed_name,
                        'type': data['type'],
                        'density': avg_density,
                        'height': avg_height,
                        'age': avg_age,
                        'diameter': avg_diameter  # ДОБАВЛЯЕМ ДИАМЕТР!
                    }

                    if data['type'] == 'coniferous':
                        zones = data.get('coniferous_zones', {})
                        breed_data.update({
                            'do_05': zones.get('do_05', 0) / len(data['plots']) if data['plots'] else 0,
                            '_05_15': zones.get('05_15', 0) / len(data['plots']) if data['plots'] else 0,
                            'bolee_15': zones.get('bolee_15', 0) / len(data['plots']) if data['plots'] else 0
                        })

                    total_data['breeds'].append(breed_data)

            return total_data

        except Exception as e:
            print(f"Ошибка получения данных из меню Итого: {e}")
            import traceback
            traceback.print_exc()
            return {}

    def parse_care_subject_density(self, care_text):
        """Парсит предмет ухода и возвращает оставляемую густоту на гектар"""
        if not care_text:
            return 0

        care_text = care_text.strip().upper()

        # Регулярное выражение для поиска чисел и букв
        # Примеры: "3С", "2Б1С", "1Е0.5С" и т.д.
        matches = re.findall(r'(\d+(?:\.\d+)?)([А-ЯA-Z]+)', care_text)

        if not matches:
            return 0

        total_density = 0
        for number_str, breed_code in matches:
            try:
                density = float(number_str)
                total_density += density
            except ValueError:
                continue

        # Предмет ухода показывает сколько деревьев оставить на гектар
        # Например, "3С" значит оставить 3000 сосен на гектар
        return total_density * 1000  # Умножаем на 1000, так как числа обычно означают тысячи деревьев

    def parse_care_subject_by_breeds(self, care_text):
        """Парсит предмет ухода и возвращает словарь {порода: густота в тыс. шт/га}"""
        if not care_text:
            return {}

        care_text = care_text.strip().upper()

        matches = re.findall(r'(\d+(?:\.\d+)?)([А-ЯA-Z]+)', care_text)

        if not matches:
            return {}

        breed_densities = {}
        for number_str, breed_code in matches:
            try:
                density = float(number_str)
                if breed_code not in breed_densities:
                    breed_densities[breed_code] = 0
                breed_densities[breed_code] += density
            except ValueError:
                continue

        return breed_densities

    def _get_current_plot_area_input(self):
        """Получить текущее значение площади участка"""
        # If stored in instance variable
        if hasattr(self, 'plot_area_input') and self.plot_area_input:
            return self.plot_area_input
        return ''

    def get_breed_letter(self, breed_name):
        """Получение первой буквы для коэффициента состава породы"""
        breed_letters = {
            'Сосна': 'С',
            'Ель': 'Е',
            'Пихта': 'П',
            'Кедр': 'К',
            'Лиственница': 'Л',
            'Берёза': 'Б',
            'Осина': 'Ос',
            'Ольха чёрная': 'ОЧ',
            'Ольха серая': 'ОС',
            'Ива': 'И',
            'Ива кустарниковая': 'ИК'
        }

        for full_name, letter in breed_letters.items():
            if full_name.lower() in breed_name.lower():
                return letter

        # Возвращаем первую букву имени породы, если не найдено
        return breed_name[0].upper() if breed_name else 'Н'

    def show_edit_plots_popup(self, instance):
        """Показать popup со списком площадок для управления породами"""
        content = MDBoxLayout(orientation='vertical', spacing=15, padding=15, md_bg_color=Colors.DARK_SURFACE)

        title_label = MDLabel(
            text="Площадки - Управление породами",
            font_size='20sp',
            bold=True,
            theme_text_color='Custom',
            text_color=Colors.GREEN,
            size_hint=(1, None),
            height=50
        )
        content.add_widget(title_label)

        # ScrollView для списка площадок
        scroll = ScrollView(size_hint=(1, 1))
        plots_layout = MDGridLayout(cols=1, spacing=10, size_hint_y=None, md_bg_color=Colors.DARK_SURFACE)
        plots_layout.bind(minimum_height=plots_layout.setter('height'))

        # Получаем все площадки из page_data
        all_plots = []
        for page_num in sorted(self.page_data.keys()):
            page_rows = self.page_data[page_num]
            for row_idx, row in enumerate(page_rows):
                if len(row) >= 6 and any(row[i] for i in range(6)):  # Есть данные
                    all_plots.append({
                        'page': page_num,
                        'row': row_idx,
                        'data': row
                    })

        if not all_plots:
            no_plots_label = MDLabel(
                text="Нет сохраненных площадок.\nДобавьте площадки через меню 'Файл' -> 'Создать'",
                font_size='16sp',
                theme_text_color='Custom',
                text_color=[0.5,0.5,0.5,1],
                size_hint=(1, None),
                height=100,
                halign='center',
                valign='middle'
            )
            no_plots_label.bind(size=lambda *args: setattr(no_plots_label, 'text_size', (no_plots_label.width, None)))
            plots_layout.add_widget(no_plots_label)
        else:
            # Создаем боксы для каждой площадки
            for plot_info in all_plots:
                plot_num = plot_info['row'] + 1
                gps_point = plot_info['data'][1] if len(plot_info['data']) > 1 else ''
                predmet_uhoda = plot_info['data'][2] if len(plot_info['data']) > 2 else ''
                breeds_data = plot_info['data'][3] if len(plot_info['data']) > 3 else ''
                forest_type = plot_info['data'][5] if len(plot_info['data']) > 5 else ''

                # Парсим данные о породах
                breeds_list = self.parse_breeds_data(breeds_data) if breeds_data else []
                breeds_count = len(breeds_list)

                # Создаем бокс площадки
                plot_box = MDBoxLayout(orientation='vertical', spacing=5, size_hint_y=None, height=120 if breeds_count > 0 else 80, md_bg_color=Colors.CARD_BG)

                # Заголовок площадки
                plot_header = MDBoxLayout(orientation='horizontal', size_hint_y=None, height=35, padding=[10, 5], md_bg_color=Colors.CARD_BG)
                plot_num_label = MDLabel(
                    text=f"Площадка №{plot_num}",
                    font_size='16sp',
                    bold=True,
                    theme_text_color='Custom',
                    text_color=Colors.GREEN,
                    size_hint=(0.3, 1),
                    halign='left'
                )
                plot_num_label.bind(size=lambda *args: setattr(plot_num_label, 'text_size', (plot_num_label.width, None)))
                plot_header.add_widget(plot_num_label)

                if gps_point:
                    gps_label = MDLabel(
                        text=f"GPS: {gps_point}",
                        font_size='14sp',
                        theme_text_color='Custom',
                        text_color=[0.7,0.7,0.7,1],
                        size_hint=(0.35, 1),
                        halign='left'
                    )
                    gps_label.bind(size=lambda *args: setattr(gps_label, 'text_size', (gps_label.width, None)))
                    plot_header.add_widget(gps_label)

                if forest_type:
                    type_label = MDLabel(
                        text=f"Тип: {forest_type}",
                        font_size='14sp',
                        theme_text_color='Custom',
                        text_color=[0.7,0.7,0.7,1],
                        size_hint=(0.35, 1),
                        halign='left'
                    )
                    type_label.bind(size=lambda *args: setattr(type_label, 'text_size', (type_label.width, None)))
                    plot_header.add_widget(type_label)

                plot_box.add_widget(plot_header)

                # Информация о породах
                if breeds_count > 0:
                    breeds_info_text = f"Пород: {breeds_count} | "
                    breed_names = [b.get('name', 'Неизвестная') for b in breeds_list[:3]]
                    breeds_info_text += ", ".join(breed_names)
                    if breeds_count > 3:
                        breeds_info_text += f" и еще {breeds_count - 3}"

                    breeds_label = MDLabel(
                        text=breeds_info_text,
                        font_size='13sp',
                        theme_text_color='Custom',
                        text_color=[0.7,0.7,0.7,1],
                        size_hint=(1, None),
                        height=30,
                        halign='left'
                    )
                    breeds_label.bind(size=lambda *args: setattr(breeds_label, 'text_size', (breeds_label.width, None)))
                    plot_box.add_widget(breeds_label)

                # Кнопка управления
                manage_btn = MDButton(style='filled', size_hint=(1, None), height=40)
                manage_btn.add_widget(MDButtonText(text='Редакция', theme_text_color='Custom', text_color=Colors.GREEN))
                manage_btn.bind(on_release=lambda x, p=plot_info: self.show_plot_breed_management(p))
                plot_box.add_widget(manage_btn)

                plots_layout.add_widget(plot_box)

        scroll.add_widget(plots_layout)
        content.add_widget(scroll)

        # Кнопка закрытия
        close_btn = MDButton(style='filled', size_hint=(1, None), height=60)
        close_btn.add_widget(MDButtonText(text='Закрыть', theme_text_color='Custom', text_color=[1,0.3,0.3,1]))
        content.add_widget(close_btn)

        popup = Popup(
            title="",
            content=content,
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3],
            size_hint=(0.9, 0.9)
        )

        close_btn.bind(on_release=popup.dismiss)
        popup.open()

    def show_plot_breed_management(self, plot_info):
        """Показать меню управления породами для конкретной площадки"""
        page_num = plot_info['page']
        row_idx = plot_info['row']
        plot_data = plot_info['data']
        plot_num = row_idx + 1

        content = MDBoxLayout(orientation='vertical', spacing=15, padding=15, md_bg_color=Colors.DARK_SURFACE)

        # Заголовок
        title_label = MDLabel(
            text=f"Управление породами - Площадка №{plot_num}",
            font_size='20sp',
            bold=True,
            theme_text_color='Custom',
            text_color=Colors.GREEN,
            size_hint=(1, None),
            height=50
        )
        content.add_widget(title_label)

        # ScrollView для содержимого
        scroll = ScrollView(size_hint=(1, 1))
        scroll_content = MDGridLayout(cols=1, spacing=15, size_hint_y=None, md_bg_color=Colors.DARK_SURFACE)
        scroll_content.bind(minimum_height=scroll_content.setter('height'))

        # Информация о площадке
        info_box = MDBoxLayout(orientation='vertical', spacing=8, size_hint=(1, None), height=150, padding=[15, 15], md_bg_color=Colors.CARD_BG)

        info_title = MDLabel(
            text='Данные площадки:',
            font_size='16sp',
            bold=True,
            theme_text_color='Custom',
            text_color=Colors.TEXT_ON_DARK,
            size_hint=(1, None),
            height=30,
            halign='left'
        )
        info_box.add_widget(info_title)

        gps_point = plot_data[1] if len(plot_data) > 1 else ''
        predmet_uhoda = plot_data[2] if len(plot_data) > 2 else ''
        breeds_data = plot_data[3] if len(plot_data) > 3 else ''
        forest_type = plot_data[5] if len(plot_data) > 5 else ''

        info_text = f"GPS точка: {gps_point or 'Не указана'}\n"
        info_text += f"Предмет ухода: {predmet_uhoda or 'Не указан'}\n"
        info_text += f"Тип леса: {forest_type or 'Не указан'}"

        info_details = MDLabel(
            text=info_text,
            font_size='14sp',
            theme_text_color='Custom',
            text_color=[0.7,0.7,0.7,1],
            size_hint=(1, None),
            height=80,
            halign='left',
            valign='top'
        )
        info_details.bind(size=lambda *args: setattr(info_details, 'text_size', (info_details.width, None)))
        info_box.add_widget(info_details)

        # Кнопка "Редакция" для данных площадки
        edit_plot_btn = MDButton(style='filled', size_hint=(1, None), height=45)
        edit_plot_btn.add_widget(MDButtonText(text='Редакция', theme_text_color='Custom', text_color=Colors.GREEN, bold=True))
        edit_plot_btn.bind(on_release=lambda x, p=page_num, r=row_idx: self.edit_plot_data_from_management(p, r))
        info_box.add_widget(edit_plot_btn)

        scroll_content.add_widget(info_box)

        # Список пород с боксами
        breeds_box = MDBoxLayout(orientation='vertical', spacing=10, size_hint=(1, None), height=400, padding=[15, 15], md_bg_color=Colors.CARD_BG)

        breeds_title = MDLabel(
            text='Сохраненные породы:',
            font_size='16sp',
            bold=True,
            theme_text_color='Custom',
            text_color=Colors.TEXT_ON_DARK,
            size_hint=(1, None),
            height=30,
            halign='left'
        )
        breeds_box.add_widget(breeds_title)

        # ScrollView для списка пород
        breeds_scroll = ScrollView(size_hint=(1, 1))
        breeds_list_layout = MDGridLayout(cols=1, spacing=8, size_hint_y=None, md_bg_color=Colors.CARD_BG)
        breeds_list_layout.bind(minimum_height=breeds_list_layout.setter('height'))

        # Получаем породы из данных
        breeds_list = self.parse_breeds_data(breeds_data) if breeds_data else []

        if not breeds_list:
            no_breeds_label = MDLabel(
                text="Нет сохраненных пород для этой площадки",
                font_size='14sp',
                theme_text_color='Custom',
                text_color=[0.5,0.5,0.5,1],
                size_hint=(1, None),
                height=50,
                halign='center'
            )
            breeds_list_layout.add_widget(no_breeds_label)
        else:
            # Создаем боксы для каждой породы
            for i, breed_info in enumerate(breeds_list):
                breed_name = breed_info.get('name', 'Неизвестная')
                breed_type = breed_info.get('type', 'unknown')
                breed_type_text = 'Хвойная' if breed_type == 'coniferous' else 'Лиственная'

                # Создаем бокс породы
                breed_box = MDBoxLayout(orientation='vertical', spacing=5, size_hint_y=None, height=140, padding=[10, 10], md_bg_color=Colors.CARD_BG)

                # Заголовок породы
                breed_header = f"№{i+1}: {breed_name} ({breed_type_text})"
                breed_header_label = MDLabel(
                    text=breed_header,
                    font_size='15sp',
                    bold=True,
                    theme_text_color='Custom',
                    text_color=Colors.GREEN,
                    size_hint=(1, None),
                    height=30,
                    halign='left'
                )
                breed_header_label.bind(size=lambda *args: setattr(breed_header_label, 'text_size', (breed_header_label.width, None)))
                breed_box.add_widget(breed_header_label)

                # Параметры породы
                params = []
                if breed_type == 'coniferous':
                    do_05 = breed_info.get('do_05', 0)
                    _05_15 = breed_info.get('05_15', 0)
                    bolee_15 = breed_info.get('bolee_15', 0)
                    total_density = do_05 + _05_15 + bolee_15
                    if total_density > 0:
                        params.append(f"Густота: {total_density}")
                    if do_05 > 0:
                        params.append(f"До 0.5м: {do_05}")
                    if _05_15 > 0:
                        params.append(f"0.5-1.5м: {_05_15}")
                    if bolee_15 > 0:
                        params.append(f">1.5м: {bolee_15}")
                else:
                    density = breed_info.get('density', 0)
                    if density > 0:
                        params.append(f"Густота: {density}")

                height = breed_info.get('height', 0)
                if height > 0:
                    params.append(f"Высота: {height}м")

                diameter = breed_info.get('diameter', 0)
                if diameter > 0:
                    params.append(f"Диаметр: {diameter}см")

                age = breed_info.get('age', 0)
                if age > 0:
                    params.append(f"Возраст: {age}л")

                params_text = ", ".join(params) if params else "Нет параметров"
                params_label = MDLabel(
                    text=params_text,
                    font_size='13sp',
                    theme_text_color='Custom',
                    text_color=[0.7,0.7,0.7,1],
                    size_hint=(1, None),
                    height=40,
                    halign='left',
                    valign='top'
                )
                params_label.bind(size=lambda *args: setattr(params_label, 'text_size', (params_label.width, None)))
                breed_box.add_widget(params_label)

                # Кнопка Редакция
                edit_breed_btn = MDButton(style='filled', size_hint=(1, None), height=35)
                edit_breed_btn.add_widget(MDButtonText(text='Редакция', theme_text_color='Custom', text_color=Colors.GREEN))
                edit_breed_btn.bind(on_release=lambda x, p=page_num, r=row_idx, b=i: self.edit_breed_from_plot(p, r, b))
                breed_box.add_widget(edit_breed_btn)

                breeds_list_layout.add_widget(breed_box)

        breeds_scroll.add_widget(breeds_list_layout)
        breeds_box.add_widget(breeds_scroll)
        scroll_content.add_widget(breeds_box)

        # Кнопки управления
        btn_layout = MDBoxLayout(orientation='horizontal', spacing=15, size_hint=(1, None), height=60, md_bg_color=Colors.DARK_SURFACE)

        add_breed_btn = MDButton(style='filled', size_hint=(0.33, 1), height=50)
        add_breed_btn.add_widget(MDButtonText(text='Добавить', theme_text_color='Custom', text_color=Colors.GREEN))
        add_breed_btn.bind(on_release=lambda x, p=page_num, r=row_idx: self.add_breed_from_plot(p, r))
        btn_layout.add_widget(add_breed_btn)

        edit_breed_main_btn = MDButton(style='filled', size_hint=(0.33, 1), height=50)
        edit_breed_main_btn.add_widget(MDButtonText(text='Изменить', theme_text_color='Custom', text_color=[0.7,0.7,0.7,1]))
        edit_breed_main_btn.bind(on_release=lambda x, p=page_num, r=row_idx: self.edit_breed_list_from_plot(p, r))
        btn_layout.add_widget(edit_breed_main_btn)

        close_btn = MDButton(style='filled', size_hint=(0.34, 1), height=50)
        close_btn.add_widget(MDButtonText(text='Закрыть', theme_text_color='Custom', text_color=[1,0.3,0.3,1]))
        close_btn.bind(on_release=lambda x: self.close_breed_management())
        btn_layout.add_widget(close_btn)

        scroll_content.add_widget(btn_layout)
        scroll.add_widget(scroll_content)
        content.add_widget(scroll)

        self.breed_management_popup = Popup(
            title="",
            content=content,
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3],
            size_hint=(0.95, 0.95)
        )

        self.breed_management_popup.open()

    def close_breed_management(self):
        """Закрыть меню управления породами"""
        if hasattr(self, 'breed_management_popup'):
            self.breed_management_popup.dismiss()

    def edit_breed_list_from_plot(self, page_num, row_idx):
        """Редактирование списка пород (выбор конкретной породы для изменения)"""
        if page_num not in self.page_data:
            self.show_error("Страница не найдена!")
            return

        page_data = self.page_data[page_num]
        if row_idx >= len(page_data):
            self.show_error("Строка не найдена!")
            return

        row = page_data[row_idx]
        breeds_data = row[3] if len(row) > 3 else ''
        breeds_list = self.parse_breeds_data(breeds_data)

        if not breeds_list:
            self.show_error("Нет пород для редактирования!")
            return

        # Закрываем popup управления породами
        self.close_breed_management()

        # Открываем popup выбора породы для редактирования
        content = MDBoxLayout(orientation='vertical', spacing=10, padding=10, md_bg_color=Colors.DARK_SURFACE)

        title_label = MDLabel(
            text=f"Площадка №{row_idx + 1} - Выберите породу для редактирования:",
            font_size='16sp',
            bold=True,
            theme_text_color='Custom',
            text_color=Colors.GREEN,
            size_hint=(1, None),
            height=40
        )
        content.add_widget(title_label)

        scroll = ScrollView(size_hint=(1, None), height=300)
        breeds_layout = MDGridLayout(cols=1, spacing=5, size_hint_y=None, md_bg_color=Colors.DARK_SURFACE)
        breeds_layout.bind(minimum_height=breeds_layout.setter('height'))

        for i, breed_info in enumerate(breeds_list):
            breed_name = breed_info.get('name', 'Неизвестная')
            breed_type = breed_info.get('type', 'unknown')
            breed_type_text = 'Хвойная' if breed_type == 'coniferous' else 'Лиственная'

            btn = MDButton(style='filled', size_hint=(1, None), height=50)
            btn.add_widget(MDButtonText(text=f"№{i+1}: {breed_name} ({breed_type_text})", theme_text_color='Custom', text_color=[0.7,0.7,0.7,1]))
            btn.bind(on_release=lambda x, p=page_num, r=row_idx, b=i: self.edit_breed_from_plot(p, r, b))
            breeds_layout.add_widget(btn)

        scroll.add_widget(breeds_layout)
        content.add_widget(scroll)

        cancel_btn = MDButton(style='filled', size_hint=(1, None), height=50)
        cancel_btn.add_widget(MDButtonText(text='Отмена', theme_text_color='Custom', text_color=[1,0.3,0.3,1]))
        cancel_btn.bind(on_release=lambda x: self.close_breed_selection_popup())
        content.add_widget(cancel_btn)

        self.breed_selection_popup = Popup(
            title="",
            content=content,
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3],
            size_hint=(0.9, 0.8)
        )
        self.breed_selection_popup.open()

    def close_breed_selection_popup(self):
        """Закрыть popup выбора породы"""
        if hasattr(self, 'breed_selection_popup'):
            self.breed_selection_popup.dismiss()

    def parse_breeds_data(self, breeds_text):
        """Парсинг данных о породах из JSON строки"""
        if not breeds_text:
            return []
        try:
            breeds_list = json.loads(breeds_text) if isinstance(breeds_text, str) else []
            return breeds_list if isinstance(breeds_list, list) else []
        except (json.JSONDecodeError, TypeError):
            return []

    def edit_breed_from_plot(self, page_num, row_idx, breed_idx):
        """Редактирование породы из меню площадок"""
        # Получаем данные породы
        if page_num not in self.page_data:
            self.show_error("Страница не найдена!")
            return

        page_data = self.page_data[page_num]
        if row_idx >= len(page_data):
            self.show_error("Строка не найдена!")
            return

        row = page_data[row_idx]
        breeds_data = row[3] if len(row) > 3 else ''
        breeds_list = self.parse_breeds_data(breeds_data)

        if breed_idx >= len(breeds_list):
            self.show_error("Порода не найдена!")
            return

        breed_info = breeds_list[breed_idx]

        # Закрываем popup управления породами
        self.close_breed_management()

        # Открываем popup редактирования породы (используем существующий метод)
        # Создаем временный input для передачи в show_breed_details_popup
        temp_input = TextInput(text=breeds_data)
        temp_input.row_index = row_idx

        breed_type = breed_info.get('type', 'coniferous')
        breed_name = breed_info.get('name', 'Неизвестная')

        # Открываем popup с параметрами породы
        self.show_breed_details_popup(temp_input, breed_type, breed_name)

    def add_breed_from_plot(self, page_num, row_idx):
        """Добавление породы из меню площадок"""
        # Закрываем popup управления породами
        self.close_breed_management()

        # Открываем popup выбора типа породы
        # Создаем временный input для передачи в show_breed_popup
        temp_input = TextInput(text='')
        temp_input.row_index = row_idx

        # Открываем popup выбора типа породы
        self.show_breed_popup(temp_input, True)

    def edit_plot_data_from_management(self, page_num, row_idx):
        """Редактирование данных площадки из меню управления породами"""
        # Закрываем popup управления породами
        self.close_breed_management()

        # Открываем popup редактирования площадки
        MolodnikiTreeDataInputPopup(self, row_idx).open()

    def edit_plot_popup(self, row_index):
        """Открыть popup редактирования для выбранной площадки"""
        MolodnikiTreeDataInputPopup(self, row_index).open()

    def save_to_json(self, instance=None):
        """Сохранение данных в JSON формате"""
        # Получаем итоговые данные (total_data) из меню Итого
        total_data = self.get_total_data_from_db()

        data = {
            'page_data': self.page_data,
            'section': self.current_section,
            'quarter': self.current_quarter,
            'plot': self.current_plot,
            'forestry': self.current_forestry,
            'radius': self.current_radius,
            'project_data': self.project_data,  # Данные проекта
            'total_data': total_data,  # ✅ ДОБАВЛЕНО: итоговые данные с породами
            'export_date': datetime.datetime.now().isoformat()
        }

        timestamp = datetime.datetime.now().strftime('%M%S')  # Только минуты и секунды
        document_name = self.project_data.get('document_name', 'Проект')
        # Очень короткое имя файла на основе названия проекта (макс 10 символов)
        short_name = document_name.replace(' ', '').replace('/', '_').replace('.', '')[:10]
        filename = f"{short_name}_{self.current_section}_{timestamp}.json"
        full_path = os.path.join(self.reports_dir, filename)

        try:
            with open(full_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            return f"JSON: {filename}", None
        except Exception as e:
            return None, f"Ошибка сохранения JSON: {str(e)}"

    def save_to_excel_without_dialog(self):
        """Сохранение в Excel без диалога"""
        timestamp = datetime.datetime.now().strftime('%M%S')  # Только минуты и секунды
        document_name = self.project_data.get('document_name', 'Проект')
        # Очень короткое имя файла на основе названия проекта (макс 10 символов)
        short_name = document_name.replace(' ', '').replace('/', '_').replace('.', '')[:10]
        filename = f"{short_name}_{self.current_section}_{timestamp}.xlsx"
        full_path = os.path.join(self.reports_dir, filename)

        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Молодняки"

            address_parts = []
            if self.current_quarter:
                address_parts.append(f"Квартал: {self.current_quarter}")
            if self.current_plot:
                address_parts.append(f"Выдел: {self.current_plot}")
            if self.current_forestry:
                address_parts.append(f"Лесничество: {self.current_forestry}")
            if self.current_radius:
                address_parts.append(f"Радиус: {self.current_radius} м")

            address_text = " | ".join(address_parts) if address_parts else "Адрес не указан"
            ws['A1'] = f"Адрес: {address_text}"
            ws['A1'].font = openpyxl.styles.Font(bold=True, size=12)

            # Расчет площади перечета
            current_radius = float(self.current_radius) if self.current_radius else 5.64
            plot_area_m2 = 3.14159 * (current_radius ** 2)
            plot_area_ha = plot_area_m2 / 10000
            plot_count = len([row for page in self.page_data.values() for row in page if any(cell for cell in row[:3] if cell)])
            total_plot_area_ha = plot_count * plot_area_ha

            ws['A2'] = f"Площадь перечета: {total_plot_area_ha:.4f} га ({total_plot_area_ha*10000:.0f} м²) - {plot_count} площадок по {plot_area_ha:.4f} га каждая"
            ws['A2'].font = openpyxl.styles.Font(bold=True, size=10)

            ws.append([])

            headers = [
                '№ППР', 'GPS точка', 'Предмет ухода', 'Порода', 'Густота', 'До 0.5м', '0.5-1.5м', '>1.5м', 'Высота', 'Возраст', 'Примечания', 'Тип Леса'
            ]
            for col_num, header in enumerate(headers, 1):
                cell = ws.cell(row=3, column=col_num, value=header)
                cell.font = openpyxl.styles.Font(bold=True)
                cell.fill = openpyxl.styles.PatternFill(start_color="00FF00", end_color="00FF00", fill_type="solid")

            all_data = []
            for page in sorted(self.page_data.keys()):
                all_data.extend(self.page_data[page])

            current_row = 4
            for row in all_data:
                if any(cell for cell in row[:3] if cell):  # Проверяем, что основные столбцы не пустые
                    try:
                        breeds_data = json.loads(row[3]) if row[3] else []
                    except (json.JSONDecodeError, TypeError):
                        breeds_data = []

                    if isinstance(breeds_data, list) and breeds_data:
                        for breed_info in breeds_data:
                            if isinstance(breed_info, dict):
                                breed_name = breed_info.get('name', 'Неизвестная')
                                density = breed_info.get('density', '')
                                height = breed_info.get('height', '')
                                age = breed_info.get('age', '')

                                # Инициализируем градации
                                do_05 = ''
                                _05_15 = ''
                                bolee_15 = ''

                                if breed_info.get('type') == 'coniferous':
                                    # Для хвойных заполняем градации
                                    do_05 = str(breed_info.get('do_05', ''))
                                    _05_15 = str(breed_info.get('05_15', ''))
                                    bolee_15 = str(breed_info.get('bolee_15', ''))
                                    # Густота оставляем пустой для хвойных
                                    density = ''
                                else:
                                    # Для лиственных оставляем густоту, градации пустые
                                    pass

                                processed_row = [
                                    row[0],  # №ППР
                                    row[1],  # GPS точка
                                    row[2],  # Предмет ухода
                                    breed_name,  # Порода
                                    str(density) if density else '',  # Густота
                                    do_05,  # До 0.5м
                                    _05_15,  # 0.5-1.5м
                                    bolee_15,  # >1.5м
                                    str(height) if height else '',  # Высота
                                    str(age) if age else '',  # Возраст
                                    row[4],  # Примечания
                                    row[5],  # Тип Леса
                                ]
                                ws.append(processed_row)
                                current_row += 1
                    else:
                        # Если нет пород, добавить строку без данных
                        processed_row = [row[0], row[1], row[2], '', '', '', '', '', '', '', row[4], row[5]]
                        ws.append(processed_row)
                        current_row += 1

            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width

            wb.save(full_path)
            return f"Excel: {filename}", None
        except Exception as e:
            return None, f"Ошибка сохранения Excel: {str(e)}"

    def save_to_word_without_dialog(self):
        """Сохранение в Word без диалога"""
        try:
            from docx import Document

            timestamp = datetime.datetime.now().strftime('%M%S')  # Только минуты и секунды
            document_name = self.project_data.get('document_name', 'Проект')
            # Очень короткое имя файла на основе названия проекта (макс 10 символов)
            short_name = document_name.replace(' ', '').replace('/', '_').replace('.', '')[:10]
            filename = f"{short_name}_{self.current_section}_{timestamp}.docx"
            full_path = os.path.join(self.reports_dir, filename)

            doc = Document()
            doc.add_heading(f'Расширенный отчет по молоднякам - Участок {self.current_section}', 0)

            # Расчет площади перечета
            current_radius = float(self.current_radius) if self.current_radius else 5.64
            plot_area_m2 = 3.14159 * (current_radius ** 2)
            plot_area_ha = plot_area_m2 / 10000
            plot_count = len([row for page in self.page_data.values() for row in page if any(cell for cell in row[:3] if cell)])
            total_plot_area_ha = plot_count * plot_area_ha

            # Добавляем информацию о площади перечета
            doc.add_paragraph(f"Площадь перечета: {total_plot_area_ha:.4f} га ({total_plot_area_ha*10000:.0f} м²) - {plot_count} площадок по {plot_area_ha:.4f} га каждая")

            all_data = []
            for page in sorted(self.page_data.keys()):
                all_data.extend(self.page_data[page])

            table = doc.add_table(rows=1, cols=12)
            table.style = 'Table Grid'

            headers = [
                '№ППР', 'GPS точка', 'Предмет ухода', 'Порода', 'Густота', 'До 0.5м', '0.5-1.5м', '>1.5м', 'Высота', 'Возраст', 'Примечания', 'Тип Леса'
            ]
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header

            for row in all_data:
                if any(cell for cell in row[:3] if cell):  # Проверяем, что основные столбцы не пустые
                    try:
                        breeds_data = json.loads(row[3]) if row[3] else []
                    except (json.JSONDecodeError, TypeError):
                        breeds_data = []

                    if isinstance(breeds_data, list) and breeds_data:
                        for breed_info in breeds_data:
                            if isinstance(breed_info, dict):
                                breed_name = breed_info.get('name', 'Неизвестная')
                                density = breed_info.get('density', '')
                                height = breed_info.get('height', '')
                                age = breed_info.get('age', '')

                                # Инициализируем градации
                                do_05 = ''
                                _05_15 = ''
                                bolee_15 = ''

                                if breed_info.get('type') == 'coniferous':
                                    # Для хвойных заполняем градации
                                    do_05 = str(breed_info.get('do_05', ''))
                                    _05_15 = str(breed_info.get('05_15', ''))
                                    bolee_15 = str(breed_info.get('bolee_15', ''))
                                    # Густота оставляем пустой для хвойных
                                    density = ''
                                else:
                                    # Для лиственных оставляем густоту, градации пустые
                                    pass

                                row_cells = table.add_row().cells
                                row_cells[0].text = str(row[0]) if row[0] else ""  # №ППР
                                row_cells[1].text = str(row[1]) if row[1] else ""  # GPS точка
                                row_cells[2].text = str(row[2]) if row[2] else ""  # Предмет ухода
                                row_cells[3].text = breed_name  # Порода
                                row_cells[4].text = str(density) if density else ""  # Густота
                                row_cells[5].text = do_05  # До 0.5м
                                row_cells[6].text = _05_15  # 0.5-1.5м
                                row_cells[7].text = bolee_15  # >1.5м
                                row_cells[8].text = str(height) if height else ""  # Высота
                                row_cells[9].text = str(age) if age else ""  # Возраст
                                row_cells[10].text = str(row[4]) if row[4] else ""  # Примечания
                                row_cells[11].text = str(row[5]) if row[5] else ""  # Тип Леса
                    else:
                        # Если нет пород, добавить строку без данных
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(row[0]) if row[0] else ""
                        row_cells[1].text = str(row[1]) if row[1] else ""
                        row_cells[2].text = str(row[2]) if row[2] else ""
                        row_cells[3].text = ""
                        row_cells[4].text = ""
                        row_cells[5].text = ""
                        row_cells[6].text = ""
                        row_cells[7].text = ""
                        row_cells[8].text = ""
                        row_cells[9].text = ""
                        row_cells[10].text = str(row[4]) if row[4] else ""
                        row_cells[11].text = str(row[5]) if row[5] else ""

            doc.save(full_path)
            return f"Word: {filename}", None
        except ImportError:
            return None, "Для сохранения в Word установите библиотеку python-docx: pip install python-docx"
        except Exception as e:
            return None, f"Ошибка сохранения Word: {str(e)}"

    def save_all_formats(self, instance=None):
        """Сохранить данные во всех форматах сразу"""
        success_messages = []
        error_messages = []

        # Проверка наличия данных
        if not self.page_data:
            error_messages.append("Нет данных для сохранения!")
            self.show_error("Ошибки сохранения:\n" + "\n".join(error_messages))
            return

        # Проверка наличия участка
        if not self.current_section:
            error_messages.append("Не указан номер участка!")
            self.show_error("Ошибки сохранения:\n" + "\n".join(error_messages))
            return

        # Валидация радиуса
        try:
            radius = float(self.current_radius) if self.current_radius else 5.64
            if radius <= 0:
                error_messages.append("Радиус должен быть положительным числом!")
                self.show_error("Ошибки сохранения:\n" + "\n".join(error_messages))
                return
        except (ValueError, TypeError):
            error_messages.append("Некорректное значение радиуса!")
            self.show_error("Ошибки сохранения:\n" + "\n".join(error_messages))
            return

        # Автоматически сохраняем текущую страницу перед сохранением отчёта
        if not self.save_current_page():
            error_messages.append("Не удалось сохранить текущую страницу в базу данных!")
            self.show_error("Ошибки сохранения:\n" + "\n".join(error_messages))
            return

        # Проверка существования папки reports
        if not os.path.exists(self.reports_dir):
            try:
                os.makedirs(self.reports_dir, exist_ok=True)
            except Exception as e:
                error_messages.append(f"Не удалось создать папку reports: {str(e)}")
                self.show_error("Ошибки сохранения:\n" + "\n".join(error_messages))
                return

        try:
            # Сохранение в JSON
            result, error = self.save_to_json()
            if result:
                success_messages.append(result)
            else:
                error_messages.append(error)

            # Сохранение в Excel
            result, error = self.save_to_excel_without_dialog()
            if result:
                success_messages.append(result)
            else:
                error_messages.append(error)

            # Сохранение в Word
            result, error = self.save_to_word_without_dialog()
            if result:
                success_messages.append(result)
            else:
                error_messages.append(error)

        except Exception as e:
            import traceback
            error_messages.append(f"Общая ошибка: {str(e)}\n{traceback.format_exc()}")

        if success_messages:
            self.show_success("Файлы сохранены:\n" + "\n".join(success_messages))
        if error_messages:
            self.show_error("Ошибки сохранения:\n" + "\n".join(error_messages))

    def show_edit_breed_popup(self, instance, breed_index, breed_info):
        """Показать popup для редактирования породы"""
        content = MDBoxLayout(orientation='vertical', spacing=10, padding=10, md_bg_color=Colors.DARK_SURFACE)

        title_label = MDLabel(
            text=f"Редактирование породы: {breed_info.get('name', '')}",
            bold=True,
            size_hint=(1, None),
            height=30,
            theme_text_color='Custom',
            text_color=Colors.TEXT_ON_DARK
        )
        content.add_widget(title_label)

        fields_layout = MDGridLayout(cols=2, spacing=5, size_hint=(1, None), height=200, md_bg_color=Colors.DARK_SURFACE)

        breed_type = breed_info.get('type', 'deciduous')
        if breed_type == 'coniferous':
            fields = [
                ('До 0.5м:', 'do_05'),
                ('0.5-1.5м:', '05_15'),
                ('>1.5м:', 'bolee_15'),
                ('Высота (м):', 'height'),
                ('Густота:', 'density'),
                ('Возраст (лет):', 'age')
            ]
        else:
            fields = [
                ('Густота:', 'density'),
                ('Высота (м):', 'height'),
                ('Возраст (лет):', 'age')
            ]

        self.edit_inputs = {}
        for label_text, field_key in fields:
            lbl = MDLabel(text=label_text, size_hint=(None, None), size=(100, 30), theme_text_color='Custom', text_color=Colors.TEXT_ON_DARK)
            inp = MDTextField(
                mode="outlined",
                size_hint=(None, None),
                size=(100, 30),
                line_color_focus=Colors.GREEN,
                text=str(breed_info.get(field_key, ''))
            )
            fields_layout.add_widget(lbl)
            fields_layout.add_widget(inp)
            self.edit_inputs[field_key] = inp

        content.add_widget(fields_layout)

        # Кнопки
        btn_layout = MDBoxLayout(orientation='horizontal', spacing=10, size_hint=(1, None), height=50, md_bg_color=Colors.DARK_SURFACE)
        save_btn = MDButton(style='filled', size_hint=(0.5, 1), height=50)
        save_btn.add_widget(MDButtonText(text='Сохранить', theme_text_color='Custom', text_color=Colors.GREEN))
        cancel_btn = MDButton(style='filled', size_hint=(0.5, 1), height=50)
        cancel_btn.add_widget(MDButtonText(text='Отмена', theme_text_color='Custom', text_color=[1,0.3,0.3,1]))
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3],
            size_hint=(0.9, 0.8)
        )

        def save_edit(btn):
            for key, inp in self.edit_inputs.items():
                if inp.text.strip():
                    try:
                        if key in ['density', 'age']:
                            breed_info[key] = int(inp.text)
                        elif key == 'height':
                            breed_info[key] = float(inp.text)
                        else:
                            breed_info[key] = float(inp.text)
                    except ValueError:
                        breed_info[key] = inp.text
                else:
                    breed_info[key] = 0 if key in ['density', 'age', 'do_05', '05_15', 'bolee_15'] else 0.0

            breeds_data = self.parse_breeds_data(instance.text)
            if 0 <= breed_index < len(breeds_data):
                breeds_data[breed_index] = breed_info
                instance.text = json.dumps(breeds_data, ensure_ascii=False, indent=2)
                self.update_totals()
                self.show_success("Порода обновлена!")
                popup.dismiss()

        save_btn.bind(on_release=save_edit)
        cancel_btn.bind(on_release=popup.dismiss)

        popup.open()

    def show_custom_breed_popup(self, instance, breed_type):
        """Показать popup для ввода названия другой породы"""
        content = MDBoxLayout(orientation='vertical', spacing=10, padding=10, md_bg_color=Colors.DARK_SURFACE)

        title_label = MDLabel(
            text="Введите название другой породы",
            bold=True,
            size_hint=(1, None),
            height=30,
            theme_text_color='Custom',
            text_color=Colors.TEXT_ON_DARK
        )
        content.add_widget(title_label)

        self.custom_breed_input = MDTextField(
            hint_text="Название породы",
            mode="outlined",
            size_hint=(1, None),
            height=40,
            line_color_focus=Colors.GREEN
        )
        content.add_widget(self.custom_breed_input)

        btn_layout = MDBoxLayout(orientation='horizontal', spacing=10, size_hint=(1, None), height=50, md_bg_color=Colors.DARK_SURFACE)
        save_btn = MDButton(style='filled', size_hint=(0.5, 1), height=50)
        save_btn.add_widget(MDButtonText(text='Сохранить', theme_text_color='Custom', text_color=Colors.GREEN))
        cancel_btn = MDButton(style='filled', size_hint=(0.5, 1), height=50)
        cancel_btn.add_widget(MDButtonText(text='Отмена', theme_text_color='Custom', text_color=[1,0.3,0.3,1]))
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3],
            size_hint=(0.8, 0.6)
        )

        def save_custom_breed(btn):
            breed_name = self.custom_breed_input.text.strip()
            if breed_name:
                # Проверяем, не является ли порода запрещенной
                forbidden_breeds = ['семенная', 'культуры', 'подрост']
                if any(forbidden.lower() in breed_name.lower() for forbidden in forbidden_breeds):
                    self.show_error("Эта порода не разрешена для использования!")
                    return
                
                # Сохраняем новую породу в базу данных
                self.save_custom_breed_to_db(breed_name, breed_type)
                
                instance.text = breed_name
                self.show_breed_details_popup(instance, breed_type, breed_name)
                popup.dismiss()
            else:
                self.show_error("Название породы не может быть пустым!")

        save_btn.bind(on_release=save_custom_breed)
        cancel_btn.bind(on_release=popup.dismiss)

        popup.open()

    def update_coniferous_density(self, instance, value):
        """Автоматический расчет густоты для хвойных пород"""
        if 'density' in self.breed_inputs:
            density_input = self.breed_inputs['density']
            try:
                do_05 = int(self.breed_inputs.get('do_05', TextInput(text='0')).text or '0')
                _05_15 = int(self.breed_inputs.get('05_15', TextInput(text='0')).text or '0')
                bolee_15 = int(self.breed_inputs.get('bolee_15', TextInput(text='0')).text or '0')

                total_density = do_05 + _05_15 + bolee_15
                density_input.text = str(total_density) if total_density > 0 else ''
            except (ValueError, AttributeError):
                pass

    def update_plot_breeds_display(self, plot_breeds_list, existing_breeds_for_plot):
        """Обновить отображение списка пород на площадке"""
        # Очищаем текущий список
        plot_breeds_list.clear_widgets()

        if existing_breeds_for_plot:
            for i, breed_info in enumerate(existing_breeds_for_plot):
                if not isinstance(breed_info, dict):
                    continue
                breed_name = breed_info.get('name', 'Неизвестная')

                # Рассчитываем густоту (для хвойных - сумма градаций)
                density = breed_info.get('density', 0)
                if breed_info.get('type') == 'coniferous':
                    conif_density = (breed_info.get('do_05', 0) + breed_info.get('05_15', 0) + breed_info.get('bolee_15', 0))
                    if conif_density > 0:
                        density = conif_density

                height = breed_info.get('height', 0)
                age = breed_info.get('age', 0)
                diameter = breed_info.get('diameter', 0)

                # Формируем строку с параметрами породы
                params_parts = []
                if density:
                    params_parts.append(f"Густота: {density}")
                if height:
                    params_parts.append(f"Высота: {height}м")
                if diameter:
                    params_parts.append(f"Диаметр: {diameter}см")
                if age:
                    params_parts.append(f"Возраст: {age}л")

                params_text = ", ".join(params_parts) if params_parts else "Нет данных"

                # Для хвойных пород добавляем градации по высоте
                gradation_text = ""
                if breed_info.get('type') == 'coniferous':
                    gradations = []
                    if breed_info.get('do_05', 0) > 0:
                        gradations.append(f"До 0.5м: {breed_info['do_05']}")
                    if breed_info.get('05_15', 0) > 0:
                        gradations.append(f"0.5-1.5м: {breed_info['05_15']}")
                    if breed_info.get('bolee_15', 0) > 0:
                        gradations.append(f">1.5м: {breed_info['bolee_15']}")
                    if gradations:
                        gradation_text = " | " + ", ".join(gradations)

                breed_item = Label(
                    text=f'{i+1}. {breed_name} - {params_text}{gradation_text}',
                    font_name='Roboto',
                    font_size='12sp',
                    color=(0.2, 0.2, 0.2, 1),
                    size_hint=(1, None),
                    height=35 if gradation_text else 30,
                    halign='left',
                    valign='middle'
                )
                breed_item.bind(size=lambda *args: setattr(breed_item, 'text_size', (breed_item.width, None)))
                plot_breeds_list.add_widget(breed_item)
        else:
            no_breeds_label = Label(
                text='Породы еще не добавлены',
                font_name='Roboto',
                font_size='12sp',
                color=(0.5, 0.5, 0.5, 1),
                size_hint=(1, None),
                height=30,
                halign='left'
            )
            no_breeds_label.bind(size=lambda *args: setattr(no_breeds_label, 'text_size', (no_breeds_label.width, None)))
            plot_breeds_list.add_widget(no_breeds_label)

    def update_address_display(self):
        """Обновить отображение адреса"""
        # Обновляем адресную строку в интерфейсе
        address_parts = []
        if self.current_quarter:
            address_parts.append(f"К: {self.current_quarter}")
        if self.current_plot:
            address_parts.append(f"В: {self.current_plot}")
        if self.current_forestry:
            address_parts.append(f"Л: {self.current_forestry}")
        if self.current_radius:
            address_parts.append(f"Р: {self.current_radius}")
        if hasattr(self, 'plot_area_input') and self.plot_area_input:
            address_parts.append(f"П: {self.plot_area_input} га")

        # Добавляем название документа
        doc_name = self.project_data.get('document_name', 'Проект')
        if doc_name:
            address_parts.append(f"Д: {doc_name}")

        address_text = " | ".join(address_parts) if address_parts else ""
        self.address_label.text = address_text

        # Обновляем текст для бокса "Настройки Адреса"
        address_display_text = (
            f"Квартал: {self.project_data['address'].get('quarter', 'Не указан')}\n"
            f"Выдел: {self.project_data['address'].get('plot', 'Не указан')}\n"
            f"Лесничество: {self.project_data['address'].get('forestry', 'Не указано')}\n"
            f"Участковое лесничество: {self.project_data['address'].get('district_forestry', 'Не указано')}\n"
            f"Радиус: {self.project_data['address'].get('radius', 'Не указан')} м\n"
            f"Площадь участка: {self.project_data['address'].get('plot_area', 'Не указана')} га"
        )
        self.current_address_display_text = address_display_text

        # Если popup открыт, обновляем и там
        self.update_address_popup_display()

    def update_quarter_display(self, value):
        """Обновить отображение квартала"""
        self.current_quarter = value.strip() if value else ""
        self.project_data['address']['quarter'] = self.current_quarter
        self.update_address_display()

    def update_plot_area_display(self, instance, value):
        """Обновить отображение площади участка"""
        self.plot_area_input = value.strip() if value else ""
        self.project_data['address']['plot_area'] = self.plot_area_input
        self.update_address_display()

    def update_plot_display(self, instance, value):
        """Обновить отображение выдела"""
        self.current_plot = value.strip() if value else ""
        self.project_data['address']['plot'] = self.current_plot
        self.update_address_display()

    def update_forestry_display(self, instance, value):
        """Обновить отображение лесничества"""
        self.current_forestry = value.strip() if value else ""
        self.project_data['address']['forestry'] = self.current_forestry
        self.update_address_display()

    def update_district_forestry_display(self, instance, value):
        """Обновить отображение участкового лесничества"""
        self.current_district_forestry = value.strip() if value else ""
        self.project_data['address']['district_forestry'] = self.current_district_forestry
        self.update_address_display()

    def update_radius_display(self, instance, value):
        """Обновить отображение радиуса"""
        self.current_radius = value.strip() if value else "5.64"
        self.project_data['address']['radius'] = self.current_radius
        self.update_address_display()

    def update_address_label(self):
        """Устаревший метод - используем update_address_display"""
        self.update_address_display()

    def update_address_popup_display(self):
        """Обновить отображение адреса в открытом popup Настройки адреса"""
        if hasattr(self, 'current_address_info') and self.current_address_info:
            self.current_address_info.text = (
                f"Название проекта: {self.project_data.get('document_name', 'Проект')}\n"
                f"Квартал: {self.project_data['address'].get('quarter', 'Не указан')}\n"
                f"Выдел: {self.project_data['address'].get('plot', 'Не указан')}\n"
                f"Лесничество: {self.project_data['address'].get('forestry', 'Не указано')}\n"
                f"Участковое лесничество: {self.project_data['address'].get('district_forestry', 'Не указано')}\n"
                f"Радиус: {self.project_data['address'].get('radius', 'Не указан')} м\n"
                f"Площадь участка: {self.project_data['address'].get('plot_area', 'Не указана')} га"
            )
        if hasattr(self, 'address_popup') and self.address_popup:
            self.address_popup.title = "Настройки адреса (обновлено)"

    def update_details_display(self):
        """Обновляет отображение деталей проекта"""
        if hasattr(self, 'current_details_info') and self.current_details_info:
            care_queue_val = self.project_data['details'].get('care_queue', '') or self.care_queue or 'Не указана'
            characteristics_val = self.project_data['details'].get('characteristics', '') or self.characteristics or 'Не указана'
            care_date_val = self.project_data['details'].get('care_date', '') or self.care_date or 'Не указана'
            technology_val = self.project_data['details'].get('technology', '') or self.technology or 'Не указана'
            forest_purpose_val = self.project_data['details'].get('forest_purpose', '') or self.forest_purpose or 'Не указано'

            self.current_details_info.text = (
                f"Очередь рубки: {care_queue_val}\n"
                f"Характеристика молодняков: {characteristics_val}\n"
                f"Дата рубки: {care_date_val}\n"
                f"Технология ухода: {technology_val}\n"
                f"Назначение лесов: {forest_purpose_val}"
            )

    def load_existing_data(self):
        """Загружаем существующие данные из базы данных"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()

        try:
            # Загружаем настройки участка
            cursor.execute('''
                SELECT radius, plot_area FROM molodniki_settings
                WHERE section_name = ?
            ''', (self.current_section,))

            settings_row = cursor.fetchone()
            if settings_row:
                self.current_radius = str(settings_row[0]) if settings_row[0] else "5.64"
                self.plot_area_input = str(settings_row[1]) if settings_row[1] else ""

            cursor.execute('''
                SELECT DISTINCT page_number FROM molodniki_data
                WHERE section_name = ?
                ORDER BY page_number
            ''', (self.current_section,))

            page_numbers = [row[0] for row in cursor.fetchall()]

            if page_numbers:
                for page_num in page_numbers:
                    cursor.execute('''
                        SELECT row_index, nn, gps_point, predmet_uhoda, poroda, primechanie, radius
                        FROM molodniki_data
                        WHERE page_number = ? AND section_name = ?
                        ORDER BY row_index
                    ''', (page_num, self.current_section))

                    page_data = []
                    rows_data = cursor.fetchall()
                    max_row = max((r[0] for r in rows_data), default=-1)

                    for row_idx in range(max_row + 1):
                        row_data = ['', '', '', '', '', '']
                        page_data.append(row_data)

                    for row_data in rows_data:
                        row_idx, nn, gps_point, predmet_uhoda, poroda, primechanie, radius = row_data
                        if row_idx < len(page_data):
                            page_data[row_idx] = [
                                str(nn) if nn is not None else '',
                                str(gps_point) if gps_point is not None else '',
                                str(predmet_uhoda) if predmet_uhoda is not None else '',
                                str(poroda) if poroda is not None else '',
                                str(primechanie) if primechanie is not None else '',
                                str(radius) if radius is not None else '',
                            ]

                    self.page_data[page_num] = page_data

                self.current_page = min(page_numbers)
                self.load_page_data()

        except Exception as e:
            print(f"Error loading existing data: {e}")
            self.show_error(f"Ошибка загрузки данных из базы: {str(e)}")
        finally:
            conn.close()

    def load_page_data(self):
        # Данные загружаются напрямую из page_data, таблица не используется
        pass

    def clear_table_data(self, instance=None):
        self.page_data.clear()
        self.show_success("Данные очищены!")

    def open_excel_file(self, instance):
        if os.path.exists(self.reports_dir):
            os.startfile(self.reports_dir)
        else:
            self.show_error("Папка reports не найдена!")

    def go_back(self, instance):
        App.get_running_app().root.current = 'main'

    def show_error(self, message):
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)
        content.add_widget(MDIcon(icon='alert-circle', theme_text_color='Custom', text_color=Colors.DANGER,
                                   size_hint=(None, None), size=(dp(48), dp(48)), halign='center'))
        content.add_widget(MDLabel(text=message, theme_text_color='Custom', text_color=[1,0.3,0.3,1],
                                    halign='center', adaptive_height=True))
        Popup(
            title="",
            content=content,
            size_hint=(0.7, None),
            height=dp(180),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3]
        ).open()

    def show_success(self, message):
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)
        content.add_widget(MDIcon(icon='check-circle', theme_text_color='Custom', text_color=Colors.GREEN,
                                   size_hint=(None, None), size=(dp(48), dp(48)), halign='center'))
        content.add_widget(MDLabel(text=message, theme_text_color='Custom', text_color=[0.3,0.8,0.3,1],
                                    halign='center', adaptive_height=True))
        Popup(
            title="",
            content=content,
            size_hint=(0.7, None),
            height=dp(180),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3]
        ).open()

    def show_quarter_popup(self, instance):
        """Показать popup для ввода квартала"""
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)

        title_label = MDLabel(
            text="Введите номер квартала",
            font_style='Title', role='medium',
            theme_text_color='Custom', text_color=Colors.GREEN,
            size_hint_y=None, height=dp(40)
        )
        content.add_widget(title_label)

        self.quarter_input = MDTextField(
            hint_text="Номер квартала",
            mode='outlined',
            size_hint_y=None, height=dp(52),
            input_filter='int',
            text=self.project_data['address'].get('quarter', self.current_quarter)
        )
        self.quarter_input.bind(text=lambda instance, value: self.update_quarter_display(value))
        content.add_widget(self.quarter_input)

        btn_layout = MDBoxLayout(orientation='horizontal', spacing=Spacing.MD, adaptive_height=True)
        save_btn = MDButton(style='filled', md_bg_color=Colors.BTN_SUCCESS,
                            size_hint=(0.5, None), height=dp(48))
        save_btn.add_widget(MDButtonText(text='Сохранить'))
        cancel_btn = MDButton(style='outlined',
                              size_hint=(0.5, None), height=dp(48),
                              line_color=Colors.DANGER)
        cancel_btn.add_widget(MDButtonText(text='Отмена',
                              theme_text_color='Custom', text_color=Colors.DANGER))
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            size_hint=(0.85, None),
            height=dp(300),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3]
        )

        def save_quarter(btn):
            quarter = self.quarter_input.text.strip()
            if quarter:
                self.current_quarter = quarter
                self.project_data['address']['quarter'] = quarter
                self.update_address_display()
                self.show_success(f"Квартал установлен: {quarter}")
                popup.dismiss()
            else:
                self.show_error("Номер квартала не может быть пустым!")

        save_btn.bind(on_release=save_quarter)
        cancel_btn.bind(on_release=popup.dismiss)

        popup.open()

    def show_plot_popup(self, instance):
        """Показать popup для ввода выдела"""
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)

        title_label = MDLabel(
            text="Введите номер выдела",
            font_style='Title', role='medium',
            theme_text_color='Custom', text_color=Colors.GREEN,
            size_hint_y=None, height=dp(40)
        )
        content.add_widget(title_label)

        self.plot_input = MDTextField(
            hint_text="Номер выдела",
            mode='outlined',
            size_hint_y=None, height=dp(52),
            input_filter='int',
            text=self.current_plot
        )
        content.add_widget(self.plot_input)

        btn_layout = MDBoxLayout(orientation='horizontal', spacing=Spacing.MD, adaptive_height=True)
        save_btn = MDButton(style='filled', md_bg_color=Colors.BTN_SUCCESS,
                            size_hint=(0.5, None), height=dp(48))
        save_btn.add_widget(MDButtonText(text='Сохранить'))
        cancel_btn = MDButton(style='outlined',
                              size_hint=(0.5, None), height=dp(48),
                              line_color=Colors.DANGER)
        cancel_btn.add_widget(MDButtonText(text='Отмена',
                              theme_text_color='Custom', text_color=Colors.DANGER))
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            size_hint=(0.85, None),
            height=dp(300),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3]
        )

        def save_plot(btn):
            plot = self.plot_input.text.strip()
            if plot:
                self.current_plot = plot
                self.project_data['address']['plot'] = plot
                self.update_address_label()
                self.show_success(f"Выдел установлен: {plot}")
                popup.dismiss()
            else:
                self.show_error("Номер выдела не может быть пустым!")

        save_btn.bind(on_release=save_plot)
        cancel_btn.bind(on_release=popup.dismiss)

        popup.open()

    def show_forestry_popup(self, instance):
        """Показать popup для ввода лесничества и участкового лесничества"""
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)

        title_label = MDLabel(
            text="Введите название лесничества и участкового лесничества",
            font_style='Title', role='medium',
            theme_text_color='Custom', text_color=Colors.GREEN,
            size_hint_y=None, height=dp(40)
        )
        content.add_widget(title_label)

        forestry_label = MDLabel(
            text="Лесничество:",
            theme_text_color='Custom', text_color=[0.8,0.8,0.8,1],
            size_hint_y=None, height=dp(25)
        )
        content.add_widget(forestry_label)

        self.forestry_input = MDTextField(
            hint_text="Название лесничества",
            mode='outlined',
            size_hint_y=None, height=dp(52),
            text=self.current_forestry
        )
        content.add_widget(self.forestry_input)

        district_forestry_label = MDLabel(
            text="Участковое лесничество:",
            theme_text_color='Custom', text_color=[0.8,0.8,0.8,1],
            size_hint_y=None, height=dp(25)
        )
        content.add_widget(district_forestry_label)

        self.district_forestry_input = MDTextField(
            hint_text="Название участкового лесничества",
            mode='outlined',
            size_hint_y=None, height=dp(52),
            text=getattr(self, 'current_district_forestry', '')
        )
        content.add_widget(self.district_forestry_input)

        btn_layout = MDBoxLayout(orientation='horizontal', spacing=Spacing.MD, adaptive_height=True)
        save_btn = MDButton(style='filled', md_bg_color=Colors.BTN_SUCCESS,
                            size_hint=(0.5, None), height=dp(48))
        save_btn.add_widget(MDButtonText(text='Сохранить'))
        cancel_btn = MDButton(style='outlined',
                              size_hint=(0.5, None), height=dp(48),
                              line_color=Colors.DANGER)
        cancel_btn.add_widget(MDButtonText(text='Отмена',
                              theme_text_color='Custom', text_color=Colors.DANGER))
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            size_hint=(0.85, None),
            height=dp(380),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3]
        )

        def save_forestry(btn):
            forestry = self.forestry_input.text.strip()
            district_forestry = self.district_forestry_input.text.strip()
            if forestry:
                self.current_forestry = forestry
                self.current_district_forestry = district_forestry
                self.project_data['address']['forestry'] = forestry
                self.project_data['address']['district_forestry'] = district_forestry
                self.update_address_label()
                self.show_success(f"Лесничество установлено: {forestry}" + (f", участковое: {district_forestry}" if district_forestry else ""))
                popup.dismiss()
            else:
                self.show_error("Название лесничества не может быть пустым!")

        save_btn.bind(on_release=save_forestry)
        cancel_btn.bind(on_release=popup.dismiss)

        popup.open()

    def show_district_forestry_popup(self, instance=None):
        """Показать popup для ввода участкового лесничества"""
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)

        title_label = MDLabel(
            text="Введите название участкового лесничества",
            font_style='Title', role='medium',
            theme_text_color='Custom', text_color=Colors.GREEN,
            size_hint_y=None, height=dp(40)
        )
        content.add_widget(title_label)

        self.district_forestry_input = MDTextField(
            hint_text="Название участкового лесничества",
            mode='outlined',
            size_hint_y=None, height=dp(52),
            text=getattr(self, 'current_district_forestry', '')
        )
        content.add_widget(self.district_forestry_input)

        btn_layout = MDBoxLayout(orientation='horizontal', spacing=Spacing.MD, adaptive_height=True)
        save_btn = MDButton(style='filled', md_bg_color=Colors.BTN_SUCCESS,
                            size_hint=(0.5, None), height=dp(48))
        save_btn.add_widget(MDButtonText(text='Сохранить'))
        cancel_btn = MDButton(style='outlined',
                              size_hint=(0.5, None), height=dp(48),
                              line_color=Colors.DANGER)
        cancel_btn.add_widget(MDButtonText(text='Отмена',
                              theme_text_color='Custom', text_color=Colors.DANGER))
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            size_hint=(0.85, None),
            height=dp(300),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3]
        )

        def save_district_forestry(btn):
            district_forestry = self.district_forestry_input.text.strip()
            self.current_district_forestry = district_forestry
            self.update_address_label()
            if district_forestry:
                self.show_success(f"Участковое лесничество установлено: {district_forestry}")
            else:
                self.show_success("Участковое лесничество очищено")
            popup.dismiss()

        save_btn.bind(on_release=save_district_forestry)
        cancel_btn.bind(on_release=popup.dismiss)

        popup.open()

    def calculate_section_totals(self):
        """Расчет итогов по всему разделу (все страницы)"""
        breed_composition = {}
        total_stats = {'density': [], 'height': [], 'age': []}
        coniferous_stats = {'do_05': [], '05_15': [], 'bolee_15': [], 'height': [], 'age': []}

        for page_num, page_data in self.page_data.items():
            for row_data in page_data:
                # row_data имеет 6 элементов: nn, gps_point, predmet_uhoda, poroda, primechanie, tip_lesa
                if row_data[2]:  # predmet_uhoda
                    composition = self.parse_composition(row_data[2])
                    for breed, count in composition.items():
                        if breed not in breed_composition:
                            breed_composition[breed] = []
                        breed_composition[breed].append(count)

                breeds_text = row_data[3]  # poroda
                if breeds_text:
                    breeds_data = self.parse_breeds_data(breeds_text)
                    for breed_info in breeds_data:
                        if breed_info.get('type') == 'coniferous':
                            # Густота хвойных = сумма градаций
                            coniferous_density = (breed_info.get('do_05', 0) +
                                                breed_info.get('05_15', 0) +
                                                breed_info.get('bolee_15', 0))
                            if coniferous_density > 0:
                                total_stats['density'].append(coniferous_density)
                            else:
                                if 'density' in breed_info and breed_info['density']:
                                    total_stats['density'].append(breed_info['density'])

                        elif 'density' in breed_info and breed_info['density']:
                            total_stats['density'].append(breed_info['density'])

                        if 'height' in breed_info and breed_info['height']:
                            total_stats['height'].append(breed_info['height'])
                        if 'age' in breed_info and breed_info['age']:
                            total_stats['age'].append(breed_info['age'])

        # Рассчитываем остальные итоги

        current_radius = float(self.current_radius) if self.current_radius else 5.64
        plot_area_m2 = 3.14159 * (current_radius ** 2)  # Площадь пробной площади в м²

        # Расчет средних по градациям для хвойных по формулам лесного хозяйства на гектар
        coniferous_stats_ha = []
        for row_data in [row for page in self.page_data.values() for row in page]:
            breeds_text = row_data[3]
            if breeds_text:
                breeds_data = self.parse_breeds_data(breeds_text)
                if breeds_data and any(b.get('type') == 'coniferous' for b in breeds_data):
                    coniferous_density_ha = 0
                    height_sum = 0
                    age_sum = 0
                    count = 0
                    for breed_info in breeds_data:
                        if breed_info.get('type') == 'coniferous':
                            do_05 = breed_info.get('do_05', 0)
                            _05_15 = breed_info.get('05_15', 0)
                            bolee_15 = breed_info.get('bolee_15', 0)
                            coniferous_density_ha += (do_05 * 10000 / plot_area_m2) + (_05_15 * 10000 / plot_area_m2) + (bolee_15 * 10000 / plot_area_m2)
                            if breed_info.get('height'):
                                height_sum += breed_info['height']
                                count += 1
                            if breed_info.get('age'):
                                age_sum += breed_info['age']

                    coniferous_stats_ha.append({
                        'density_ha': coniferous_density_ha if coniferous_density_ha > 0 else 0,
                        'height': height_sum / count if count > 0 else 0,
                        'age': age_sum / count if count > 0 else 0
                    })

        # Итоги по лиственным
        deciduous_stats = []
        for row_data in [row for page in self.page_data.values() for row in page]:
            breeds_text = row_data[3]
            if breeds_text:
                breeds_data = self.parse_breeds_data(breeds_text)
                if breeds_data:
                    deciduous_density_total = 0
                    deciduous_height = []
                    deciduous_age = []
                    for breed_info in breeds_data:
                        if breed_info.get('type') == 'deciduous':
                            deciduous_density_total += breed_info.get('density', 0)
                            if breed_info.get('height', 0) > 0:
                                deciduous_height.append(breed_info['height'])
                            if breed_info.get('age', 0) > 0:
                                deciduous_age.append(breed_info['age'])
                    if deciduous_height or deciduous_age:
                        deciduous_density_ha = deciduous_density_total
                        avg_height = sum(deciduous_height) / len(deciduous_height) if deciduous_height else 0
                        avg_age = sum(deciduous_age) / len(deciduous_age) if deciduous_age else 0
                        deciduous_stats.append({'density': deciduous_density_ha, 'height': avg_height, 'age': avg_age})

        # Сводные итоги
        avg_composition = {}
        for breed, counts in breed_composition.items():
            if counts:
                avg_composition[breed] = sum(counts) / len(counts)

        composition_text = ""
        for breed in sorted(avg_composition.keys()):
            count = avg_composition[breed]
            if count > 0:
                composition_text += f"{int(count)}{breed}"

        # Обновляем итоговую строку или возвращаем данные
        # В зависимости от логики приложения
        return {
            'composition_text': composition_text,
            'forestry_formulas_text': forestry_formulas_text if 'forestry_formulas_text' in locals() else "",
            'total_plots': total_plots if 'total_plots' in locals() else 0
        }

        forestry_formulas_text = ""

        # Расчет градаций для хвойных пород
        coniferous_gradiations_stats = {'do_05_ha': [], '05_15_ha': [], 'bolee_15_ha': [], 'height': [], 'age': []}

        current_radius = float(self.current_radius) if self.current_radius else 5.64
        plot_area_m2 = 3.14159 * (current_radius ** 2)

        for row_data in [row for page in self.page_data.values() for row in page]:
            breeds_text = row_data[3]
            if breeds_text:
                breeds_data = self.parse_breeds_data(breeds_text)
                for breed_info in breeds_data:
                    if breed_info.get('type') == 'coniferous':
                        do_05_ha = breed_info.get('do_05', 0) * 10000 / plot_area_m2 if plot_area_m2 > 0 else 0
                        _05_15_ha = breed_info.get('05_15', 0) * 10000 / plot_area_m2 if plot_area_m2 > 0 else 0
                        bolee_15_ha = breed_info.get('bolee_15', 0) * 10000 / plot_area_m2 if plot_area_m2 > 0 else 0
                        height = breed_info.get('height', 0)
                        age = breed_info.get('age', 0)

                        coniferous_gradiations_stats['do_05_ha'].append(do_05_ha)
                        coniferous_gradiations_stats['05_15_ha'].append(_05_15_ha)
                        coniferous_gradiations_stats['bolee_15_ha'].append(bolee_15_ha)
                        if height > 0:
                            coniferous_gradiations_stats['height'].append(height)
                        if age > 0:
                            coniferous_gradiations_stats['age'].append(age)

        if coniferous_gradiations_stats['do_05_ha'] or coniferous_gradiations_stats['05_15_ha'] or coniferous_gradiations_stats['bolee_15_ha']:
            forestry_formulas_text += "Хвойные: "
            gradiations = []
            avg_do_05 = sum(coniferous_gradiations_stats['do_05_ha']) / len(coniferous_gradiations_stats['do_05_ha']) if coniferous_gradiations_stats['do_05_ha'] else 0
            gradiations.append(f"до 0.5м: {avg_do_05:.1f} шт/га")

            avg_05_15 = sum(coniferous_gradiations_stats['05_15_ha']) / len(coniferous_gradiations_stats['05_15_ha']) if coniferous_gradiations_stats['05_15_ha'] else 0
            gradiations.append(f"0.5-1.5м: {avg_05_15:.1f} шт/га")

            avg_bolee_15 = sum(coniferous_gradiations_stats['bolee_15_ha']) / len(coniferous_gradiations_stats['bolee_15_ha']) if coniferous_gradiations_stats['bolee_15_ha'] else 0
            gradiations.append(f">1.5м: {avg_bolee_15:.1f} шт/га")

            forestry_formulas_text += ", ".join(gradiations)

            if coniferous_gradiations_stats['height']:
                avg_height = sum(coniferous_gradiations_stats['height']) / len(coniferous_gradiations_stats['height'])
                forestry_formulas_text += f", высота: {avg_height:.1f}м"
            if coniferous_gradiations_stats['age']:
                avg_age = sum(coniferous_gradiations_stats['age']) / len(coniferous_gradiations_stats['age'])
                forestry_formulas_text += f", возраст: {avg_age:.1f} лет"

        # Лиственные итоги
        if deciduous_stats:
            if forestry_formulas_text:
                forestry_formulas_text += "; "
            forestry_formulas_text += "Лиственные: "
            avg_deciduous_density = sum(d['density'] for d in deciduous_stats) / len(deciduous_stats) if deciduous_stats else 0
            avg_deciduous_height = sum(d['height'] for d in deciduous_stats) / len(deciduous_stats) if deciduous_stats else 0
            avg_deciduous_age = sum(d['age'] for d in deciduous_stats) / len(deciduous_stats) if deciduous_stats else 0

            if avg_deciduous_density > 0:
                forestry_formulas_text += f"густота: {avg_deciduous_density:.1f} шт/га "
            if avg_deciduous_height > 0:
                forestry_formulas_text += f"высота: {avg_deciduous_height:.1f}м "
            if avg_deciduous_age > 0:
                forestry_formulas_text += f"возраст: {avg_deciduous_age:.1f} лет"

        return {
            'composition_text': composition_text,
            'forestry_formulas_text': forestry_formulas_text,
            'total_plots': sum(1 for page in self.page_data.values() for row in page if any(cell for cell in row[:3] if cell))
        }

    def show_total_summary_popup(self, *args, **kwargs):
        """Показать popup со сводными итогами и таксационными расчетами - 10 отдельных цветных боксов"""
        try:
            default_radius = float(self.current_radius) if self.current_radius else 1.78
            plot_area_m2_default = 3.14159 * (default_radius ** 2)
            plot_area_ha_default = plot_area_m2_default / 10000
            trees_per_ha = 10000 / plot_area_m2_default if plot_area_m2_default > 0 else 0

            # Словарь для сбора данных по породам
            breeds_data = {}
            forest_types_set = set()
            
            # Подсчитываем общее количество площадок
            total_plots_count = sum(1 for page in self.page_data.values() for row in page if any(cell for cell in row[:3] if cell))

            # Обрабатываем все страницы для сбора данных
            for page_num, page_rows in self.page_data.items():
                for row in page_rows:
                    if len(row) < 4:
                        continue

                    plot_radius = default_radius
                    plot_area_m2 = 3.14159 * (plot_radius ** 2)
                    plot_area_ha = plot_area_m2 / 10000

                    # Собираем тип леса из row[5] (индекс 5 в расширенной таблице)
                    if len(row) > 5 and row[5]:
                        forest_types_set.add(str(row[5]).strip())

                    breeds_text = row[3]
                    if not breeds_text:
                        continue

                    try:
                        breeds_list = json.loads(breeds_text) if isinstance(breeds_text, str) else []
                    except json.JSONDecodeError:
                        continue

                    for breed_info in breeds_list:
                        if not isinstance(breed_info, dict):
                            continue

                        breed_name = breed_info.get('name', '').strip()
                        if not breed_name:
                            continue

                        breed_type = breed_info.get('type', 'deciduous')
                        density = 0
                        height = None
                        age = None

                        if breed_type == 'coniferous':
                            do_05 = breed_info.get('do_05', 0)
                            _05_15 = breed_info.get('05_15', 0)
                            bolee_15 = breed_info.get('bolee_15', 0)
                            density = (do_05 + _05_15 + bolee_15) / plot_area_ha if plot_area_ha > 0 else 0

                            if any([do_05, _05_15, bolee_15]):
                                if bolee_15 > 0:
                                    height = 2.0
                                elif _05_15 > 0:
                                    height = 1.0
                                elif do_05 > 0:
                                    height = 0.3
                                else:
                                    height = 0.0
                            else:
                                height = breed_info.get('height', 0) or 0
                        else:
                            density_value = breed_info.get('density', 0)
                            density = density_value / plot_area_ha if plot_area_ha > 0 else 0
                            height = breed_info.get('height', 0) or 0

                        age = breed_info.get('age', 0) or 0
                        diameter = breed_info.get('diameter', 0) or 0

                        if breed_name not in breeds_data:
                            breeds_data[breed_name] = {
                                'type': breed_type,
                                'plots': [],
                                'coniferous_zones': {'do_05': 0, '05_15': 0, 'bolee_15': 0} if breed_type == 'coniferous' else None,
                                'diameters': []
                            }

                        plot_data = {
                            'density': density,  # в шт/га
                            'density_raw': breed_info.get('density', 0) if breed_type == 'deciduous' else (do_05 + _05_15 + bolee_15),  # исходное кол-во деревьев
                            'height': height,
                            'age': age,
                            'diameter': diameter  # ✅ ДОБАВЛЕНО: диаметр
                        }

                        if breed_type == 'coniferous':
                            plot_data.update({
                                'do_05_density': do_05 / plot_area_ha if plot_area_ha > 0 else 0,
                                '05_15_density': _05_15 / plot_area_ha if plot_area_ha > 0 else 0,
                                'bolee_15_density': bolee_15 / plot_area_ha if plot_area_ha > 0 else 0,
                                'do_05': do_05,  # исходное кол-во
                                '05_15': _05_15,  # исходное кол-во
                                'bolee_15': bolee_15  # исходное кол-во
                            })

                        breeds_data[breed_name]['plots'].append(plot_data)
                        breeds_data[breed_name]['diameters'].append(diameter)

                        if breed_type == 'coniferous':
                            breeds_data[breed_name]['coniferous_zones']['do_05'] += plot_data['do_05_density']
                            breeds_data[breed_name]['coniferous_zones']['05_15'] += plot_data['05_15_density']
                            breeds_data[breed_name]['coniferous_zones']['bolee_15'] += plot_data['bolee_15_density']

            # Создаем popup с результатами
            content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                                  md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)

            scroll = ScrollView(size_hint=(1, None), height=600)
            results_layout = MDGridLayout(cols=1, spacing=Spacing.MD, adaptive_height=True, size_hint_y=None)
            results_layout.bind(minimum_height=results_layout.setter('height'))

            # ============================================================
            # БОКС 1: АДРЕСНАЯ ИНФОРМАЦИЯ (#E6F3FF)
            # ============================================================
            address_block = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, adaptive_height=True, padding=Spacing.MD)
            address_title = MDLabel(
                text='АДРЕСНАЯ ИНФОРМАЦИЯ',
                theme_text_color='Custom',
                text_color=Colors.GREEN,
                font_style='Title',
                role='medium',
                adaptive_height=True
            )
            address_block.add_widget(address_title)

            address_inner = MDBoxLayout(orientation='vertical', spacing=5, adaptive_height=True, padding=[10, 10])
            plot_area_from_address = self.project_data['address'].get('plot_area', '')
            address_area_text = f"Площадь участка: {plot_area_from_address} га" if plot_area_from_address else "Площадь участка: не указана"
            address_text = MDLabel(
                text=f"Название проекта: {self.project_data.get('document_name', 'Проект')}\n"
                     f"Квартал: {self.current_quarter or 'Не указан'}\n"
                     f"Выдел: {getattr(self, 'current_plot', '') or 'Не указан'}\n"
                     f"Лесничество: {self.current_forestry or 'Не указано'}\n"
                     f"Участковое лесничество: {self.project_data['address'].get('district_forestry', '') or 'Не указано'}\n"
                     f"Радиус: {default_radius:.2f} м | Площадь: {plot_area_m2_default:.1f} м² | {address_area_text}",
                theme_text_color='Custom',
                text_color=[1, 1, 1, 1],
                adaptive_height=True
            )
            address_inner.add_widget(address_text)
            address_block.add_widget(address_inner)
            results_layout.add_widget(address_block)

            # ============================================================
            # БОКС 2: ДЕТАЛИ УХОДА (#FFF8E6)
            # ============================================================
            details_block = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, adaptive_height=True, padding=Spacing.MD)
            details_title = MDLabel(
                text='ДЕТАЛИ УХОДА',
                theme_text_color='Custom',
                text_color=Colors.GREEN,
                font_style='Title',
                role='medium',
                adaptive_height=True
            )
            details_block.add_widget(details_title)

            details_inner = MDBoxLayout(orientation='vertical', spacing=5, adaptive_height=True, padding=[10, 10])

            care_queue_val = self.project_data['details'].get('care_queue', '') or self.care_queue or 'Не указана'
            characteristics_val = self.project_data['details'].get('characteristics', '') or self.characteristics or 'Не указана'
            technology_val = self.project_data['details'].get('technology', '') or self.technology or 'Не указана'
            care_date_val = self.project_data['details'].get('care_date', '') or self.care_date or 'Не указана'
            forest_purpose_val = self.project_data['details'].get('forest_purpose', '') or self.forest_purpose or 'Не указано'

            details_text = MDLabel(
                text=f"Очередь рубки: {care_queue_val}\n"
                     f"Дата рубки: {care_date_val}\n"
                     f"Назначение лесов: {forest_purpose_val}\n"
                     f"Характеристика: {characteristics_val[:80]}{'...' if len(characteristics_val) > 80 else ''}\n"
                     f"Технология: {technology_val[:60]}{'...' if len(technology_val) > 60 else ''}",
                theme_text_color='Custom',
                text_color=[1, 1, 1, 1],
                adaptive_height=True
            )
            details_inner.add_widget(details_text)
            details_block.add_widget(details_inner)
            results_layout.add_widget(details_block)

            # ============================================================
            # БОКС 3: КОЭФФИЦИЕНТ СОСТАВА (#F3E5F5)
            # ============================================================
            composition_block = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, adaptive_height=True, padding=Spacing.MD)
            composition_title = MDLabel(
                text='КОЭФФИЦИЕНТ СОСТАВА',
                theme_text_color='Custom',
                text_color=Colors.GREEN,
                font_style='Title',
                role='medium',
                adaptive_height=True
            )
            composition_block.add_widget(composition_title)

            composition_inner = MDBoxLayout(orientation='vertical', spacing=5, adaptive_height=True, padding=[10, 10])

            total_densities = {}
            total_density_all = 0  # Общая густота всех пород (сумма средних)
            
            for breed_name, data in breeds_data.items():
                if data['plots']:
                    # Рассчитываем СРЕДНЮЮ густоту породы (сумма плотностей / кол-во площадок)
                    # Плотности уже в шт/га, делим на 1000 для перевода в тыс.шт/га
                    avg_density = sum(p.get('density', 0) for p in data['plots']) / len(data['plots'])
                    avg_density_thousands = avg_density / 1000  # переводим в тыс.шт/га
                    total_density_all += avg_density_thousands
                    if avg_density_thousands > 0:
                        total_densities[breed_name] = avg_density_thousands

            if total_densities:
                # Метод наибольшего остатка для распределения коэффициентов
                # Сначала рассчитываем точные коэффициенты
                exact_coeffs = []
                for breed_name, density in sorted(total_densities.items(), key=lambda x: x[1], reverse=True):
                    exact_coeff = (density / total_density_all * 10) if total_density_all > 0 else 1
                    exact_coeffs.append(exact_coeff)
                
                # Округляем вниз
                coeffs_floor = [int(coeff) for coeff in exact_coeffs]
                # Считаем остатки
                remainders = [(exact_coeffs[i] - coeffs_floor[i], i) for i in range(len(exact_coeffs))]
                # Сортируем по убыванию остатков
                remainders.sort(reverse=True)
                
                # Распределяем единицы начиная с наибольших остатков
                coeffs = coeffs_floor[:]
                total = sum(coeffs)
                i = 0
                while total < 10 and i < len(remainders):
                    idx = remainders[i][1]
                    coeffs[idx] += 1
                    total += 1
                    i += 1
                
                # Формируем формулу состава
                sorted_breeds = sorted(total_densities.items(), key=lambda x: x[1], reverse=True)
                composition_parts = []
                for i, (breed_name, _) in enumerate(sorted_breeds):
                    if i < len(coeffs):
                        breed_letter = self.get_breed_letter(breed_name)
                        composition_parts.append(f"{coeffs[i]}{breed_letter}")

                composition_text = ''.join(composition_parts) + "Др"
                composition_label = MDLabel(
                    text=f"Формула состава: {composition_text}",
                    theme_text_color='Custom',
                    text_color=[1, 1, 1, 1],
                    adaptive_height=True
                )
                composition_inner.add_widget(composition_label)
            else:
                no_composition = MDLabel(
                    text="Коэффициент состава не определен (недостаточно данных)",
                    theme_text_color='Custom',
                    text_color=[1, 0.3, 0.3, 1],
                    adaptive_height=True
                )
                composition_inner.add_widget(no_composition)

            composition_block.add_widget(composition_inner)
            results_layout.add_widget(composition_block)

            # ============================================================
            # БОКС 4: ХВОЙНЫЕ ПОРОДЫ (#E8F5E9)
            # ============================================================
            coniferous_block = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, adaptive_height=True, padding=Spacing.MD)
            coniferous_title = MDLabel(
                text='ХВОЙНЫЕ ПОРОДЫ',
                theme_text_color='Custom',
                text_color=Colors.GREEN,
                font_style='Title',
                role='medium',
                adaptive_height=True
            )
            coniferous_block.add_widget(coniferous_title)

            coniferous_inner = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, adaptive_height=True, padding=[5, 5])

            has_coniferous = False
            for breed_name, data in sorted(breeds_data.items()):
                if data['type'] == 'coniferous' and data['plots']:
                    has_coniferous = True
                    zones = data.get('coniferous_zones', {})
                    plot_area_ha = 3.14159 * (float(self.current_radius) if self.current_radius else 1.78) ** 2 / 10000
                    total_area_ha = plot_area_ha * total_plots_count
                    total_do_05 = sum(p.get('do_05', 0) for p in data['plots'])
                    total_05_15 = sum(p.get('05_15', 0) for p in data['plots'])
                    total_bolee_15 = sum(p.get('bolee_15', 0) for p in data['plots'])
                    avg_do_05 = total_do_05 / total_area_ha if total_area_ha > 0 else 0
                    avg_05_15 = total_05_15 / total_area_ha if total_area_ha > 0 else 0
                    avg_bolee_15 = total_bolee_15 / total_area_ha if total_area_ha > 0 else 0

                    avg_heights_over_15 = [p['height'] for p in data['plots'] if p['height'] > 1.5]
                    avg_height_total = sum(avg_heights_over_15) / len(avg_heights_over_15) if avg_heights_over_15 else 0
                    avg_diameters = [d for d in data['diameters'] if d > 0]
                    avg_diameter = sum(avg_diameters) / len(avg_diameters) if avg_diameters else 0
                    avg_ages = [p['age'] for p in data['plots'] if p['age'] > 0]
                    avg_age = sum(avg_ages) / len(avg_ages) if avg_ages else 0

                    breed_label = MDLabel(
                        text=f"{breed_name}:\nдо 0.5м: {avg_do_05:.1f} шт/га | 0.5-1.5м: {avg_05_15:.1f} шт/га | >1.5м: {avg_bolee_15:.1f} шт/га\nСр. высота: {avg_height_total:.1f}м | Возраст: {avg_age:.1f} лет | Диаметр: {avg_diameter:.1f} см",
                        theme_text_color='Custom',
                        text_color=Colors.GREEN,
                        adaptive_height=True
                    )
                    coniferous_inner.add_widget(breed_label)

            if not has_coniferous:
                no_coniferous = MDLabel(
                    text="Хвойные породы не найдены",
                    theme_text_color='Custom',
                    text_color=[0.5, 0.5, 0.5, 1],
                    adaptive_height=True
                )
                coniferous_inner.add_widget(no_coniferous)

            coniferous_block.add_widget(coniferous_inner)
            results_layout.add_widget(coniferous_block)

            # ============================================================
            # БОКС 5: ЛИСТВЕННЫЕ ПОРОДЫ (#E3F2FD)
            # ============================================================
            deciduous_block = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, adaptive_height=True, padding=Spacing.MD)
            deciduous_title = MDLabel(
                text='ЛИСТВЕННЫЕ ПОРОДЫ',
                theme_text_color='Custom',
                text_color=Colors.GREEN,
                font_style='Title',
                role='medium',
                adaptive_height=True
            )
            deciduous_block.add_widget(deciduous_title)

            deciduous_inner = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, adaptive_height=True, padding=[5, 5])

            has_deciduous = False
            for breed_name, data in sorted(breeds_data.items()):
                if data['type'] == 'deciduous' and data['plots']:
                    has_deciduous = True
                    total_trees = sum(p.get('density_raw', 0) for p in data['plots'])
                    plot_area_ha = 3.14159 * (float(self.current_radius) if self.current_radius else 1.78) ** 2 / 10000
                    total_area_ha = plot_area_ha * total_plots_count
                    avg_density = total_trees / total_area_ha if total_area_ha > 0 else 0
                    avg_heights = [p['height'] for p in data['plots'] if p['height'] > 0]
                    avg_height = sum(avg_heights) / len(avg_heights) if avg_heights else 0
                    avg_ages = [p['age'] for p in data['plots'] if p['age'] > 0]
                    avg_age = sum(avg_ages) / len(avg_ages) if avg_ages else 0
                    avg_diameters = [d for d in data['diameters'] if d > 0]
                    avg_diameter = sum(avg_diameters) / len(avg_diameters) if avg_diameters else 0

                    breed_label = MDLabel(
                        text=f"{breed_name}:\nГустота: {avg_density:.1f} шт/га | Высота: {avg_height:.1f}м | Возраст: {avg_age:.1f} лет | Диаметр: {avg_diameter:.1f} см",
                        theme_text_color='Custom',
                        text_color=[1, 1, 1, 1],
                        adaptive_height=True
                    )
                    deciduous_inner.add_widget(breed_label)

            if not has_deciduous:
                no_deciduous = MDLabel(
                    text="Лиственные породы не найдены",
                    theme_text_color='Custom',
                    text_color=[0.5, 0.5, 0.5, 1],
                    adaptive_height=True
                )
                deciduous_inner.add_widget(no_deciduous)

            deciduous_block.add_widget(deciduous_inner)
            results_layout.add_widget(deciduous_block)

            # ============================================================
            # БОКС 6: СРЕДНИЕ ДАННЫЕ (#FFF3E0)
            # ============================================================
            overall_block = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, adaptive_height=True, padding=Spacing.MD)
            overall_title = MDLabel(
                text='СРЕДНИЕ ДАННЫЕ',
                theme_text_color='Custom',
                text_color=Colors.GREEN,
                font_style='Title',
                role='medium',
                adaptive_height=True
            )
            overall_block.add_widget(overall_title)

            overall_inner = MDBoxLayout(orientation='vertical', spacing=5, adaptive_height=True, padding=[10, 10])

            all_densities = []
            all_heights = []
            all_ages = []
            all_diameters = []

            total_trees_all = 0
            for breed_name, data in breeds_data.items():
                if data['plots']:
                    total_trees_all += sum(p.get('density_raw', 0) for p in data['plots'])
                    all_heights.extend([p['height'] for p in data['plots'] if p['height'] > 0])
                    all_ages.extend([p['age'] for p in data['plots'] if p['age'] > 0])
                    all_diameters.extend([d for d in data['diameters'] if d > 0])

            plot_area_ha = 3.14159 * (float(self.current_radius) if self.current_radius else 1.78) ** 2 / 10000
            total_area_ha = plot_area_ha * total_plots_count
            avg_overall_density = total_trees_all / total_area_ha if total_area_ha > 0 else 0
            avg_overall_height = sum(all_heights) / len(all_heights) if all_heights else 0
            avg_overall_age = sum(all_ages) / len(all_ages) if all_ages else 0
            avg_overall_diameter = sum(all_diameters) / len(all_diameters) if all_diameters else 0

            plot_area_from_address = self.project_data['address'].get('plot_area', '')
            total_area_ha_input = float(plot_area_from_address) if plot_area_from_address else total_area_ha
            total_density_entire = avg_overall_density * total_area_ha_input if total_area_ha_input > 0 else 0

            overall_text = MDLabel(
                text=f"Средняя густота: {avg_overall_density:.1f} шт/га\n"
                     f"Общая густота на всю площадь ({total_area_ha_input:.2f} га): {total_density_entire:.0f} шт\n"
                     f"Средняя высота: {avg_overall_height:.1f} м\n"
                     f"Средний возраст: {avg_overall_age:.1f} лет\n"
                     f"Средний диаметр: {avg_overall_diameter:.1f} см",
                theme_text_color='Custom',
                text_color=[1, 1, 1, 1],
                adaptive_height=True
            )
            overall_inner.add_widget(overall_text)
            overall_block.add_widget(overall_inner)
            results_layout.add_widget(overall_block)

            # ============================================================
            # БОКС 7: ПРЕДМЕТ УХОДА (#FCE4EC)
            # ============================================================
            care_block = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, adaptive_height=True, padding=Spacing.MD)
            care_title = MDLabel(
                text='ПРЕДМЕТ УХОДА',
                theme_text_color='Custom',
                text_color=Colors.GREEN,
                font_style='Title',
                role='medium',
                adaptive_height=True
            )
            care_block.add_widget(care_title)

            care_inner = MDBoxLayout(orientation='vertical', spacing=5, adaptive_height=True, padding=[10, 10])

            care_data = []
            total_density_all = 0
            plot_count_all = 0
            total_remaining_density = 0
            plot_count_with_care = 0

            for page_num, page_rows in self.page_data.items():
                for row in page_rows:
                    if len(row) >= 4 and row[3]:
                        plot_density = 0
                        breeds_text = row[3]
                        if breeds_text:
                            try:
                                breeds_list = json.loads(breeds_text) if isinstance(breeds_text, str) else []
                                for breed_info in breeds_list:
                                    if isinstance(breed_info, dict):
                                        if breed_info.get('type') == 'coniferous':
                                            do_05 = breed_info.get('do_05', 0)
                                            _05_15 = breed_info.get('05_15', 0)
                                            bolee_15 = breed_info.get('bolee_15', 0)
                                            plot_density += (do_05 + _05_15 + bolee_15) / plot_area_ha_default if plot_area_ha_default > 0 else 0
                                        else:
                                            density = breed_info.get('density', 0)
                                            plot_density += density / plot_area_ha_default if plot_area_ha_default > 0 else 0
                            except (json.JSONDecodeError, TypeError):
                                pass

                        if plot_density > 0:
                            total_density_all += plot_density
                            plot_count_all += 1

                    if len(row) >= 4 and row[2]:
                        care_text = row[2].strip()
                        if care_text:
                            plot_density = 0
                            breeds_text = row[3]
                            if breeds_text:
                                try:
                                    breeds_list = json.loads(breeds_text) if isinstance(breeds_text, str) else []
                                    for breed_info in breeds_list:
                                        if isinstance(breed_info, dict):
                                            if breed_info.get('type') == 'coniferous':
                                                do_05 = breed_info.get('do_05', 0)
                                                _05_15 = breed_info.get('05_15', 0)
                                                bolee_15 = breed_info.get('bolee_15', 0)
                                                plot_density += (do_05 + _05_15 + bolee_15) / plot_area_ha_default if plot_area_ha_default > 0 else 0
                                            else:
                                                density = breed_info.get('density', 0)
                                                plot_density += density / plot_area_ha_default if plot_area_ha_default > 0 else 0
                                except (json.JSONDecodeError, TypeError):
                                    pass

                            if plot_density > 0:
                                remaining_density = self.parse_care_subject_density(care_text)
                                if remaining_density > 0:
                                    care_data.append({
                                        'care_text': care_text,
                                        'plot_density': plot_density,
                                        'remaining_density': remaining_density
                                    })
                                    total_remaining_density += remaining_density
                                    plot_count_with_care += 1

            if care_data:
                care_breed_totals = {}
                care_plot_count = 0
                for item in care_data:
                    care_text = item['care_text']
                    breed_densities = self.parse_care_subject_by_breeds(care_text)
                    for breed, density in breed_densities.items():
                        if breed not in care_breed_totals:
                            care_breed_totals[breed] = 0
                        care_breed_totals[breed] += density
                    care_plot_count += 1

                if care_breed_totals and care_plot_count > 0:
                    avg_care_parts = []
                    short_parts = []
                    for breed, total_density in sorted(care_breed_totals.items()):
                        avg_density = total_density / care_plot_count
                        avg_care_parts.append(f"{avg_density * 1000:.0f}шт/га{breed}")
                        short_parts.append(f"{avg_density:.1f}{breed}")
                    avg_care_text = ''.join(avg_care_parts)
                    short_text = ''.join(short_parts).replace('.', ',')

                    care_label = MDLabel(
                        text=f"Средний предмет ухода: {avg_care_text} = {short_text}",
                        theme_text_color='Custom',
                        text_color=[1, 1, 1, 1],
                        adaptive_height=True
                    )
                    care_inner.add_widget(care_label)

            if not care_data:
                no_care = MDLabel(
                    text="Предмет ухода не указан или недостаточно данных",
                    theme_text_color='Custom',
                    text_color=[0.5, 0.5, 0.5, 1],
                    adaptive_height=True
                )
                care_inner.add_widget(no_care)

            care_block.add_widget(care_inner)
            results_layout.add_widget(care_block)

            # ============================================================
            # БОКС 8: ИНТЕНСИВНОСТЬ РУБКИ
            # ============================================================
            intensity_block = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, adaptive_height=True, padding=Spacing.MD)
            intensity_title = MDLabel(
                text='ИНТЕНСИВНОСТЬ РУБКИ',
                theme_text_color='Custom',
                text_color=Colors.GREEN,
                font_style='Title',
                role='medium',
                adaptive_height=True
            )
            intensity_block.add_widget(intensity_title)

            intensity_inner = MDBoxLayout(orientation='vertical', spacing=5, adaptive_height=True, padding=[10, 10])

            if plot_count_with_care > 0 and plot_count_all > 0 and avg_overall_density > 0:
                avg_remaining_density = total_remaining_density / plot_count_with_care
                intensity = ((avg_overall_density - avg_remaining_density) / avg_overall_density) * 100

                intensity_label = MDLabel(
                    text=f"Интенсивность рубки: {intensity:.1f}%\n(было {avg_overall_density:.0f} шт/га, останется {avg_remaining_density:.0f} шт/га)",
                    theme_text_color='Custom',
                    text_color=[1, 1, 1, 1],
                    adaptive_height=True
                )
                intensity_inner.add_widget(intensity_label)
            else:
                no_intensity = MDLabel(
                    text="Недостаточно данных для расчета интенсивности",
                    theme_text_color='Custom',
                    text_color=[0.5, 0.5, 0.5, 1],
                    adaptive_height=True
                )
                intensity_inner.add_widget(no_intensity)

            intensity_block.add_widget(intensity_inner)
            results_layout.add_widget(intensity_block)

            # ============================================================
            # БОКС 9: ИНФОРМАЦИЯ О ПЛОЩАДИ
            # ============================================================
            plot_area_block = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, adaptive_height=True, padding=Spacing.MD)
            plot_area_title = MDLabel(
                text='ИНФОРМАЦИЯ О ПЛОЩАДИ',
                theme_text_color='Custom',
                text_color=Colors.GREEN,
                font_style='Title',
                role='medium',
                adaptive_height=True
            )
            plot_area_block.add_widget(plot_area_title)

            plot_area_inner = MDBoxLayout(orientation='vertical', spacing=5, adaptive_height=True, padding=[10, 10])

            plot_count = len([row for page in self.page_data.values() for row in page if any(cell for cell in row[:3] if cell)])
            total_plot_area_ha = plot_count * plot_area_ha_default

            plot_area_label = MDLabel(
                text=f"Радиус пробной площади: {default_radius:.2f} м\n"
                     f"Площадь одной площадки: {plot_area_ha_default:.4f} га ({plot_area_m2_default:.2f} м²)\n"
                     f"Всего площадок: {plot_count}\n"
                     f"Совокупная площадь перечета: {total_plot_area_ha:.4f} га ({total_plot_area_ha*10000:.0f} м²)",
                theme_text_color='Custom',
                text_color=[1, 1, 1, 1],
                adaptive_height=True
            )
            plot_area_inner.add_widget(plot_area_label)
            plot_area_block.add_widget(plot_area_inner)
            results_layout.add_widget(plot_area_block)

            # ============================================================
            # БОКС 10: ТИП ЛЕСА
            # ============================================================
            forest_type_block = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, adaptive_height=True, padding=Spacing.MD)
            forest_type_title = MDLabel(
                text='ТИП ЛЕСА',
                theme_text_color='Custom',
                text_color=Colors.GREEN,
                font_style='Title',
                role='medium',
                adaptive_height=True
            )
            forest_type_block.add_widget(forest_type_title)

            forest_type_inner = MDBoxLayout(orientation='vertical', spacing=5, adaptive_height=True, padding=[10, 10])

            if forest_types_set:
                forest_types_text = ', '.join(sorted(forest_types_set))
                forest_type_label = MDLabel(
                    text=f"Встречающиеся типы леса:\n{forest_types_text}",
                    theme_text_color='Custom',
                    text_color=[1, 1, 1, 1],
                    adaptive_height=True
                )
                forest_type_inner.add_widget(forest_type_label)
            else:
                no_forest_type = MDLabel(
                    text="Тип леса не указан",
                    theme_text_color='Custom',
                    text_color=[0.5, 0.5, 0.5, 1],
                    adaptive_height=True
                )
                forest_type_inner.add_widget(no_forest_type)

            forest_type_block.add_widget(forest_type_inner)
            results_layout.add_widget(forest_type_block)

            scroll.add_widget(results_layout)
            content.add_widget(scroll)

            close_btn = MDButton(style='outlined', line_color=Colors.DANGER, size_hint=(1, None), height=50)
            close_btn.add_widget(MDButtonText(text='Закрыть', theme_text_color='Custom', text_color=Colors.DANGER))
            content.add_widget(close_btn)

            popup = Popup(
                title='',
                content=content,
                size_hint=(0.98, 0.98),
                separator_height=0,
                background_color=[0,0,0,0.3],
                overlay_color=[0,0,0,0.3]
            )

            close_btn.bind(on_press=popup.dismiss)
            popup.open()

        except Exception as e:
            import traceback
            self.show_error(f"Ошибка расчета таксационных показателей: {str(e)}\n{traceback.format_exc()}")

    def update_totals(self, update_global=True):
        """Обновление строки итогов с поддержкой множественных пород"""
        breed_composition = {}  # Initialize at top
        total_stats = {'density': [], 'height': [], 'age': []}  # Initialize at top

        # Calculate radius and area once
        current_radius = float(self.current_radius) if self.current_radius else 5.64
        plot_area_m2 = 3.14159 * (current_radius ** 2)
        if update_global:
            section_data = self.calculate_section_totals()
            composition_text = section_data['composition_text']
            forestry_formulas_text = section_data['forestry_formulas_text']
            total_plots = section_data['total_plots']
        else:
            # Старая логика по странице
            breed_composition = {}
            total_stats = {'density': [], 'height': [], 'age': []}
            coniferous_stats = {'do_05': [], '05_15': [], 'bolee_15': [], 'height': [], 'age': []}

            for row in self.inputs:
                predmet_text = row[2].text
                if predmet_text:
                    composition = self.parse_composition(predmet_text)
                    for breed, count in composition.items():
                        if breed not in breed_composition:
                            breed_composition[breed] = []
                        breed_composition[breed].append(count)

                breeds_text = row[3].text
                if breeds_text:
                    breeds_data = self.parse_breeds_data(breeds_text)
                    for breed_info in breeds_data:
                        if breed_info.get('type') == 'coniferous':
                            # Густота хвойных = сумма градаций
                            coniferous_density = (breed_info.get('do_05', 0) +
                                                breed_info.get('05_15', 0) +
                                                breed_info.get('bolee_15', 0))
                            if coniferous_density > 0:
                                total_stats['density'].append(coniferous_density)
                            else:
                                if 'density' in breed_info and breed_info['density']:
                                    total_stats['density'].append(breed_info['density'])

                            # Сбор данных по градациям для хвойных
                            if breed_info.get('do_05', 0) > 0:
                                coniferous_stats['do_05'].append(breed_info['do_05'])
                            if breed_info.get('05_15', 0) > 0:
                                coniferous_stats['05_15'].append(breed_info['05_15'])
                            if breed_info.get('bolee_15', 0) > 0:
                                coniferous_stats['bolee_15'].append(breed_info['bolee_15'])
                            if breed_info.get('height', 0) > 0:
                                coniferous_stats['height'].append(breed_info['height'])
                            if breed_info.get('age', 0) > 0:
                                coniferous_stats['age'].append(breed_info['age'])
                        elif 'density' in breed_info and breed_info['density']:
                            total_stats['density'].append(breed_info['density'])

                        if 'height' in breed_info and breed_info['height']:
                            total_stats['height'].append(breed_info['height'])
                        if 'age' in breed_info and breed_info['age']:
                            total_stats['age'].append(breed_info['age'])

            # Рассчитываем итоги по странице

            current_radius = float(self.current_radius) if self.current_radius else 1.78
            plot_area_m2 = 3.14159 * (current_radius ** 2)  # Площадь пробной площади в м²

            # Расчет средних по градациям для хвойных по формулам лесного хозяйства на гектар
            coniferous_stats_ha = []
            for row in range(len(self.inputs)):
                row_do_05 = coniferous_stats['do_05'][row] if row < len(coniferous_stats['do_05']) and row < len(coniferous_stats['do_05']) else 0
                row_05_15 = coniferous_stats['05_15'][row] if row < len(coniferous_stats['05_15']) else 0
                row_bolee_15 = coniferous_stats['bolee_15'][row] if row < len(coniferous_stats['bolee_15']) else 0
                row_height = coniferous_stats['height'][row] if row < len(coniferous_stats['height']) else 0
                row_age = coniferous_stats['age'][row] if row < len(coniferous_stats['age']) else 0

                # Рассчитываем густоту на гектар для градаций
                do_05_ha = (row_do_05 * 10000) / plot_area_m2 if plot_area_m2 > 0 else 0
                _05_15_ha = (row_05_15 * 10000) / plot_area_m2 if plot_area_m2 > 0 else 0
                bolee_15_ha = (row_bolee_15 * 10000) / plot_area_m2 if plot_area_m2 > 0 else 0

                coniferous_stats_ha.append({
                    'do_05_ha': do_05_ha,
                    '05_15_ha': _05_15_ha,
                    'bolee_15_ha': bolee_15_ha,
                    'height': row_height,
                    'age': row_age
                })

            avg_coniferous_do_05_ha = sum(d['do_05_ha'] for d in coniferous_stats_ha) / len(coniferous_stats_ha) if coniferous_stats_ha else 0
            avg_coniferous_05_15_ha = sum(d['05_15_ha'] for d in coniferous_stats_ha) / len(coniferous_stats_ha) if coniferous_stats_ha else 0
            avg_coniferous_bolee_15_ha = sum(d['bolee_15_ha'] for d in coniferous_stats_ha) / len(coniferous_stats_ha) if coniferous_stats_ha else 0
            avg_coniferous_height_ha = sum(d['height'] for d in coniferous_stats_ha) / len(coniferous_stats_ha) if coniferous_stats_ha else 0
            avg_coniferous_age_ha = sum(d['age'] for d in coniferous_stats_ha) / len(coniferous_stats_ha) if coniferous_stats_ha else 0

            # Формирование текста для столбца Порода в строке итогов с формулами лесного хозяйства
            forestry_formulas_text = ""

            # Хвойные породы - средние значения по градациям на га
            if coniferous_stats_ha:
                forestry_formulas_text += "Хвойные: "
                gradiations = []
                if avg_coniferous_do_05_ha > 0:
                    gradiations.append(f"до 0.5м: {avg_coniferous_do_05_ha:.1f} шт/га")
                if avg_coniferous_05_15_ha > 0:
                    gradiations.append(f"0.5-1.5м: {avg_coniferous_05_15_ha:.1f} шт/га")
                if avg_coniferous_bolee_15_ha > 0:
                    gradiations.append(f">1.5м: {avg_coniferous_bolee_15_ha:.1f} шт/га")
                if gradiations:
                    forestry_formulas_text += ", ".join(gradiations)
                if avg_coniferous_height_ha > 0:
                    forestry_formulas_text += f", высота: {avg_coniferous_height_ha:.1f}м"
                if avg_coniferous_age_ha > 0:
                    forestry_formulas_text += f", возраст: {avg_coniferous_age_ha:.1f} лет"

            # Лиственные породы - средние значения без градаций
            deciduous_density = []
            deciduous_height = []
            deciduous_age = []

            for row in self.inputs:
                breeds_text = row[3].text
                if breeds_text:
                    breeds_data = self.parse_breeds_data(breeds_text)
                    for breed_info in breeds_data:
                        if breed_info.get('type') == 'deciduous':
                            if breed_info.get('density'):
                                deciduous_density.append(breed_info['density'] * (10000 / plot_area_m2) if plot_area_m2 > 0 else breed_info['density'])
                            if breed_info.get('height'):
                                deciduous_height.append(breed_info['height'])
                            if breed_info.get('age'):
                                deciduous_age.append(breed_info['age'])

            # Рассчитываем средние по лиственным на га
            if deciduous_density or deciduous_height or deciduous_age:
                avg_deciduous_density = sum(deciduous_density) / len(deciduous_density) if deciduous_density else 0
                avg_deciduous_height = sum(deciduous_height) / len(deciduous_height) if deciduous_height else 0
                avg_deciduous_age = sum(deciduous_age) / len(deciduous_age) if deciduous_age else 0

                if forestry_formulas_text:
                    forestry_formulas_text += "; "
                forestry_formulas_text += "Лиственные: "
                parts = []
                if avg_deciduous_density > 0:
                    parts.append(f"густота: {avg_deciduous_density:.1f} шт/га")
                if avg_deciduous_height > 0:
                    parts.append(f"высота: {avg_deciduous_height:.1f}м")
                if avg_deciduous_age > 0:
                    parts.append(f"возраст: {avg_deciduous_age:.1f} лет")
                forestry_formulas_text += ", ".join(parts)

        avg_composition = {}
        for breed, counts in breed_composition.items():
            if counts:
                avg_composition[breed] = sum(counts) / len(counts)

        composition_text = ""
        for breed in sorted(avg_composition.keys()):
            count = avg_composition[breed]
            if count > 0:
                composition_text += f"{int(count)}{breed}"

    def parse_composition(self, text):
        """Парсит текстовое представление состава пород"""
        composition = {}
        if isinstance(text, str):
            matches = re.findall(r'(\d+)([А-ЯA-Z])', text.upper())
            for count, breed in matches:
                try:
                    composition[breed] = int(count)
                except ValueError:
                    pass
        return composition

    def parse_breeds_data(self, breeds_text):
        """Парсит данные пород из текстового поля"""
        if not breeds_text or not isinstance(breeds_text, str):
            return []

        try:
            if isinstance(breeds_text, str) and breeds_text.startswith('['):
                return json.loads(breeds_text)
            elif isinstance(breeds_text, str) and breeds_text.startswith('{'):
                return [json.loads(breeds_text)]
        except (json.JSONDecodeError, TypeError):
            pass

        return []

    def calculate_page_totals(self):
        """Вычисляет итоговые значения для текущей страницы"""
        totals = {
            'composition': '',
            'total_area': 0.0,
            'avg_age': 0.0,
            'avg_density': 0.0,
            'avg_height': 0.0
        }

        breed_composition = {}
        total_stats = {'density': [], 'height': [], 'age': []}
        total_area = 0.0

        for row in self.inputs:
            predmet_text = row[2].text
            if predmet_text:
                composition = self.parse_composition(predmet_text)
                for breed, count in composition.items():
                    if breed not in breed_composition:
                        breed_composition[breed] = []
                    breed_composition[breed].append(count)

            radius = 5.64
            try:
                if row[5].text:
                    radius = float(row[5].text)
            except (ValueError, IndexError):
                pass

            area = 3.14159 * (radius ** 2)
            total_area += area

            breeds_text = row[3].text
            if breeds_text:
                breeds_data = self.parse_breeds_data(breeds_text)
                for breed_info in breeds_data:
                    if breed_info.get('type') == 'coniferous':
                        coniferous_density = (breed_info.get('do_05', 0) +
                                            breed_info.get('05_15', 0) +
                                            breed_info.get('bolee_15', 0))
                        if coniferous_density > 0:
                            total_stats['density'].append(coniferous_density)
                    elif 'density' in breed_info and breed_info['density']:
                        total_stats['density'].append(breed_info['density'])

                    if 'height' in breed_info and breed_info['height']:
                        total_stats['height'].append(breed_info['height'])
                    if 'age' in breed_info and breed_info['age']:
                        total_stats['age'].append(breed_info['age'])

        avg_composition = {}
        for breed, counts in breed_composition.items():
            if counts:
                avg_composition[breed] = sum(counts) / len(counts)

        composition_text = ""
        for breed in sorted(avg_composition.keys()):
            count = avg_composition[breed]
            if count > 0:
                composition_text += f"{int(count)}{breed}"

        totals['composition'] = composition_text
        totals['total_area'] = total_area
        totals['avg_density'] = sum(total_stats['density']) / len(total_stats['density']) if total_stats['density'] else 0.0
        totals['avg_height'] = sum(total_stats['height']) / len(total_stats['height']) if total_stats['height'] else 0.0
        totals['avg_age'] = sum(total_stats['age']) / len(total_stats['age']) if total_stats['age'] else 0.0

        return totals

    def save_current_page(self, instance=None):
        """Сохраняем текущую страницу в базу данных"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()

        try:
            cursor.execute('''
                DELETE FROM molodniki_data
                WHERE page_number = ? AND section_name = ?
            ''', (self.current_page, self.current_section))

            # Сохраняем данные из page_data
            if self.current_page in self.page_data:
                for row_idx, row_data in enumerate(self.page_data[self.current_page]):
                    if any(row_data[:5]):
                        radius = 5.64
                        try:
                            if row_data[5]:
                                radius = float(row_data[5])
                        except (ValueError, IndexError):
                            pass

                        cursor.execute('''
                            INSERT INTO molodniki_data
                            (page_number, row_index, nn, gps_point, predmet_uhoda, poroda, primechanie, radius, section_name)
                            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                        ''', (
                            self.current_page,
                            row_idx,
                            row_data[0] or None,
                            row_data[1] or None,
                            row_data[2] or None,
                            row_data[3] or None,
                            row_data[4] or None,
                            radius,
                            self.current_section
                        ))

                        molodniki_data_id = cursor.lastrowid

                        if row_data[3]:
                            breeds_data = self.parse_breeds_data(row_data[3])
                            for breed_info in breeds_data:
                                try:
                                    # Validate and convert data types
                                    density = int(breed_info.get('density', 0) or 0)
                                    height = float(breed_info.get('height', 0.0) or 0.0)
                                    age = int(breed_info.get('age', 0) or 0)
                                    do_05 = int(breed_info.get('do_05', 0) or 0)
                                    _05_15 = int(breed_info.get('05_15', 0) or 0)
                                    bolee_15 = int(breed_info.get('bolee_15', 0) or 0)

                                    composition_coeff = 0.0
                                    if density and radius:
                                        area = 3.14159 * (radius ** 2)
                                        composition_coeff = (density * area) / 10000

                                    diameter = float(breed_info.get('diameter', 0.0) or 0.0)

                                    cursor.execute('''
                                        INSERT INTO molodniki_breeds
                                        (molodniki_data_id, breed_name, breed_type, do_05, _05_15, bolee_15,
                                         density, height, diameter, age, composition_coefficient)
                                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                                    ''', (
                                        molodniki_data_id,
                                        breed_info.get('name', ''),
                                        breed_info.get('type', 'deciduous'),
                                        do_05,
                                        _05_15,
                                        bolee_15,
                                        density,
                                        height,
                                        diameter,
                                        age,
                                        composition_coeff
                                    ))
                                except Exception as e:
                                    print(f"Error inserting breed: {e}, skipping this breed")
                                    continue

            conn.commit()
            self.show_success("Страница сохранена в базу данных!")
            success = True

        except Exception as e:
            conn.rollback()
            self.show_error(f"Ошибка сохранения: {str(e)}")
            success = False
        finally:
            conn.close()

        return success

    def show_save_dialog(self, instance=None):
        content = MDBoxLayout(orientation='vertical', spacing=10, padding=10, md_bg_color=Colors.DARK_SURFACE)

        title_label = MDLabel(
            text="Введите имя файла",
            bold=True,
            size_hint=(1, None),
            height=30,
            theme_text_color='Custom',
            text_color=Colors.TEXT_ON_DARK
        )
        content.add_widget(title_label)

        self.filename_input = MDTextField(
            hint_text="Имя файла",
            mode="outlined",
            size_hint=(1, None),
            height=40,
            line_color_focus=Colors.GREEN
        )
        timestamp = datetime.datetime.now().strftime('%M%S')  # Только минуты и секунды
        document_name = self.project_data.get('document_name', 'Проект')
        # Очень короткое имя файла на основе названия проекта (макс 10 символов)
        short_name = document_name.replace(' ', '').replace('/', '_').replace('.', '')[:10]
        default_name = f"{short_name}_{self.current_section}_{timestamp}"
        self.filename_input.text = default_name
        content.add_widget(self.filename_input)

        btn_layout = MDBoxLayout(orientation='horizontal', spacing=10, size_hint=(1, None), height=50, md_bg_color=Colors.DARK_SURFACE)
        save_btn = MDButton(style='filled', size_hint=(0.5, 1), height=50)
        save_btn.add_widget(MDButtonText(text='Сохранить', theme_text_color='Custom', text_color=Colors.GREEN))
        cancel_btn = MDButton(style='filled', size_hint=(0.5, 1), height=50)
        cancel_btn.add_widget(MDButtonText(text='Отмена', theme_text_color='Custom', text_color=[1,0.3,0.3,1]))
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        self.save_popup = Popup(
            title="",
            content=content,
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3],
            size_hint=(0.7, 0.5)
        )
        save_btn.bind(on_release=self.save_to_excel)
        cancel_btn.bind(on_release=self.save_popup.dismiss)
        self.save_popup.open()

    def save_to_excel(self, instance):
        filename = self.filename_input.text.strip()
        if not filename:
            self.show_error("Имя файла не может быть пустым!")
            return
        filename = re.sub(r'[\\/*?:"<>|]', "", filename)
        filename = f"{filename}.xlsx" if not filename.endswith(".xlsx") else filename
        full_path = os.path.join(self.reports_dir, filename)

        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Молодняки"

            address_parts = []
            if self.current_quarter:
                address_parts.append(f"Квартал: {self.current_quarter}")
            if self.current_plot:
                address_parts.append(f"Выдел: {self.current_plot}")
            if self.current_forestry:
                address_parts.append(f"Лесничество: {self.current_forestry}")
            if self.current_radius:
                address_parts.append(f"Радиус: {self.current_radius} м")

            address_text = " | ".join(address_parts) if address_parts else "Адрес не указан"
            ws['A1'] = f"Адрес: {address_text}"
            ws['A1'].font = openpyxl.styles.Font(bold=True, size=12)

            ws.append([])

            headers = [
                '№ППР', 'GPS точка', 'Предмет ухода', 'Порода', 'Густота', 'Высота', 'Возраст', 'Примечания', 'Тип Леса'
            ]
            for col_num, header in enumerate(headers, 1):
                cell = ws.cell(row=3, column=col_num, value=header)
                cell.font = openpyxl.styles.Font(bold=True)
                cell.fill = openpyxl.styles.PatternFill(start_color="00FF00", end_color="00FF00", fill_type="solid")

            all_data = []
            for page in sorted(self.page_data.keys()):
                all_data.extend(self.page_data[page])

            current_row = 4
            for row in all_data:
                if any(cell for cell in row[:3] if cell):  # Проверяем, что основные столбцы не пустые
                    try:
                        breeds_data = json.loads(row[3]) if row[3] else []
                    except (json.JSONDecodeError, TypeError):
                        breeds_data = []

                    if isinstance(breeds_data, list) and breeds_data:
                        for breed_info in breeds_data:
                            if isinstance(breed_info, dict):
                                breed_name = breed_info.get('name', 'Неизвестная')
                                density = breed_info.get('density', '')
                                height = breed_info.get('height', '')
                                age = breed_info.get('age', '')

                                # Для хвойных рассчитываем густоту по градациям
                                if breed_info.get('type') == 'coniferous':
                                    coniferous_density = (breed_info.get('do_05', 0) +
                                                        breed_info.get('05_15', 0) +
                                                        breed_info.get('bolee_15', 0))
                                    if coniferous_density > 0:
                                        density = str(coniferous_density)

                                processed_row = [
                                    row[0],  # №ППР
                                    row[1],  # GPS точка
                                    row[2],  # Предмет ухода
                                    breed_name,  # Порода
                                    str(density) if density else '',  # Густота
                                    str(height) if height else '',  # Высота
                                    str(age) if age else '',  # Возраст
                                    row[4],  # Примечания
                                    row[5],  # Тип Леса
                                ]
                                ws.append(processed_row)
                                current_row += 1
                    else:
                        # Если нет пород, добавить строку без данных
                        processed_row = [row[0], row[1], row[2], '', '', '', '', row[4], row[5]]
                        ws.append(processed_row)
                        current_row += 1

            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width

            wb.save(full_path)
            self.save_popup.dismiss()
            self.show_success(f"Файл сохранен: {filename}")
        except Exception as e:
            self.show_error(f"Ошибка: {str(e)}")

    def save_to_word(self, instance):
        try:
            from docx import Document

            timestamp = datetime.datetime.now().strftime('%M%S')  # Только минуты и секунды
            document_name = self.project_data.get('document_name', 'Проект')
            # Очень короткое имя файла на основе названия проекта (макс 10 символов)
            short_name = document_name.replace(' ', '').replace('/', '_').replace('.', '')[:10]
            filename = f"{short_name}_{self.current_section}_{timestamp}.docx"
            full_path = os.path.join(self.reports_dir, filename)

            doc = Document()
            doc.add_heading(f'Расширенный отчет по молоднякам - Участок {self.current_section}', 0)

            all_data = []
            for page in sorted(self.page_data.keys()):
                all_data.extend(self.page_data[page])

            table = doc.add_table(rows=1, cols=9)
            table.style = 'Table Grid'

            headers = [
                '№ППР', 'GPS точка', 'Предмет ухода', 'Порода', 'Густота', 'Высота', 'Возраст', 'Примечания', 'Тип Леса'
            ]
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header

            for row in all_data:
                if any(cell for cell in row[:3] if cell):  # Проверяем, что основные столбцы не пустые
                    try:
                        breeds_data = json.loads(row[3]) if row[3] else []
                    except (json.JSONDecodeError, TypeError):
                        breeds_data = []

                    if isinstance(breeds_data, list) and breeds_data:
                        for breed_info in breeds_data:
                            if isinstance(breed_info, dict):
                                breed_name = breed_info.get('name', 'Неизвестная')
                                density = breed_info.get('density', '')
                                height = breed_info.get('height', '')
                                age = breed_info.get('age', '')

                                # Для хвойных рассчитываем густоту по градациям
                                if breed_info.get('type') == 'coniferous':
                                    coniferous_density = (breed_info.get('do_05', 0) +
                                                        breed_info.get('05_15', 0) +
                                                        breed_info.get('bolee_15', 0))
                                    if coniferous_density > 0:
                                        density = str(coniferous_density)

                                row_cells = table.add_row().cells
                                row_cells[0].text = str(row[0]) if row[0] else ""  # №ППР
                                row_cells[1].text = str(row[1]) if row[1] else ""  # GPS точка
                                row_cells[2].text = str(row[2]) if row[2] else ""  # Предмет ухода
                                row_cells[3].text = breed_name  # Порода
                                row_cells[4].text = str(density) if density else ""  # Густота
                                row_cells[5].text = str(height) if height else ""  # Высота
                                row_cells[6].text = str(age) if age else ""  # Возраст
                                row_cells[7].text = str(row[4]) if row[4] else ""  # Примечания
                                row_cells[8].text = str(row[5]) if row[5] else ""  # Тип Леса
                    else:
                        # Если нет пород, добавить строку без данных
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(row[0]) if row[0] else ""
                        row_cells[1].text = str(row[1]) if row[1] else ""
                        row_cells[2].text = str(row[2]) if row[2] else ""
                        row_cells[3].text = ""
                        row_cells[4].text = ""
                        row_cells[5].text = ""
                        row_cells[6].text = ""
                        row_cells[7].text = str(row[4]) if row[4] else ""
                        row_cells[8].text = str(row[5]) if row[5] else ""

            doc.save(full_path)
            self.show_success(f"Word документ сохранен: {filename}")
        except ImportError:
            self.show_error("Для сохранения в Word установите библиотеку python-docx: pip install python-docx")
        except Exception as e:
            self.show_error(f"Ошибка сохранения Word: {str(e)}")

    def aggregate_breeds_data(self, df):
        """Агрегирует данные пород по площадкам из Excel файла с учетом заголовков"""
        # Получаем заголовки из первой строки
        headers = df.iloc[0] if not df.empty else []

        # Находим индексы нужных столбцов
        nn_idx = None
        gps_idx = None
        predmet_idx = None
        breed_name_idx = None
        density_idx = None
        do_05_idx = None
        _05_15_idx = None
        bolee_15_idx = None
        height_idx = None
        age_idx = None
        primechanie_idx = None
        tip_lesa_idx = None

        for i, header in enumerate(headers):
            header_str = str(header).strip().lower()
            if '№ппр' in header_str:
                nn_idx = i
            elif 'gps' in header_str:
                gps_idx = i
            elif 'предмет ухода' in header_str:
                predmet_idx = i
            elif 'порода' in header_str:
                breed_name_idx = i
            elif 'густота' in header_str:
                density_idx = i
            elif 'до 0.5м' in header_str:
                do_05_idx = i
            elif '0.5-1.5м' in header_str:
                _05_15_idx = i
            elif '>1.5м' in header_str or 'выше' in header_str:
                bolee_15_idx = i
            elif 'высота' in header_str:
                height_idx = i
            elif 'возраст' in header_str:
                age_idx = i
            elif 'примечания' in header_str:
                primechanie_idx = i
            elif 'тип леса' in header_str:
                tip_lesa_idx = i

        # Группировка по данным площадки (GPS, Предмет ухода, Примечания, Тип Леса)
        grouped = {}

        print(f"DEBUG: aggregate_breeds_data starting, df shape: {df.shape}")

        # Начинаем с второй строки (после заголовков)
        for index in range(1, len(df)):
            row = df.iloc[index]

            # Извлекаем данные по найденным индексам
            gps = str(row.iloc[gps_idx]) if gps_idx is not None and gps_idx < len(row) else ''
            predmet = str(row.iloc[predmet_idx]) if predmet_idx is not None and predmet_idx < len(row) else ''
            breed_name = str(row.iloc[breed_name_idx]) if breed_name_idx is not None and breed_name_idx < len(row) else ''
            density = str(row.iloc[density_idx]) if density_idx is not None and density_idx < len(row) else ''
            do_05 = str(row.iloc[do_05_idx]) if do_05_idx is not None and do_05_idx < len(row) else ''
            _05_15 = str(row.iloc[_05_15_idx]) if _05_15_idx is not None and _05_15_idx < len(row) else ''
            bolee_15 = str(row.iloc[bolee_15_idx]) if bolee_15_idx is not None and bolee_15_idx < len(row) else ''
            height = str(row.iloc[height_idx]) if height_idx is not None and height_idx < len(row) else ''
            age = str(row.iloc[age_idx]) if age_idx is not None and age_idx < len(row) else ''
            primechanie = str(row.iloc[primechanie_idx]) if primechanie_idx is not None and primechanie_idx < len(row) else ''
            tip_lesa = str(row.iloc[tip_lesa_idx]) if tip_lesa_idx is not None and tip_lesa_idx < len(row) else ''

            print(f"DEBUG: Processing row {index}: breed_name='{breed_name}', primechanie='{primechanie}', tip_lesa='{tip_lesa}'")

            # Ключ группы по уникальной комбинации данных площадки
            key = (str(gps), str(predmet), str(primechanie), str(tip_lesa))

            if key not in grouped:
                grouped[key] = {
                    'gps': gps,
                    'predmet': predmet,
                    'primechanie': primechanie,
                    'tip_lesa': tip_lesa,
                    'breeds': []
                }

            # Пропускаем строки без породы
            if not breed_name.strip() or breed_name in ['nan', 'NaN', '']:
                continue

            # Определяем тип породы
            breed_type = self.determine_breed_type(breed_name)

            # Создаем объект породы
            breed_data = {
                'name': breed_name,
                'type': breed_type
            }

            # Добавляем параметры с проверкой типов
            if density and density not in ['nan', 'NaN', '']:
                try:
                    breed_data['density'] = int(float(density))
                except (ValueError, TypeError):
                    pass

            if do_05 and do_05 not in ['nan', 'NaN', '']:
                try:
                    breed_data['do_05'] = int(float(do_05))
                except (ValueError, TypeError):
                    pass

            if _05_15 and _05_15 not in ['nan', 'NaN', '']:
                try:
                    breed_data['05_15'] = int(float(_05_15))
                except (ValueError, TypeError):
                    pass

            if bolee_15 and bolee_15 not in ['nan', 'NaN', '']:
                try:
                    breed_data['bolee_15'] = int(float(bolee_15))
                except (ValueError, TypeError):
                    pass

            if height and height not in ['nan', 'NaN', '']:
                try:
                    breed_data['height'] = float(height)
                except (ValueError, TypeError):
                    pass

            if age and age not in ['nan', 'NaN', '']:
                try:
                    breed_data['age'] = int(float(age))
                except (ValueError, TypeError):
                    pass

            # Если порода уже есть в списке, добавляем/обновляем параметры
            existing_breed = None
            for b in grouped[key]['breeds']:
                if b['name'] == breed_name:
                    existing_breed = b
                    break

            if existing_breed:
                # Обновляем существующую породу
                for k, v in breed_data.items():
                    if k not in existing_breed or not existing_breed.get(k):
                        existing_breed[k] = v
            else:
                grouped[key]['breeds'].append(breed_data)

        # Формируем финальный список данных с автоматической нумерацией площадок
        result = []
        nn_counter = 1
        for group_key, group_data in grouped.items():
            # Создаем JSON строку для пород
            breeds_json = json.dumps(group_data['breeds'], ensure_ascii=False, indent=2) if group_data['breeds'] else ''

            result.append([
                str(nn_counter),  # Автоматическая нумерация от 1
                group_data['gps'],
                group_data['predmet'],
                breeds_json,  # Данные пород в JSON формате
                group_data['primechanie'],
                group_data['tip_lesa']
            ])
            nn_counter += 1

        return result

    def determine_breed_type(self, breed_name):
        """Определяет тип породы по названию"""
        coniferous_breeds = ['Сосна', 'Ель', 'Пихта', 'Кедр', 'Лиственница']
        deciduous_breeds = ['Берёза', 'Осина', 'Ольха чёрная', 'Ольха серая', 'Ива', 'Ива кустарниковая']

        if any(coniferous.lower() in breed_name.lower() for coniferous in coniferous_breeds):
            return 'coniferous'
        elif any(deciduous.lower() in breed_name.lower() for deciduous in deciduous_breeds):
            return 'deciduous'
        else:
            # По умолчанию считаем лиственными
            return 'deciduous'

    def load_section(self, instance):
        """Показать popup для выбора JSON файла"""
        if not os.path.exists(self.reports_dir):
            self.show_error("Папка reports не найдена!")
            return

        content = MDBoxLayout(orientation='vertical', spacing=10, padding=10, md_bg_color=Colors.DARK_SURFACE)

        title_label = MDLabel(
            text='Выберите файл JSON данных приложения:',
            font_size='18sp',
            theme_text_color='Custom',
            text_color=Colors.TEXT_ON_DARK,
            size_hint=(1, None),
            height=50
        )
        content.add_widget(title_label)

        scroll = ScrollView(size_hint=(1, 1))
        files_layout = MDGridLayout(cols=1, spacing=5, size_hint_y=None, md_bg_color=Colors.DARK_SURFACE)
        files_layout.bind(minimum_height=files_layout.setter('height'))

        # Добавляем кнопку для ручного ввода пути
        manual_input_layout = MDBoxLayout(orientation='vertical', size_hint_y=None, height=80, spacing=5, md_bg_color=Colors.DARK_SURFACE)
        manual_label = MDLabel(
            text="Или введите полный путь к файлу:",
            size_hint=(1, None),
            height=30,
            theme_text_color='Custom',
            text_color=Colors.TEXT_ON_DARK
        )
        self.manual_file_input = MDTextField(
            hint_text="Полный путь к JSON файлу",
            mode="outlined",
            size_hint=(1, None),
            height=40,
            line_color_focus=Colors.GREEN
        )
        manual_input_layout.add_widget(manual_label)
        manual_input_layout.add_widget(self.manual_file_input)
        files_layout.add_widget(manual_input_layout)

        uploader_label = MDLabel(
            text="Доступные JSON файлы:",
            bold=True,
            size_hint=(1, None),
            height=30,
            theme_text_color='Custom',
            text_color=Colors.TEXT_ON_DARK
        )
        files_layout.add_widget(uploader_label)

        # Получаем список JSON файлов
        json_files = [f for f in os.listdir(self.reports_dir) if f.endswith('.json')]
        if not json_files:
            no_files_label = MDLabel(
                text="JSON файлы не найдены в папке reports\nИспользуйте ручной ввод пути выше",
                size_hint=(1, None),
                height=50,
                theme_text_color='Custom',
                text_color=[0.8,0.8,0.8,1],
                valign='top'
            )
            no_files_label.bind(size=lambda *args: setattr(no_files_label, 'text_size', (no_files_label.width, None)))
            files_layout.add_widget(no_files_label)
        else:
            for filename in sorted(json_files):
                btn = MDButton(style='filled', size_hint=(1, None), height=50)
                btn.add_widget(MDButtonText(text=filename, theme_text_color='Custom', text_color=[0.7,0.7,0.7,1]))
                btn.bind(on_release=lambda x, f=filename: self.select_json_file(os.path.join(self.reports_dir, f)))
                files_layout.add_widget(btn)

        scroll.add_widget(files_layout)
        content.add_widget(scroll)

        btn_layout = MDBoxLayout(
            orientation='horizontal',
            spacing=10,
            size_hint=(1, None),
            height=60,
            md_bg_color=Colors.DARK_SURFACE
        )
        load_manual_btn = MDButton(style='filled', size_hint=(0.35, 1), height=60)
        load_manual_btn.add_widget(MDButtonText(text='Загрузить', theme_text_color='Custom', text_color=Colors.GREEN))
        load_manual_btn.bind(on_release=self.load_manual_json)
        cancel_btn = MDButton(style='filled', size_hint=(0.35, 1), height=60)
        cancel_btn.add_widget(MDButtonText(text='Отмена', theme_text_color='Custom', text_color=[1,0.3,0.3,1]))
        cancel_btn.bind(on_release=self.dismiss_json_popup)
        btn_layout.add_widget(load_manual_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        self.json_popup = Popup(
            title="",
            content=content,
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3],
            size_hint=(0.9, 0.9)
        )
        self.json_popup.open()

    def load_section_popup(self):
        """Показать popup для выбора JSON файла (вызывается из главного меню)"""
        return self.load_section(None)

    def select_json_file(self, file_path):
        """Обработка выбора JSON файла из списка"""
        try:
            self.load_json_data(file_path)
            self.json_popup.dismiss()
        except Exception as e:
            self.show_error(f"Ошибка загрузки: {str(e)}")

    def load_manual_json(self, instance):
        """Загрузка JSON файла по указанному пути"""
        file_path = self.manual_file_input.text.strip()
        if not file_path:
            self.show_error("Укажите путь к файлу!")
            return

        if not os.path.exists(file_path):
            self.show_error("Файл не найден!")
            return

        try:
            self.load_json_data(file_path)
            self.json_popup.dismiss()
        except Exception as e:
            self.show_error(f"Ошибка загрузки: {str(e)}")

    def dismiss_json_popup(self, instance=None):
        """Закрыть popup выбора файла"""
        if hasattr(self, 'json_popup'):
            self.json_popup.dismiss()

    def load_json_data(self, file_path):
        """Загрузка данных из JSON файла"""
        print(f"DEBUG: Loading JSON file: {file_path}")
        try:
            # Попытка загрузить с UTF-8
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
            except UnicodeDecodeError:
                # Если UTF-8 не работает, пробуем Windows-1251
                try:
                    with open(file_path, 'r', encoding='cp1251') as f:
                        data = json.load(f)
                except UnicodeDecodeError:
                    # Если и это не работает, пробуем автоопределение
                    import chardet
                    with open(file_path, 'rb') as f:
                        raw_data = f.read()
                        detected_encoding = chardet.detect(raw_data)['encoding']
                        if detected_encoding:
                            data = json.loads(raw_data.decode(detected_encoding))
                        else:
                            raise UnicodeDecodeError("Не удалось определить кодировку файла")

            self.current_section = os.path.splitext(os.path.basename(file_path))[0].replace('.json', '').replace('_приложение', '')
            self.page_data.clear()

            # Загружаем настройки адреса
            if isinstance(data, dict):
                if 'radius' in data:
                    self.current_radius = str(data['radius']) if data['radius'] else "5.64"
                if 'plot_area' in data:
                    self.plot_area_input = str(data['plot_area']) if data['plot_area'] else ""

            # Загружаем данные проекта
            if isinstance(data, dict) and 'project_data' in data:
                loaded_project_data = data['project_data']
                if isinstance(loaded_project_data, dict):
                    # Обновляем адресные данные
                    if 'address' in loaded_project_data and isinstance(loaded_project_data['address'], dict):
                        for key, value in loaded_project_data['address'].items():
                            if key in self.project_data['address']:
                                self.project_data['address'][key] = str(value) if value else ''
                                # Также обновляем текущие переменные
                                if key == 'quarter':
                                    self.current_quarter = str(value) if value else ''
                                elif key == 'plot':
                                    self.current_plot = str(value) if value else ''
                                elif key == 'forestry':
                                    self.current_forestry = str(value) if value else ''
                                elif key == 'radius':
                                    self.current_radius = str(value) if value else '5.64'

                    # Обновляем данные деталей проекта
                    if 'details' in loaded_project_data and isinstance(loaded_project_data['details'], dict):
                        for key, value in loaded_project_data['details'].items():
                            if key in self.project_data['details']:
                                self.project_data['details'][key] = str(value) if value else ''

                    # Загружаем название документа
                    if 'document_name' in loaded_project_data:
                        self.project_data['document_name'] = str(loaded_project_data['document_name']) if loaded_project_data['document_name'] else 'Проект'

            # Ожидаем, что JSON содержит page_data как словарь
            if isinstance(data, dict) and 'page_data' in data:
                self.page_data = data['page_data']
                print(f"DEBUG: Loaded page_data: {len(self.page_data)} pages")
            else:
                # Старый формат или простая структура
                self.page_data = data if isinstance(data, dict) else {}
                print(f"DEBUG: Loaded data as dict: {len(self.page_data) if isinstance(self.page_data, dict) else 'not dict'}")

            # Проверяем и исправляем формат страницы
            corrected_page_data = {}
            for page_key, page_rows in self.page_data.items():
                if isinstance(page_key, str):
                    try:
                        page_num = int(page_key)
                    except ValueError:
                        continue
                else:
                    page_num = page_key

                if isinstance(page_rows, list):
                    # Убеждаемся, что каждая строка - список из 6 элементов
                    corrected_rows = []
                    for row in page_rows:
                        if isinstance(row, list) and len(row) == 6:
                            corrected_rows.append(row)
                        elif isinstance(row, list):
                            # Дополняем до 6 элементов пустыми строками
                            corrected_row = row + [''] * (6 - len(row))
                            corrected_rows.append(corrected_row[:6])
                        else:
                            continue
                    corrected_page_data[page_num] = corrected_rows
                else:
                    continue

            self.page_data = corrected_page_data

            if self.page_data:
                self.current_page = min(self.page_data.keys())
            else:
                self.current_page = 0

            self.load_page_data()

            total_plots = sum(1 for rows in self.page_data.values() for row in rows if any(row[:3]))
            Clock.schedule_once(lambda dt: self.show_create_popup(), 0.2)
            self.show_success(f"Загружено {total_plots} площадок в {len(self.page_data)} страницах")
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            self.show_error(f"Ошибка загрузки JSON файла: {str(e)}\n{error_details}")

    def save_settings_to_db(self):
        """Сохранить настройки участка в базу данных"""
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()

        cursor.execute('''
            INSERT OR REPLACE INTO molodniki_settings (section_name, radius, plot_area, updated_at)
            VALUES (?, ?, ?, CURRENT_TIMESTAMP)
        ''', (self.current_section, self.current_radius, self.plot_area_input))

        conn.commit()
        conn.close()

    def show_radius_popup(self, instance):
        """Показать popup для установки радиуса"""
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)

        title_label = MDLabel(
            text="Установка радиуса для расчета коэффициента состава",
            font_style='Title', role='medium',
            theme_text_color='Custom', text_color=Colors.GREEN,
            adaptive_height=True
        )
        content.add_widget(title_label)

        self.radius_input = MDTextField(
            hint_text="Радиус (метры)",
            mode='outlined',
            size_hint_y=None, height=dp(52),
            input_filter='float',
            text=self.current_radius
        )
        self.radius_input.bind(text=self.update_radius_display)
        content.add_widget(self.radius_input)

        info_label = MDLabel(
            text="Радиус используется для расчета площади круга:\n"
                 "Площадь = π × радиус²\n"
                 "Коэффициент состава = (густота × площадь) / 10000\n"
                 "Радиус применяется автоматически ко всем площадкам",
            theme_text_color='Custom', text_color=[0.8,0.8,0.8,1],
            adaptive_height=True
        )
        content.add_widget(info_label)

        btn_layout = MDBoxLayout(orientation='horizontal', spacing=Spacing.MD, adaptive_height=True)
        save_btn = MDButton(style='filled', md_bg_color=Colors.BTN_SUCCESS,
                            size_hint=(0.5, None), height=dp(48))
        save_btn.add_widget(MDButtonText(text='Сохранить'))
        cancel_btn = MDButton(style='outlined',
                              size_hint=(0.5, None), height=dp(48),
                              line_color=Colors.DANGER)
        cancel_btn.add_widget(MDButtonText(text='Отмена',
                              theme_text_color='Custom', text_color=Colors.DANGER))
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            size_hint=(0.85, None),
            height=dp(360),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3]
        )

        def apply_radius(btn):
            try:
                radius = float(self.radius_input.text.strip())
                if radius <= 0:
                    self.show_error("Радиус должен быть положительным числом!")
                    return
                self.current_radius = str(radius)
                self.project_data['address']['radius'] = str(radius)
                self.save_settings_to_db()
                self.update_totals()
                self.show_success(f"Радиус {radius} м сохранен для всех расчетов")
                popup.dismiss()
                self.update_address_popup_display()
            except ValueError:
                self.show_error("Введите корректное числовое значение радиуса!")

        save_btn.bind(on_release=apply_radius)
        cancel_btn.bind(on_release=popup.dismiss)

        popup.open()

    def show_breed_choice_popup(self, instance, selected_breed):
        """Показать popup с выбором после добавления первой породы"""
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)

        title_label = MDLabel(
            text=f"Порода '{selected_breed}' добавлена!",
            font_style='Title', role='medium',
            theme_text_color='Custom', text_color=Colors.GREEN,
            adaptive_height=True
        )
        content.add_widget(title_label)

        info_label = MDLabel(
            text="Автоматически присвоен номер: 1 порода",
            theme_text_color='Custom', text_color=[0.8,0.8,0.8,1],
            adaptive_height=True
        )
        content.add_widget(info_label)

        btn_layout = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, adaptive_height=True)
        add_more_btn = MDButton(style='filled', md_bg_color=Colors.BTN_SUCCESS,
                                size_hint=(1, None), height=dp(48))
        add_more_btn.add_widget(MDButtonText(text='Добавить еще породу'))
        save_exit_btn = MDButton(style='filled', md_bg_color=Colors.BTN_SUCCESS,
                                 size_hint=(1, None), height=dp(48))
        save_exit_btn.add_widget(MDButtonText(text='Сохранить и выйти'))
        btn_layout.add_widget(add_more_btn)
        btn_layout.add_widget(save_exit_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            size_hint=(0.85, None),
            height=dp(300),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3]
        )

        def add_more_breed(btn):
            popup.dismiss()
            self.show_breed_popup(instance, True)

        def save_and_exit(btn):
            popup.dismiss()
            self.table_screen.show_success("Данные по площадке сохранены!")

        add_more_btn.bind(on_release=add_more_breed)
        save_exit_btn.bind(on_release=save_and_exit)

        popup.open()

    def update_row_total(self, instance, value):
        """Обновляем итоги по строке"""
        # Обновляем общие итоги страницы при изменении данных
        self.update_totals()

    def update_plot_total(self, instance, value):
        """Обновляем итог по площадке при изменении данных"""
        row_idx = instance.row_index
        breeds_text = value
        breeds_data = self.parse_breeds_data(breeds_text)

        if not breeds_data:
            return

        total_density = 0
        total_height = 0.0
        total_age = 0
        breed_count = 0
        breed_names = []

        for breed_info in breeds_data:
            breed_count += 1
            breed_name = breed_info.get('name', 'Неизвестная')
            breed_names.append(breed_name)

            if breed_info.get('type') == 'coniferous':
                coniferous_density = (breed_info.get('do_05', 0) +
                                    breed_info.get('05_15', 0) +
                                    breed_info.get('bolee_15', 0))
                if coniferous_density > 0:
                    total_density += coniferous_density
            elif 'density' in breed_info and breed_info['density']:
                total_density += breed_info['density']

            if 'height' in breed_info and breed_info['height']:
                total_height += breed_info['height']
            if 'age' in breed_info and breed_info['age']:
                total_age += breed_info['age']

        # Обновляем общие итоги
        self.update_totals()

    def create_new_plot(self, instance=None):
        """Создать новую площадку с помощью всплывающего окна"""
        if self.current_page not in self.page_data:
            self.page_data[self.current_page] = []
        row_idx = len(self.page_data[self.current_page])
        MolodnikiTreeDataInputPopup(self, row_idx).open()

    def restore_last_plot(self, instance=None):
        """Восстановить последнюю заполненную площадку"""
        last_plot = None
        for page_num in sorted(self.page_data.keys(), reverse=True):
            for row_idx in range(len(self.page_data[page_num]) - 1, -1, -1):
                row = self.page_data[page_num][row_idx]
                if any(row):
                    last_plot = {'page': page_num, 'row': row_idx, 'data': row.copy()}
                    break
            if last_plot:
                break
        if not last_plot:
            self.show_error("Нет данных для восстановления!")
            return
        if self.current_page not in self.page_data:
            self.page_data[self.current_page] = []
        new_row_idx = len(self.page_data[self.current_page])
        self.page_data[self.current_page].append(last_plot['data'])
        self.save_current_page()
        self.load_page_data()
        MolodnikiTreeDataInputPopup(self, new_row_idx).open()

    def show_create_popup(self, instance=None):
        """Показать popup управления площадками (современный интерфейс)"""
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)

        title_label = MDLabel(
            text="Управление площадками",
            font_style='Headline', role='medium',
            theme_text_color='Custom', text_color=Colors.GREEN,
            size_hint_y=None, height=dp(44)
        )
        content.add_widget(title_label)

        scroll = ScrollView(size_hint_y=None, height=dp(450))
        plots_grid = MDGridLayout(cols=1, spacing=Spacing.SM, adaptive_height=True,
                                  size_hint_y=None, padding=[0, 0])

        all_plots = []
        for page_num in sorted(self.page_data.keys()):
            for row_idx, row in enumerate(self.page_data[page_num]):
                if any(row):
                    all_plots.append({'page': page_num, 'row': row_idx, 'data': row})

        if not all_plots:
            empty_label = MDLabel(
                text="Нет добавленных площадок.\nНажмите «Добавить» чтобы создать новую.",
                theme_text_color='Custom', text_color=[0.6,0.6,0.6,1],
                halign='center', size_hint_y=None, height=dp(80)
            )
            plots_grid.add_widget(empty_label)
        else:
            for idx, plot in enumerate(all_plots):
                row = plot['data']
                gps = row[1] if len(row) > 1 else ''
                predmet = row[2] if len(row) > 2 else ''
                breeds_raw = row[3] if len(row) > 3 else ''
                notes = row[4] if len(row) > 4 else ''
                forest_type = row[5] if len(row) > 5 else ''

                breeds_list = self.parse_breeds_data(breeds_raw)
                breeds_str = ', '.join(b.get('name', '?') for b in breeds_list[:3])
                if len(breeds_list) > 3:
                    breeds_str += f' … +{len(breeds_list) - 3}'

                card = MDCard(style='elevated', size_hint=(1, None), height=dp(110),
                              md_bg_color=get_color_from_hex('#2A2A2A'),
                              padding=[Spacing.SM, Spacing.XS], spacing=Spacing.XS)
                card_content = MDBoxLayout(orientation='vertical', spacing=Spacing.XS,
                                           adaptive_height=True)

                header = MDBoxLayout(orientation='horizontal', size_hint_y=None, height=dp(22))
                header.add_widget(MDLabel(
                    text=f"Площадка №{idx + 1}", font_style='Label', role='medium',
                    theme_text_color='Custom', text_color=Colors.GREEN,
                    size_hint=(0.4, 1), valign='middle'))
                if gps:
                    header.add_widget(MDLabel(
                        text=f"GPS: {gps}", font_style='Body', role='small',
                        theme_text_color='Custom', text_color=[0.7,0.7,0.7,1],
                        size_hint=(0.6, 1), valign='middle'))
                card_content.add_widget(header)

                info_text = ''
                if predmet:
                    info_text += f"Предмет: {predmet}  "
                if forest_type:
                    info_text += f"Тип: {forest_type}  "
                if breeds_str:
                    info_text += f"Породы: {breeds_str}"
                if not info_text and not notes:
                    info_text = 'Нет данных'
                card_content.add_widget(MDLabel(
                    text=info_text, font_style='Body', role='small',
                    theme_text_color='Custom', text_color=[0.8,0.8,0.8,1],
                    adaptive_height=True))

                if notes:
                    card_content.add_widget(MDLabel(
                        text=f"Прим: {notes[:60]}", font_style='Body', role='small',
                        theme_text_color='Custom', text_color=[0.6,0.6,0.6,1],
                        adaptive_height=True))

                btn_row = MDBoxLayout(orientation='horizontal', spacing=Spacing.SM,
                                      size_hint_y=None, height=dp(30))
                edit_btn = MDButton(style='text', size_hint=(0.5, None), height=dp(28))
                edit_btn.add_widget(MDButtonText(text='Редактировать', font_size='12sp'))
                edit_btn.bind(on_release=lambda x, p=plot: self._edit_plot(p))
                delete_btn = MDButton(style='text', size_hint=(0.5, None), height=dp(28))
                delete_btn.add_widget(MDButtonText(text='Удалить', font_size='12sp',
                                      theme_text_color='Custom', text_color=Colors.DANGER))
                delete_btn.bind(on_release=lambda x, p=plot: self._delete_plot(p))
                btn_row.add_widget(edit_btn)
                btn_row.add_widget(delete_btn)
                card_content.add_widget(btn_row)

                card.add_widget(card_content)
                plots_grid.add_widget(card)

        scroll.add_widget(plots_grid)
        content.add_widget(scroll)

        btn_layout = MDBoxLayout(orientation='horizontal', spacing=Spacing.MD,
                                 adaptive_height=True)
        restore_btn = MDButton(style='filled', md_bg_color=Colors.BTN_WARNING,
                               size_hint=(1/3, None), height=dp(48))
        restore_btn.add_widget(MDButtonText(text='Восстановить'))
        restore_btn.bind(on_release=self.restore_last_plot)
        add_btn = MDButton(style='filled', md_bg_color=Colors.BTN_SUCCESS,
                           size_hint=(1/3, None), height=dp(48))
        add_btn.add_widget(MDButtonText(text='Добавить'))
        add_btn.bind(on_release=self.create_new_plot)
        close_btn = MDButton(style='outlined', size_hint=(1/3, None), height=dp(48),
                             line_color=Colors.DANGER)
        close_btn.add_widget(MDButtonText(text='Закрыть',
                             theme_text_color='Custom', text_color=Colors.DANGER))
        btn_layout.add_widget(restore_btn)
        btn_layout.add_widget(add_btn)
        btn_layout.add_widget(close_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="", content=content, size_hint=(0.9, None),
            height=dp(600), separator_height=0,
            background_color=[0,0,0,0.3], overlay_color=[0,0,0,0.3]
        )
        close_btn.bind(on_release=popup.dismiss)
        popup.open()

    def _edit_plot(self, plot_info):
        """Открыть редактирование площадки"""
        self.current_page = plot_info['page']
        MolodnikiTreeDataInputPopup(self, plot_info['row']).open()

    def _delete_plot(self, plot_info, popup_ref=None):
        """Удалить площадку с подтверждением"""
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)
        content.add_widget(MDLabel(
            text=f"Удалить площадку №{plot_info['row'] + 1}?",
            theme_text_color='Custom', text_color=[1,1,1,1], halign='center',
            size_hint_y=None, height=dp(50)))
        btn_layout = MDBoxLayout(orientation='horizontal', spacing=Spacing.MD, adaptive_height=True)
        yes_btn = MDButton(style='filled', md_bg_color=Colors.BTN_DANGER,
                           size_hint=(0.5, None), height=dp(44))
        yes_btn.add_widget(MDButtonText(text='Да, удалить'))
        no_btn = MDButton(style='outlined', size_hint=(0.5, None), height=dp(44),
                          line_color=Colors.GREEN)
        no_btn.add_widget(MDButtonText(text='Отмена',
                          theme_text_color='Custom', text_color=Colors.GREEN))
        btn_layout.add_widget(yes_btn)
        btn_layout.add_widget(no_btn)
        content.add_widget(btn_layout)
        popup = Popup(title="", content=content, size_hint=(0.7, None), height=dp(180),
                      separator_height=0, background_color=[0,0,0,0.3], overlay_color=[0,0,0,0.3])
        def do_delete(btn):
            page, row = plot_info['page'], plot_info['row']
            if page in self.page_data and row < len(self.page_data[page]):
                self.page_data[page][row] = [''] * 6
                self.save_current_page()
                self.load_page_data()
            self.show_success(f"Площадка №{plot_info['row'] + 1} удалена")
            popup.dismiss()
        yes_btn.bind(on_release=do_delete)
        no_btn.bind(on_release=popup.dismiss)
        popup.open()

    def show_address_popup(self, instance):
        """Показать popup с настройками адреса"""
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE)

        title_label = MDLabel(
            text="Настройки адреса",
            font_style='Headline', role='medium',
            theme_text_color='Custom', text_color=Colors.GREEN,
            size_hint_y=None, height=dp(50)
        )
        content.add_widget(title_label)

        scroll = ScrollView(size_hint=(1, 1))
        scroll_content = MDGridLayout(cols=1, spacing=Spacing.MD, size_hint_y=None, adaptive_height=True)
        scroll_content.bind(minimum_height=scroll_content.setter('height'))

        current_values = MDBoxLayout(orientation='vertical', spacing=Spacing.SM, adaptive_height=True,
                                     padding=Spacing.MD)

        current_title = MDLabel(
            text='Текущие значения:',
            font_style='Title', role='small',
            theme_text_color='Custom', text_color=Colors.GREEN,
            adaptive_height=True
        )
        current_values.add_widget(current_title)

        self.current_address_info = MDLabel(
            text=f"Название проекта: {self.project_data.get('document_name', 'Проект')}\n"
                 f"Квартал: {self.project_data['address'].get('quarter', 'Не указан')}\n"
                 f"Выдел: {self.project_data['address'].get('plot', 'Не указан')}\n"
                 f"Лесничество: {self.project_data['address'].get('forestry', 'Не указано')}\n"
                 f"Участковое лесничество: {self.project_data['address'].get('district_forestry', 'Не указано')}\n"
                 f"Радиус: {self.project_data['address'].get('radius', 'Не указан')} м\n"
                 f"Площадь участка: {self.project_data['address'].get('plot_area', 'Не указана')} га",
            theme_text_color='Custom', text_color=[0.8,0.8,0.8,1],
            adaptive_height=True
        )
        current_values.add_widget(self.current_address_info)
        scroll_content.add_widget(current_values)

        buttons_layout = MDGridLayout(cols=2, spacing=Spacing.MD, adaptive_height=True, size_hint_y=None)

        def make_addr_card(text, callback):
            card = MDCard(style='elevated', size_hint=(1, None), height=dp(60),
                          md_bg_color=Colors.CARD_BG, ripple_behavior=True,
                          on_release=callback, focus_behavior=True)
            card.add_widget(MDLabel(text=text, font_style='Label', role='medium',
                                    theme_text_color='Custom', text_color=[1,1,1,1],
                                    halign='center', valign='middle'))
            return card

        buttons_layout.add_widget(make_addr_card('Квартал', lambda x: self.show_quarter_popup(x)))
        buttons_layout.add_widget(make_addr_card('Выдел', lambda x: self.show_plot_popup(x)))
        buttons_layout.add_widget(make_addr_card('Лесничество', lambda x: self.show_forestry_popup(x)))
        buttons_layout.add_widget(make_addr_card('Радиус', lambda x: self.show_radius_popup(x)))
        buttons_layout.add_widget(make_addr_card('Площадь', lambda x: self.show_plot_area_input_popup(x)))
        buttons_layout.add_widget(make_addr_card('Название', lambda x: self.show_document_name_popup(x)))

        scroll_content.add_widget(buttons_layout)
        scroll.add_widget(scroll_content)
        content.add_widget(scroll)

        cancel_btn = MDButton(style='outlined', size_hint=(1, None), height=dp(48),
                              line_color=Colors.DANGER)
        cancel_btn.add_widget(MDButtonText(text='Закрыть',
                              theme_text_color='Custom', text_color=Colors.DANGER))
        content.add_widget(cancel_btn)

        self.address_popup = Popup(
            title="",
            content=content,
            size_hint=(0.9, 0.9),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3]
        )

        self.update_address_display()
        self.update_details_display()

        cancel_btn.bind(on_release=self.address_popup.dismiss)
        self.address_popup.open()

    def show_document_name_popup(self, instance):
        """Показать popup для ввода названия документа"""
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)

        title_label = MDLabel(
            text="Введите название документа",
            font_style='Title', role='medium',
            theme_text_color='Custom', text_color=Colors.GREEN,
            adaptive_height=True
        )
        content.add_widget(title_label)

        self.document_name_input = MDTextField(
            hint_text="Название документа",
            mode='outlined',
            size_hint_y=None, height=dp(52),
            text=self.project_data.get('document_name', 'Проект')
        )
        content.add_widget(self.document_name_input)

        info_label = MDLabel(
            text="Это название будет использоваться в имени файла при сохранении",
            theme_text_color='Custom', text_color=[0.8,0.8,0.8,1],
            adaptive_height=True
        )
        content.add_widget(info_label)

        btn_layout = MDBoxLayout(orientation='horizontal', spacing=Spacing.MD, adaptive_height=True)
        save_btn = MDButton(style='filled', md_bg_color=Colors.BTN_SUCCESS,
                            size_hint=(0.5, None), height=dp(48))
        save_btn.add_widget(MDButtonText(text='Сохранить'))
        cancel_btn = MDButton(style='outlined',
                              size_hint=(0.5, None), height=dp(48),
                              line_color=Colors.DANGER)
        cancel_btn.add_widget(MDButtonText(text='Отмена',
                              theme_text_color='Custom', text_color=Colors.DANGER))
        btn_layout.add_widget(save_btn)
        btn_layout.add_widget(cancel_btn)
        content.add_widget(btn_layout)

        popup = Popup(
            title="",
            content=content,
            size_hint=(0.85, None),
            height=dp(300),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3]
        )

        def save_document_name(btn):
            doc_name = self.document_name_input.text.strip()
            if doc_name:
                self.project_data['document_name'] = doc_name
                self.show_success(f"Название документа установлено: {doc_name}")
                popup.dismiss()
            else:
                self.show_error("Название документа не может быть пустым!")

        save_btn.bind(on_release=save_document_name)
        cancel_btn.bind(on_release=popup.dismiss)

        popup.open()

    def show_file_popup(self, instance):
        """Показать popup с операциями над файлами"""
        content = MDBoxLayout(orientation='vertical', spacing=Spacing.MD, padding=Spacing.MD,
                              md_bg_color=Colors.DARK_SURFACE, adaptive_height=True)

        title_label = MDLabel(
            text="Операции с файлами",
            font_style='Headline', role='medium',
            theme_text_color='Custom', text_color=Colors.GREEN,
            size_hint_y=None, height=dp(50)
        )
        content.add_widget(title_label)

        buttons_layout = MDGridLayout(cols=2, spacing=Spacing.MD, adaptive_height=True, size_hint_y=None)

        create_plot_btn = MDButton(style='filled', md_bg_color=Colors.BTN_SUCCESS,
                                   size_hint=(1, None), height=dp(48))
        create_plot_btn.add_widget(MDButtonText(text='Создать'))
        create_plot_btn.bind(on_release=self.show_create_popup)
        buttons_layout.add_widget(create_plot_btn)

        save_btn = MDButton(style='filled', md_bg_color=Colors.WARNING,
                            size_hint=(1, None), height=dp(48))
        save_btn.add_widget(MDButtonText(text='Сохранить'))
        save_btn.bind(on_release=self.save_all_formats)
        buttons_layout.add_widget(save_btn)

        load_btn = MDButton(style='filled', md_bg_color=Colors.PRIMARY,
                            size_hint=(1, None), height=dp(48))
        load_btn.add_widget(MDButtonText(text='Загрузить'))
        load_btn.bind(on_release=self.load_section)
        buttons_layout.add_widget(load_btn)

        edit_btn = MDButton(style='outlined', size_hint=(1, None), height=dp(48),
                            line_color=Colors.DANGER)
        edit_btn.add_widget(MDButtonText(text='Изменить',
                              theme_text_color='Custom', text_color=Colors.DANGER))
        edit_btn.bind(on_release=self.show_edit_plots_popup)
        buttons_layout.add_widget(edit_btn)

        open_folder_btn = MDButton(style='filled', md_bg_color=Colors.INFO,
                                   size_hint=(1, None), height=dp(48))
        open_folder_btn.add_widget(MDButtonText(text='Открыть'))
        open_folder_btn.bind(on_release=self.open_excel_file)
        buttons_layout.add_widget(open_folder_btn)

        clear_btn = MDButton(style='filled', md_bg_color=Colors.DARK,
                             size_hint=(1, None), height=dp(48))
        clear_btn.add_widget(MDButtonText(text='Очистить'))
        clear_btn.bind(on_release=self.clear_table_data)
        buttons_layout.add_widget(clear_btn)

        content.add_widget(buttons_layout)

        cancel_btn = MDButton(style='outlined', size_hint=(1, None), height=dp(48),
                              line_color=Colors.DANGER)
        cancel_btn.add_widget(MDButtonText(text='Закрыть',
                              theme_text_color='Custom', text_color=Colors.DANGER))
        content.add_widget(cancel_btn)

        popup = Popup(
            title="",
            content=content,
            size_hint=(0.85, None),
            height=dp(420),
            separator_height=0,
            background_color=[0,0,0,0.3],
            overlay_color=[0,0,0,0.3]
        )

        cancel_btn.bind(on_release=popup.dismiss)
        popup.open()

    def format_breeds_cell(self, breeds_text):
        """Форматирует содержимое ячейки пород для отображения"""
        if not breeds_text or not isinstance(breeds_text, str):
            return ""

        try:
            breeds_data = json.loads(breeds_text) if breeds_text.startswith('[') else []
        except (json.JSONDecodeError, TypeError):
            return breeds_text[:50] + "..." if len(breeds_text) > 50 else breeds_text

        if not breeds_data:
            return ""

        formatted_breeds = []
        for breed in breeds_data:
            if isinstance(breed, dict):
                name = breed.get('name', '')
                breed_type = breed.get('type', '')
                density = breed.get('density', '')
                height = breed.get('height', '')
                age = breed.get('age', '')

                # Для хвойных показываем сумму градаций
                if breed_type == 'coniferous':
                    conif_density = (breed.get('do_05', 0) + breed.get('05_15', 0) + breed.get('bolee_15', 0))
                    if conif_density > 0:
                        density = conif_density

                parts = [name]
                if density:
                    parts.append(f"гус.{density}")
                if height:
                    parts.append(f"выс.{height}м")
                if age:
                    parts.append(f"воз.{age}л")

                formatted_breeds.append(" ".join(parts))

        return "\n".join(formatted_breeds)
