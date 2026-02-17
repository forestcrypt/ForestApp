#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для чтения содержимого шаблона проекта ухода
"""

from docx import Document

def read_template():
    """Читаем шаблон и показываем его содержимое"""
    try:
        doc = Document('reports/Шаблон проект.docx')

        print("=== ВСЕ ПАРАГРАФЫ ШАБЛОНА ===")
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:  # Показываем все непустые параграфы
                print(f'{i+1}: "{text}"')

        print("\n=== ВСЕ ТАБЛИЦЫ ШАБЛОНА ===")
        for table_idx, table in enumerate(doc.tables):
            print(f'\nТаблица {table_idx+1}:')
            for row_idx, row in enumerate(table.rows):
                cells_text = [cell.text.strip() for cell in row.cells]
                print(f'  Строка {row_idx+1}: {cells_text}')

    except Exception as e:
        print(f"Ошибка: {e}")

if __name__ == '__main__':
    read_template()
