"""
Профессиональный экран таблицы деревьев
KivyMD + единая система стилей ForestApp
"""
from kivy.app import App
from kivy.core.window import Window
from kivy.clock import Clock
from kivy.properties import (NumericProperty, BooleanProperty, ObjectProperty,
                             ListProperty, StringProperty)
from kivy.metrics import dp
from kivy.utils import get_color_from_hex
from kivy.animation import Animation
from kivy.core.text import LabelBase

from kivymd.uix.screen import MDScreen
from kivymd.uix.boxlayout import MDBoxLayout
from kivymd.uix.scrollview import MDScrollView
from kivymd.uix.gridlayout import MDGridLayout
from kivymd.uix.label import MDLabel
from kivymd.uix.textfield import MDTextField
from kivymd.uix.button import MDButton, MDButtonText, MDIconButton
from kivymd.uix.dialog import (
    MDDialog,
    MDDialogHeadlineText,
    MDDialogSupportingText,
    MDDialogButtonContainer,
    MDDialogContentContainer,
)
from kivymd.uix.snackbar import MDSnackbar, MDSnackbarText
from kivymd.uix.card import MDCard
from kivymd.uix.menu import MDDropdownMenu
from kivymd.uix.appbar import MDTopAppBar, MDTopAppBarLeadingButtonContainer, MDTopAppBarTrailingButtonContainer, MDTopAppBarTitle, MDActionTopAppBarButton

import sqlite3
import pandas as pd
import os
import datetime
import json
import glob
import openpyxl
from openpyxl import Workbook
from openpyxl import styles

from ui_styles import Colors, Spacing, Fonts


class ModernTableTextInput(MDTextField):
    """Современное текстовое поле для таблицы"""
    row_index = NumericProperty(0)
    col_index = NumericProperty(0)
    next_widget = ObjectProperty(None)
    prev_widget = ObjectProperty(None)

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.bind(on_text_validate=self.on_enter)

    def on_enter(self, instance=None):
        if instance and self.next_widget:
            self.next_widget.focus = True


class TableScreen(MDScreen):
    """Профессиональный экран таблицы деревьев"""
    current_page = NumericProperty(0)
    total_pages = NumericProperty(1)
    current_section = StringProperty('')
    rows_per_page = NumericProperty(50)
    page_data = {}
    column_names = ListProperty([])
    inputs = []

    def _snack(self, message, duration=2.5):
        snack = MDSnackbar(duration=duration)
        snack.add_widget(MDSnackbarText(text=message))
        snack.open()

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.name = 'table'
        self.db_name = 'forest_data.db'
        self.reports_dir = 'reports'
        os.makedirs(self.reports_dir, exist_ok=True)

        self.default_column_names = [
            '№ дерева', 'Порода', 'ж/ф', 'шт/либо лет',
            'D, см', 'H, м', 'Сост-е', 'Модель', 'Примечания'
        ]
        self.column_names = self.default_column_names.copy()
        self.page_data = {}
        self.inputs = []
        self._edit_mode = False

        self.load_column_config()
        if not self.column_names:
            self.column_names = self.default_column_names.copy()
        self.setup_database()
        Clock.schedule_once(lambda dt: self.create_ui(), 0)
        self.load_existing_data()

    def load_column_config(self):
        try:
            with open('table_config.json', 'r', encoding='utf-8') as f:
                config = json.load(f)
                self.column_names = config.get('column_names', self.default_column_names.copy())
        except (FileNotFoundError, json.JSONDecodeError):
            self.column_names = self.default_column_names.copy()

    def save_column_config(self):
        with open('table_config.json', 'w', encoding='utf-8') as f:
            json.dump({'column_names': self.column_names}, f, ensure_ascii=False, indent=4)

    def setup_database(self):
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        cursor.execute('''CREATE TABLE IF NOT EXISTS trees (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            tree_number INTEGER, species TEXT, age TEXT, count TEXT,
            diameter REAL, height REAL, condition TEXT, model TEXT, notes TEXT,
            section_id INTEGER, FOREIGN KEY(section_id) REFERENCES sections(id))
        ''')
        cursor.execute('''CREATE TABLE IF NOT EXISTS sections (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            section_number TEXT UNIQUE, quarter TEXT, plot TEXT,
            forestry TEXT, district_forestry TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)
        ''')
        cursor.execute('''CREATE TABLE IF NOT EXISTS suggestions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            column_index INTEGER, value TEXT, UNIQUE(column_index, value))
        ''')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_suggestions ON suggestions (column_index, value)')
        conn.commit()
        conn.close()

    def create_ui(self):
        self.clear_widgets()
        main_layout = MDBoxLayout(orientation='vertical', spacing=0)

        # Верхняя панель инструментов
        toolbar = MDTopAppBar(
            type='small',
            elevation=2,
            md_bg_color=Colors.SECONDARY,
        )
        toolbar.add_widget(MDTopAppBarLeadingButtonContainer(
            MDActionTopAppBarButton(icon='arrow-left', on_release=self.go_back),
        ))
        toolbar.add_widget(MDTopAppBarTitle(text='Перечётная ведомость'))
        toolbar.add_widget(MDTopAppBarTrailingButtonContainer(
            MDActionTopAppBarButton(icon='counter', on_release=self.auto_fill_numbers),
            MDActionTopAppBarButton(icon='content-save', on_release=self.show_save_dialog),
        ))
        main_layout.add_widget(toolbar)

        # Информационная панель участка
        info_bar = MDCard(
            size_hint_y=None,
            height=dp(48),
            orientation='horizontal',
            padding=[Spacing.LG, 0],
            spacing=Spacing.SM,
            md_bg_color=Colors.SURFACE_ALT,
            radius=[0]
        )
        section_icon = MDIconButton(
            icon='map-marker',
            theme_text_color='Custom',
            text_color=Colors.PRIMARY,
            size_hint_x=None,
            size_hint_y=None,
            width=dp(32),
            height=dp(32),
            font_size=dp(20),
            pos_hint={'center_y': 0.5}
        )
        self.section_label = MDLabel(
            text=f'Участок: {self.current_section or "Не выбран"}',
            font_style='Body', role='large',
            theme_text_color='Custom',
            text_color=Colors.TEXT,
            halign='left',
            adaptive_height=True
        )
        info_bar.add_widget(section_icon)
        info_bar.add_widget(self.section_label)
        main_layout.add_widget(info_bar)

        # Пагинация
        pagination = self._create_pagination()
        main_layout.add_widget(pagination)

        # Таблица
        scroll = MDScrollView(do_scroll_x=True, do_scroll_y=True, bar_width=dp(6))
        self.table = MDGridLayout(cols=9, size_hint=(None, None), spacing=dp(1),
                                   padding=[Spacing.SM, 0])
        self.table.bind(
            minimum_height=self.table.setter('height'),
            minimum_width=self.table.setter('width')
        )

        # Заголовки таблицы
        header_colors = [Colors.PRIMARY_DIM if i == 4 else Colors.PRIMARY for i in range(9)]
        for i, column_name in enumerate(self.column_names):
            header = MDCard(
                size_hint_y=None, height=dp(36),
                size_hint_x=None, width=dp(110),
                md_bg_color=header_colors[min(i, len(header_colors)-1)],
                radius=[Spacing.RADIUS_XS],
                padding=[Spacing.XS, 0],
                orientation='vertical',
                on_release=lambda x, idx=i: self.edit_column_name(idx)
            )
            header_text = MDLabel(
                text=column_name,
                font_style='Label', role='small',
                theme_text_color='Custom',
                text_color=Colors.TEXT_ON_PRIMARY,
                bold=True,
                halign='center',
                valign='middle',
                size_hint_y=None,
                height=dp(36)
            )
            header.add_widget(header_text)
            self.table.add_widget(header)

        # Строки таблицы
        self.create_table_rows()
        scroll.add_widget(self.table)
        main_layout.add_widget(scroll)

        # Нижняя панель действий
        bottom_bar = MDBoxLayout(
            orientation='horizontal',
            size_hint_y=None,
            height=Spacing.HEADER_HEIGHT,
            padding=[Spacing.SM, Spacing.XS],
            spacing=dp(2),
            md_bg_color=Colors.SURFACE
        )

        actions = [
            ('chart-box-outline', 'Итого', Colors.SECONDARY, self.show_total_summary),
            ('folder-open', 'Загрузить', Colors.SECONDARY, self.load_section),
            ('delete', 'Очистить', Colors.DANGER, self.clear_table_data),
            ('home', 'Меню', Colors.PRIMARY_DIM, self.go_back),
        ]
        for icon, text, color, callback in actions:
            btn = MDCard(
                orientation='vertical',
                size_hint_x=0.25, size_hint_y=None, height=dp(48),
                padding=[dp(4), dp(2)], spacing=dp(1), radius=[dp(8)],
                elevation=1, md_bg_color=color,
                on_release=callback,
            )
            btn.add_widget(MDIconButton(
                icon=icon, font_size='20sp',
                theme_text_color='Custom', text_color=Colors.TEXT_ON_PRIMARY,
                pos_hint={'center_x': 0.5},
                size_hint=(None, None), size=(dp(28), dp(24)),
                on_release=lambda: None,
            ))
            btn.add_widget(MDLabel(
                text=text, font_size='10sp', bold=True,
                theme_text_color='Custom', text_color=Colors.TEXT_ON_PRIMARY,
                halign='center', size_hint_y=None, height=dp(16),
            ))
            bottom_bar.add_widget(btn)

        main_layout.add_widget(bottom_bar)
        self.add_widget(main_layout)

    def _create_pagination(self):
        layout = MDBoxLayout(
            orientation='horizontal',
            size_hint_y=None,
            height=dp(44),
            spacing=Spacing.SM,
            padding=[Spacing.LG, Spacing.XS]
        )
        prev_btn = MDButton(style='outlined', size_hint_x=None, width=dp(90),
                              on_release=lambda x: self.change_page(-1))
        prev_btn.add_widget(MDButtonText(text='← Пред.'))
        layout.add_widget(prev_btn)

        page_info = MDBoxLayout(orientation='vertical', adaptive_height=True, pos_hint={'center_y': 0.5})
        self.page_label = MDLabel(
            text=f'Стр. {self.current_page + 1} из {self.total_pages}',
            font_style='Body', role='small',
            halign='center',
            theme_text_color='Primary',
            adaptive_height=True
        )
        self.total_label = MDLabel(
            text=f'Всего: {self.rows_per_page} записей',
            font_style='Body', role='small',
            halign='center',
            theme_text_color='Secondary',
            adaptive_height=True
        )
        page_info.add_widget(self.page_label)
        page_info.add_widget(self.total_label)
        layout.add_widget(page_info)

        next_btn = MDButton(style='outlined', size_hint_x=None, width=dp(90),
                              on_release=lambda x: self.change_page(1))
        next_btn.add_widget(MDButtonText(text='След. →'))
        layout.add_widget(next_btn)
        return layout

    def create_table_rows(self):
        self.inputs = []
        for row_idx in range(self.rows_per_page):
            row = []
            for col_idx in range(9):
                inp = ModernTableTextInput(
                    multiline=False,
                    size_hint_y=None,
                    height=dp(34),
                    size_hint_x=None,
                    width=dp(110),
                    mode='outlined',
                    row_index=row_idx,
                    col_index=col_idx,
                    hint_text='' if col_idx == 0 else (['№', 'Порода', 'ж/ф', 'шт/лет', 'D, см', 'H, м', 'Сост.', 'Модель', 'Прим.'][col_idx]),
                    font_size=Fonts.BODY_XS,
                )
                if col_idx == 0:
                    inp.input_filter = 'int'
                elif col_idx in (4, 5):
                    inp.input_filter = 'float'

                if row:
                    inp.prev_widget = row[col_idx - 1]
                    row[col_idx - 1].next_widget = inp
                if row_idx > 0 and col_idx < len(self.inputs[row_idx - 1]):
                    inp.up_widget = self.inputs[row_idx - 1][col_idx]
                    self.inputs[row_idx - 1][col_idx].down_widget = inp

                inp.bind(on_text_validate=self.navigate_on_enter)
                row.append(inp)
                self.table.add_widget(inp)
            self.inputs.append(row)

    def navigate_on_enter(self, instance):
        if hasattr(instance, 'next_widget') and instance.next_widget:
            instance.next_widget.focus = True

    def change_page(self, direction):
        self.save_current_page()
        new_page = self.current_page + direction
        if 0 <= new_page < max(self.total_pages, 1):
            self.current_page = new_page
            self.load_page_data()
            self._update_pagination()

    def _update_pagination(self):
        self.total_pages = max(len(self.page_data), 1)
        self.page_label.text = f'Стр. {self.current_page + 1} из {self.total_pages}'

    def load_page_data(self):
        if self.current_page in self.page_data:
            page_data = self.page_data[self.current_page]
            for row_idx, row in enumerate(self.inputs):
                if row_idx < len(page_data):
                    for col_idx, cell in enumerate(row):
                        if col_idx < len(page_data[row_idx]):
                            cell.text = str(page_data[row_idx][col_idx])
                        else:
                            cell.text = ''
                else:
                    for cell in row:
                        cell.text = ''

    def save_current_page(self, instance=None):
        page_data = []
        for row in self.inputs:
            page_data.append([inp.text for inp in row])
        self.page_data[self.current_page] = page_data
        if instance:
            self._snack('✅ Страница сохранена!', duration=2)

    def show_save_dialog(self, instance):
        content = MDBoxLayout(orientation='vertical', spacing=dp(12), padding=dp(8), adaptive_height=True)
        content.add_widget(MDLabel(
            text='Сохранение отчёта', font_style='Title', role='medium', bold=True,
            halign='center', size_hint_y=None, height=dp(32),
        ))
        filename_field = MDTextField(
            hint_text='Имя файла', mode='outlined',
            text=f'{self.current_section}_{datetime.datetime.now().strftime("%Y-%m-%d_%H-%M")}',
            size_hint_y=None, height=dp(52), font_size='15sp',
        )
        content.add_widget(filename_field)

        format_row = MDBoxLayout(orientation='horizontal', spacing=dp(8), size_hint_y=None, height=dp(48))
        formats = [
            ('file-excel', 'Excel', Colors.PRIMARY, lambda x: self.save_to_excel()),
            ('file-word', 'Word', Colors.SECONDARY, lambda x: self.save_to_word()),
            ('file-pdf-box', 'PDF', Colors.BTN_DANGER, lambda x: self.save_to_pdf()),
            ('code-json', 'JSON', Colors.ACCENT, lambda x: self.save_to_json()),
            ('file-multiple', 'Все', Colors.BTN_PURPLE, lambda x: self.save_all_formats()),
        ]
        for icon, label, color, cb in formats:
            btn = MDCard(
                orientation='vertical', size_hint_x=0.25, size_hint_y=None, height=dp(48),
                padding=[dp(2), dp(2)], spacing=dp(1), radius=[dp(8)],
                elevation=1, md_bg_color=color, on_release=cb,
            )
            btn.add_widget(MDIconButton(
                icon=icon, font_size='18sp',
                theme_text_color='Custom', text_color=[1,1,1,1],
                size_hint=(None, None), size=(dp(24), dp(20)),
                pos_hint={'center_x': 0.5}, on_release=lambda: None,
            ))
            btn.add_widget(MDLabel(
                text=label, font_size='9sp', bold=True,
                theme_text_color='Custom', text_color=[1,1,1,1],
                halign='center', size_hint_y=None, height=dp(16),
            ))
            format_row.add_widget(btn)
        content.add_widget(format_row)

        cancel_save_btn = MDButton(style='outlined', on_release=lambda x: dialog.dismiss())
        cancel_save_btn.add_widget(MDButtonText(text='Отмена'))
        content.add_widget(cancel_save_btn)
        dialog = MDDialog(
            MDDialogContentContainer(content),
        )
        self.save_dialog = dialog
        dialog.open()

    def save_all_formats(self, instance):
        self.save_dialog.dismiss()
        try:
            self.save_to_excel()
            self.save_to_word()
            self.save_to_pdf()
            self.save_to_json()
            self._snack('✅ Все файлы сохранены!', duration=3)
        except Exception as e:
            self._snack(f'❌ Ошибка: {str(e)}', duration=3)

    def save_to_excel(self, instance=None):
        try:
            filename = f'Перечетная_ведомость_{self.current_section}_{datetime.datetime.now().strftime("%Y-%m-%d_%H-%M")}.xlsx'
            full_path = os.path.join(self.reports_dir, filename)
            wb = Workbook()
            ws = wb.active
            ws.title = 'Перечетная ведомость'
            ws['A1'] = f'ПЕРЕЧЕТНАЯ ВЕДОМОСТЬ - УЧАСТОК {self.current_section}'
            ws['A1'].font = openpyxl.styles.Font(bold=True, size=14)
            ws.merge_cells('A1:I1')

            all_data = []
            for page in sorted(self.page_data.keys()):
                all_data.extend(self.page_data[page])

            for col_num, header in enumerate(self.column_names, 1):
                cell = ws.cell(row=3, column=col_num, value=header)
                cell.font = openpyxl.styles.Font(bold=True)
                cell.fill = openpyxl.styles.PatternFill(start_color='2E7D32', end_color='2E7D32', fill_type='solid')
                cell.font = openpyxl.styles.Font(bold=True, color='FFFFFF')

            for row_num, row_data in enumerate(all_data, 4):
                for col_num, cell_value in enumerate(row_data, 1):
                    ws.cell(row=row_num, column=col_num, value=cell_value)

            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(cell.value)
                    except:
                        pass
                ws.column_dimensions[column_letter].width = min(max_length + 2, 50)

            wb.save(full_path)
            self._snack(f'✅ Excel: {filename}', duration=3)
        except Exception as e:
            self._snack(f'❌ Ошибка Excel: {str(e)}', duration=3)

    def save_to_word(self, instance=None):
        try:
            from docx import Document
            filename = f'Перечетная_ведомость_{self.current_section}_{datetime.datetime.now().strftime("%Y-%m-%d_%H-%M")}.docx'
            full_path = os.path.join(self.reports_dir, filename)
            doc = Document()
            doc.add_heading(f'Перечетная ведомость - Участок {self.current_section}', 0)

            all_data = []
            for page in sorted(self.page_data.keys()):
                all_data.extend(self.page_data[page])

            table = doc.add_table(rows=1, cols=len(self.column_names))
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            for i, col_name in enumerate(self.column_names):
                header_cells[i].text = col_name
                header_cells[i].paragraphs[0].runs[0].bold = True

            for row_data in all_data:
                row = table.add_row().cells
                for i, cell_value in enumerate(row_data):
                    row[i].text = str(cell_value) if cell_value else ''

            totals = self.calculate_totals()
            if totals and totals.get('total_trees', 0) > 0:
                doc.add_page_break()
                doc.add_heading('Итоги по перечетной ведомости', 1)
                doc.add_paragraph(f'Всего деревьев: {totals["total_trees"]}')
                doc.add_paragraph(f'Средний диаметр: {totals["avg_diameter"]:.1f} см (измерено: {totals["diameter_count"]})')
                doc.add_paragraph(f'Средняя высота: {totals["avg_height"]:.1f} м (измерено: {totals["height_count"]})')
                species_summary = totals.get('species_summary', {})
                if species_summary:
                    doc.add_heading('Распределение по породам', 2)
                    species_table = doc.add_table(rows=1, cols=4)
                    species_table.style = 'Table Grid'
                    hdr = species_table.rows[0].cells
                    hdr[0].text = 'Порода'
                    hdr[1].text = 'Количество'
                    hdr[2].text = 'Средний диаметр'
                    hdr[3].text = 'Средняя высота'
                    for species, data in sorted(species_summary.items()):
                        row_cells = species_table.add_row().cells
                        row_cells[0].text = species
                        row_cells[1].text = str(data['count'])
                        d = data.get('diameters', [])
                        h = data.get('heights', [])
                        row_cells[2].text = f'{sum(d)/len(d):.1f} см' if d else '-'
                        row_cells[3].text = f'{sum(h)/len(h):.1f} м' if h else '-'

            doc.save(full_path)
            self._snack(f'✅ Word: {filename}', duration=3)
        except Exception as e:
            self._snack(f'❌ Ошибка Word: {str(e)}', duration=3)

    def save_to_pdf(self, instance=None):
        try:
            from core.pdf_export import export_table_to_pdf
            filename = f'reports/Перечетная_ведомость_{self.current_section}.pdf'
            result = export_table_to_pdf(self, filename)
            self._snack(f'✅ PDF: {os.path.basename(result)}', duration=3)
        except Exception as e:
            self._snack(f'❌ Ошибка PDF: {str(e)}', duration=3)

    def save_to_json(self, instance=None):
        try:
            filename = f'Перечетная_ведомость_{self.current_section}_{datetime.datetime.now().strftime("%Y-%m-%d_%H-%M")}.json'
            full_path = os.path.join(self.reports_dir, filename)
            data = {
                'page_data': self.page_data,
                'section': self.current_section,
                'column_names': self.column_names,
                'export_date': datetime.datetime.now().isoformat()
            }
            with open(full_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            self._snack(f'✅ JSON: {filename}', duration=3)
        except Exception as e:
            self._snack(f'❌ Ошибка JSON: {str(e)}', duration=3)

    def load_section(self, instance):
        try:
            conn = sqlite3.connect('forest_data.db')
            cursor = conn.cursor()
            cursor.execute('SELECT section_number FROM sections WHERE section_number IS NOT NULL ORDER BY id DESC')
            sections = cursor.fetchall()
            conn.close()
            if not sections:
                self._snack('⚠️ Нет сохранённых участков!', duration=3)
                return
            items = []
            for section in sections:
                items.append({
                'leading_icon': 'folder',
                'text': section[0],
                'on_release': lambda x, s=section[0]: self.select_section(s)
            })
            self.menu = MDDropdownMenu(items=items)
            self.menu.open()
        except Exception as e:
            self._snack(f'❌ Ошибка: {str(e)}', duration=3)

    def select_section(self, section_number):
        self.menu.dismiss()
        files = glob.glob(os.path.join(self.reports_dir, f'{section_number}_*.xlsx'))
        if files:
            latest_file = max(files, key=os.path.getctime)
            try:
                df = pd.read_excel(latest_file)
                self.current_section = section_number
                self.page_data.clear()
                for page_num in range(0, len(df), self.rows_per_page):
                    page = page_num // self.rows_per_page
                    page_data = df.iloc[page_num:page_num + self.rows_per_page].values.tolist()
                    self.page_data[page] = page_data
                self.current_page = 0
                self.load_page_data()
                self._update_pagination()
                self.update_section_label()
                self._snack('✅ Данные загружены!', duration=3)
            except Exception as e:
                self._snack(f'❌ Ошибка загрузки: {str(e)}', duration=3)
        else:
            self._snack('❌ Файл не найден!', duration=3)

    def update_section_label(self):
        self.section_label.text = f'Участок: {self.current_section or "Не выбран"}'

    def clear_table_data(self, instance):
        clear_yes_btn = MDButton(style='outlined', on_release=lambda x: self.clear_data())
        clear_yes_btn.add_widget(MDButtonText(text='Да, очистить'))
        clear_no_btn = MDButton(style='outlined', on_release=lambda x: self.confirm_dialog.dismiss())
        clear_no_btn.add_widget(MDButtonText(text='Отмена'))
        self.confirm_dialog = MDDialog(
            MDDialogHeadlineText(text='🗑️ Очистка данных'),
            MDDialogSupportingText(text='Вы уверены, что хотите очистить все данные?'),
            MDDialogButtonContainer(
                clear_yes_btn,
                clear_no_btn,
                spacing='8dp',
            ),
        )
        self.confirm_dialog.open()

    def clear_data(self, instance=None):
        self.page_data.clear()
        for row in self.inputs:
            for cell in row:
                cell.text = ''
        self.current_page = 0
        self.total_pages = 1
        self._update_pagination()
        if instance:
            self.confirm_dialog.dismiss()
            self._snack('🗑️ Данные очищены!', duration=2)

    def go_back(self, instance):
        App.get_running_app().root.current = 'main'

    def _get_address_info(self):
        try:
            conn = sqlite3.connect(self.db_name)
            cursor = conn.cursor()
            cursor.execute('''
                SELECT quarter, plot, forestry, district_forestry
                FROM sections WHERE section_number = ?
                ORDER BY id DESC LIMIT 1
            ''', (self.current_section,))
            row = cursor.fetchone()
            conn.close()
            if row:
                parts = []
                if row[0]: parts.append(f'Квартал: {row[0]}')
                if row[1]: parts.append(f'Выдел: {row[1]}')
                if row[2]: parts.append(f'Лесничество: {row[2]}')
                if row[3]: parts.append(f'Участковое лесничество: {row[3]}')
                return '\n'.join(parts) if parts else 'Адрес не указан'
        except Exception:
            pass
        return 'Адрес не указан'

    def show_total_summary(self, instance):
        totals = self.calculate_totals()
        if not totals or totals.get('total_trees', 0) == 0:
            self._snack('⚠️ Нет данных для расчёта итогов', duration=3)
            return

        content = MDBoxLayout(orientation='vertical', spacing=dp(8), padding=dp(16), adaptive_height=True)
        content.add_widget(MDLabel(
            text=f'ИТОГИ — Участок {self.current_section}',
            font_style='Title', role='medium', bold=True, halign='center',
            size_hint_y=None, height=dp(36),
        ))

        address_info = self._get_address_info()
        if address_info:
            addr_card = MDCard(
                orientation='vertical', size_hint=(1, None),
                padding=dp(12), spacing=dp(4), radius=[dp(8)],
                elevation=1, md_bg_color=Colors.INFO_BG,
            )
            addr_card.add_widget(MDLabel(
                text='АДРЕС УЧАСТКА', font_style='Label', role='small', bold=True,
                theme_text_color='Custom', text_color=Colors.INFO,
                size_hint_y=None, height=dp(20),
            ))
            addr_card.add_widget(MDLabel(
                text=address_info, font_size='12sp',
                theme_text_color='Custom', text_color=Colors.TEXT_SECONDARY,
                adaptive_height=True,
            ))
            content.add_widget(addr_card)

        summary = (
            f'Всего деревьев: {totals.get("total_trees", 0)}\n'
            f'Средний диаметр: {totals.get("avg_diameter", 0):.1f} см (изм.: {totals.get("diameter_count", 0)})\n'
            f'Средняя высота: {totals.get("avg_height", 0):.1f} м (изм.: {totals.get("height_count", 0)})'
        )
        content.add_widget(MDLabel(
            text=summary, font_size='14sp',
            size_hint_y=None, height=dp(72),
        ))

        species_summary = totals.get('species_summary', {})
        if species_summary:
            content.add_widget(MDLabel(
                text='РАСПРЕДЕЛЕНИЕ ПО ПОРОДАМ:',
                font_style='Label', role='small', bold=True, size_hint_y=None, height=dp(24),
            ))
            for species, data in sorted(species_summary.items()):
                diameters = data.get('diameters', [])
                heights = data.get('heights', [])
                ages = data.get('ages', [])
                avg_d = sum(diameters) / len(diameters) if diameters else 0
                avg_h = sum(heights) / len(heights) if heights else 0
                ages_text = f', возрасты: {", ".join(ages[:5])}' if ages else ''
                species_text = f'{species}: {data["count"]} шт, Dср={avg_d:.1f}см, Hср={avg_h:.1f}м{ages_text}'
                content.add_widget(MDLabel(
                    text=species_text, font_size='12sp',
                    theme_text_color='Secondary', size_hint_y=None, height=dp(22),
                ))

        close_total_btn = MDButton(style='filled', md_bg_color=Colors.PRIMARY,
                                      size_hint=(1, None), height=dp(44),
                                      on_release=lambda x: dialog.dismiss())
        close_total_btn.add_widget(MDButtonText(text='Закрыть'))
        content.add_widget(close_total_btn)
        dialog = MDDialog(
            MDDialogContentContainer(content),
            size_hint=(0.7, None),
        )
        dialog.open()

    def calculate_totals(self):
        try:
            all_data = []
            for page in sorted(self.page_data.keys()):
                all_data.extend(self.page_data[page])

            coniferous_breeds = ['Сосна', 'Ель', 'Пихта', 'Кедр', 'Лиственница']
            deciduous_breeds = ['Берёза', 'Осина', 'Ольха чёрная', 'Ольха серая', 'Ива', 'Ива кустарниковая']

            def get_breed_type(species):
                for c in coniferous_breeds:
                    if c.lower() in species.lower():
                        return 'coniferous'
                for d in deciduous_breeds:
                    if d.lower() in species.lower():
                        return 'deciduous'
                return 'unknown'

            species_summary = {}
            total_trees = 0
            total_diameter = 0.0
            total_height = 0.0
            diameter_count = 0
            height_count = 0

            for row in all_data:
                if len(row) >= 6:
                    tree_num = str(row[0]).strip() if row[0] else ''
                    species = str(row[1]).strip() if row[1] else ''
                    age_str = str(row[3]).strip() if row[3] else ''
                    diameter_str = str(row[4]).strip() if row[4] else ''
                    height_str = str(row[5]).strip() if row[5] else ''

                    if tree_num and species:
                        total_trees += 1
                        if species not in species_summary:
                            species_summary[species] = {'count': 0, 'diameters': [], 'heights': [], 'ages': [], 'type': get_breed_type(species)}
                        species_summary[species]['count'] += 1
                        if age_str:
                            species_summary[species]['ages'].append(age_str)
                        try:
                            d = float(diameter_str)
                            species_summary[species]['diameters'].append(d)
                            total_diameter += d
                            diameter_count += 1
                        except (ValueError, TypeError):
                            pass
                        try:
                            h = float(height_str)
                            species_summary[species]['heights'].append(h)
                            total_height += h
                            height_count += 1
                        except (ValueError, TypeError):
                            pass

            return {
                'total_trees': total_trees,
                'avg_diameter': total_diameter / diameter_count if diameter_count else 0,
                'avg_height': total_height / height_count if height_count else 0,
                'diameter_count': diameter_count,
                'height_count': height_count,
                'species_summary': species_summary,
            }
        except Exception as e:
            return {'total_trees': 0, 'avg_diameter': 0, 'avg_height': 0, 'species_summary': {}}

    def edit_column_name(self, column_index):
        current_name = self.column_names[column_index] if column_index < len(self.column_names) else ''
        textfield = MDTextField(
            text=current_name, mode='outlined',
            size_hint_y=None, height=dp(52), font_size='16sp',
        )
        save_col_btn = MDButton(style='outlined',
                                  on_release=lambda x: self._save_column_name(column_index, textfield.text, dialog))
        save_col_btn.add_widget(MDButtonText(text='Сохранить'))
        cancel_col_btn = MDButton(style='outlined', on_release=lambda x: dialog.dismiss())
        cancel_col_btn.add_widget(MDButtonText(text='Отмена'))
        dialog = MDDialog(
            MDDialogHeadlineText(text='✏️ Редактировать название столбца'),
            MDDialogContentContainer(textfield),
            MDDialogButtonContainer(
                save_col_btn,
                cancel_col_btn,
                spacing='8dp',
            ),
        )
        dialog.open()

    def _save_column_name(self, col_idx, new_name, dialog):
        if new_name.strip():
            self.column_names[col_idx] = new_name.strip()
            dialog.dismiss()
            self.create_ui()
            self._snack('✅ Название столбца изменено', duration=2)
        else:
            self._snack('⚠️ Название не может быть пустым', duration=2)

    def auto_fill_numbers(self, instance=None):
        count = 0
        for row_idx, row in enumerate(self.inputs):
            current_val = row[0].text.strip()
            if not current_val:
                tree_num = self.current_page * self.rows_per_page + row_idx + 1
                row[0].text = str(tree_num)
                count += 1
        if count > 0:
            self._snack(f'✅ Добавлено номеров: {count}', duration=2)
        else:
            self._snack('ℹ️ Все строки уже пронумерованы', duration=2)

    def load_existing_data(self):
        pass

    def save_suggestion(self, col_index, value):
        try:
            conn = sqlite3.connect(self.db_name)
            cursor = conn.cursor()
            cursor.execute('''
                INSERT OR IGNORE INTO suggestions (column_index, value)
                VALUES (?, ?)
            ''', (col_index, value))
            conn.commit()
            conn.close()
        except Exception as e:
            print(f"Error saving suggestion: {e}")

    def validate_page_data(self):
        warnings = []
        for row_idx, row in enumerate(self.page_data.get(self.current_page, [])):
            if len(row) >= 6:
                tree_num = str(row[0]).strip() if row[0] else ''
                species = str(row[1]).strip() if row[1] else ''
                diameter_str = str(row[4]).strip() if row[4] else ''
                height_str = str(row[5]).strip() if row[5] else ''

                if tree_num and not species:
                    warnings.append(f"Строка {row_idx+1}: Нет породы для дерева №{tree_num}")
                if species and not tree_num:
                    warnings.append(f"Строка {row_idx+1}: Нет номера дерева для породы {species}")
                if diameter_str:
                    try:
                        d = float(diameter_str)
                        if d <= 0 or d > 500:
                            warnings.append(f"Строка {row_idx+1}: Подозрительный диаметр {d} см")
                    except ValueError:
                        warnings.append(f"Строка {row_idx+1}: Некорректный диаметр '{diameter_str}'")
                if height_str:
                    try:
                        h = float(height_str)
                        if h <= 0 or h > 100:
                            warnings.append(f"Строка {row_idx+1}: Подозрительная высота {h} м")
                    except ValueError:
                        warnings.append(f"Строка {row_idx+1}: Некорректная высота '{height_str}'")
        if warnings:
            self._snack('⚠️ ' + '\n'.join(warnings[:3]), duration=4)
