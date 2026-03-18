#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Чтение документа участка для анализа
"""

import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document

doc = Document('reports/Молодняки_3_20260309_1402.docx')

print("=" * 60)
print("ЧТЕНИЕ ДОКУМЕНТА УЧАСТКА")
print("=" * 60)

# Читаем параграфы
print("\n=== ПАРАГРАФЫ ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        print(f"\nПараграф {i}:")
        print(f"  {text}")

# Читаем таблицы
print("\n=== ТАБЛИЦЫ ===")
for i, table in enumerate(doc.tables):
    print(f"\nТаблица {i+1} (колонок: {len(table.columns)}, строк: {len(table.rows)}):")
    
    for j, row in enumerate(table.rows):
        row_text = []
        for k, cell in enumerate(row.cells):
            cell_text = cell.text.strip().replace('\n', ' ')
            if cell_text:
                row_text.append(f"[{k}] {cell_text}")
        
        if row_text:
            print(f"  Строка {j}: {' | '.join(row_text)}")

print("\n" + "=" * 60)
print("ЧТЕНИЕ ЗАВЕРШЕНО")
print("=" * 60)
