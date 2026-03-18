#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для заполнения нашего шаблона проекта ухода
ИСПРАВЛЕНИЯ ВСЕХ ПРОБЛЕМ:
1. care_activity_text - не дублировать, если care_queue уже содержит "- РУМ"
2. forest_type - брать из address_data.forest_type (из участка)
3. Диаметры - заполнять из breed.get('diameter'), рассчитывать проектные
4. Густота проектируемая - парсинг предмета ухода (3000шт/га = 3.0 тыс.шт/га)
5. Характеристики деревьев - парсить из details_data.characteristics построчно
6. Коэффициент состава - рассчитывать правильно по породам
7. Интенсивность рубки - из total_data.get('intensity')
"""

import os
import sys
import json
import re
import datetime
from docx import Document

# БАЗА ДАННЫХ ПОРОД для предмета ухода
BREED_DATABASE = {
    'С': 'Сосна',
    'Б': 'Берёза',
    'Ос': 'Осина',
    'Е': 'Ель',
    'П': 'Пихта',
    'Ол': 'Ольха',
    'Д': 'Дуб',
    'Я': 'Ясень',
    'Г': 'Граб',
    'Лц': 'Лиственница',
    'Лп': 'Липа',
    'К': 'Кедр',
    'Мж': 'Можжевельник',
    'Кл': 'Клён',
    'Р': 'Рябина'
}

# Обратный маппинг: название породы -> код
BREED_TO_CODE = {
    'Сосна': 'С',
    'Берёза': 'Б',
    'Осина': 'Ос',
    'Ель': 'Е',
    'Пихта': 'П',
    'Ольха': 'Ол',
    'Дуб': 'Д',
    'Ясень': 'Я',
    'Граб': 'Г',
    'Лиственница': 'Лц',
    'Липа': 'Лп',
    'Кедр': 'К',
    'Можжевельник': 'Мж',
    'Клён': 'Кл',
    'Рябина': 'Р'
}

class OurTemplateFiller:
    def __init__(self, data_file=None):
        self.document_path = 'reports/Шаблон проект_наш.docx'
        self.data_file = data_file
        self.address_data = {}
        self.total_data = {}
        self.details_data = {}
        self.breeds_data = []

    def load_data_from_json(self, file_path):
        """Загружаем данные из JSON файла"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            self.address_data = data.get('address_data', {})
            self.total_data = data.get('total_data', {})
            self.breeds_data = self.total_data.get('breeds', [])

            # Извлекаем детали из total_data
            self.details_data = {
                'care_queue': self.total_data.get('care_queue', ''),
                'characteristics': self.total_data.get('characteristics', ''),
                'care_date': self.total_data.get('care_date', ''),
                'technology': self.total_data.get('technology', ''),
                'forest_purpose': self.total_data.get('forest_purpose', ''),
                'care_subject': self.total_data.get('care_subject', ''),
                'forest_type': self.address_data.get('forest_type', '')
            }

            print(f"[OK] Данные загружены")
            print(f"  Пород: {len(self.breeds_data)}")
            for breed in self.breeds_data:
                print(f"    - {breed.get('name', 'Н/Д')}: d={breed.get('diameter', 0):.1f}, h={breed.get('height', 0):.1f}, density={breed.get('density', 0):.1f}")
            print(f"  Коэффициент состава: {self.total_data.get('composition', 'Н/Д')}")
            print(f"  Интенсивность: {self.total_data.get('intensity', 'Н/Д')}")
            print(f"  Предмет ухода: {self.details_data.get('care_subject', 'Н/Д')}")
            print(f"  Тип леса: {self.details_data.get('forest_type', 'Н/Д')}")
            print(f"  Характеристики: {self.details_data.get('characteristics', 'Н/Д')}")
            return True

        except Exception as e:
            print(f"[ERROR] Ошибка загрузки данных: {e}")
            import traceback
            traceback.print_exc()
            return False

    def format_number(self, value, default=''):
        """Форматирует число с одной десятичной точкой"""
        if value is None or value == 'Н/Д' or value == 0:
            return default
        try:
            num = float(value)
            return f"{num:.1f}"
        except (ValueError, TypeError):
            return str(value)

    def inflect_forestry(self, name):
        """Склоняет название лесничества (упрощённо)"""
        if not name:
            return ''

        name = name.strip()

        # Простые правила склонения
        if name.endswith('ое'):
            return name[:-2] + 'ом'  # Красное -> Красном
        elif name.endswith('ее'):
            return name[:-2] + 'ем'  # Сегежское -> Сегежском
        elif name.endswith('ий'):
            return name[:-2] + 'ем'  # Лесной -> Лесном
        elif name.endswith('ый'):
            return name[:-2] + 'ом'  # Борвый -> Борвом
        elif name.endswith('о'):
            return name + 'м'  # Волом -> Воломом
        else:
            return name + 'е'  # По умолчанию

    def parse_care_subject_density(self, care_text):
        """
        Парсит предмет ухода и возвращает густоту по породам {порода: густоту тыс. шт/га}
        Примеры: "3000шт/гаС", "3С2Б", "300шт/гаС + 50шт/гаБ"
        Возвращает: {порода: густота в тыс. шт/га}
        """
        if not care_text:
            return {}

        result = {}
        care_text = care_text.strip()

        # Паттерн 1: "3000шт/гаС" или "3000 шт/га С"
        # Ищем числа с возможными "шт/га" и буквой породы
        matches = re.findall(r'(\d+(?:\.\d+)?)\s*(?:шт/га)?\s*([А-ЯA-Z][а-яa-z]*)', care_text)

        for density_str, breed_str in matches:
            density_val = float(density_str)
            
            # Если число большое (>100), значит это шт/га, делим на 1000 для тыс.шт/га
            if density_val > 100:
                density = density_val / 1000  # Переводим в тыс. шт/га
            else:
                # Это коэффициент состава (1-10)
                density = density_val

            # Определяем породу по коду (первая заглавная буква)
            breed_code = breed_str[0].upper() if breed_str else ''

            # Ищем полное название породы по коду
            breed_name = BREED_DATABASE.get(breed_code, None)

            # Если не нашли по коду, ищем по названию
            if not breed_name:
                for code, name in BREED_DATABASE.items():
                    if code.lower() in breed_str.lower():
                        breed_name = name
                        break

            if breed_name:
                result[breed_name] = density

        # Паттерн 2: "3С2Б" (коэффициенты состава)
        if not result:
            matches = re.findall(r'(\d+(?:\.\d+)?)([А-ЯA-Z][а-яa-z]*)', care_text)
            for coeff_str, breed_str in matches:
                coeff = float(coeff_str)
                # Коэффициент состава переводим в густоту (примерно)
                density = coeff  # тыс. шт/га

                breed_code = breed_str[0].upper() if breed_str else ''
                breed_name = BREED_DATABASE.get(breed_code, None)

                if not breed_name:
                    for code, name in BREED_DATABASE.items():
                        if code.lower() in breed_str.lower():
                            breed_name = name
                            break

                if breed_name:
                    result[breed_name] = density

        print(f"[INFO] Распарсен предмет ухода: {result}")
        return result

    def parse_characteristics(self):
        """
        Парсит характеристики деревьев из details_data.characteristics
        Формат: "Лучшие: Сосна\nВспомогательные: Берёза\nНежелательные: Осина"
        Возвращает: {'best': 'Сосна', 'auxiliary': 'Берёза', 'undesirable': 'Осина'}
        """
        characteristics_text = self.details_data.get('characteristics', '')
        
        result = {
            'best': '',
            'auxiliary': '',
            'undesirable': ''
        }
        
        if not characteristics_text:
            # Значения по умолчанию
            result['best'] = 'здоровая, хорошо укорененная сосна, с хорошо сформированной кроной'
            result['auxiliary'] = 'деревья всех пород обеспечивающие сохранение целостности и устойчивости насаждения'
            result['undesirable'] = 'деревья мешающие росту и формированию крон отобранных лучших и вспомогательных деревьев; деревья неудовлетворительного состояния'
            return result
        
        # Парсим по строкам
        lines = characteristics_text.strip().split('\n')
        for line in lines:
            line = line.strip()
            if line.lower().startswith('лучшие:'):
                result['best'] = line.split(':', 1)[1].strip()
            elif line.lower().startswith('вспомогательные:'):
                result['auxiliary'] = line.split(':', 1)[1].strip()
            elif line.lower().startswith('нежелательные:'):
                result['undesirable'] = line.split(':', 1)[1].strip()
        
        # Если не распарсилось, используем значения по умолчанию
        if not result['best']:
            result['best'] = 'здоровая, хорошо укорененная сосна, с хорошо сформированной кроной'
        if not result['auxiliary']:
            result['auxiliary'] = 'деревья всех пород обеспечивающие сохранение целостности и устойчивости насаждения'
        if not result['undesirable']:
            result['undesirable'] = 'деревья мешающие росту и формированию крон отобранных лучших и вспомогательных деревьев; деревья неудовлетворительного состояния'
        
        print(f"[INFO] Распарсены характеристики: {result}")
        return result

    def calculate_project_values(self, breed, intensity):
        """Рассчитывает проектируемые значения (после рубки)"""
        # Исходные данные
        age = breed.get('age', 0)
        height = breed.get('height', 0)
        diameter = breed.get('diameter', 0)

        print(f"[DEBUG] breed={breed.get('name', 'Н/Д')}, d={diameter}, h={height}, intensity={intensity}")

        # После рубки остаются лучшие деревья, поэтому:
        # - высота увеличивается (мелочь убирается)
        # - диаметр увеличивается (мелочь убирается)
        # - густота уменьшается согласно интенсивности

        # Коэффициент увеличения для высоты и диаметра (зависит от интенсивности)
        # Чем выше интенсивность, тем больше остаются крупные деревья
        growth_factor = 1 + (intensity / 100) * 0.3  # Увеличение на 30% при 100% интенсивности

        project_height = height * growth_factor if height > 0 else 0
        project_diameter = diameter * growth_factor if diameter > 0 else 0

        print(f"[DEBUG] project: d={project_diameter:.2f}, h={project_height:.2f}")

        return {
            'age': age,  # Возраст не меняется
            'height': project_height,
            'diameter': project_diameter,
            'density': None  # Будет заполнено из предмета ухода
        }

    def calculate_breed_composition(self, breed_name, density, total_density):
        """Рассчитывает коэффициент состава для породы"""
        if total_density > 0:
            coeff = round((density / total_density) * 10)
            if coeff < 1:
                coeff = 1
        else:
            coeff = 1

        # Получаем букву породы из БД - ищем по полному названию или по части
        letter = 'Др'  # По умолчанию
        
        # Сначала пробуем точное совпадение
        if breed_name in BREED_TO_CODE:
            letter = BREED_TO_CODE[breed_name]
        else:
            # Ищем по первому слову или аббревиатуре
            breed_name_lower = breed_name.lower()
            for full_name, code in BREED_TO_CODE.items():
                if full_name.lower() in breed_name_lower or breed_name_lower.startswith(full_name.lower()[:3]):
                    letter = code
                    break
            
            # Если не нашли, берём первую букву
            if letter == 'Др':
                letter = breed_name[0].upper()
                # Особые случаи
                if breed_name.startswith('Ос'):
                    letter = 'Ос'
                elif breed_name.startswith('Ол'):
                    letter = 'Ол'
                elif breed_name.startswith('Мож'):
                    letter = 'Мж'

        return f"{coeff}{letter}"

    def fill_document(self):
        """Заполняем документ данными"""
        if not os.path.exists(self.document_path):
            print(f"[ERROR] Файл {self.document_path} не найден!")
            return False

        try:
            doc = Document(self.document_path)

            # Получаем характеристики деревьев (парсим из details_data)
            characteristics = self.parse_characteristics()

            # Рассчитываем параметры площадок
            current_radius = float(self.address_data.get('radius', 1.78))
            plot_area_m2 = 3.14159 * current_radius ** 2
            total_plots = self.total_data.get('total_plots', 0)

            # Получаем ИНТЕНСИВНОСТЬ ИЗ МЕНЮ ИТОГО (рассчитанную)
            intensity = self.total_data.get('intensity', 25)
            print(f"[INFO] Интенсивность рубки из меню Итого: {intensity}%")

            # Формируем вид рубки
            activity_name = self.total_data.get('activity_name', 'осветление')
            care_queue = self.details_data.get('care_queue', 'первая')

            # ИСПРАВЛЕНИЕ 1: Не дублируем "- РУМ", если care_queue уже содержит
            care_queue_clean = care_queue.replace('- РУМ', '').replace('(осветление)', '').strip()
            # Удаляем лишние дефисы и пробелы
            care_queue_clean = re.sub(r'\s*,\s*', ', ', care_queue_clean)
            care_queue_clean = re.sub(r'\s+', ' ', care_queue_clean).strip()
            
            # Формируем правильный текст
            care_activity_text = f"{activity_name}, {care_queue_clean}"

            # Склоняем названия лесничеств
            forestry = self.address_data.get('forestry', '')
            district_forestry = self.address_data.get('district_forestry', '')
            forestry_inflected = self.inflect_forestry(forestry)
            district_forestry_inflected = self.inflect_forestry(district_forestry)

            # ИСПРАВЛЕНИЕ 2: forest_type берём из address_data (из участка)
            forest_type = self.address_data.get('forest_type', '')
            if not forest_type:
                forest_type = self.details_data.get('forest_type', 'Смешанный лес')

            # Парсим предмет ухода для густоты
            care_subject = self.details_data.get('care_subject', '')
            care_density_by_breed = self.parse_care_subject_density(care_subject)
            
            # Общая густота из предмета ухода (в тыс. шт/га)
            total_care_density = sum(care_density_by_breed.values())

            # Словарь общих замен
            replacements = {
                # Адресные данные (со склонением)
                '{forestry}': forestry_inflected,
                '{district_forestry}': district_forestry_inflected,
                '{quarter}': self.address_data.get('quarter', ''),
                '{plot}': self.address_data.get('plot', ''),
                '{plot_area}': self.address_data.get('plot_area', ''),

                # Детали ухода
                '{care_activity_text}': care_activity_text,
                '{care_queue}': care_queue,
                '{care_date}': self.details_data.get('care_date', ''),
                '{technology}': self.details_data.get('technology', ''),
                '{forest_purpose}': self.details_data.get('forest_purpose', ''),
                '{best}': characteristics.get('best', ''),
                '{auxiliary}': characteristics.get('auxiliary', ''),
                '{undesirable}': characteristics.get('undesirable', ''),
                '{care_subject}': care_subject,

                # Тип леса из участка
                '{forest_type}': forest_type,

                # ОБЩИЕ ДАННЫЕ ИЗ МЕНЮ ИТОГО
                '{total_composition_isx}': self.total_data.get('composition', ''),
                '{total_composition_project}': self.total_data.get('composition', ''),
                '{total_age_isx}': self.format_number(self.total_data.get('avg_age')),
                '{total_age_project}': self.format_number(self.total_data.get('avg_age')),
                '{total_height_isx}': self.format_number(self.total_data.get('avg_height')),
                '{total_height_project}': self.format_number(
                    self.total_data.get('avg_height', 0) * (1 + intensity/100 * 0.3) if self.total_data.get('avg_height') else 0
                ),
                '{total_diameter_isx}': self.format_number(self.total_data.get('avg_diameter', 0)),
                '{total_diameter_project}': self.format_number(
                    self.total_data.get('avg_diameter', 0) * (1 + intensity/100 * 0.3) if self.total_data.get('avg_diameter') else 0
                ),
                '{total_density_isx}': self.format_number(self.total_data.get('avg_density')),
                # ИСПРАВЛЕНИЕ 4: Густота проектируемая из предмета ухода (в тыс. шт/га)
                '{total_density_project}': self.format_number(total_care_density) if total_care_density > 0 else '',
                '{intensity}': f"{intensity:.1f}%",

                # Другие
                '{radius_info}': f"{total_plots} шт. {plot_area_m2:.0f}м²(R-{current_radius:.2f}м)",
            }

            # Заполняем параграфы общими данными
            for paragraph in doc.paragraphs:
                for old_text, new_text in replacements.items():
                    if old_text in paragraph.text:
                        paragraph.text = paragraph.text.replace(old_text, str(new_text))

            # Заполняем таблицы
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for old_text, new_text in replacements.items():
                                if old_text in paragraph.text:
                                    paragraph.text = paragraph.text.replace(old_text, str(new_text))

            # === ТЕПЕРЬ ЗАПОЛНЯЕМ ДАННЫМИ ПО ПОРОДАМ ===
            if self.breeds_data:
                # Находим таблицу с породами
                breeds_table = None
                for table in doc.tables:
                    if len(table.columns) == 11 and 'Порода' in table.rows[0].cells[0].text:
                        breeds_table = table
                        break

                if breeds_table:
                    # Удаляем строку-шаблон (первую после заголовка)
                    if len(breeds_table.rows) > 1:
                        breeds_table._tbl.remove(breeds_table.rows[1]._tr)

                    # Рассчитываем общую густоту для коэффициента состава
                    total_density = sum(b.get('density', 0) for b in self.breeds_data)

                    # Добавляем строки для каждой породы
                    for breed in self.breeds_data:
                        breed_name = breed.get('name', '')
                        density = breed.get('density', 0)
                        age = breed.get('age', 0)
                        height = breed.get('height', 0)
                        diameter = breed.get('diameter', 0)

                        print(f"[DEBUG] Заполняем породу: {breed_name}, d={diameter}, h={height}, density={density}")

                        # ИСПРАВЛЕНИЕ 6: Рассчитываем состав правильно
                        composition_isx = self.calculate_breed_composition(breed_name, density, total_density)
                        composition_project = composition_isx

                        # Рассчитываем проектируемые значения
                        project_values = self.calculate_project_values(breed, intensity)

                        # ИСПРАВЛЕНИЕ 4: Густота проектируемая из предмета ухода
                        # Ищем густоту по частичному совпадению названия породы
                        project_density = None
                        for care_breed, care_dens in care_density_by_breed.items():
                            # Если название породы содержит код из предмета ухода
                            if care_breed in breed_name or breed_name.startswith(care_breed):
                                project_density = care_dens
                                break
                            # Или если код породы совпадает с первым словом
                            breed_name_first_word = breed_name.split()[0].lower() if ' ' in breed_name else breed_name.lower()
                            if care_breed.lower() == breed_name_first_word:
                                project_density = care_dens
                                break
                        
                        project_density_str = self.format_number(project_density) if project_density is not None else ''

                        # Добавляем новую строку
                        row = breeds_table.add_row()
                        row.cells[0].text = breed_name
                        row.cells[1].text = composition_isx  # Состав исх
                        row.cells[2].text = composition_project  # Состав проект
                        row.cells[3].text = self.format_number(age)  # Возраст исх
                        row.cells[4].text = self.format_number(age)  # Возраст проект
                        row.cells[5].text = self.format_number(diameter)  # ИСХОДНЫЙ диаметр
                        row.cells[6].text = self.format_number(project_values['diameter'])  # ПРОЕКТНЫЙ диаметр
                        row.cells[7].text = self.format_number(height)  # ИСХОДНАЯ высота
                        row.cells[8].text = self.format_number(project_values['height'])  # ПРОЕКТНАЯ высота
                        row.cells[9].text = self.format_number(density)  # ИСХОДНАЯ густота
                        row.cells[10].text = project_density_str  # ПРОЕКТНАЯ густота из предмета ухода

                        print(f"[DEBUG] Заполнено: d_isx={diameter}, d_prj={project_values['diameter']:.2f}, density_prj={project_density_str}")

            # Сохраняем документ
            timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = self.document_path.replace('.docx', f'_заполненный_{timestamp}.docx')
            doc.save(output_path)

            print(f"[OK] Документ заполнен и сохранён: {output_path}")
            return True

        except Exception as e:
            print(f"[ERROR] Ошибка при заполнении документа: {e}")
            import traceback
            traceback.print_exc()
            return False

    def run(self):
        """Основной метод выполнения"""
        print("=" * 60)
        print("ЗАПОЛНЕНИЕ НАШЕГО ШАБЛОНА ПРОЕКТА УХОДА")
        print("=" * 60)

        if self.data_file:
            if not self.load_data_from_json(self.data_file):
                return False
        else:
            print("[ERROR] Не указан файл данных!")
            return False

        success = self.fill_document()

        if success:
            print("=" * 60)
            print("[OK] Задача выполнена успешно!")
            print("=" * 60)
        else:
            print("=" * 60)
            print("[ERROR] Произошла ошибка при выполнении задачи")
            print("=" * 60)

        return success

def fill_document_from_json(json_file_path):
    """Функция для заполнения документа из JSON"""
    try:
        filler = OurTemplateFiller(data_file=json_file_path)
        return filler.run()
    except Exception as e:
        print(f"[ERROR] Ошибка при заполнении документа: {e}")
        return False

if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(description='Заполнение нашего шаблона проекта ухода')
    parser.add_argument('--data-file', type=str, help='Путь к JSON файлу с данными')
    args = parser.parse_args()

    if args.data_file:
        fill_document_from_json(args.data_file)
    else:
        print("[ERROR] Не указан файл данных! Используйте --data-file <путь>")
