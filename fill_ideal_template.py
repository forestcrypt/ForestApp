#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для заполнения идеального шаблона проекта ухода
"""

import os
import sys
import json
import argparse
from docx import Document
import sqlite3
import datetime

class IdealWordFiller:
    def __init__(self, db_name='forest_data.db', data_file=None, address_data=None, total_data=None):
        self.db_name = db_name
        self.document_path = 'reports/Шаблон проект_идеальный.docx'
        self.data_file = data_file
        self.address_data = address_data or {}
        self.total_data = total_data or {}

    def load_data_from_json(self, file_path):
        """Загружаем данные из JSON файла"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            self.address_data = data.get('address_data', self.address_data)
            self.total_data = data.get('total_data', self.total_data)

            # Добавляем данные деталей ухода
            for key in ['care_queue', 'characteristics', 'care_date', 'technology', 'forest_purpose']:
                if key in self.total_data:
                    setattr(self, key, self.total_data[key])

            print(f"[OK] Данные загружены из JSON")
            print(f"  Квартал: {self.address_data.get('quarter', 'Н/Д')}")
            print(f"  Выдел: {self.address_data.get('plot', 'Н/Д')}")
            print(f"  Очередь рубки: {self.total_data.get('care_queue', 'Н/Д')}")
            print(f"  Коэффициент состава: {self.total_data.get('composition', 'Н/Д')}")
            return True

        except Exception as e:
            print(f"[ERROR] Ошибка загрузки данных: {e}")
            return False

    def format_number(self, value, default='Н/Д'):
        """Форматирует число с одной десятичной точкой"""
        if value is None or value == 'Н/Д':
            return default
        try:
            num = float(value)
            return f"{num:.1f}"
        except (ValueError, TypeError):
            return str(value)

    def get_characteristics(self):
        """Получает характеристики деревьев"""
        characteristics = self.total_data.get('characteristics', '')
        
        if isinstance(characteristics, dict):
            return {
                'best': characteristics.get('best', 'здоровая, хорошо укорененная сосна, с хорошо сформированной кроной'),
                'auxiliary': characteristics.get('auxiliary', 'деревья всех пород обеспечивающие сохранение целостности и устойчивости насаждения'),
                'undesirable': characteristics.get('undesirable', 'деревья мешающие росту и формированию крон отобранных лучших и вспомогательных деревьев; деревья неудовлетворительного состояния')
            }
        else:
            return {
                'best': 'здоровая, хорошо укорененная сосна, с хорошо сформированной кроной',
                'auxiliary': 'деревья всех пород обеспечивающие сохранение целостности и устойчивости насаждения',
                'undesirable': 'деревья мешающие росту и формированию крон отобранных лучших и вспомогательных деревьев; деревья неудовлетворительного состояния'
            }

    def fill_document(self):
        """Заполняем документ данными"""
        if not os.path.exists(self.document_path):
            print(f"[ERROR] Файл {self.document_path} не найден!")
            return False

        try:
            doc = Document(self.document_path)
            
            # Получаем данные
            care_queue = self.total_data.get('care_queue', 'первая')
            care_date = self.total_data.get('care_date', 'сент 2025 года')
            technology = self.total_data.get('technology', '')
            forest_purpose = self.total_data.get('forest_purpose', 'Эксплуатационные леса')
            characteristics = self.get_characteristics()
            
            # Формируем мероприятие с очередью
            activity_name = self.total_data.get('activity_name', 'осветление')
            care_activity_text = f"{activity_name}, {care_queue} очередь"
            
            # Рассчитываем параметры площадок
            current_radius = float(self.address_data.get('radius', 1.78))
            plot_area_m2 = 3.14159 * current_radius ** 2
            
            # Словарь замен
            replacements = {
                # Адресные данные
                '(Адрес-лесничество)': self.address_data.get('forestry', ''),
                '(Адрес-участковое лесничество)': self.address_data.get('district_forestry', ''),
                '(Адрес-квартал)': self.address_data.get('quarter', ''),
                '(Адрес-выдел)': self.address_data.get('plot', ''),
                '(Адрес-площадь)': self.address_data.get('plot_area', ''),
                
                # Детали ухода
                '(Детали-очередь рубки-тип мероприятий)': care_activity_text,
                '(Детали-назначение лесов)': forest_purpose,
                '(Детали-дата рубки)': care_date,
                '(Детали-технология ухода)': technology,
                '(Детали-характеристика-лучшие)': characteristics['best'],
                '(Детали-характеристика-вспомогательные)': characteristics['auxiliary'],
                '(Детали-характеристика-нежелательные)': characteristics['undesirable'],
                
                # Данные из Итого
                '(Итого-состав исх)': self.total_data.get('composition', 'Н/Д'),
                '(Итого-состав проект)': self.total_data.get('composition', 'Н/Д'),
                '(Итого-возраст)': self.format_number(self.total_data.get('avg_age')),
                '(Итого-диаметр)': self.format_number(self.total_data.get('avg_diameter')),
                '(Итого-высота)': self.format_number(self.total_data.get('avg_height')),
                '(Итого-густота)': self.format_number(self.total_data.get('avg_density')),
                '(Итого-интенсивность)': self.total_data.get('intensity', '25%'),
                
                # Другие данные
                '(Функции-очередь рубки)': care_queue,
                '(Итого и Адрес-радиус)': f"{self.total_data.get('total_plots', 0)} шт. {plot_area_m2:.0f}м²(R-{current_radius:.2f}м)",
                '(Данных по площадкам-тип леса)': self.address_data.get('forest_type', 'Смешанный лес'),
            }
            
            # Заполняем параграфы
            for paragraph in doc.paragraphs:
                for old_text, new_text in replacements.items():
                    if old_text in paragraph.text:
                        paragraph.text = paragraph.text.replace(old_text, str(new_text))
            
            # Заполняем таблицы
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for old_text, new_text in replacements.items():
                                if old_text in paragraph.text:
                                    paragraph.text = paragraph.text.replace(old_text, str(new_text))
            
            # Сохраняем документ
            timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = self.document_path.replace('.docx', f'_заполненный_{timestamp}.docx')
            doc.save(output_path)
            
            print(f"[OK] Документ заполнен и сохранён: {output_path}")
            return True

        except Exception as e:
            print(f"[ERROR] Ошибка при заполнении документа: {e}")
            import traceback
            traceback.print_exc()
            return False

    def run(self):
        """Основной метод выполнения"""
        print("=" * 60)
        print("ЗАПОЛНЕНИЕ ИДЕАЛЬНОГО ШАБЛОНА ПРОЕКТА УХОДА")
        print("=" * 60)
        
        if self.data_file:
            if not self.load_data_from_json(self.data_file):
                return False
        else:
            print("[ERROR] Не указан файл данных!")
            return False
        
        success = self.fill_document()
        
        if success:
            print("=" * 60)
            print("[OK] Задача выполнена успешно!")
            print("=" * 60)
        else:
            print("=" * 60)
            print("[ERROR] Произошла ошибка при выполнении задачи")
            print("=" * 60)
        
        return success

def fill_document_from_json(json_file_path):
    """Функция для заполнения документа из JSON"""
    try:
        filler = IdealWordFiller(data_file=json_file_path)
        return filler.run()
    except Exception as e:
        print(f"[ERROR] Ошибка при заполнении документа: {e}")
        return False

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Заполнение идеального шаблона проекта ухода')
    parser.add_argument('--data-file', type=str, help='Путь к JSON файлу с данными')
    args = parser.parse_args()
    
    if args.data_file:
        fill_document_from_json(args.data_file)
    else:
        print("[ERROR] Не указан файл данных! Используйте --data-file <путь>")
